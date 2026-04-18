"""
=============================================================================
FAA RF Interference Analysis Tool
For ITU-R Policy Support — Aeronautical Spectrum Protection
=============================================================================
Purpose : Help RF engineers supporting FAA/NTIA policy assess interference
          threats to protected aeronautical frequencies before and during
          ITU-R WP 5D / 5B proceedings.
Libraries: itur (ITU-R propagation models), numpy, pandas, matplotlib
=============================================================================
"""

import streamlit as st
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import itur
import itur.models.itu676 as itu676
import warnings
import anthropic
import math
import io
from auth import (show_login_page, show_admin_panel, is_authenticated,
                  is_admin, current_user, logout)
warnings.filterwarnings("ignore")

# ─────────────────────────────────────────────────────────────────────────────
# PAGE CONFIG
# ─────────────────────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="FAA RF Interference Analysis Tool",
    page_icon="✈️",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ─────────────────────────────────────────────────────────────────────────────
# AUTH GATE — must pass before anything else renders
# ─────────────────────────────────────────────────────────────────────────────
if not is_authenticated():
    show_login_page()
    st.stop()


# ─────────────────────────────────────────────────────────────────────────────
# GLOBAL STYLES
# ─────────────────────────────────────────────────────────────────────────────
st.markdown("""
<style>
.explainer {
    background: #1a2a3a;
    border-left: 4px solid #00aaff;
    padding: 6px 12px;
    border-radius: 4px;
    color: #aaddff;
    font-size: 0.85em;
    margin-bottom: 10px;
}
.warn-box {
    background: #3a1a1a;
    border-left: 4px solid #ff4444;
    padding: 6px 12px;
    border-radius: 4px;
    color: #ffaaaa;
    font-size: 0.85em;
    margin-bottom: 10px;
}
.ok-box {
    background: #1a3a1a;
    border-left: 4px solid #44ff44;
    padding: 6px 12px;
    border-radius: 4px;
    color: #aaffaa;
    font-size: 0.85em;
    margin-bottom: 10px;
}
.metric-label { font-size: 0.75em; color: #888; }
</style>
""", unsafe_allow_html=True)

def ex(text):
    """Render a one-line blue explainer box."""
    st.markdown(f'<div class="explainer">💡 {text}</div>', unsafe_allow_html=True)

def warn(text):
    st.markdown(f'<div class="warn-box">⚠️ {text}</div>', unsafe_allow_html=True)

def ok(text):
    st.markdown(f'<div class="ok-box">✅ {text}</div>', unsafe_allow_html=True)

# ─────────────────────────────────────────────────────────────────────────────
# FAA PROTECTED BAND DATABASE
# ─────────────────────────────────────────────────────────────────────────────
FAA_BANDS = {
    "VOR / ILS Localizer": {
        "f_low_mhz": 108.0, "f_high_mhz": 117.975,
        "system": "VOR / ILS Localizer", "allocation": "ARNS",
        "service_category": "AM(R)S + ARNS",
        "in_threshold_db": -6, "aviation_safety_factor_db": 6, "effective_threshold_db": -12,
        "noise_floor_dbm": -120, "safety_of_life": True, "rr_1_59": True,
        "protection_basis": "I/N ≤ −6 dB + 6 dB aviation safety factor for precision approach",
        "spr_source": "Max EIRP toward airport; worst-case azimuth",
        "spr_path": "FSPL; distances >20 km; free-space worst-case",
        "spr_victim": "Receiver susceptibility mask per DO-196; −120 dBm noise floor",
        "notes": "En-route navigation and precision approach guidance. ILS CAT III approach — 6 dB additional safety factor applies.",
        "rtca_doc": "DO-196",
    },
    "ILS Glide Slope": {
        "f_low_mhz": 328.6, "f_high_mhz": 335.4,
        "system": "ILS Glide Slope", "allocation": "ARNS",
        "service_category": "ARNS",
        "in_threshold_db": -6, "aviation_safety_factor_db": 6, "effective_threshold_db": -12,
        "noise_floor_dbm": -115, "safety_of_life": True, "rr_1_59": True,
        "protection_basis": "I/N ≤ −6 dB + 6 dB aviation safety factor; CAT III precision approach",
        "spr_source": "Max EIRP interferer; worst-case signal characteristics",
        "spr_path": "FSPL worst-case; free-space attenuation above 300 MHz",
        "spr_victim": "Receiver susceptibility per DO-148; noise floor −115 dBm",
        "notes": "Vertical guidance for precision landings (CAT I/II/III). Safety factor mandatory for CAT III.",
        "rtca_doc": "DO-148",
    },
    "DME / TACAN": {
        "f_low_mhz": 960.0, "f_high_mhz": 1215.0,
        "system": "DME / TACAN / SSR / TCAS", "allocation": "ARNS + ANS",
        "service_category": "ARNS + ANS",
        "in_threshold_db": -6, "aviation_safety_factor_db": 0, "effective_threshold_db": -6,
        "epfd_threshold_dbw_m2_mhz": -121.5,
        "noise_floor_dbm": -106, "safety_of_life": True, "rr_1_59": True,
        "protection_basis": "I/N ≤ −6 dB; epfd ≤ −121.5 dBW/m² in any 1 MHz band",
        "spr_source": "Max Tx power; worst-case antenna gain toward aircraft",
        "spr_path": "FSPL; distance separation >20 km worst-case",
        "spr_victim": "Receiver susceptibility mask; antenna gain; noise power per DO-189",
        "notes": "Distance measuring, ATC surveillance, collision avoidance. epfd limit applies to satellite downlinks.",
        "rtca_doc": "DO-189 / DO-185B",
    },
    "ADS-B / Mode-S (1090 MHz)": {
        "f_low_mhz": 1085.0, "f_high_mhz": 1095.0,
        "system": "ADS-B / Mode-S Transponder", "allocation": "ARNS + ANS",
        "service_category": "ANS (Aeronautical Navigation Service)",
        "in_threshold_db": -10, "aviation_safety_factor_db": 0, "effective_threshold_db": -10,
        "noise_floor_dbm": -100, "safety_of_life": True, "rr_1_59": True,
        "protection_basis": "I/N ≤ −10 dB; ASR protection level per system protection table",
        "spr_source": "Max EIRP; worst-case signal characteristics",
        "spr_path": "FSPL worst-case; distances >20 km",
        "spr_victim": "Receiver susceptibility per DO-260B; noise floor −100 dBm",
        "notes": "1090 MHz squitter — global ATC surveillance backbone. ASR threshold −10 dB applies.",
        "rtca_doc": "DO-260B",
    },
    "GNSS L5 / ARNS": {
        "f_low_mhz": 1164.0, "f_high_mhz": 1215.0,
        "system": "GNSS L5 / Galileo E5", "allocation": "ARNS + RNSS",
        "service_category": "RNSS + ARNS",
        "in_threshold_db": -10, "aviation_safety_factor_db": 6, "effective_threshold_db": -16,
        "delta_t_t_pct_aggregate": 6.0,
        "noise_floor_dbm": -130, "safety_of_life": True, "rr_1_59": True,
        "protection_basis": "I/N ≤ −10 dB wideband; ΔT/T ≤ 6% single-entry; +6 dB aviation safety factor",
        "spr_source": "Max EIRP; worst-case signal characteristics",
        "spr_path": "FSPL worst-case; distances >20 km; aggregate sources",
        "spr_victim": "Receiver susceptibility mask per DO-292; noise floor −130 dBm",
        "notes": "Safety-of-life GNSS signal; aviation approach procedures. ΔT/T = 6% for RNSS feeder links.",
        "rtca_doc": "DO-292",
    },
    "GPS L1 / GNSS": {
        "f_low_mhz": 1559.0, "f_high_mhz": 1610.0,
        "system": "GPS L1 / GLONASS / Galileo E1", "allocation": "RNSS + ARNS",
        "service_category": "RNSS + ARNS",
        "in_threshold_db": -10, "aviation_safety_factor_db": 6, "effective_threshold_db": -16,
        "psd_threshold_dbw_mhz": -146.5,
        "noise_floor_dbm": -130, "safety_of_life": True, "rr_1_59": True,
        "protection_basis": "I < −146.5 dBW/MHz (L1 SBAS Type 1); I/N ≈ −5 dB + 6 dB safety margin (wideband RFI)",
        "spr_source": "Max EIRP; worst-case antenna gain; signal characteristics",
        "spr_path": "FSPL worst-case; aggregate effect of multiple sources",
        "spr_victim": "Receiver susceptibility per DO-235B/DO-253; noise floor −130 dBm",
        "notes": "Primary GNSS band. SBAS/WAAS critical. PSD limit −146.5 dBW/MHz for wideband RFI.",
        "rtca_doc": "DO-235B / DO-253",
    },
    "En-Route Radar": {
        "f_low_mhz": 2700.0, "f_high_mhz": 2900.0,
        "system": "ATC En-Route Surveillance Radar (ARSR / ASR)", "allocation": "ARNS + RN",
        "service_category": "ARNS (Safety Service per RR 1.59)",
        "in_threshold_db": -6, "aviation_safety_factor_db": 0, "effective_threshold_db": -6,
        "noise_floor_dbm": -100, "safety_of_life": True, "rr_1_59": True,
        "protection_basis": "ARSR: I/N ≤ −6 dB; ASR: I/N ≤ −10 dB (per system protection levels table)",
        "spr_source": "Max EIRP; worst-case azimuth toward radar",
        "spr_path": "FSPL; worst-case for distances >20 km",
        "spr_victim": "Radar receiver susceptibility; noise power; antenna gain",
        "notes": "ARSR (long-range) I/N = −6 dB; ASR (short-range, airport) I/N = −10 dB.",
        "rtca_doc": "N/A (ITU-R M.1849)",
    },
    "Radio Altimeter": {
        "f_low_mhz": 4200.0, "f_high_mhz": 4400.0,
        "system": "Radio Altimeter (Rad Alt)", "allocation": "ARNS",
        "service_category": "ARNS (Safety Service per RR 1.59)",
        "in_threshold_db": -6, "aviation_safety_factor_db": 6, "effective_threshold_db": -12,
        "noise_floor_dbm": -90, "safety_of_life": True, "rr_1_59": True,
        "protection_basis": "I/N ≤ −6 dB + 6 dB aviation safety factor for CAT III precision approach",
        "spr_source": "Max EIRP interferer (e.g. 5G base station); OOB/spurious emissions",
        "spr_path": "FSPL + possible blocking; distances >20 km worst-case; aggregate effect",
        "spr_victim": "LNA susceptibility mask; blocking threshold; noise floor −90 dBm",
        "notes": "Critical for CAT III landings, TAWS, GPWS, helicopter ops.",
        "rtca_doc": "DO-155 / ETSO-C87",
    },
    "ARNS 5 GHz": {
        "f_low_mhz": 5000.0, "f_high_mhz": 5150.0,
        "system": "ARNS / Future Aeronautical Systems", "allocation": "ARNS",
        "service_category": "ARNS",
        "in_threshold_db": -6, "aviation_safety_factor_db": 6, "effective_threshold_db": -12,
        "noise_floor_dbm": -95, "safety_of_life": True, "rr_1_59": True,
        "protection_basis": "I/N ≤ −6 dB + 6 dB aviation safety factor; under pressure from IMT-2030 WP 5D studies",
        "spr_source": "Potential IMT base station EIRP; OOB emissions from adjacent IMT bands",
        "spr_path": "FSPL worst-case; P.528 for airborne victim",
        "spr_victim": "Future ARNS receiver; noise floor −95 dBm",
        "notes": "Protected for future aeronautical use; under pressure from IMT. WRC-27 WP 5D AI 1.2 threat.",
        "rtca_doc": "N/A",
    },
    "Airborne Weather Radar": {
        "f_low_mhz": 9000.0, "f_high_mhz": 9500.0,
        "system": "Airborne / Surface Movement Radar", "allocation": "ARNS + RN",
        "service_category": "ARNS (Safety Service per RR 1.59)",
        "in_threshold_db": -6, "aviation_safety_factor_db": 0, "effective_threshold_db": -6,
        "noise_floor_dbm": -95, "safety_of_life": True, "rr_1_59": True,
        "protection_basis": "I/N ≤ −6 dB; safety service per RR 1.59",
        "spr_source": "Max EIRP; worst-case signal characteristics",
        "spr_path": "FSPL worst-case; gaseous absorption significant at X-band",
        "spr_victim": "Radar receiver susceptibility; noise floor −95 dBm",
        "notes": "X-band weather radar and airport surface detection.",
        "rtca_doc": "DO-220",
    },
    "MLS (Microwave Landing System)": {
        "f_low_mhz": 5030.0, "f_high_mhz": 5091.0,
        "system": "Microwave Landing System (MLS)", "allocation": "ARNS",
        "service_category": "ARNS (Safety Service per RR 1.59)",
        "in_threshold_db": -6, "aviation_safety_factor_db": 6, "effective_threshold_db": -12,
        "pfd_threshold_dbw_m2_khz": -124.5,
        "noise_floor_dbm": -110, "safety_of_life": True, "rr_1_59": True,
        "protection_basis": "pfd ≤ −124.5 dBW/m² in 150 kHz band; I/N ≤ −6 dB + 6 dB safety factor",
        "spr_source": "Max EIRP; signal characteristics",
        "spr_path": "FSPL; worst-case approach geometry",
        "spr_victim": "MLS receiver susceptibility; noise floor −110 dBm",
        "notes": "Precision approach system. PFD limit −124.5 dBW/m² in 150 kHz band.",
        "rtca_doc": "N/A",
    },
    "L-band AMS(R)S": {
        "f_low_mhz": 1525.0, "f_high_mhz": 1559.0,
        "system": "L-band Aeronautical Mobile Satellite (Route) Service", "allocation": "AMS(R)S",
        "service_category": "AMS(R)S — Aeronautical Mobile Satellite (Route) Service",
        "in_threshold_db": -6, "aviation_safety_factor_db": 0, "effective_threshold_db": -6,
        "delta_t_t_pct_aggregate": 20.0, "delta_t_t_pct_single": 6.0,
        "noise_floor_dbm": -120, "safety_of_life": True, "rr_1_59": True,
        "protection_basis": "ΔT/T ≤ 20% aggregate, ΔT/T ≤ 6% single-entry (from FAA system protection table)",
        "spr_source": "Satellite downlink EIRP; terrestrial co-frequency EIRP",
        "spr_path": "Slant path (P.619); FSPL worst-case for terrestrial",
        "spr_victim": "Aircraft terminal; noise temperature; antenna gain",
        "notes": "INMARSAT/IRIDIUM safety comms. ΔT/T metric used (not I/N). 20% aggregate = 6% single-entry.",
        "rtca_doc": "N/A",
    },
}

# ─────────────────────────────────────────────────────────────────────────────
# WRC-27 AGENDA ITEMS DATABASE
# ─────────────────────────────────────────────────────────────────────────────
WRC27_AGENDA_ITEMS = {
    "AI 1.7":  {"ref":"AI 1.7","title":"IMT in 4.4–4.8 GHz, 7.125–8.4 GHz, 14.8–15.35 GHz","service":"IMT","working_party":"WP 5D","threat_level":"HIGH","faa_systems_at_risk":["Radio Altimeter/WAICS — 4.2–4.4 GHz","FAA fixed links — 7.125–8.4 GHz","FAA fixed links — 14.8–15.35 GHz"],"faa_bands_mhz":[(4200,4400),(7125,8400),(14800,15350)],"mechanism":"OOB emissions/blocking from IMT base stations into RA band","key_concern":"4.4–4.8 GHz IMT is 0 MHz from Radio Altimeter band","citations":["RR No. 4.10","ITU-R SM.1540","ITU-R SM.1541","RR Appendix 3","RTCA DO-155"],"us_position":"Oppose IMT identification without OOB compliance per SM.1541 and coordination zones","notes":"RA is safety-of-life, 6 dB aviation safety factor applies."},
    "AI 1.13": {"ref":"AI 1.13","title":"MSS 694–2700 MHz for DC-MSS-IMT space-to-Earth connectivity","service":"MSS+IMT","working_party":"WP 4C","threat_level":"HIGH","faa_systems_at_risk":["ARNS/AM(R)S/AMS(R)S — 960–1215 MHz","MSS SatCom DL — 1525–1559 MHz","ASR — 2700–2900 MHz"],"faa_bands_mhz":[(960,1215),(1525,1559),(2700,2900)],"candidate_bands_mhz":[(925,960),(1475,1518),(2620,2690)],"mechanism":"OOB/spurious from satellite downlinks; aggregate interference","key_concern":"Candidate bands immediately adjacent to DME, AMS(R)S, and ASR","citations":["RR No. 4.10","RR No. 5.444","ITU-R SM.2028","ITU-R M.1642","ITU-R SM.1540"],"us_position":"Require aggregate analysis per SM.2028 for all three candidate bands","notes":"960–1215 MHz contains DME — critical navigation."},
    "AI 1.15": {"ref":"AI 1.15","title":"SRS (space-to-space) for lunar surface communications","service":"SRS","working_party":"WP 7B","threat_level":"MEDIUM","faa_systems_at_risk":["ASR — 2700–2900 MHz","3600–4200 MHz","ARNS 5 GHz — 5350–5470 MHz","Fixed — 7190–7235 MHz","Fixed — 8450–8500 MHz"],"faa_bands_mhz":[(2700,2900),(3600,4200),(5350,5470),(7190,7235),(8450,8500)],"mechanism":"Novel use case — no established Earth-Moon methodology","key_concern":"No established ITU-R methodology for lunar SRS vs terrestrial ARNS","citations":["RR No. 4.10","ITU-R SM.2028","ITU-R P.528","RR Appendix 3"],"us_position":"Require methodology before allocation","notes":"Methodology gap is the primary FAA policy argument."},
    "AI 1.17": {"ref":"AI 1.17","title":"EESS passive space weather sensors","service":"EESS","working_party":"WP 7C","threat_level":"LOW-MEDIUM","faa_systems_at_risk":["HF comms — 2.1–29.89 MHz","ILS-related — 74.8–75.2 MHz"],"faa_bands_mhz":[(2100,29890),(74800,75200)],"mechanism":"Passive — allocation policy concern","key_concern":"Coordination obligations on FAA HF comms","citations":["RR No. 4.10","ICAO Annex 10"],"us_position":"Monitor; ensure no coordination burden on FAA","notes":"Low direct threat; procedural concern."},
    "AI 1.19": {"ref":"AI 1.19","title":"EESS (passive) in 4.2–4.4 GHz and 8.4–8.5 GHz","service":"EESS","working_party":"WP 7C","threat_level":"MEDIUM","faa_systems_at_risk":["Radio Altimeter — 4.2–4.4 GHz","Fixed — 8.4–8.5 GHz"],"faa_bands_mhz":[(4200,4400),(8400,8500)],"mechanism":"Allocation precedent — weakens exclusive ARNS status","key_concern":"EESS co-primary in RA band weakens AI 1.7 defense","citations":["RR No. 4.10","RR No. 1.59","RTCA DO-155","ITU-R M.1477"],"us_position":"Oppose co-primary; demand secondary status only","notes":"Strategic link to AI 1.7 — this is the allocation table fight."},
}

# ─────────────────────────────────────────────────────────────────────────────
# WP ANALYSIS PROFILES
# ─────────────────────────────────────────────────────────────────────────────
WP_ANALYSIS_PROFILES = {
    "WP 5D (IMT/Mobile)": {"label":"WP 5D — IMT/Mobile (5G/6G)","interferer_type":"Terrestrial IMT base station or UE","victim_type":"Airborne or ground-based aeronautical receiver","primary_threat":"IMT identification in bands adjacent to or overlapping aeronautical allocations","propagation_models":["FSPL (worst-case)","ITU-R P.452 (terrestrial)","ITU-R P.528 (airborne victim)"],"interference_metrics":["I/N (dB) — primary metric","pfd (dBW/m²) for field strength limits"],"key_recommendations":["ITU-R M.1642","ITU-R SM.2028","ITU-R P.452","ITU-R P.528","ITU-R M.1477","ITU-R SM.1540","ITU-R SM.1541"],"protection_criteria":"I/N thresholds from FAA system protection table; +6 dB aviation safety factor for precision approach","aggregate_method":"Monte Carlo per SM.2028","specific_checks":["Propagation model P.452/P.528?","EIRP worst-case?","OOB mask per SM.1541?","Coordination zone adequate?","Aviation safety factor applied?","Aggregate interference per SM.2028?"],"common_proponent_tactics":["Using median EIRP instead of maximum","Using P.452 urban clutter for airborne victim","Claiming OOB compliance without providing mask","Single base station instead of aggregate"],"wrc27_items":["AI 1.7"],"policy_levers":["SM.1541 OOB mask compliance","Coordination zones","RR No. 4.10","RR Resolution 750"]},
    "WP 5B (Maritime/Radiodetermination)": {"label":"WP 5B — Maritime/Radiodetermination","interferer_type":"Ship-borne, coastal, or radiodetermination transmitter","victim_type":"Airborne or ground aeronautical receiver","primary_threat":"Radiolocation/maritime allocations near DME, ATC radar, or ARNS bands","propagation_models":["FSPL","ITU-R P.452","ITU-R P.528 for airborne victim"],"interference_metrics":["I/N (dB)","pfd (dBW/m²)"],"key_recommendations":["ITU-R M.1849","ITU-R P.528","ITU-R SM.2028"],"protection_criteria":"ARSR I/N = −6 dB; ASR I/N = −10 dB","aggregate_method":"Monte Carlo per SM.2028","specific_checks":["Maritime allocation adjacent to ATC radar bands?","Ship/coastal station EIRP and emission mask?","Airborne geometry modeled?"],"common_proponent_tactics":["Sea-surface propagation only — ignores airborne victim","Average power instead of peak for pulsed systems"],"wrc27_items":[],"policy_levers":["I/N thresholds for ASR/ARSR","Coordination distances from airports","RR No. 4.10"]},
    "WP 4C (MSS / DC-MSS-IMT)": {"label":"WP 4C — MSS/DC-MSS-IMT","interferer_type":"Satellite downlink (space-to-Earth) — LEO/MEO/GEO constellation","victim_type":"Ground-based or airborne aeronautical receiver","primary_threat":"MSS satellite downlinks in candidate bands adjacent to DME, AMS(R)S, ASR","propagation_models":["ITU-R P.619 (Earth-space) — CORRECT model for satellite downlinks","NOT P.452 — P.452 is for terrestrial links only","ITU-R P.676 for atmospheric absorption on slant path"],"interference_metrics":["epfd (dBW/m²/MHz) — PRIMARY metric for satellite downlinks","ΔT/T (%) — noise temperature rise for RNSS and AMS(R)S","I/N (dB) — for ASR"],"key_recommendations":["ITU-R P.619","ITU-R SM.2028","ITU-R S.1586","ITU-R M.1477","RR No. 5.444"],"protection_criteria":"DME: epfd ≤ −121.5 dBW/m²/MHz; AMS(R)S: ΔT/T ≤ 6% single-entry; ASR: I/N ≤ −10 dB","aggregate_method":"epfd Monte Carlo per S.1586/SM.2028 — all visible satellites contribute simultaneously","specific_checks":["P.619 used (not P.452)?","epfd calculated for full constellation?","All three candidate bands analyzed separately?","Airborne victim analyzed?","ΔT/T single-entry compliance?"],"common_proponent_tactics":["Single-satellite pfd instead of constellation epfd","P.452 instead of P.619","Ground-only victim","Average EIRP instead of nadir worst-case"],"wrc27_items":["AI 1.13"],"policy_levers":["epfd limits","ΔT/T limits for AMS(R)S","ASR I/N = −10 dB","RR No. 5.444","RR No. 4.10","ITU-R S.1586"]},
    "WP 7B (Space Radiocommunication / Lunar SRS)": {"label":"WP 7B — Space Research/Lunar","interferer_type":"SRS transmitter — Earth-based uplink or lunar surface transmitter","victim_type":"Terrestrial aeronautical receivers","primary_threat":"Novel use case — no established ITU-R methodology for lunar SRS vs ARNS","propagation_models":["Novel geometry — no established ITU-R model for lunar-to-Earth interference","P.452/FSPL for Earth-side SRS uplinks"],"interference_metrics":["pfd (dBW/m²) at Earth surface","I/N (dB) if interference reaches terrestrial receiver"],"key_recommendations":["ITU-R SA.509","NOTE: No ITU-R Recommendation for lunar surface SRS vs terrestrial ARNS"],"protection_criteria":"ASR I/N = −10 dB; ARNS 5350–5470 MHz I/N = −6 dB","aggregate_method":"Not yet established — FAA should argue methodology must be developed BEFORE allocation","specific_checks":["Does contribution propose a coordination methodology?","Is pfd from lunar transmitters calculated?","Are Earth-side SRS uplinks analyzed for co-frequency impact?"],"common_proponent_tactics":["Claiming lunar SRS is low power without quantitative analysis","Proposing allocation before methodology is established"],"wrc27_items":["AI 1.15"],"policy_levers":["Methodology gap argument — oppose allocation before methodology","Precautionary principle","RR No. 4.10"]},
    "WP 7C (EESS / Space Weather Sensors)": {"label":"WP 7C — EESS/Science (Passive)","interferer_type":"PASSIVE — receive-only sensors, NO transmission","victim_type":"FAA concern is ALLOCATION PRECEDENT not interference","primary_threat":"NOT interference — ALLOCATION POLICY. Co-primary in RA band weakens ARNS exclusivity for AI 1.7","propagation_models":["NOT APPLICABLE — passive sensors do not interfere with FAA systems"],"interference_metrics":["NOT APPLICABLE for direct interference","Assess allocation policy implications only"],"key_recommendations":["ITU Radio Regulations — Table of Frequency Allocations","RR Resolution 750"],"protection_criteria":"N/A — assess regulatory allocation table consequences","aggregate_method":"N/A — passive sensors do not transmit","specific_checks":["Does EESS passive allocation propose co-primary status in 4.2–4.4 GHz?","Does it weaken FAA ARNS exclusivity for AI 1.7?","Does it create coordination obligations on FAA transmitters?","Is this a stepping stone to future active allocation?"],"common_proponent_tactics":["Arguing passive = harmless = should be allowed","Using passive as foot-in-door for future active allocation"],"wrc27_items":["AI 1.17","AI 1.19"],"policy_levers":["Allocation table exclusivity","Strategic linkage to AI 1.7","Demand secondary status not co-primary"]},
    "WP 4A (Fixed Satellite Service)": {"label":"WP 4A — Fixed Satellite Service","interferer_type":"FSS satellite downlink or Earth station uplink","victim_type":"FAA fixed microwave links, ARNS","primary_threat":"FSS downlinks/uplinks in bands shared with FAA fixed microwave links","propagation_models":["ITU-R P.619 (Earth-space)","ITU-R P.452 (Earth station to terrestrial)"],"interference_metrics":["pfd (dBW/m²)","epfd","I/N (dB)"],"key_recommendations":["ITU-R S.1586","ITU-R P.619","ITU-R SM.2028"],"protection_criteria":"FAA fixed links: coordination per national frequency assignment","aggregate_method":"epfd per S.1586","specific_checks":["FSS downlink pfd comply with RR Appendix 5?","FAA fixed links included in victim analysis?"],"common_proponent_tactics":["Using coordination distance from populated areas only"],"wrc27_items":[],"policy_levers":["RR Appendix 5 pfd limits","RR No. 4.10"]},
}

WP_PROFILE_MAP = {
    "WP 5D (IMT/Mobile)":                         "WP 5D (IMT/Mobile)",
    "WP 5B (Maritime/Radiodetermination)":         "WP 5B (Maritime/Radiodetermination)",
    "WP 4C (MSS / DC-MSS-IMT)":                   "WP 4C (MSS / DC-MSS-IMT)",
    "WP 7B (Space Radiocommunication / Lunar SRS)":"WP 7B (Space Radiocommunication / Lunar SRS)",
    "WP 7C (EESS / Space Weather Sensors)":        "WP 7C (EESS / Space Weather Sensors)",
    "WP 4A (Fixed Satellite Service)":             "WP 4A (Fixed Satellite Service)",
    "WP 4B (Satellite News Gathering / ESIM)":     "WP 4A (Fixed Satellite Service)",
}



def _make_analysis_docx(analysis_md, meta):
    """
    Convert a markdown analysis string to a formatted Word document.
    Returns bytes. Robust — no OxmlElement calls that may fail on cloud.
    """
    from docx import Document as _D
    from docx.shared import Pt, Inches, RGBColor
    from docx.oxml.ns import qn as _qn
    from docx.oxml import OxmlElement as _OE
    import re as _re
    import io as _io2
    from datetime import date as _dt

    FAA_BLUE = RGBColor(0x1F, 0x4E, 0x79)
    MED_BLUE = RGBColor(0x2E, 0x75, 0xB6)
    RED      = RGBColor(0xC0, 0x00, 0x00)
    AMBER    = RGBColor(0x7F, 0x60, 0x00)
    GREEN    = RGBColor(0x37, 0x56, 0x23)
    DARK     = RGBColor(0x20, 0x20, 0x20)
    GRAY     = RGBColor(0x60, 0x60, 0x60)
    WHITE    = RGBColor(0xFF, 0xFF, 0xFF)

    doc = _D()
    sec = doc.sections[0]
    sec.page_width  = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = sec.bottom_margin = Inches(1.0)
    sec.left_margin = sec.right_margin = Inches(1.25)

    def add_heading(text, level=1):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14 if level == 1 else 8)
        p.paragraph_format.space_after  = Pt(3)
        r = p.add_run(text)
        r.bold = True
        # Section G gets a distinct amber color to flag it as the actionable draft
        if 'Draft U.S. Response' in text or 'Draft US Response' in text:
            r.font.size = Pt(13)
            r.font.color.rgb = RGBColor(0x7F, 0x40, 0x00)  # dark amber
        else:
            r.font.size = Pt(14 if level == 1 else 11)
            r.font.color.rgb = FAA_BLUE if level == 1 else MED_BLUE
        return p

    def add_body(text, color=None, bold=False, size=10):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(3)
        # Handle **bold** inline markers
        parts = _re.split(r'\*\*(.+?)\*\*', text)
        for idx, part in enumerate(parts):
            if not part:
                continue
            r = p.add_run(part)
            r.font.size = Pt(size)
            r.bold = bold or (idx % 2 == 1)
            c = color or DARK
            if "REQUIRES HUMAN REVIEW" in part: c = RED
            elif "NOT RELEVANT" in part:        c = GREEN
            elif "CLARIFICATION" in part:       c = AMBER
            elif "⚠️" in part or "UNVERIFIED" in part or "Cannot confirm" in part: c = RED
            r.font.color.rgb = c
        return p

    def add_bullet(text):
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(1)
        parts = _re.split(r'\*\*(.+?)\*\*', text)
        for idx, part in enumerate(parts):
            if not part: continue
            r = p.add_run(part)
            r.font.size = Pt(9.5)
            r.bold = (idx % 2 == 1)
            c = DARK
            if "⚠️" in part or "UNVERIFIED" in part or "Cannot confirm" in part: c = RED
            elif "High" == part.strip(): c = RED
            elif "Medium" == part.strip(): c = AMBER
            elif "Low" == part.strip(): c = GREEN
            r.font.color.rgb = c
        return p

    def add_md_table(rows_data):
        """
        Add a formatted Word table from list of lists.
        Applies column width heuristics, header shading, and semantic cell coloring.
        """
        if not rows_data: return
        import re as _tr
        from docx.shared import Inches as _In
        from docx.oxml.ns import qn as _tqn
        from docx.oxml import OxmlElement as _TOE

        n_cols = max(len(r) for r in rows_data)
        headers = [c.strip().lower() for c in rows_data[0]] if rows_data else []

        # ── Column width heuristics based on header name ──────────────────────
        # Total usable width ≈ 6.0 inches (8.5" - 2 × 1.25" margins)
        WIDTH_HINTS = {
            "proposed band": 1.2, "proposed bw": 0.6, "faa band": 1.1,
            "faa system": 1.1, "overlap": 1.0, "overlap or gap": 1.0, "gap": 0.9,
            "relationship": 0.8, "study type": 0.8, "verdict": 1.3,
            "severity": 0.6, "mechanism": 1.0, "required by": 0.9,
            "priority": 0.7, "owner": 0.7, "action": 1.5, "mitigation": 1.2,
            "title": 2.0, "description": 2.0, "finding": 1.5, "recommendation": 1.5,
        }
        col_widths = []
        for h in headers:
            w = next((v for k, v in WIDTH_HINTS.items() if k in h), None)
            if w is None: w = 1.0  # default
            col_widths.append(w)
        # Pad to n_cols
        while len(col_widths) < n_cols:
            col_widths.append(1.0)
        # Scale so total ≈ 6.0 inches
        total = sum(col_widths[:n_cols])
        scale = 6.0 / total if total > 0 else 1.0
        col_widths = [w * scale for w in col_widths[:n_cols]]

        def _shd(cell, hex_color):
            try:
                tcPr = cell._element.get_or_add_tcPr()
                # Remove existing shd elements
                for s in tcPr.findall(_tqn('w:shd')):
                    tcPr.remove(s)
                shd = _TOE('w:shd')
                shd.set(_tqn('w:fill'), hex_color)
                shd.set(_tqn('w:val'), 'clear')
                shd.set(_tqn('w:color'), 'auto')
                tcPr.append(shd)
            except Exception:
                pass

        def _cell_color(val, header_hint=""):
            """Return (text_color_rgb, bg_hex or None) based on cell value."""
            v = str(val).strip().upper()
            h = header_hint.lower()
            # Relationship column
            if "IN-BAND" in v:       return RED,   "FCE4D6"
            if "ADJACENT" in v:      return AMBER,  "FFF2CC"
            if "NEARBY" in v:        return AMBER,  "FFF2CC"
            if "NOT RELEVANT" in v:  return GREEN,  "E2EFDA"
            # Verdict
            if "REQUIRES HUMAN REVIEW" in v: return RED,   "FCE4D6"
            if "LIKELY NOT RELEVANT"  in v:  return GREEN, "E2EFDA"
            if "FLAG FOR CLARIFICATION" in v: return AMBER, "FFF2CC"
            # Severity
            if "severity" in h or v in ("HIGH","HIGH ","HIGH\n"):
                if "HIGH" in v: return RED, None
                if "MEDIUM" in v: return AMBER, None
                if "LOW" in v: return GREEN, None
            if v == "HIGH":    return RED,   None
            if v == "MEDIUM":  return AMBER,  None
            if v == "LOW":     return GREEN,  None
            # Overlap/gap
            if "OVERLAP" in v and "overlap" in h: return RED, "FCE4D6"
            if "GAP: 0" in v:  return AMBER, "FFF2CC"
            # Methodology
            if "NON-COMPLIANT" in v: return RED, None
            if "COMPLIANT" in v:     return GREEN, None
            # Study type
            if "SHARING" in v:       return AMBER, None
            if "COMPATIBILITY" in v: return MED_BLUE, None
            # Unverified
            if "⚠️" in val or "UNVERIFIED" in v: return RED, None
            return DARK, None

        tbl = doc.add_table(rows=len(rows_data), cols=n_cols)
        tbl.style = "Table Grid"

        # Apply column widths
        for ci, width_in in enumerate(col_widths):
            for ri in range(len(rows_data)):
                try:
                    tbl.rows[ri].cells[ci].width = _In(width_in)
                except Exception:
                    pass

        for ri, row_cells in enumerate(rows_data):
            is_header = (ri == 0)
            for ci in range(n_cols):
                cell = tbl.rows[ri].cells[ci]
                cell.paragraphs[0].clear()
                val = row_cells[ci].strip() if ci < len(row_cells) else ""
                hdr_hint = headers[ci] if ci < len(headers) else ""

                r = cell.paragraphs[0].add_run(val)
                r.font.size = Pt(8.5 if n_cols > 4 else 9)
                cell.paragraphs[0].paragraph_format.space_after  = Pt(2)
                cell.paragraphs[0].paragraph_format.space_before = Pt(2)

                if is_header:
                    r.bold = True
                    r.font.color.rgb = WHITE
                    _shd(cell, "1F4E79")  # FAA blue header
                    cell.paragraphs[0].alignment = 1  # center
                else:
                    txt_color, bg = _cell_color(val, hdr_hint)
                    r.font.color.rgb = txt_color
                    # Bold for special values
                    if txt_color in (RED, AMBER, GREEN):
                        r.bold = True
                    if bg:
                        _shd(cell, bg)

        doc.add_paragraph()


    # ── HEADER ───────────────────────────────────────────────────────────────
    p_title = doc.add_paragraph()
    rt = p_title.add_run("FEDERAL AVIATION ADMINISTRATION")
    rt.bold = True; rt.font.size = Pt(9); rt.font.color.rgb = MED_BLUE

    p_main = doc.add_paragraph()
    rm = p_main.add_run("ITU-R Contribution Analysis Report")
    rm.bold = True; rm.font.size = Pt(18); rm.font.color.rgb = FAA_BLUE

    # Contribution identity — document number + admin prominent on own line
    doc_id_parts = [meta.get("doc_number"), meta.get("submitting_admin")]
    doc_id = "  |  ".join(p for p in doc_id_parts if p)
    if doc_id:
        p_id = doc.add_paragraph()
        ri = p_id.add_run(doc_id)
        ri.bold = True; ri.font.size = Pt(13); ri.font.color.rgb = MED_BLUE

    # Metadata summary line
    meta_line = "  |  ".join(filter(None, [
        meta.get("working_party"),
        meta.get("meeting_date"),
        meta.get("agenda_item"),
        meta.get("doc_type"),
    ]))
    p_meta = doc.add_paragraph()
    pm = p_meta.add_run(meta_line or "—")
    pm.font.size = Pt(9); pm.font.color.rgb = GRAY

    p_date = doc.add_paragraph()
    pd2 = p_date.add_run(f"Analysis Depth: {meta.get('analysis_depth','N/A')}   |   Generated: {_dt.today()}")
    pd2.font.size = Pt(9); pd2.italic = True; pd2.font.color.rgb = GRAY

    doc.add_paragraph()

    # ── PARSE MARKDOWN ────────────────────────────────────────────────────────
    lines = analysis_md.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith('## '):
            add_heading(stripped[3:], level=1)

        elif stripped.startswith('### '):
            add_heading(stripped[4:], level=2)

        elif stripped.startswith('| ') and '|' in stripped[2:]:
            # Collect table rows
            tbl_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                tbl_lines.append(lines[i].strip())
                i += 1
            i -= 1
            # Strip separator rows
            data = []
            for tl in tbl_lines:
                if _re.match(r'^\|[-:\s|]+\|$', tl):
                    continue
                cells = [c.strip() for c in tl.split('|')]
                cells = [c for c in cells if c != '']
                if cells:
                    data.append(cells)
            if data:
                add_md_table(data)

        elif stripped.startswith('- ') or stripped.startswith('* '):
            text_b = stripped[2:]
            # If bullet contains " | " separators (e.g. "Severity: High | Confidence: High")
            # render as a mini structured row with each part on same paragraph
            if ' | ' in text_b:
                p = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.space_after = Pt(1)
                parts_pipe = text_b.split(' | ')
                for pi, part in enumerate(parts_pipe):
                    if pi > 0:
                        sep_r = p.add_run('  |  ')
                        sep_r.font.size = Pt(9); sep_r.font.color.rgb = GRAY
                    # Bold labels before ':'
                    if ':' in part:
                        label, _, value = part.partition(':')
                        lr = p.add_run(label + ':')
                        lr.bold = True; lr.font.size = Pt(9); lr.font.color.rgb = MED_BLUE
                        # Color-code the value
                        v = value.strip()
                        vc = DARK
                        if v in ('High','High\n'):   vc = RED
                        elif v in ('Medium',):        vc = AMBER
                        elif v in ('Low',):           vc = GREEN
                        elif '⚠️' in v or 'UNVERIFIED' in v: vc = RED
                        vr = p.add_run(' ' + v)
                        vr.font.size = Pt(9); vr.font.color.rgb = vc
                    else:
                        pr = p.add_run(part)
                        pr.font.size = Pt(9); pr.font.color.rgb = DARK
            else:
                add_bullet(text_b)

        elif _re.match(r'^\d+\. ', stripped):
            text_n = _re.sub(r'^\d+\. ', '', stripped)
            # Numbered action items with " | " separators — render structured
            if ' | ' in text_n:
                p = doc.add_paragraph(style="List Number")
                p.paragraph_format.space_after = Pt(2)
                parts_pipe = text_n.split(' | ')
                for pi, part in enumerate(parts_pipe):
                    if pi > 0:
                        sep_r = p.add_run('   ')
                        sep_r.font.size = Pt(9)
                    parts_bold = _re.split(r'\*\*(.+?)\*\*', part)
                    for bi, bp in enumerate(parts_bold):
                        if not bp: continue
                        br = p.add_run(bp)
                        br.font.size = Pt(9.5)
                        br.bold = (bi % 2 == 1)
                        # Color priority labels
                        if 'Immediate' in bp: br.font.color.rgb = RED
                        elif 'Before next meeting' in bp: br.font.color.rgb = AMBER
                        else: br.font.color.rgb = DARK
            else:
                p = doc.add_paragraph(style="List Number")
                p.paragraph_format.space_after = Pt(1)
                parts = _re.split(r'\*\*(.+?)\*\*', text_n)
                for idx, part in enumerate(parts):
                    if not part: continue
                    r = p.add_run(part)
                    r.font.size = Pt(9.5)
                    r.bold = (idx % 2 == 1)
                    r.font.color.rgb = DARK

        elif stripped == '---':
            doc.add_paragraph()

        elif stripped == '':
            pass

        elif stripped:
            # Check for REVIEW VERDICT line
            if 'REVIEW VERDICT' in stripped:
                add_body(stripped, bold=True, size=11)
            elif stripped.startswith('📋 STATUS') or stripped.startswith('📋 **STATUS'):
                add_body(stripped, bold=True, size=10)
            # Section G: Draft US Response — render with a highlighted box
            elif stripped.startswith('**Intervention Header') or stripped.startswith('Intervention Header'):
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after  = Pt(2)
                r_h = p.add_run(_re.sub(r'\*+', '', stripped))
                r_h.bold = True; r_h.font.size = Pt(10); r_h.font.color.rgb = FAA_BLUE
            else:
                add_body(stripped)

        i += 1

    # ── FOOTER ────────────────────────────────────────────────────────────────
    doc.add_paragraph()
    p_foot = doc.add_paragraph()
    rf = p_foot.add_run(
        "Generated by FAA RF Interference Analysis Tool  |  "
        "REVIEW ALL AI-GENERATED FINDINGS BEFORE OPERATIONAL USE"
    )
    rf.font.size = Pt(7.5); rf.italic = True; rf.font.color.rgb = GRAY

    buf = _io2.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()




def _extract_analysis_fields(analysis_text, meta=None):
    """
    Parse a single analysis text into a flat dict for Excel.
    Uses broad, flexible patterns that match the AI's varied output formats.
    """
    import re as _rx
    t = analysis_text or ""
    m = meta or {}

    def _first(patterns, default="—"):
        """Try multiple regex patterns, return first match."""
        for pat in patterns:
            hit = _rx.search(pat, t, _rx.IGNORECASE | _rx.DOTALL)
            if hit:
                return hit.group(1).strip()[:200]
        return default

    def _all(pattern):
        return list(dict.fromkeys(
            x.strip() for x in _rx.findall(pattern, t, _rx.IGNORECASE)
            if x.strip()
        ))

    # ── Source / admin ────────────────────────────────────────────────────────
    _SKIP_ADMIN = {"document","working party","wp","n/a","—","not stated","unknown","tbd",""}
    admin = (m.get("submitting_admin") or "").strip()
    if not admin or admin.lower() in _SKIP_ADMIN:
        # Look for "Submitting Administration: X" line in analysis
        _am = _rx.search(
            r'(?:Submitting Administration|Source / Admin|Administration|Submitted by|Source)[:\s*|]+([^\n|*]{3,60})',
            t, _rx.IGNORECASE)
        if _am:
            _raw = _rx.sub(r'\*+', '', _am.group(1)).strip()
            # Accept if not a generic word
            if _raw.lower() not in _SKIP_ADMIN:
                admin = _raw[:60]
    # Final strip of markdown and punctuation
    admin = _rx.sub(r'[*_`]+', '', admin).strip()
    admin = _rx.sub(r'^[\s\-|:]+|[\s\-|:]+$', '', admin) or "—"

    # ── Doc number ────────────────────────────────────────────────────────────
    doc_num = (m.get("doc_number") or "").strip() or _first([
        # Explicit "Document No." label
        r'(?:Doc(?:ument)?(?:\s+No\.?|#|:\s*))([A-Z0-9]{1,3}[A-Z]?/\d[\d\w\-]{1,15})',
        # ITU format: 5D/123-E, 5B/456, 4C/789-E  (digit+letter / number)
        r'\b(\d[A-Z]{1,2}/\d[\d\w\-]{1,15})\b',
        # R23-WP5B-C-0435 style
        r'\b(R\d{2}-WP[\w]+-[\w\-]+)\b',
    ])
    # Sanity check: must contain slash or hyphen followed by digit
    if doc_num and doc_num != "—":
        import re as _dn_re
        if not _dn_re.search(r'[/\-]\d', doc_num):
            doc_num = "—"

    # ── Working party — scan document text first, meta is fallback only ────────
    # WP codes in the document itself are authoritative
    _WP_PATTERNS = [
        # Explicit "Working Party 5B" / "WP 5B" in text or title
        r'\b(WP\s*5B)\b',   r'\b(WP\s*5D)\b',  r'\b(WP\s*4C)\b',
        r'\b(WP\s*7B)\b',   r'\b(WP\s*7C)\b',  r'\b(WP\s*4A)\b',
        r'\b(WP\s*7D)\b',   r'\b(WP\s*5A)\b',
        r'Working Party\s+(5B|5D|4C|7B|7C|4A|7D|5A)\b',
    ]
    wp_from_text = "—"
    for pat in _WP_PATTERNS:
        _wm = _rx.search(pat, t, _rx.IGNORECASE)
        if _wm:
            raw = _wm.group(1).strip().upper().replace(' ','')
            # Normalise to "WP 5B" format
            wp_from_text = "WP " + raw.replace('WP','').strip() if raw.startswith('WP') else "WP " + raw
            break

    # Cross-check admin field — filenames and admins carry WP cues too
    _admin_wp_map = {
        "5B": "WP 5B", "5D": "WP 5D", "4C": "WP 4C",
        "7B": "WP 7B", "7C": "WP 7C", "4A": "WP 4A",
    }
    if wp_from_text == "—":
        # Check admin field and doc number for WP code embedded in WP-style strings
        _check_str = str(admin) + " " + str(doc_num) + " " + t[:500]
        for code, label in _admin_wp_map.items():
            # Match "WP 4C", "WP4C", "WP-4C" patterns (not bare "4C" which is too ambiguous)
            if _rx.search(rf'WP[\s\-_]?{code}\b', _check_str, _rx.IGNORECASE):
                wp_from_text = label
                break

    # Use text-extracted WP if found; fall back to meta (the dropdown selection)
    if wp_from_text != "—":
        wp = wp_from_text
    else:
        _meta_wp = (m.get("working_party") or "").strip()
        # Extract just "WP 5D" from "WP 5D (IMT/Mobile)" style strings
        _mwm = _rx.search(r'(WP\s*(?:5[ABCD]|4[ABC]|7[ABCD]))', _meta_wp, _rx.IGNORECASE)
        wp = _mwm.group(1).strip() if _mwm else (_meta_wp or "—")

    # ── Agenda items ─ scan broadly, then filter to FAA-relevant WRC-27 AIs ────
    # FAA-relevant WRC-27 AIs: 1.7 (RA/5G), 1.13 (MSS), 1.15 (lunar), 1.17, 1.19 (EESS)
    FAA_RELEVANT_AIS = {"1.7","1.13","1.15","1.17","1.19"}

    # Collect all AI references from text
    _raw_ai = _all(r'(?:AI|Agenda Item|WRC-27 AI)[:\s]*(1\.\d+)')
    if not _raw_ai:
        _raw_ai = _all(r'\b(?:AI\s*)?(1\.\d+)\b')   # bare numbers like "1.7"

    # Filter to FAA-relevant AIs first; fall back to all detected if none are relevant
    _faa_ai = [x for x in _raw_ai if x.replace("AI ","").replace("AI","").strip() in FAA_RELEVANT_AIS]
    ai_hits = _faa_ai if _faa_ai else _raw_ai[:6]   # cap at 6 to avoid joint-meeting noise

    # Meta fallback
    if not ai_hits and m.get("agenda_item"):
        _ai_m = _rx.search(r'(1\.\d+)', m["agenda_item"])
        if _ai_m: ai_hits = [_ai_m.group(1)]

    ai_hits = list(dict.fromkeys(ai_hits))
    if ai_hits:
        agenda_items = ", ".join(
            f"AI {x}" if not x.upper().startswith("AI") else x
            for x in ai_hits
        )
    else:
        agenda_items = m.get("agenda_item") or "—"

    # ── Review verdict ────────────────────────────────────────────────────────
    verdict = _first([
        r'REVIEW VERDICT[*:\s]+([A-Z][A-Z ]+(?:HUMAN REVIEW|NOT RELEVANT|CLARIFICATION|ANALYSIS))',
        r'\*\*REVIEW VERDICT[^*]*\*\*[:\s]*([A-Z ]+)',
        r'VERDICT[:\s]+([A-Z][A-Z ]+)',
    ], "SEE FULL ANALYSIS")
    # US contribution override
    if _rx.search(r'US contribution.*summary only|summary only.*review policy', t, _rx.IGNORECASE):
        verdict = "U.S. CONTRIBUTION — SUMMARY ONLY"

    # ── Doc status ────────────────────────────────────────────────────────────
    doc_status = _first([
        r'STATUS[*:\s]+(NEW DOCUMENT|REVISION|UNCLEAR)',
        r'📋\s*STATUS[*:\s]+(NEW DOCUMENT|REVISION|UNCLEAR)',
        r'\b(NEW DOCUMENT|REVISION)\b.*confidence',
    ], "—").upper() if "NEW DOCUMENT" in t.upper() or "REVISION" in t.upper() else "—"
    if doc_status == "—":
        if _rx.search(r'\bNEW DOCUMENT\b', t, _rx.IGNORECASE):
            doc_status = "NEW"
        elif _rx.search(r'\bREVISION\b', t, _rx.IGNORECASE):
            doc_status = "REVISION"

    # ── Proposed frequency bands ──────────────────────────────────────────────
    # Pull only from column 1 of the frequency table (first | cell per row)
    # Table format: | proposed | bw | FAA band | system | gap | rel | type |
    prop_col = []
    faa_col  = []

    # Pattern A: 7-col table — | proposed | BW | FAA band | ...
    for row in _rx.finditer(
        r'^\|\s*(\d[\d,\.]*\s*[–\-]\s*\d[\d,\.]*\s*(?:MHz|GHz))\s*\|'   # col1: proposed
        r'[^|]+\|'                                                          # col2: BW (any content)
        r'\s*(\d[\d,\.]*\s*[–\-]\s*\d[\d,\.]*\s*(?:MHz|GHz))\s*\|',      # col3: FAA band
        t, _rx.IGNORECASE | _rx.MULTILINE):
        prop_col.append(row.group(1).strip())
        faa_col.append(row.group(2).strip())

    # Pattern B: 6-col table — | proposed | FAA band | ... (no BW column)
    if not prop_col:
        for row in _rx.finditer(
            r'^\|\s*(\d[\d,\.]*\s*[–\-]\s*\d[\d,\.]*\s*(?:MHz|GHz))\s*\|'  # col1: proposed
            r'\s*(\d[\d,\.]*\s*[–\-]\s*\d[\d,\.]*\s*(?:MHz|GHz))\s*\|',    # col2: FAA band (no BW)
            t, _rx.IGNORECASE | _rx.MULTILINE):
            prop_col.append(row.group(1).strip())
            faa_col.append(row.group(2).strip())

    # Pattern C: extract just proposed band from any table row
    # (fallback for NOT RELEVANT tables where FAA col is "None applicable" / "N/A")
    if not prop_col:
        for row in _rx.finditer(
            r'^\|\s*(\d[\d,\.]*\s*[–\-]\s*\d[\d,\.]*\s*(?:MHz|GHz))\s*\|',
            t, _rx.IGNORECASE | _rx.MULTILINE):
            prop_col.append(row.group(1).strip())

    # Fallback: Section D "Proposed band:" / "FAA band:" labels
    prop_label = _all(r'Proposed band[:\s]+(\d[\d,\.]*\s*[–\-]\s*\d[\d,\.]*\s*(?:MHz|GHz))')
    faa_label  = _all(r'FAA band[:\s]+(\d[\d,\.]*\s*[–\-]\s*\d[\d,\.]*\s*(?:MHz|GHz))')

    all_prop = list(dict.fromkeys(prop_col + prop_label))
    all_faa  = list(dict.fromkeys(faa_col  + faa_label))

    # ── Prose fallback: extract bands from Section D and running text ──────────
    if not all_prop:
        # "proposed 4.4–4.8 GHz" / "band 925–960 MHz" patterns in prose
        all_prop = _all(r'(?:proposed|band|allocation)[:\s]+(\d[\d,\.]*\s*[–\-]\s*\d[\d,\.]*\s*(?:MHz|GHz))')[:4]
    if not all_faa:
        # "FAA band/protected band X–Y MHz" in prose
        all_faa  = _all(r'(?:FAA\s+band|protected\s+band)[:\s]+(\d[\d,\.]*\s*[–\-]\s*\d[\d,\.]*\s*(?:MHz|GHz))')[:4]
    # Last resort: any frequency range — deduplicate proposed vs FAA by position
    if not all_prop and not all_faa:
        all_ranges = _all(r'(\d[\d,\.]*\s*[–\-]\s*\d[\d,\.]*\s*(?:MHz|GHz))')
        if all_ranges:
            all_prop = all_ranges[:3]

    proposed_bands = "; ".join(all_prop[:5]) if all_prop else "—"
    faa_bands      = "; ".join(all_faa[:5])  if all_faa  else "—"

    # ── FAA systems ───────────────────────────────────────────────────────────
    SYSTEMS = {
        "Radio Altimeter": r"radio alt(?:imeter)?|RA\b|WAICS|4[,\.]?2.{1,5}4[,\.]?4\s*GHz",
        "DME / TACAN":     r"\bDME\b|\bTACAN\b|960.{1,10}1215",
        "GPS L1 / GNSS":   r"GPS L1|GNSS L1|1575|L1\s+SBAS",
        "GNSS L5":         r"GNSS L5|GPS L5|1164.{1,10}1215",
        "ADS-B / Mode-S":  r"ADS-?B|Mode-?S\b|1090\s*MHz",
        "ASR":             r"\bASR\b|airport surv|short.range.radar",
        "ARSR":            r"\bARSR\b|en.?route.+radar|long.range.radar",
        "ILS / VOR":       r"\bILS\b|\bVOR\b|localizer|glide slope",
        "L-band AMS(R)S":  r"AMS\(R\)S|L.band.+sat|1525.{1,10}1559",
        "MLS":             r"\bMLS\b|microwave landing",
        "ARNS 5 GHz":      r"ARNS.{1,10}5\s*GHz|5000.{1,10}5150",
        "En-Route Radar":  r"en.?route.+radar|ARSR|2700.{1,10}2900",
    }
    sys_hits = [s for s, p in SYSTEMS.items() if _rx.search(p, t, _rx.IGNORECASE)]
    faa_systems = "; ".join(sys_hits) if sys_hits else "—"

    # ── Overlap / Gap ─────────────────────────────────────────────────────────
    og_m = _rx.search(
        r'((?:OVERLAP|GAP)[:\s]*\d[\d,\.]*\s*(?:MHz|GHz)(?:[^\n|]{0,40})?)',
        t, _rx.IGNORECASE)
    if og_m:
        # Normalise label to uppercase and strip trailing punctuation
        raw_og = og_m.group(1).strip()[:60]
        overlap_gap = _rx.sub(r'^(overlap|gap)', lambda m2: m2.group(1).upper(), raw_og, flags=_rx.IGNORECASE)
        overlap_gap = _rx.sub(r'[\)\.\,\s]+$', '', overlap_gap).strip()
    elif _rx.search(r'immediately adjacent|touching at|0\s*MHz\s*gap|0\s*MHz.*adjacent', t, _rx.IGNORECASE):
        overlap_gap = "GAP: 0 MHz (adjacent)"
    else:
        # Try prose: "overlaps X by N MHz", "N MHz overlap", "N MHz gap"
        prose_og = _rx.search(r'(\d[\d,\.]*\s*MHz)\s*(overlap|gap)', t, _rx.IGNORECASE)
        if prose_og:
            overlap_gap = f"{'OVERLAP' if 'overlap' in prose_og.group(2).lower() else 'GAP'}: {prose_og.group(1)}"
        else:
            overlap_gap = "—"

    # ── Relationship ──────────────────────────────────────────────────────────
    relationship = _first([
        r'\|\s*(IN-BAND|ADJACENT|NEARBY|NOT RELEVANT)\s*\|',
        r'\b(IN-BAND|ADJACENT|NEARBY|NOT RELEVANT)\b',
        r'(in.band|adjacent|nearby)',
    ]).upper()

    # ── Study type ────────────────────────────────────────────────────────────
    study_type = _first([
        r'\|\s*(SHARING|COMPATIBILITY)\s*\|',           # table cell
        r'Study Type[:\s|]+\**(SHARING|COMPATIBILITY)', # table or label
        r'\b(SHARING|COMPATIBILITY)\b.{0,30}study',     # "compatibility study"
        r'(sharing|compatibility)\s+study',
        r'study\s+type[:\s]+(sharing|compatibility)',   # "study type: compatibility"
        r'requires\s+a\s+(SHARING|COMPATIBILITY)\s+study',
    ]).upper()

    # If still empty but Relationship tells us — IN-BAND → SHARING, ADJACENT → COMPATIBILITY
    if study_type in ("—", ""):
        if relationship == "IN-BAND":
            study_type = "SHARING"
        elif relationship in ("ADJACENT", "NEARBY"):
            study_type = "COMPATIBILITY"

    # Infer relationship from study type or overlap when not explicitly stated
    if relationship in ("—",""):
        if study_type == "SHARING":
            relationship = "IN-BAND"
        elif study_type == "COMPATIBILITY":
            relationship = "ADJACENT"
        elif _rx.search(r'OVERLAP:\s*\d', t, _rx.IGNORECASE):
            relationship = "IN-BAND"
        elif _rx.search(r'GAP:\s*0\s*MHz', t, _rx.IGNORECASE):
            relationship = "ADJACENT"

    # ── Proposal summary ─────────────────────────────────────────────────────
    # Try the Document Overview section first
    summ = _first([
        r'##\s*A\).*?\n+[-•]\s*(?:Title[^:]*:)?\s*([^\n]{30,200})',
        r'Document Overview.*?\n+[-•]\s*([^\n]{30,200})',
        r'(?:summary|purpose|proposes?)[:\s]+([^\n]{30,200})',
    ])
    if summ == "—":
        paras = [p.strip() for p in _rx.split(r'\n{2,}', t)
                 if len(p.strip()) > 60 and not p.strip().startswith('#')
                 and not p.strip().startswith('|') and not p.strip().startswith('━')]
        summ = paras[0][:200] if paras else "—"
    # Strip markdown bold/italic markers from summary
    summary = _rx.sub(r'\*+', '', summ).strip()
    summary = _rx.sub(r'__?(.+?)__?', r'\1', summary)
    summary = _rx.sub(r'#+\s*', '', summary).strip()[:200]

    # ── US stance ─────────────────────────────────────────────────────────────
    # Search F section first, then US Position label, then full text
    _stance_areas = []
    _sf = _rx.search(r'(?:##\s*F\)|Recommended Actions?)[^\n]*\n([\s\S]{0,600}?)(?:\n##|\Z)',
                     t, _rx.IGNORECASE)
    if _sf: _stance_areas.append(_sf.group(1))
    _sp = _rx.search(r'(?:Recommended US Position|US Position|US\s+Stance)[:\s#*]*([\s\S]{0,300}?)(?:\n##|\n\d+\.|\Z)',
                     t, _rx.IGNORECASE)
    if _sp: _stance_areas.append(_sp.group(1))
    _stance_areas.append(t)
    stance = "—"
    for _area in _stance_areas:
        for kw in ("Oppose","Support","Propose amendments","Neutral",
                   "Flag for clarification","Monitor","Abstain"):
            # Use full word boundary on both sides to avoid matching "Proposed" as "Propose"
            _kw_pat = r'\b' + _rx.escape(kw) + r'\b'
            if _rx.search(_kw_pat, _area, _rx.IGNORECASE):
                stance = kw; break
        if stance != "—": break

    # ── Severity ─────────────────────────────────────────────────────────────
    # Explicit: look for "Severity: High/Medium/Low" anywhere in text
    for sev in ("High", "Medium", "Low"):
        if _rx.search(rf'Severity[:\|*\s]+{sev}', t, _rx.IGNORECASE):
            severity = sev; break
    else:
        # Infer from relationship and verdict when no explicit severity stated
        if relationship == "IN-BAND" and "REQUIRES HUMAN REVIEW" in verdict:
            severity = "High"
        elif relationship == "ADJACENT" and "REQUIRES HUMAN REVIEW" in verdict:
            severity = "Medium"
        elif "REQUIRES HUMAN REVIEW" in verdict and faa_systems != "—":
            severity = "Medium"
        else:
            severity = "—"

    # ── Methodology ──────────────────────────────────────────────────────────
    # Check Section E content and broader keywords
    _sec_e = _rx.search(
        r'##\s*E\)[^\n]*\n([\s\S]{0,800}?)(?:\n##|\Z)', t, _rx.IGNORECASE)
    _meth_text = (_sec_e.group(1) if _sec_e else "") + " " + t

    if _rx.search(r'Non-compliant|non.compliance|WRONG|FUNDAMENTAL ERROR|FLAG.*ERROR'
                  r'|absent|omit|missing.*requir|not.*analyz|not.*provid', _meth_text, _rx.IGNORECASE):
        methodology = "Non-compliant"
    elif _rx.search(r'No methodology|no study provided|methodology.*absent'
                    r'|no.*study.*present|study.*not.*found|no.*analysis.*provid', _meth_text, _rx.IGNORECASE):
        methodology = "No study"
    elif _rx.search(r'Compliant|methodology.*sound|appears.*compliant'
                    r'|correct.*model|appropriate.*method|properly.*appli', _meth_text, _rx.IGNORECASE):
        methodology = "Compliant"
    elif _sec_e:
        # Section E exists — infer from presence of P.452/M.1642/P.619 citations
        if _rx.search(r'P\.452|P\.528|P\.619|M\.1642|SM\.2028', _meth_text, _rx.IGNORECASE):
            methodology = "Compliant"  # cited methodology = at minimum attempting compliance
        else:
            methodology = "Non-compliant"
    else:
        methodology = "—"

    # ── Top action ────────────────────────────────────────────────────────────
    action_block = _rx.search(
        r'(?:##\s*F\)|Recommended Actions?)[^\n]*\n([\s\S]{0,600}?)(?:\n##|\Z)',
        t, _rx.IGNORECASE)
    if action_block:
        first_action = _rx.search(r'(?:1\.|\*\*Action\*\*:?)\s*([^\n]{10,120})',
                                   action_block.group(1), _rx.IGNORECASE)
        top_action = first_action.group(1).strip()[:120] if first_action else action_block.group(1)[:120]
    else:
        top_action = "—"
    # Strip markdown from all text fields
    def _strip_md(s, max_len=200):
        s = _rx.sub(r'\*+', '', s)
        s = _rx.sub(r'__?(.+?)__?', r'\1', s)
        s = _rx.sub(r'#+\s*', '', s)
        # Strip common label prefixes
        s = _rx.sub(r'^(?:Title|Summary|Purpose|Description)[:\s]+', '', s, flags=_rx.IGNORECASE)
        s = _rx.sub(r'\s+', ' ', s)
        return s.strip()[:max_len] if s.strip() else "—"
    top_action = _strip_md(top_action, 120)
    summary    = _strip_md(summary, 200)
    admin      = _strip_md(admin, 60)

    # ── Review track (routing path) ───────────────────────────────────────────
    if _rx.search(r'US contribution|PATH 1|United States|NTIA|FCC', t, _rx.IGNORECASE) and \
       _rx.search(r'summary only|US contribution.*summary', t, _rx.IGNORECASE):
        review_track = "U.S. contribution"
        review_track_justification = "Document submitted by US/NTIA — summary only per review policy."
    elif _rx.search(r'NOT RELEVANT|no further analysis|no FAA band affected', t, _rx.IGNORECASE):
        review_track = "Screened out after FAA relevance review"
        _rtj_raw = _first([
            r'NOT RELEVANT[^\n]*\n([^\n]{20,200})',
            r'no FAA band affected[^\n]*\n?([^\n]{20,150})',
            r'(NOT RELEVANT[^\n]{0,100})',
        ], "No FAA band overlap detected.")[:150]
        review_track_justification = _rx.sub(r'[*_`]+', '', _rtj_raw).strip()
    else:
        review_track = "FAA-relevant foreign document"
        review_track_justification = _first([
            r'C\) Relevance Screen\n+([^\n]{30,300})',
            r'(?:IN-BAND|ADJACENT)[^\n]{0,20}(FAA[^\n]{0,100})',
        ], "Band overlap or adjacency detected with FAA protected band.")[:200]

    # Fill stance for special paths (must be after review_track is set)
    if review_track == "U.S. contribution":
        stance = "N/A — U.S. contribution"
    elif stance == "—" and _rx.search(r'NOT RELEVANT|no further analysis', t, _rx.IGNORECASE):
        stance = "Monitor"

    # ── Stance (nuanced 4-way taxonomy) ──────────────────────────────────────
    # risky / mixed / protective / neutral_but_check
    if review_track == "U.S. contribution":
        nuanced_stance = "neutral_but_check"   # US docs: no adversarial stance
    elif review_track == "Screened out after FAA relevance review":
        nuanced_stance = "neutral_but_check"
    elif severity == "High" or stance == "Oppose":
        nuanced_stance = "risky"
    elif _rx.search(r'protective|supports.*FAA|FAA.*protection criteria.*met', t, _rx.IGNORECASE):
        if severity in ("Medium","High") or _rx.search(r'Non-compliant|missing|unresolved', t, _rx.IGNORECASE):
            nuanced_stance = "mixed"
        else:
            nuanced_stance = "protective"
    elif stance in ("Propose amendments","Flag for clarification") or severity == "Medium":
        nuanced_stance = "mixed"
    else:
        nuanced_stance = "neutral_but_check"

    # Stance justification
    _stance_just_m = _rx.search(
        r'(?:Recommended US Position|US Position|## F\))[^\n]*\n([\s\S]{0,400}?)(?:\n##|\Z)',
        t, _rx.IGNORECASE)
    stance_justification = _strip_md(
        _stance_just_m.group(1) if _stance_just_m else stance, 250)

    # ── Track change fields ────────────────────────────────────────────────────
    if doc_status == "REVISION" and _rx.search(r'TRACK CHANGES[:\s]+\d+ insert', t, _rx.IGNORECASE):
        track_change_summary = "revision"
    elif doc_status == "REVISION":
        track_change_summary = "revision"
    elif doc_status == "NEW DOCUMENT" or doc_status == "NEW":
        track_change_summary = "new_document"
    elif doc_status == "UNCLEAR":
        track_change_summary = "uncertain"
    else:
        track_change_summary = "uncertain"

    tc_just_m = _rx.search(r'(?:Track Changes|Document Status|📋 STATUS)[^\n]*\n([^\n]{20,200})', t, _rx.IGNORECASE)
    track_changes_justification = _strip_md(
        tc_just_m.group(1) if tc_just_m else "No revision markup detected.", 200)

    has_track_changes = _rx.search(r'TRACK CHANGES.*(?:insert|delet)', t, _rx.IGNORECASE) is not None
    tracked_change_count = 0
    tc_count_m = _rx.search(r'(\d+) insertions?.*?(\d+) deletions?', t, _rx.IGNORECASE)
    if tc_count_m:
        tracked_change_count = int(tc_count_m.group(1)) + int(tc_count_m.group(2))

    # ── Risk counts (from Section D) ──────────────────────────────────────────
    high_risk_count   = len(_rx.findall(r'Severity[:\s|*]+High',   t, _rx.IGNORECASE))
    medium_risk_count = len(_rx.findall(r'Severity[:\s|*]+Medium', t, _rx.IGNORECASE))
    low_risk_count    = len(_rx.findall(r'Severity[:\s|*]+Low',    t, _rx.IGNORECASE))

    # ── All frequency mentions (pipe-delimited) ────────────────────────────────
    all_freq_raw = _all(r'(\d[\d,\.]*\s*[–\-]\s*\d[\d,\.]*\s*(?:MHz|GHz))')
    frequency_mentions = " | ".join(dict.fromkeys(all_freq_raw[:20])) if all_freq_raw else "—"

    # ── Agenda item justification ──────────────────────────────────────────────
    ai_just_parts = []
    for ai_ref in (agenda_items.split(", ") if agenda_items != "—" else []):
        num = _rx.search(r'(1\.\d+)', ai_ref)
        if num:
            ai_just_parts.append(f"Explicit agenda-item cue detected: {ai_ref}")
    ai_justification = "; ".join(ai_just_parts) if ai_just_parts else (
        "No explicit WRC-27 agenda item reference detected." if agenda_items == "—"
        else f"Detected via document text: {agenda_items}")

    # ── Full recommended actions text ─────────────────────────────────────────
    rec_block = _rx.search(
        r'(?:## F\)|Recommended Actions?)[^\n]*\n([\s\S]{0,800}?)(?:\n##|\Z)',
        t, _rx.IGNORECASE)
    recommended_actions = _strip_md(
        rec_block.group(1) if rec_block else top_action, 400)

    return {
        # ── Core identification ──────────────────────────────────────────────
        "Document No.":              doc_num,
        "Source / Admin":            admin,
        "Working Party":             wp,
        "Agenda Item(s)":            agenda_items,
        "Doc Status":                doc_status,
        "Review Verdict":            verdict,
        # ── Frequency analysis ───────────────────────────────────────────────
        "Proposed Band(s)":          proposed_bands,
        "FAA Band(s)":               faa_bands,
        "FAA System(s)":             faa_systems,
        "Overlap / Gap":             overlap_gap,
        "Relationship":              relationship,
        "Study Type":                study_type,
        "All Frequency Mentions":    frequency_mentions,
        # ── Risk assessment ──────────────────────────────────────────────────
        "Highest Severity":          severity,
        "High Risk Count":           high_risk_count,
        "Medium Risk Count":         medium_risk_count,
        "Low Risk Count":            low_risk_count,
        "Methodology":               methodology,
        # ── Review routing (matches reference tool schema) ───────────────────
        "Review Track":              review_track,
        "Review Track Justification":review_track_justification,
        # ── Stance (nuanced 4-way + justification) ───────────────────────────
        "Stance":                    nuanced_stance,
        "Stance Justification":      stance_justification,
        "US Stance":                 stance,
        # ── Track changes ────────────────────────────────────────────────────
        "Track Change Summary":      track_change_summary,
        "Track Changes Justification":track_changes_justification,
        "Has Track Changes":         has_track_changes,
        "Tracked Change Count":      tracked_change_count,
        # ── Narrative fields ─────────────────────────────────────────────────
        "Agenda Item Justification": ai_justification,
        "Proposal Summary":          summary,
        "Recommended Actions":       recommended_actions,
        "Top Action":                top_action,
    }


def _make_summary_xlsx(rows):
    """
    Build a formatted Excel workbook from a list of dicts (one per document).
    Returns bytes of a valid .xlsx file.
    rows: list of dicts from _extract_analysis_fields()
    """
    from openpyxl import Workbook as _WB
    from openpyxl.styles import (Font, PatternFill, Alignment, Border, Side,
                                  GradientFill)
    from openpyxl.utils import get_column_letter as _gcl
    import io as _xio

    # ── Colours ────────────────────────────────────────────────────────────────
    FAA_BLUE   = "1F4E79"
    MED_BLUE   = "2E75B6"
    LIGHT_BLUE = "D6E4F0"
    RED_FILL   = "C00000"
    AMBER_FILL = "FF8C00"
    GREEN_FILL = "375623"
    YELLOW_BG  = "FFF2CC"
    RED_BG     = "FCE4D6"
    GREEN_BG   = "E2EFDA"
    WHITE      = "FFFFFF"
    HEADER_FG  = "FFFFFF"
    LIGHT_GRAY = "F2F2F2"

    def hdr_font(bold=True):
        return Font(name="Arial", bold=bold, color=HEADER_FG, size=9)

    def body_font(bold=False, color="000000", size=9):
        return Font(name="Arial", bold=bold, color=color, size=size)

    def fill(hex_color):
        return PatternFill("solid", fgColor=hex_color)

    def thin_border():
        s = Side(style="thin", color="BFBFBF")
        return Border(left=s, right=s, top=s, bottom=s)

    def wrap_align(h="left", v="top"):
        return Alignment(horizontal=h, vertical=v, wrap_text=True)

    wb = _WB()

    # ══════════════════════════════════════════════════════════════════════════
    # SHEET 1 — TRIAGE SUMMARY (one row per document, high-level)
    # ══════════════════════════════════════════════════════════════════════════
    ws1 = wb.active
    ws1.title = "Triage Summary"
    ws1.sheet_view.showGridLines = False

    # Title row
    # Group colours for column section banding
    GRP_COLORS = {
        "id":       FAA_BLUE,
        "freq":     "1A5276",
        "risk":     "7B241C",
        "routing":  "154360",
        "stance":   "145A32",
        "track":    "4A235A",
        "narrative":"2C3E50",
    }

    # Column definitions: (header, field_key, width, group)
    COLS = [
        ("Doc No.",              "Document No.",               14, "id"),
        ("Source / Admin",       "Source / Admin",             18, "id"),
        ("WP",                   "Working Party",               8, "id"),
        ("Agenda Item(s)",       "Agenda Item(s)",             16, "id"),
        ("Status",               "Doc Status",                 12, "id"),
        ("Verdict",              "Review Verdict",             22, "id"),
        ("Proposed Band(s)",     "Proposed Band(s)",           22, "freq"),
        ("FAA Band(s)",          "FAA Band(s)",                20, "freq"),
        ("FAA System(s)",        "FAA System(s)",              26, "freq"),
        ("Overlap / Gap",        "Overlap / Gap",              16, "freq"),
        ("Relationship",         "Relationship",               14, "freq"),
        ("Study Type",           "Study Type",                 13, "freq"),
        ("All Freq. Mentions",   "All Frequency Mentions",     35, "freq"),
        ("Severity",             "Highest Severity",           11, "risk"),
        ("High #",               "High Risk Count",             8, "risk"),
        ("Med #",                "Medium Risk Count",           8, "risk"),
        ("Low #",                "Low Risk Count",              8, "risk"),
        ("Methodology",          "Methodology",                14, "risk"),
        ("Review Track",         "Review Track",               26, "routing"),
        ("Track Justification",  "Review Track Justification", 40, "routing"),
        ("Stance",               "Stance",                     18, "stance"),
        ("Stance Justification", "Stance Justification",       40, "stance"),
        ("US Stance",            "US Stance",                  16, "stance"),
        ("TC Summary",           "Track Change Summary",       14, "track"),
        ("TC Justification",     "Track Changes Justification",35, "track"),
        ("Has TC",               "Has Track Changes",          10, "track"),
        ("TC Count",             "Tracked Change Count",        9, "track"),
        ("AI Justification",     "Agenda Item Justification",  38, "narrative"),
        ("Proposal Summary",     "Proposal Summary",           45, "narrative"),
        ("Recommended Actions",  "Recommended Actions",        50, "narrative"),
        ("Top Action",           "Top Action",                 38, "narrative"),
    ]
    n_cols = len(COLS)

    # Title row
    ws1.merge_cells(f"A1:{_gcl(n_cols)}1")
    t_cell = ws1["A1"]
    t_cell.value = "FAA RF INTERFERENCE ANALYSIS — TRIAGE SUMMARY"
    t_cell.font  = Font(name="Arial", bold=True, color=WHITE, size=13)
    t_cell.fill  = fill(FAA_BLUE)
    t_cell.alignment = Alignment(horizontal="center", vertical="center")
    ws1.row_dimensions[1].height = 22

    # Sub-header row
    ws1.merge_cells(f"A2:{_gcl(n_cols)}2")
    from datetime import date as _xdate
    sub_cell = ws1["A2"]
    sub_cell.value = (f"Generated: {_xdate.today()}  |  {len(rows)} document(s) analyzed  |  "
                      "FAA RF Interference Analysis Tool")
    sub_cell.font  = Font(name="Arial", italic=True, color=WHITE, size=8)
    sub_cell.fill  = fill(MED_BLUE)
    sub_cell.alignment = Alignment(horizontal="center", vertical="center")
    ws1.row_dimensions[2].height = 14

    # Group label row (row 3)
    GRP_ROW = 3
    grp_spans = {}
    _last_grp = None; _grp_start = 1
    for ci, (_, _, _, grp) in enumerate(COLS, 1):
        if grp != _last_grp:
            if _last_grp is not None:
                grp_spans[_last_grp] = (grp_spans[_last_grp][0], ci-1)
            grp_spans[grp] = (ci, ci)
            _grp_start = ci; _last_grp = grp
        else:
            grp_spans[grp] = (grp_spans[grp][0], ci)
    GRP_LABELS = {"id":"IDENTIFICATION","freq":"FREQUENCY ANALYSIS","risk":"RISK ASSESSMENT",
                  "routing":"REVIEW ROUTING","stance":"STANCE","track":"TRACK CHANGES","narrative":"NARRATIVE"}
    for grp, (c1, c2) in grp_spans.items():
        if c1 != c2:
            try: ws1.merge_cells(f"{_gcl(c1)}{GRP_ROW}:{_gcl(c2)}{GRP_ROW}")
            except Exception: pass
        gc = ws1.cell(row=GRP_ROW, column=c1, value=GRP_LABELS.get(grp,""))
        gc.font = Font(name="Arial", bold=True, color=WHITE, size=8)
        gc.fill = fill(GRP_COLORS.get(grp, FAA_BLUE))
        gc.alignment = Alignment(horizontal="center", vertical="center")
        gc.border = thin_border()
    ws1.row_dimensions[GRP_ROW].height = 16

    # Column header row (row 4)
    HDR_ROW = 4
    for ci, (hdr, _, width, grp) in enumerate(COLS, 1):
        cell = ws1.cell(row=HDR_ROW, column=ci, value=hdr)
        cell.font      = hdr_font()
        cell.fill      = fill(GRP_COLORS.get(grp, FAA_BLUE))
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border    = thin_border()
        ws1.column_dimensions[_gcl(ci)].width = width
    ws1.row_dimensions[HDR_ROW].height = 32

    VERDICT_COLORS = {
        "REQUIRES HUMAN REVIEW": (RED_BG,   "C00000"),
        "FLAG FOR CLARIFICATION": (YELLOW_BG,"7F4000"),
        "LIKELY NOT RELEVANT":   (GREEN_BG,  "375623"),
    }
    SEV_COLORS    = {"High":("C00000",WHITE),"Medium":("FF8C00",WHITE),"Low":("375623",WHITE)}
    STANCE_COLORS = {"risky":("C00000",WHITE),"mixed":("FF8C00",WHITE),
                     "protective":("375623",WHITE),"neutral_but_check":("1A5276",WHITE)}
    TRACK_COLORS  = {"revision":MED_BLUE,"new_document":"1A5276",
                     "clean_update":"145A32","uncertain":"7D6608"}
    ROUTING_COLORS= {"FAA-relevant foreign document":("C00000",WHITE),
                     "Screened out after FAA relevance review":("375623",WHITE),
                     "U.S. contribution":("1A5276",WHITE)}

    for ri, row in enumerate(rows):
        excel_row = HDR_ROW + 1 + ri
        verdict   = row.get("Review Verdict","")
        bg_hex, _ = VERDICT_COLORS.get(
            next((k for k in VERDICT_COLORS if k in verdict),""), (WHITE,"000000"))

        for ci, (_, field, _, grp) in enumerate(COLS, 1):
            val = row.get(field,"—")
            if val is False: val = "No"
            if val is True:  val = "Yes"
            cell = ws1.cell(row=excel_row, column=ci, value=val)
            cell.border    = thin_border()
            cell.alignment = wrap_align()
            cell.font      = body_font()
            base = bg_hex if bg_hex != WHITE else (LIGHT_GRAY if ri%2==1 else WHITE)
            cell.fill = fill(base)

            if field == "Review Verdict":
                _, fg = VERDICT_COLORS.get(next((k for k in VERDICT_COLORS if k in str(val)),""), (WHITE,"000000"))
                cell.font = body_font(bold=True, color=fg)
            elif field == "Highest Severity":
                sc = SEV_COLORS.get(str(val))
                if sc:
                    cell.fill = fill(sc[0]); cell.font = body_font(bold=True, color=sc[1])
                    cell.alignment = Alignment(horizontal="center", vertical="center")
            elif field in ("High Risk Count","Medium Risk Count","Low Risk Count"):
                cell.alignment = Alignment(horizontal="center", vertical="center")
                if field == "High Risk Count" and isinstance(val,int) and val > 0:
                    cell.font = body_font(bold=True, color="C00000")
                elif field == "Medium Risk Count" and isinstance(val,int) and val > 0:
                    cell.font = body_font(bold=True, color="FF8C00")
            elif field == "Stance":
                sc2 = STANCE_COLORS.get(str(val))
                if sc2:
                    cell.fill = fill(sc2[0]); cell.font = body_font(bold=True, color=sc2[1])
                    cell.alignment = Alignment(horizontal="center", vertical="center")
            elif field == "Review Track":
                rc = ROUTING_COLORS.get(str(val))
                if rc:
                    cell.fill = fill(rc[0]); cell.font = body_font(bold=True, color=rc[1])
            elif field == "Track Change Summary":
                tc_color = TRACK_COLORS.get(str(val))
                if tc_color:
                    cell.fill = fill(tc_color); cell.font = body_font(bold=True, color=WHITE)
                    cell.alignment = Alignment(horizontal="center", vertical="center")
            elif field == "Has Track Changes":
                cell.alignment = Alignment(horizontal="center", vertical="center")
                if val == "Yes": cell.font = body_font(bold=True, color=MED_BLUE)
            elif field == "Methodology":
                if str(val) == "Non-compliant": cell.font = body_font(bold=True, color="C00000")
                elif str(val) == "Compliant":   cell.font = body_font(color="375623")
            elif field == "US Stance":
                if str(val) == "Oppose":   cell.font = body_font(bold=True, color="C00000")
                elif str(val) == "Support": cell.font = body_font(color="375623")
            elif field == "Relationship":
                if "IN-BAND" in str(val):
                    cell.fill = fill("FCE4D6"); cell.font = body_font(bold=True, color="C00000")
                elif "ADJACENT" in str(val):
                    cell.fill = fill(YELLOW_BG)
            elif field == "Doc Status":
                if str(val) in ("REVISION","revision"):
                    cell.font = body_font(bold=True, color=MED_BLUE)
        ws1.row_dimensions[excel_row].height = 52

    ws1.freeze_panes = f"A{HDR_ROW+1}"
    ws1.auto_filter.ref = f"A{HDR_ROW}:{_gcl(n_cols)}{HDR_ROW+len(rows)}"

    # SHEET 2 — LEGEND
    ws2 = wb.create_sheet("Legend")
    ws2.sheet_view.showGridLines = False
    ws2.column_dimensions["A"].width = 26
    ws2.column_dimensions["B"].width = 52
    ws2.column_dimensions["C"].width = 28
    legend_title = ws2.cell(row=1, column=1, value="FIELD LEGEND — FAA RF Interference Analysis Tool")
    legend_title.font = Font(name="Arial", bold=True, color=WHITE, size=11)
    legend_title.fill = fill(FAA_BLUE)
    ws2.merge_cells("A1:C1")
    legend_title.alignment = Alignment(horizontal="center")
    ws2.row_dimensions[1].height = 20
    LEGEND = [
        ("Field","Description","Values / Notes"),
        ("Doc No.","Document identifier","e.g. 5D/123-E or batch filename"),
        ("Source / Admin","Submitting country or body","e.g. China, Working Party 5B, ICAO"),
        ("Agenda Item(s)","WRC-27 agenda items referenced","AI 1.7, AI 1.13, AI 1.15, AI 1.17, AI 1.19"),
        ("Status","New submission or revision","NEW DOCUMENT / REVISION / UNCLEAR"),
        ("Review Verdict","AI triage decision","REQUIRES HUMAN REVIEW / LIKELY NOT RELEVANT / FLAG FOR CLARIFICATION"),
        ("Proposed Band(s)","Exact proposed frequency band(s)","From analysis frequency table column 1"),
        ("FAA Band(s)","FAA protected bands affected","From analysis frequency table column 3"),
        ("FAA System(s)","FAA aviation systems at risk","Radio Altimeter, DME, ASR, GPS, ADS-B, etc."),
        ("Overlap / Gap","Calculated frequency overlap or gap","GAP: 0 MHz / OVERLAP: X MHz"),
        ("Relationship","Spectral relationship","IN-BAND / ADJACENT / NEARBY / NOT RELEVANT"),
        ("Study Type","Coexistence study type","SHARING (co-band) / COMPATIBILITY (adjacent)"),
        ("All Freq. Mentions","All frequency ranges found (pipe-delimited)","Raw scan — for manual cross-check"),
        ("Severity","Highest FAA impact severity","High / Medium / Low"),
        ("High # / Med # / Low #","Issue count at each severity level","Integer — 0 means no issues at that level"),
        ("Methodology","ITU-R methodology compliance","Compliant / Non-compliant / No study"),
        ("Review Track","Document routing path","FAA-relevant foreign document / Screened out / U.S. contribution"),
        ("Track Justification","Why this routing was assigned","From Section C analysis text"),
        ("Stance","4-way stance taxonomy","risky / mixed / protective / neutral_but_check"),
        ("Stance Justification","Why this stance was assigned","From Section F analysis text"),
        ("US Stance","Recommended US delegation position","Oppose / Support / Propose amendments / Monitor / Neutral"),
        ("TC Summary","Track change classification","new_document / revision / clean_update / uncertain"),
        ("TC Justification","Evidence for track change classification","From document status section"),
        ("Has TC","Were track changes detected in the file?","Yes / No"),
        ("TC Count","Total insertions + deletions detected","Integer — 0 if no track changes"),
        ("AI Justification","Why each agenda item was flagged","Cue type and location in document"),
        ("Proposal Summary","Brief description of what the document proposes","Up to 200 characters"),
        ("Recommended Actions","Full Section F recommended actions text","Up to 400 characters"),
        ("Top Action","Single highest-priority action","Up to 120 characters"),
    ]
    for ri, (col1, col2, col3) in enumerate(LEGEND):
        is_hdr = (ri==0)
        for ci, val in enumerate([col1,col2,col3], 1):
            cell = ws2.cell(row=ri+2, column=ci, value=val)
            cell.font      = Font(name="Arial", bold=is_hdr, size=9,
                                  color=WHITE if is_hdr else "000000")
            cell.fill      = fill(MED_BLUE if is_hdr else (LIGHT_GRAY if ri%2==0 else WHITE))
            cell.border    = thin_border()
            cell.alignment = Alignment(wrap_text=True, vertical="top")
        ws2.row_dimensions[ri+2].height = 30 if not is_hdr else 18


    buf = _xio.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf.getvalue()



# ─────────────────────────────────────────────────────────────────────────────
# NEO4J INTEGRATION — Persistent knowledge graph for cross-document analysis
# ─────────────────────────────────────────────────────────────────────────────

def _neo4j_driver():
    """Return a connected Neo4j driver using Streamlit secrets, or None if not configured."""
    try:
        from neo4j import GraphDatabase
        cfg = st.secrets.get("neo4j", {})
        uri  = cfg.get("uri")
        user = cfg.get("username", "neo4j")
        pwd  = cfg.get("password")
        if not uri or not pwd:
            return None
        return GraphDatabase.driver(uri, auth=(user, pwd))
    except Exception:
        return None


def _neo4j_write_analysis(driver, analysis_text, meta, contrib_text=""):
    """
    Write one analyzed contribution into Neo4j using the existing schema:
    Document → RiskFinding, FAAProtectionBand, RecommendedAction,
               RegulatoryIssue, SourceExcerpt, TrackedChange,
               StudyAssumption, SimulationMethod, FrequencyBand
    All writes use MERGE so re-running is idempotent.
    Returns (nodes_created, rels_created) counts.
    """
    import re as _nr
    from datetime import date as _nd

    fields = _extract_analysis_fields(analysis_text, meta)
    doc_id = fields.get("Document No.") or meta.get("doc_number") or "UNKNOWN"
    if doc_id == "—":
        doc_id = f"DOC_{hash(analysis_text) % 100000}"

    nodes_c = 0
    rels_c   = 0

    with driver.session() as s:

        # ── Document node ─────────────────────────────────────────────────────
        s.run("""
            MERGE (d:Document {doc_id: $doc_id})
            SET d.admin         = $admin,
                d.working_party = $wp,
                d.agenda_items  = $ai,
                d.doc_status    = $status,
                d.verdict       = $verdict,
                d.us_stance     = $stance,
                d.severity      = $severity,
                d.methodology   = $meth,
                d.study_type    = $study,
                d.summary       = $summary,
                d.analyzed_date = $date
        """, doc_id=doc_id,
             admin   = fields.get("Source / Admin","—"),
             wp      = fields.get("Working Party","—"),
             ai      = fields.get("Agenda Item(s)","—"),
             status  = fields.get("Doc Status","—"),
             verdict = fields.get("Review Verdict","—"),
             stance  = fields.get("US Stance","—"),
             severity= fields.get("Highest Severity","—"),
             meth    = fields.get("Methodology","—"),
             study   = fields.get("Study Type","—"),
             summary = fields.get("Proposal Summary","—"),
             date    = str(_nd.today()))
        nodes_c += 1

        # ── FAA Protection Bands ──────────────────────────────────────────────
        faa_bands_str = fields.get("FAA Band(s)","")
        faa_systems   = fields.get("FAA System(s)","")
        overlap_gap   = fields.get("Overlap / Gap","—")
        relationship  = fields.get("Relationship","—")

        for band in [b.strip() for b in faa_bands_str.split(";") if b.strip() and b.strip() != "—"]:
            s.run("""
                MERGE (b:FAAProtectionBand {band_range: $band})
                MERGE (d:Document {doc_id: $doc_id})
                MERGE (d)-[r:AFFECTS_BAND]->(b)
                SET r.overlap_gap = $og, r.relationship = $rel
            """, band=band, doc_id=doc_id, og=overlap_gap, rel=relationship)
            nodes_c += 1; rels_c += 1

        for sys_name in [s2.strip() for s2 in faa_systems.split(";") if s2.strip() and s2.strip() != "—"]:
            s.run("""
                MERGE (fb:FrequencyBand {system_name: $sys})
                MERGE (d:Document {doc_id: $doc_id})
                MERGE (d)-[:MENTIONS_BAND]->(fb)
            """, sys=sys_name, doc_id=doc_id)
            nodes_c += 1; rels_c += 1

        # ── Proposed bands ────────────────────────────────────────────────────
        for pb in [b.strip() for b in fields.get("Proposed Band(s)","").split(";")
                   if b.strip() and b.strip() != "—"]:
            s.run("""
                MERGE (fb:FrequencyBand {band_range: $band})
                SET fb.proposed = true
                MERGE (d:Document {doc_id: $doc_id})
                MERGE (d)-[:ADDRESSES]->(fb)
            """, band=pb, doc_id=doc_id)
            nodes_c += 1; rels_c += 1

        # ── Risk Findings (Section D issues) ──────────────────────────────────
        risk_pattern = _nr.findall(
            r'\*\*Issue\s*(\d+)[:\*]+\*\*\s*([^\n]+)\n((?:.*\n){0,8})',
            analysis_text, _nr.IGNORECASE)
        for issue_num, issue_title, issue_body in risk_pattern:
            # Extract severity, mechanism, mitigation from the body block
            sev_m = _nr.search(r'Severity[:\s]*(High|Medium|Low)', issue_body, _nr.IGNORECASE)
            mech_m = _nr.search(r'Mechanism[:\s]*([^\n]+)', issue_body, _nr.IGNORECASE)
            mit_m  = _nr.search(r'Mitigation[:\s]*([^\n]+)', issue_body, _nr.IGNORECASE)
            req_m  = _nr.search(r'Required by[:\s]*([^\n]+)', issue_body, _nr.IGNORECASE)

            s.run("""
                MERGE (rf:RiskFinding {doc_id: $doc_id, issue_num: $num})
                SET rf.title      = $title,
                    rf.severity   = $sev,
                    rf.mechanism  = $mech,
                    rf.mitigation = $mit,
                    rf.required_by= $req
                MERGE (d:Document {doc_id: $doc_id})
                MERGE (d)-[:HAS_RISK]->(rf)
            """, doc_id=doc_id,
                 num   = issue_num,
                 title = issue_title.strip()[:200],
                 sev   = sev_m.group(1) if sev_m else "—",
                 mech  = mech_m.group(1).strip()[:100] if mech_m else "—",
                 mit   = mit_m.group(1).strip()[:200] if mit_m else "—",
                 req   = req_m.group(1).strip()[:100] if req_m else "—")
            nodes_c += 1; rels_c += 1

        # ── Recommended Actions (Section F) ───────────────────────────────────
        action_pattern = _nr.findall(
            r'^\d+\.\s*\*\*Action[:\*]+\*\*\s*([^\n|]+)',
            analysis_text, _nr.IGNORECASE | _nr.MULTILINE)
        for ai, action_text in enumerate(action_pattern[:5], 1):
            s.run("""
                MERGE (ra:RecommendedAction {doc_id: $doc_id, action_num: $num})
                SET ra.action_text = $text
                MERGE (d:Document {doc_id: $doc_id})
                MERGE (d)-[:RECOMMENDS_ACTION]->(ra)
            """, doc_id=doc_id, num=str(ai),
                 text=action_text.strip()[:300])
            nodes_c += 1; rels_c += 1

        # ── Regulatory Issues ─────────────────────────────────────────────────
        reg_pattern = _nr.findall(
            r'(?:cite|per|required by|violation of)[:\s]*(RR[^,\n;]{5,60}|SM\.\d+[^,\n;]{0,40}|M\.\d+[^,\n;]{0,40})',
            analysis_text, _nr.IGNORECASE)
        for reg_ref in list(dict.fromkeys(reg_pattern))[:8]:
            s.run("""
                MERGE (ri:RegulatoryIssue {citation: $cit})
                MERGE (d:Document {doc_id: $doc_id})
                MERGE (d)-[:HAS_REGULATORY_ISSUE]->(ri)
            """, cit=reg_ref.strip()[:100], doc_id=doc_id)
            nodes_c += 1; rels_c += 1

        # ── Source Excerpts (verbatim citations from document) ────────────────
        quote_pattern = _nr.findall(r'"([^"]{20,300})"', analysis_text)
        for qi, quote in enumerate(quote_pattern[:6], 1):
            s.run("""
                MERGE (se:SourceExcerpt {doc_id: $doc_id, excerpt_num: $num})
                SET se.text = $text
                MERGE (d:Document {doc_id: $doc_id})
                MERGE (d)-[:SUPPORTED_BY]->(se)
            """, doc_id=doc_id, num=str(qi),
                 text=quote.strip()[:400])
            nodes_c += 1; rels_c += 1

        # ── Tracked Changes ───────────────────────────────────────────────────
        tc_summary = meta.get("tc_summary","")
        if tc_summary and "TRACK CHANGES DETECTED" in tc_summary:
            ins_m = _nr.search(r'(\d+) insertions', tc_summary)
            del_m = _nr.search(r'(\d+) deletions', tc_summary)
            s.run("""
                MERGE (tc:TrackedChange {doc_id: $doc_id})
                SET tc.insertions = $ins,
                    tc.deletions  = $dels,
                    tc.summary    = $summ
                MERGE (d:Document {doc_id: $doc_id})
                MERGE (d)-[:HAS_CHANGE]->(tc)
            """, doc_id=doc_id,
                 ins  = int(ins_m.group(1)) if ins_m else 0,
                 dels = int(del_m.group(1)) if del_m else 0,
                 summ = tc_summary[:500])
            nodes_c += 1; rels_c += 1

        # ── Study Assumptions & Simulation Methods ────────────────────────────
        for method in ["P.452","P.528","P.619","SM.2028","M.1642","P.676"]:
            if method in analysis_text:
                s.run("""
                    MERGE (sm:SimulationMethod {method_name: $method})
                    MERGE (d:Document {doc_id: $doc_id})
                    MERGE (d)-[:USES_METHOD]->(sm)
                """, method=method, doc_id=doc_id)
                nodes_c += 1; rels_c += 1

    return nodes_c, rels_c


def _neo4j_nl_query(driver, question, api_key_val):
    """
    Convert a natural language question to Cypher via Claude,
    execute it against Neo4j, and return a formatted answer.
    """
    # Step 1: Generate Cypher from NL question
    schema_hint = """
Node labels: Document, FAAProtectionBand, FrequencyBand, RecommendedAction,
             RegulatoryIssue, RiskFinding, SimulationMethod, SourceExcerpt,
             StudyAssumption, TrackedChange

Key Document properties: doc_id, admin, working_party, agenda_items, verdict,
  us_stance, severity, methodology, study_type, summary, analyzed_date

Relationship types: ADDRESSES, AFFECTS_BAND, HAS_CHANGE, HAS_REGULATORY_ISSUE,
  HAS_RISK, MAKES_ASSUMPTION, MENTIONS_BAND, RECOMMENDS_ACTION, SUPPORTED_BY,
  USES_METHOD

Example queries:
- All docs affecting DME: MATCH (d:Document)-[:AFFECTS_BAND]->(b:FAAProtectionBand) WHERE b.band_range CONTAINS '960' RETURN d.doc_id, d.admin, d.verdict
- High severity risks: MATCH (d:Document)-[:HAS_RISK]->(r:RiskFinding {severity:'High'}) RETURN d.doc_id, r.title
- By admin: MATCH (d:Document) WHERE d.admin CONTAINS 'China' RETURN d.doc_id, d.verdict, d.agenda_items
"""

    import anthropic as _anth
    client = _anth.Anthropic(api_key=api_key_val)

    cypher_resp = client.messages.create(
        model="claude-sonnet-4-5",
        max_tokens=400,
        system=f"Convert the user's question to a Cypher query for this Neo4j schema. Return ONLY the Cypher query, no explanation, no markdown fences.\n\nSchema:\n{schema_hint}",
        messages=[{"role": "user", "content": question}]
    )
    cypher = cypher_resp.content[0].text.strip().strip("`").replace("cypher","").strip()

    # Step 2: Execute Cypher
    try:
        with driver.session() as s:
            result = s.run(cypher)
            records = [dict(r) for r in result]
    except Exception as e:
        return f"⚠️ Cypher error: {e}\n\nGenerated query:\n```\n{cypher}\n```", cypher, []

    if not records:
        return "No results found for this query.", cypher, []

    # Step 3: Format answer in natural language
    fmt_resp = client.messages.create(
        model="claude-sonnet-4-5",
        max_tokens=600,
        system="You are summarizing Neo4j query results for an FAA spectrum policy analyst. Be concise and direct. Use bullet points. State counts clearly.",
        messages=[{"role": "user", "content": f"Question: {question}\n\nResults ({len(records)} rows):\n{records[:20]}"}]
    )
    return fmt_resp.content[0].text, cypher, records


def _make_meeting_docx(info, sessions, docs, ais, actions):
    """Build a formatted Word trip report from Meeting Notes. Returns bytes."""
    from docx import Document as _MD
    from docx.shared import Pt, Inches, RGBColor
    from docx.oxml.ns import qn as _mqn
    from docx.oxml import OxmlElement as _MOE
    import io as _mio
    from datetime import date as _mdate

    FAA_BLUE = RGBColor(0x1F,0x4E,0x79); MED_BLUE = RGBColor(0x2E,0x75,0xB6)
    RED=RGBColor(0xC0,0x00,0x00); AMBER=RGBColor(0x7F,0x60,0x00)
    GREEN=RGBColor(0x37,0x56,0x23); DARK=RGBColor(0x20,0x20,0x20)
    GRAY=RGBColor(0x60,0x60,0x60); WHITE=RGBColor(0xFF,0xFF,0xFF)

    doc = _MD()
    sec = doc.sections[0]
    sec.page_width=Inches(8.5); sec.page_height=Inches(11)
    sec.top_margin=sec.bottom_margin=Inches(1.0)
    sec.left_margin=sec.right_margin=Inches(1.25)

    def h1(text):
        p=doc.add_paragraph()
        p.paragraph_format.space_before=Pt(14); p.paragraph_format.space_after=Pt(4)
        r=p.add_run(text); r.bold=True; r.font.size=Pt(13); r.font.color.rgb=FAA_BLUE

    def h2(text):
        p=doc.add_paragraph()
        p.paragraph_format.space_before=Pt(8); p.paragraph_format.space_after=Pt(2)
        r=p.add_run(text); r.bold=True; r.font.size=Pt(11); r.font.color.rgb=MED_BLUE

    def lv(label, val, lc=None):
        p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(2)
        rl=p.add_run(f"{label}: "); rl.bold=True; rl.font.size=Pt(10); rl.font.color.rgb=lc or MED_BLUE
        rv=p.add_run(str(val) if val else "—"); rv.font.size=Pt(10); rv.font.color.rgb=DARK

    def body(text, bold=False, color=None, size=10):
        p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(3)
        r=p.add_run(str(text)); r.bold=bold; r.font.size=Pt(size); r.font.color.rgb=color or DARK

    def mk_table(headers, rows):
        if not rows: return
        tbl=doc.add_table(rows=len(rows)+1,cols=len(headers)); tbl.style="Table Grid"
        for ci,h in enumerate(headers):
            cell=tbl.rows[0].cells[ci]; cell.paragraphs[0].clear()
            r2=cell.paragraphs[0].add_run(h); r2.bold=True; r2.font.size=Pt(9); r2.font.color.rgb=WHITE
            try:
                shd=_MOE("w:shd"); shd.set(_mqn("w:fill"),"1F4E79"); shd.set(_mqn("w:val"),"clear")
                cell._element.get_or_add_tcPr().append(shd)
            except Exception: pass
        for ri,row in enumerate(rows):
            for ci,val in enumerate(row):
                cell=tbl.rows[ri+1].cells[ci]; cell.paragraphs[0].clear()
                r3=cell.paragraphs[0].add_run(str(val) if val else "—"); r3.font.size=Pt(9)
                c=DARK
                if str(val) in ("HIGH",): c=RED
                elif str(val) in ("MEDIUM",): c=AMBER
                elif str(val) in ("LOW",): c=GREEN
                elif str(val).startswith("\u2705"): c=GREEN
                r3.font.color.rgb=c
        doc.add_paragraph()

    # Title
    p_o=doc.add_paragraph(); r_o=p_o.add_run("FEDERAL AVIATION ADMINISTRATION")
    r_o.bold=True; r_o.font.size=Pt(9); r_o.font.color.rgb=MED_BLUE
    p_t=doc.add_paragraph(); r_t=p_t.add_run("ITU-R Working Party Meeting Record")
    r_t.bold=True; r_t.font.size=Pt(20); r_t.font.color.rgb=FAA_BLUE
    meeting_name = info.get("meeting_name") or "ITU-R Meeting"
    p_s=doc.add_paragraph(); r_s=p_s.add_run(meeting_name)
    r_s.bold=True; r_s.font.size=Pt(14); r_s.font.color.rgb=MED_BLUE
    doc.add_paragraph()
    lv("Working Party", info.get("working_party"))
    lv("Location",      info.get("location"))
    lv("Dates",         info.get("dates"))
    lv("US Delegation Head", info.get("us_head"))
    lv("FAA Technical Lead", info.get("faa_lead"))
    p_g=doc.add_paragraph(); r_g=p_g.add_run(f"Generated: {_mdate.today()} — FAA RF Interference Analysis Tool")
    r_g.font.size=Pt(8); r_g.italic=True; r_g.font.color.rgb=GRAY

    # Section 1: Agenda Items
    doc.add_page_break(); h1("SECTION 1 — US POSITION MATRIX (Agenda Items)")
    if ais:
        for a in ais.values():
            h2(f"AI {a.get('num','?')} — {a.get('title','')}")
            lv("FAA Bands at Risk", a.get("faa_bands"))
            lv("US Position",       a.get("us_position"))
            lv("Meeting Status",    a.get("status"))
            lv("Rapporteur",        a.get("rapporteur"))
            lv("Allied Admins",     a.get("allies"))
            if a.get("faa_concerns"):  lv("FAA Concerns", a["faa_concerns"])
            if a.get("next_steps"):    lv("Next Steps", a["next_steps"], lc=RED)
            if a.get("current_text"):
                body("Draft Text:", bold=True, color=MED_BLUE)
                pb=doc.add_paragraph(style="List Bullet"); pb.add_run(a["current_text"]).font.size=Pt(9)
            doc.add_paragraph()
    else:
        body("No agenda items logged.", color=GRAY)

    # Section 2: Documents
    doc.add_page_break(); h1("SECTION 2 — DOCUMENT INDEX (FAA Flagged Documents)")
    if docs:
        for concern_key, c_color, label in [("HIGH",RED,"🔴 HIGH"), ("MEDIUM",AMBER,"🟡 MEDIUM"), ("LOW",GREEN,"🟢 LOW/MONITOR")]:
            group=[d for d in docs.values() if concern_key in str(d.get("concern",""))]
            if not group: continue
            h2(label)
            mk_table(["Doc #","Admin","AI","Session","Title","US Action"],
                [[d.get("doc_num",""),d.get("admin",""),d.get("ai",""),d.get("session",""),
                  (d.get("title","") or "")[:45],d.get("us_action","")] for d in group])
            for d in group:
                if d.get("summary"):   lv(f"  {d.get('doc_num','')} Summary", d["summary"])
                if d.get("faa_response"): lv(f"  {d.get('doc_num','')} US Response", d["faa_response"])
    else:
        body("No documents logged.", color=GRAY)

    # Section 3: Session Notes
    doc.add_page_break(); h1("SECTION 3 — SESSION NOTES")
    if sessions:
        for s in sessions.values():
            h2(f"{s.get('session','')} — {s.get('date','')}")
            lv("Chair", s.get("chair")); lv("Agenda Items", s.get("ai_context"))
            outcome=s.get("faa_outcome","")
            oc=RED if "Unfavorable" in outcome else (GREEN if "Favorable" in outcome else AMBER)
            lv("FAA Outcome", outcome, lc=oc)
            if s.get("notes"):
                body("Notes:", bold=True, color=MED_BLUE)
                for line in s["notes"].split("\n"):
                    if line.strip():
                        pb=doc.add_paragraph(style="List Bullet"); pb.add_run(line.strip()).font.size=Pt(9)
            if s.get("key_decisions"): lv("Key Decisions", s["key_decisions"], lc=RED)
            if s.get("follow_up"):     lv("Follow-up", s["follow_up"], lc=AMBER)
            doc.add_paragraph()
    else:
        body("No sessions logged.", color=GRAY)

    # Section 4: Actions
    doc.add_page_break(); h1("SECTION 4 — ACTION ITEMS")
    if actions:
        open_a=[a for a in actions if "Complete" not in str(a.get("status",""))]
        done_a=[a for a in actions if "Complete"  in str(a.get("status",""))]
        if open_a:
            h2("Open Actions")
            mk_table(["Priority","Status","Owner","Due","AI","Action"],
                [[a.get("priority",""),a.get("status",""),a.get("owner",""),
                  a.get("due",""),a.get("ai",""),str(a.get("desc",""))[:60]] for a in open_a])
        if done_a:
            h2("Completed Actions")
            mk_table(["","Owner","AI","Action"],
                [["\u2705",a.get("owner",""),a.get("ai",""),str(a.get("desc",""))[:70]] for a in done_a])
    else:
        body("No action items logged.", color=GRAY)

    # Footer
    doc.add_paragraph()
    p_f=doc.add_paragraph()
    r_f=p_f.add_run(f"END  |  {meeting_name}  |  Generated {_mdate.today()}  |  FAA RF Interference Analysis Tool  |  DISTRIBUTE PER FAA/NTIA GUIDELINES")
    r_f.font.size=Pt(7.5); r_f.italic=True; r_f.font.color.rgb=GRAY

    buf=_mio.BytesIO(); doc.save(buf); buf.seek(0)
    return buf.getvalue()


def _make_batch_docx(batch_results, triage_rows, working_party, analysis_depth):
    """
    Build a combined Word document for all batch analyses.
    Returns bytes.
    """
    from docx import Document as _BD
    from docx.shared import Pt, Inches, RGBColor
    from docx.oxml.ns import qn as _qn
    from docx.oxml import OxmlElement as _OE
    import re as _bre
    import io as _bio
    from datetime import date as _bdate

    FAA_BLUE = RGBColor(0x1F,0x4E,0x79); MED_BLUE = RGBColor(0x2E,0x75,0xB6)
    RED  = RGBColor(0xC0,0x00,0x00);  GREEN = RGBColor(0x37,0x56,0x23)
    GRAY = RGBColor(0x60,0x60,0x60);  DARK  = RGBColor(0x20,0x20,0x20)
    WHITE = RGBColor(0xFF,0xFF,0xFF)

    bdoc = _BD()
    bsec = bdoc.sections[0]
    bsec.page_width = Inches(8.5); bsec.page_height = Inches(11)
    bsec.top_margin = bsec.bottom_margin = Inches(1.0)
    bsec.left_margin = bsec.right_margin = Inches(1.25)

    # Title page
    tp = bdoc.add_paragraph()
    tr = tp.add_run("FEDERAL AVIATION ADMINISTRATION")
    tr.bold = True; tr.font.size = Pt(9); tr.font.color.rgb = MED_BLUE

    tp2 = bdoc.add_paragraph()
    tr2 = tp2.add_run("ITU-R Batch Contribution Analysis Report")
    tr2.bold = True; tr2.font.size = Pt(18); tr2.font.color.rgb = FAA_BLUE

    tp3 = bdoc.add_paragraph()
    tr3 = tp3.add_run(f"Working Party: {working_party}  |  Depth: {analysis_depth}  |  Documents: {len(batch_results)}  |  Date: {_bdate.today()}")
    tr3.font.size = Pt(9); tr3.italic = True; tr3.font.color.rgb = GRAY

    bdoc.add_paragraph()

    # Triage table
    th = bdoc.add_paragraph()
    thr = th.add_run("TRIAGE SUMMARY")
    thr.bold = True; thr.font.size = Pt(13); thr.font.color.rgb = FAA_BLUE

    cols_t = ["File", "Verdict", "Doc Status", "Proposed Freq"]
    tbl = bdoc.add_table(rows=len(triage_rows)+1, cols=len(cols_t))
    tbl.style = "Table Grid"
    for ci, ch in enumerate(cols_t):
        cell = tbl.rows[0].cells[ci]; cell.paragraphs[0].clear()
        r2 = cell.paragraphs[0].add_run(ch)
        r2.bold = True; r2.font.size = Pt(9); r2.font.color.rgb = WHITE
        try:
            shd = _OE('w:shd'); shd.set(_qn('w:fill'),'1F4E79'); shd.set(_qn('w:val'),'clear')
            cell._element.get_or_add_tcPr().append(shd)
        except Exception: pass
    for ri, row in enumerate(triage_rows):
        for ci, key in enumerate(cols_t):
            cell = tbl.rows[ri+1].cells[ci]; cell.paragraphs[0].clear()
            val = str(row.get(key, ""))
            r3 = cell.paragraphs[0].add_run(val)
            r3.font.size = Pt(9)
            if "HUMAN REVIEW" in val:  r3.font.color.rgb = RED
            elif "NOT RELEVANT" in val: r3.font.color.rgb = GREEN
            else: r3.font.color.rgb = DARK

    bdoc.add_paragraph()

    # Individual analyses
    for res in batch_results:
        bdoc.add_page_break()

        fp = bdoc.add_paragraph()
        fr = fp.add_run(f"Document: {res['file']}")
        fr.bold = True; fr.font.size = Pt(14); fr.font.color.rgb = FAA_BLUE

        if res.get("tc") and "TRACK CHANGES" in res.get("tc",""):
            tp4 = bdoc.add_paragraph()
            tr4 = tp4.add_run(res["tc"][:200])
            tr4.font.size = Pt(8); tr4.italic = True; tr4.font.color.rgb = GRAY

        bdoc.add_paragraph()

        # Render analysis lines
        for line in res["analysis"].split("\n"):
            ls = line.strip()
            if not ls: continue
            if ls.startswith("## "):
                p2 = bdoc.add_paragraph(); r4 = p2.add_run(ls[3:])
                r4.bold = True; r4.font.size = Pt(12); r4.font.color.rgb = FAA_BLUE
            elif ls.startswith("### "):
                p2 = bdoc.add_paragraph(); r4 = p2.add_run(ls[4:])
                r4.bold = True; r4.font.size = Pt(10); r4.font.color.rgb = MED_BLUE
            elif ls.startswith("| ") and "|" in ls[2:]:
                p2 = bdoc.add_paragraph()
                r4 = p2.add_run(ls); r4.font.size = Pt(8); r4.font.color.rgb = DARK
            elif ls.startswith("- ") or ls.startswith("* "):
                p2 = bdoc.add_paragraph(style="List Bullet")
                text2 = ls[2:]
                parts2 = _bre.split(r'\*\*(.+?)\*\*', text2)
                for pi, pt in enumerate(parts2):
                    if not pt: continue
                    r4 = p2.add_run(pt); r4.font.size = Pt(9.5); r4.bold = (pi%2==1)
                    c2 = DARK
                    if "⚠️" in pt or "UNVERIFIED" in pt: c2 = RED
                    r4.font.color.rgb = c2
            else:
                p2 = bdoc.add_paragraph()
                parts3 = _bre.split(r'\*\*(.+?)\*\*', ls)
                for pi, pt in enumerate(parts3):
                    if not pt: continue
                    r4 = p2.add_run(pt); r4.font.size = Pt(9.5); r4.bold = (pi%2==1)
                    c3 = DARK
                    if "REQUIRES HUMAN REVIEW" in pt: c3 = RED
                    elif "NOT RELEVANT" in pt: c3 = GREEN
                    r4.font.color.rgb = c3

    bio = _bio.BytesIO(); bdoc.save(bio); bio.seek(0)
    return bio.getvalue()


def noise_floor_dbm(bandwidth_hz: float, noise_figure_db: float) -> float:
    """Thermal noise floor: N = -174 + 10log10(BW) + NF"""
    return -174.0 + 10.0 * np.log10(bandwidth_hz) + noise_figure_db

def free_space_path_loss_db(freq_mhz: float, dist_km: float) -> float:
    """FSPL(dB) = 20log10(d_km) + 20log10(f_MHz) + 32.44"""
    return 20.0 * np.log10(dist_km) + 20.0 * np.log10(freq_mhz) + 32.44

def received_power_dbm(tx_power_dbm, tx_gain_dbi, path_loss_db,
                        rx_gain_dbi, cable_loss_db=0.0) -> float:
    """Friis: Pr = Pt + Gt - L - cable + Gr"""
    return tx_power_dbm + tx_gain_dbi - path_loss_db + rx_gain_dbi - cable_loss_db

def in_ratio_db(interference_dbm: float, noise_floor_dbm: float) -> float:
    """I/N = interference power relative to receiver noise floor."""
    return interference_dbm - noise_floor_dbm

def protection_margin_db(in_ratio: float, threshold_db: float) -> float:
    """Margin = threshold - I/N.  Positive = protected, negative = violated."""
    return threshold_db - in_ratio

def eirp_dbm(tx_power_dbm: float, tx_gain_dbi: float) -> float:
    """EIRP = transmit power + antenna gain (dBm + dBi)."""
    return tx_power_dbm + tx_gain_dbi

def pfd_dbm_per_m2(eirp_dbm_val: float, dist_km: float, freq_mhz: float) -> float:
    """
    Power flux density at receiver (dBW/m²).
    PFD(dBW/m²) = EIRP(dBW) - 10·log10(4π·d²)
    This is frequency-independent — it only depends on EIRP and distance.
    freq_mhz retained in signature for API compatibility but not used.
    """
    eirp_dbw = eirp_dbm_val - 30.0          # dBm → dBW
    d_m      = dist_km * 1000.0             # km → m
    return eirp_dbw - 10.0 * np.log10(4.0 * np.pi * d_m**2)

# ─── Simplified ITU-R P.452 terrestrial interference model ───────────────────
def p452_basic_loss_db(freq_mhz: float, dist_km: float,
                        tx_height_m: float, rx_height_m: float,
                        terrain_type: str = "suburban") -> float:
    """
    Simplified P.452 basic transmission loss (line-of-sight + diffraction).
    Full P.452 requires terrain profiles; this uses FSPL + empirical corrections.
    For rigorous analysis use the ITU-R SoftTools P.452 implementation.
    """
    fspl = free_space_path_loss_db(freq_mhz, dist_km)
    # Effective earth radius factor
    k = 4.0 / 3.0  # standard effective earth radius
    Re = 6371.0 * k  # km
    # Fresnel zone clearance check (simplified)
    d1, d2 = dist_km / 2, dist_km / 2
    f_hz = freq_mhz * 1e6
    r1 = np.sqrt(3e8 * d1 * 1e3 * d2 * 1e3 / (f_hz * dist_km * 1e3))
    # Height gain correction (log)
    ht_correction = -10.0 * np.log10(tx_height_m / 10.0) - 10.0 * np.log10(rx_height_m / 10.0)
    terrain_corr = {"open": 0, "suburban": 8, "urban": 18, "dense_urban": 26}
    tc = terrain_corr.get(terrain_type, 8)
    return fspl + tc + ht_correction * 0.5

# ─── Simplified ITU-R P.528 aeronautical propagation ─────────────────────────
def p528_aero_path_loss_db(freq_mhz: float, ground_dist_km: float,
                            aircraft_alt_km: float,
                            time_percent: float = 50.0) -> float:
    """
    Simplified P.528 aeronautical path loss (airborne receiver).
    P.528 accounts for: FSPL + atmospheric absorption + troposcatter.
    Uses itur gaseous attenuation for atmospheric component.
    Full P.528 requires detailed marine-layer / refractive corrections.
    """
    # Slant distance
    slant_km = np.sqrt(ground_dist_km**2 + aircraft_alt_km**2)
    fspl = free_space_path_loss_db(freq_mhz, slant_km)

    # Gaseous attenuation via itur P.676
    try:
        elev_angle_deg = np.degrees(np.arctan(aircraft_alt_km / max(ground_dist_km, 0.1)))
        # itur gaseous_attenuation_terrestrial_path for horizontal path approximation
        if freq_mhz >= 1000:  # Above 1 GHz, gaseous matters more
            gamma = itu676.gaseous_attenuation_terrestrial_path(
                freq_mhz / 1e3,  # GHz
                p=50,
                T=15,
                H=50,
                P=1013.25,
                d=slant_km,
                mode="approx"
            )
            gas_atten_db = float(gamma.value) if hasattr(gamma, 'value') else float(gamma)
        else:
            gas_atten_db = 0.0
    except Exception:
        gas_atten_db = 0.0

    # Time-variability factor (simplified log-normal)
    t_factor = 0.0
    if time_percent < 50:
        t_factor = -5.0 * np.log10(time_percent / 50.0)

    return fspl + gas_atten_db + t_factor

# ─── Monte Carlo Aggregate Interference ──────────────────────────────────────
def monte_carlo_aggregate(
    n_interferers: int,
    tx_power_dbm: float,
    tx_gain_dbi: float,
    freq_mhz: float,
    deployment_radius_km: float,
    rx_noise_floor_dbm: float,
    in_threshold_db: float,
    n_trials: int = 5000,
    model: str = "FSPL",
    terrain_type: str = "suburban",
    aircraft_alt_km: float = 0.0,
) -> dict:
    """
    Monte Carlo simulation of aggregate interference from N random interferers.
    Each trial: randomly place N transmitters, compute total I at victim receiver.
    Returns distribution statistics and exceedance probability.
    """
    violations = 0
    in_values = []
    agg_power_list = []

    for _ in range(n_trials):
        # Random distances: uniform in area → radius ~ sqrt(uniform) * max_radius
        r = np.sqrt(np.random.uniform(0.01**2, deployment_radius_km**2, n_interferers))
        # Small random power variation ±2 dB (log-normal transmitter variation)
        pwr_variation = np.random.normal(0, 2.0, n_interferers)
        effective_eirp = eirp_dbm(tx_power_dbm, tx_gain_dbi) + pwr_variation

        # Path loss per interferer
        path_losses = np.array([
            p528_aero_path_loss_db(freq_mhz, ri, aircraft_alt_km)
            if (model == "P.528" and aircraft_alt_km > 0)
            else p452_basic_loss_db(freq_mhz, ri, 30.0, 5.0, terrain_type)
            if model == "P.452"
            else free_space_path_loss_db(freq_mhz, ri)
            for ri in r
        ])

        rx_powers_dbm = effective_eirp - path_losses
        # Linear sum of interferers
        agg_power_mw = np.sum(10.0 ** (rx_powers_dbm / 10.0))
        agg_power_dbm = 10.0 * np.log10(max(agg_power_mw, 1e-30))

        i_n = in_ratio_db(agg_power_dbm, rx_noise_floor_dbm)
        in_values.append(i_n)
        agg_power_list.append(agg_power_dbm)
        if i_n > in_threshold_db:
            violations += 1

    in_values = np.array(in_values)
    return {
        "in_mean_db": float(np.mean(in_values)),
        "in_p50_db": float(np.percentile(in_values, 50)),
        "in_p95_db": float(np.percentile(in_values, 95)),
        "in_p99_db": float(np.percentile(in_values, 99)),
        "violation_probability": violations / n_trials,
        "in_values": in_values,
        "agg_power_list": np.array(agg_power_list),
        "n_trials": n_trials,
    }

# ─────────────────────────────────────────────────────────────────────────────
# SIDEBAR
# ─────────────────────────────────────────────────────────────────────────────
st.sidebar.image("https://upload.wikimedia.org/wikipedia/commons/thumb/1/17/FAA_seal.svg/200px-FAA_seal.svg.png", width=80)
st.sidebar.title("FAA RF Interference\nAnalysis Tool")
st.sidebar.markdown("*For ITU-R WP 5D / 5B / 4C / 7B / 7C Policy Support*")

# ── User info bar ─────────────────────────────────────────────────────────────
user = current_user()
role_badge = "🔴 Admin" if is_admin() else "🟢 User"
st.sidebar.markdown(
    f"<div style='background:#1a2a1a;border-left:3px solid #44bb44;"
    f"padding:6px 10px;border-radius:4px;font-size:0.82em;color:#aaffaa'>"
    f"👤 <b>{user.get('name','')}</b> &nbsp;|&nbsp; {role_badge}</div>",
    unsafe_allow_html=True
)
if st.sidebar.button("🚪 Sign Out", use_container_width=True):
    logout()

st.sidebar.markdown("---")

tab_names = [
    "🤖 Contribution Analyzer",
    "📋 Contribution Summary",
    "📓 Meeting Notes",
    "🔬 Contribution Code Analyzer",
    "📡 Protected Bands",
    "🔗 Link Budget",
    "📊 Noise & I/N",
    "🌐 Propagation",
    "🎲 Monte Carlo",
    "📚 Tutorial",
    "🎓 RF Training",
    "📖 Glossary",
    "📻 Microwave Link Budget",
]

# Admin-only tab
if is_admin():
    tab_names.append("⚙️ Admin Panel")
# Handle programmatic navigation (e.g. from tutorial "Go to module" buttons)
if "_nav_to" in st.session_state:
    _nav_target = st.session_state.pop("_nav_to")
    selected_tab = st.sidebar.radio("Module", tab_names, index=tab_names.index(_nav_target) if _nav_target in tab_names else 0)
else:
    selected_tab = st.sidebar.radio("Module", tab_names)

st.sidebar.markdown("---")
st.sidebar.markdown("""
**Key References**
- ITU-R P.452 — Terrestrial interference
- ITU-R P.528 — Aeronautical propagation  
- ITU-R M.1642 — IMT → ARNS methodology
- RTCA DO-235B — GNSS protection
- RR No. 4.10 — Safety-of-life protection
""")

# ─────────────────────────────────────────────────────────────────────────────
# TAB 1 — PROTECTED BANDS DATABASE
# ─────────────────────────────────────────────────────────────────────────────
if selected_tab == "📡 Protected Bands":
    st.title("📡 FAA Protected Frequency Bands")
    ex("Each band carries a primary or secondary allocation under the ITU Radio Regulations — knowing the allocation type and footnote structure determines which regulatory instruments you can invoke.")

    # Spectrum overview chart
    st.subheader("Spectrum Overview")
    ex("Frequency proximity matters: receiver front-end selectivity is finite — a new allocation 200 MHz away at C-band can still cause LNA desensitization via blocking, even with no spectral overlap.")

    band_colors = [
        "#4fc3f7","#81c784","#ffb74d","#e57373",
        "#ce93d8","#4db6ac","#fff176","#ff8a65",
        "#a5d6a7","#90caf9","#f48fb1","#80cbc4",
    ]

    band_list = list(FAA_BANDS.items())
    n_bands   = len(band_list)

    ROW_H    = 1.10
    ROW_GAP  = 0.30
    ROW_STEP = ROW_H + ROW_GAP
    FIG_H    = n_bands * ROW_STEP + 1.5

    plt.rcParams['figure.dpi']   = 180
    plt.rcParams['font.family']  = 'DejaVu Sans'
    plt.rcParams['text.antialiased'] = True

    fig, ax = plt.subplots(figsize=(20, FIG_H))
    fig.patch.set_facecolor("#0e1117")
    ax.set_facecolor("#0e1117")

    log_min, log_max = np.log10(90), np.log10(12000)
    min_log_w = (log_max - log_min) * 0.018

    for kf in [100, 200, 300, 500, 1000, 2000, 3000, 5000, 10000]:
        ax.axvline(kf, color='#1e1e2e', linewidth=0.8, zorder=0)

    for i, (name, b) in enumerate(band_list):
        col   = band_colors[i % len(band_colors)]
        row_y = (n_bands - 1 - i) * ROW_STEP

        row_bg = "#111118" if i % 2 == 0 else "#0e0e15"
        ax.fill_betweenx(
            [row_y, row_y + ROW_H],
            10**log_min, 10**log_max,
            color=row_bg, linewidth=0, zorder=1
        )

        log_fl = np.log10(b["f_low_mhz"])
        log_fh = np.log10(b["f_high_mhz"])
        if (log_fh - log_fl) < min_log_w:
            lm = (log_fl + log_fh) / 2
            log_fl = lm - min_log_w / 2
            log_fh = lm + min_log_w / 2

        bar_bot = row_y + ROW_H * 0.10
        bar_top = row_y + ROW_H * 0.90

        ax.fill_betweenx([bar_bot, bar_top], 10**log_fl, 10**log_fh,
                         color=col, alpha=0.92, linewidth=0, zorder=3)
        ax.plot([10**log_fl, 10**log_fh, 10**log_fh, 10**log_fl, 10**log_fl],
                [bar_bot, bar_bot, bar_top, bar_top, bar_bot],
                color='white', linewidth=1.0, alpha=0.55, zorder=4)

        log_mid = (log_fl + log_fh) / 2
        bar_cy  = (bar_bot + bar_top) / 2
        ax.text(10**log_mid, bar_cy, str(i+1),
                ha='center', va='center',
                fontsize=10, fontweight='bold',
                color='#0a0a0a', zorder=5)

        label_x = 10**log_min * 1.012

        # Band name — large, bold, bright
        ax.text(label_x, bar_cy + 0.13, name,
                ha='left', va='center',
                fontsize=13, fontweight='bold',
                color=col, zorder=5)

        # Frequency range — clearly readable grey
        ax.text(label_x, bar_cy - 0.16,
                f"{b['f_low_mhz']:.0f} – {b['f_high_mhz']:.0f} MHz",
                ha='left', va='center',
                fontsize=11, color='#cccccc', zorder=5)

        # I/N threshold — large, prominent
        sf = b.get("aviation_safety_factor_db", 0)
        thresh_str = f"I/N  {b['in_threshold_db']} dB"
        if sf > 0:
            thresh_str += f"   +{sf} dB safety"
        ax.text(10**log_max * 0.983, bar_cy + 0.13,
                thresh_str,
                ha='right', va='center',
                fontsize=12, fontweight='bold',
                color='#ffd966' if sf > 0 else '#77ff99', zorder=5)

        # Service category
        svc = b.get("service_category", b["allocation"]).split(" —")[0][:26]
        ax.text(10**log_max * 0.983, bar_cy - 0.16,
                svc,
                ha='right', va='center',
                fontsize=10, color='#999999', zorder=5)

        ax.axhline(row_y, color='#1a1a2a', linewidth=0.6, zorder=2)

    ax.set_xscale("log")
    ax.set_xlim(10**log_min, 10**log_max)
    ax.set_ylim(-0.2, n_bands * ROW_STEP + 0.3)
    ax.set_yticks([])

    key_ticks = [100, 200, 330, 500, 960, 1090, 1176, 1575, 2000, 2800, 4300, 5000, 9375]
    ax.set_xticks(key_ticks)
    ax.set_xticklabels([str(k) for k in key_ticks],
                       fontsize=11, fontweight='bold',
                       color='#cccccc', rotation=45, ha='right')

    ax.tick_params(axis='x', colors='#555', pad=6, length=5)
    ax.spines['bottom'].set_color('#555')
    for sp in ['top','left','right']:
        ax.spines[sp].set_visible(False)

    ax.set_title("FAA Protected Aeronautical Frequency Bands",
                 color='white', fontsize=17, fontweight='bold', pad=14, loc='left')
    ax.set_xlabel("Frequency (MHz) — log scale", color='#666', fontsize=9, labelpad=6)

    plt.tight_layout(rect=[0, 0, 1, 0.98])
    st.pyplot(fig)

    # Compact legend strip
    leg_cols = st.columns(min(6, n_bands))
    for i, (name, b_leg) in enumerate(band_list):
        col_leg = leg_cols[i % 6]
        col_leg.markdown(
            f"<div style='border-left:3px solid {band_colors[i % len(band_colors)]};padding:2px 6px;"
            f"margin:1px 0;font-size:0.75em;color:#ccc'>"
            f"<b style='color:{band_colors[i % len(band_colors)]}'>{i+1}</b> {name.split('/')[0].strip()}"
            f"</div>",
            unsafe_allow_html=True
        )
    st.markdown("")

    st.subheader("Band Details")
    ex("I/N threshold is defined such that the total noise floor rise stays below ~1 dB — at I/N = −6 dB the noise power increases by 10·log(1 + 10^(−0.6)) ≈ 0.97 dB; at −10 dB it is 0.41 dB.")

    selected_band = st.selectbox("Select a band for details:", list(FAA_BANDS.keys()))
    b = FAA_BANDS[selected_band]

    col1, col2, col3, col4 = st.columns(4)
    col1.metric("Lower Freq", f"{b['f_low_mhz']} MHz",
        help="Bottom edge of the protected band in Megahertz (MHz) — new allocations must not encroach below this frequency.")
    col2.metric("Upper Freq", f"{b['f_high_mhz']} MHz",
        help="Top edge of the protected band in Megahertz (MHz) — out-of-band emissions from systems above this frequency can still cause Low Noise Amplifier (LNA) desensitization or blocking in the protected receiver.")
    col3.metric("Bandwidth", f"{b['f_high_mhz'] - b['f_low_mhz']:.1f} MHz",
        help="Total protected bandwidth in Megahertz (MHz). Wider bandwidth raises the thermal noise floor (noise power = kTBF), making the receiver less sensitive — but also more spectrum to defend.")
    col4.metric("I/N Threshold", f"{b['in_threshold_db']} dB",
        help="Interference-to-Noise (I/N) ratio threshold — the maximum tolerable interference power relative to the receiver noise floor. Exceeding this threshold is the basis for citing harmful interference under Radio Regulations No. 4.10.")

    col5, col6, col7 = st.columns(3)
    col5.metric("Allocation", b["allocation"],
        help="ITU Radio Regulations allocation type. ARNS = Aeronautical Radionavigation Service. RNSS = Radionavigation Satellite Service. ANS = Aeronautical Navigation Service. Primary allocation has stronger protection than secondary.")
    col6.metric("Noise Floor (est.)", f"{b['noise_floor_dbm']} dBm",
        help="Estimated receiver thermal noise floor (kTBF). Verify exact value against applicable RTCA Minimum Operational Performance Standard (MOPS) before citing in a contribution.")
    col7.metric("RTCA Standard", b["rtca_doc"],
        help="Radio Technical Commission for Aeronautics (RTCA) Minimum Operational Performance Standard governing this receiver type. Primary source for protection criteria cited in ITU-R contributions.")

    # New fields from slides
    col8, col9, col10 = st.columns(3)
    safety_factor = b.get("aviation_safety_factor_db", 0)
    effective_thresh = b.get("effective_threshold_db", b["in_threshold_db"])
    col8.metric("Aviation Safety Factor", f"{safety_factor} dB" if safety_factor > 0 else "0 dB (N/A)",
        help="Per FAA slides: 'Some aeronautical applications (e.g. precision approach and landing) are critical, meriting an additional safety factor of not less than 6 dB.' Applies on top of the base I/N threshold.")
    col9.metric("Effective Threshold (incl. safety factor)", f"{effective_thresh} dB",
        help="Base I/N threshold plus the aviation safety factor. This is the operationally applied protection criterion for safety-critical precision approach and landing applications.")
    col10.metric("Service Category (ITU-R)", b.get("service_category", b["allocation"]),
        help="ITU service category. AM(R)S = Aeronautical Mobile (Route) Service — safety and regularity. AMS(R)S = satellite variant. ARNS = Aeronautical Radionavigation Service.")

    st.markdown(f"**System:** {b['system']}")
    st.markdown(f"**Notes:** {b['notes']}")
    st.markdown(f"**Protection Basis:** {b.get('protection_basis', '—')}")

    # Special protection values (ΔT/T, PFD, epfd)
    special_vals = []
    if b.get("delta_t_t_pct_aggregate"):
        special_vals.append(f"ΔT/T aggregate: {b['delta_t_t_pct_aggregate']}%")
    if b.get("delta_t_t_pct_single"):
        special_vals.append(f"ΔT/T single-entry: {b['delta_t_t_pct_single']}%")
    if b.get("psd_threshold_dbw_mhz"):
        special_vals.append(f"PSD threshold: {b['psd_threshold_dbw_mhz']} dBW/MHz")
    if b.get("epfd_threshold_dbw_m2_mhz"):
        special_vals.append(f"epfd limit: {b['epfd_threshold_dbw_m2_mhz']} dBW/m²/MHz")
    if b.get("pfd_threshold_dbw_m2_khz"):
        special_vals.append(f"pfd limit: {b['pfd_threshold_dbw_m2_khz']} dBW/m² in 150 kHz")
    if special_vals:
        st.markdown("**Additional Protection Criteria:** " + "  |  ".join(special_vals))

    if b.get("rr_1_59"):
        st.markdown("🔴 **Safety Service (RR 1.59)** — any radiocommunication service used permanently or temporarily for the safeguarding of human life and property. Protected under RR No. 4.10.")

    # Source-Path-Receiver model
    with st.expander("🔍 Source-Path-Receiver (SPR) Analysis Framework for this band"):
        ex("The SPR model is the formal FAA methodology for aviation interference analysis (per FAA slides). Worst-case limits must be applied on all three elements.")
        spr1, spr2, spr3 = st.columns(3)
        with spr1:
            st.markdown(f"""
<div style='background:#2a1a0a;border-left:4px solid #ff8844;padding:10px;border-radius:4px'>
<b style='color:#ff8844'>📡 SOURCE (Interferer)</b><br>
<span style='color:#ffddaa;font-size:0.85em'>{b.get('spr_source','Max Tx power; worst-case antenna gain; signal characteristics')}</span>
</div>""", unsafe_allow_html=True)
        with spr2:
            st.markdown(f"""
<div style='background:#0a0a2a;border-left:4px solid #4488ff;padding:10px;border-radius:4px'>
<b style='color:#4488ff'>🌐 PATH (Propagation)</b><br>
<span style='color:#aaddff;font-size:0.85em'>{b.get('spr_path','FSPL worst-case; free-space attenuation especially >1 GHz; distance separation >20 km')}</span>
</div>""", unsafe_allow_html=True)
        with spr3:
            st.markdown(f"""
<div style='background:#0a2a0a;border-left:4px solid #44bb44;padding:10px;border-radius:4px'>
<b style='color:#44bb44'>📻 VICTIM (Receiver)</b><br>
<span style='color:#aaffaa;font-size:0.85em'>{b.get('spr_victim','Receiver susceptibility mask; antenna gain; receiver noise power')}</span>
</div>""", unsafe_allow_html=True)
        st.markdown("")
        st.caption("Per FAA guidance: use worst-case limits on all aspects. Aggregate effect of multiple interference sources must be considered. Max interference threshold is used as protection criteria taking into account all environmental conditions.")

    # Full table — updated with new columns
    st.subheader("All Bands Summary Table")
    ex("Now includes aviation safety factor and effective threshold per FAA system protection levels. Noise floor estimates are representative — verify against applicable RTCA MOPS.")
    rows = []
    for name, b_row in FAA_BANDS.items():
        rows.append({
            "Band": name,
            "Low (MHz)": b_row["f_low_mhz"],
            "High (MHz)": b_row["f_high_mhz"],
            "BW (MHz)": round(b_row["f_high_mhz"] - b_row["f_low_mhz"], 3),
            "Allocation": b_row["allocation"],
            "I/N Threshold (dB)": b_row["in_threshold_db"],
            "Aviation Safety Factor (dB)": b_row.get("aviation_safety_factor_db", 0),
            "Effective Threshold (dB)": b_row.get("effective_threshold_db", b_row["in_threshold_db"]),
            "Noise Floor (dBm)": b_row["noise_floor_dbm"],
            "RR 1.59 Safety Service": "✅" if b_row.get("rr_1_59") else "—",
            "RTCA": b_row["rtca_doc"],
        })
    st.dataframe(pd.DataFrame(rows), use_container_width=True)

# ─────────────────────────────────────────────────────────────────────────────
# ─────────────────────────────────────────────────────────────────────────────
# TAB 2 — LINK BUDGET
# ─────────────────────────────────────────────────────────────────────────────
elif selected_tab == "🔗 Link Budget":
    st.title("🔗 Link Budget Calculator")
    ex("The Friis transmission equation: Pr = Pt + Gt − L_tx_total − FSPL − L_rx_total + Gr. Every loss field on both Tx and Rx sides defaults to zero — enter only what applies to your scenario.")

    col_l, col_r = st.columns([1, 1])

    with col_l:
        st.subheader("📡 Transmitter (Interferer)")
        ex("For worst-case analysis use maximum authorized EIRP. EIRP = Pt + Gt − total Tx losses.")
        tx_power_dbm     = st.number_input("Tx Power (dBm)", value=43.0, step=1.0,
            help="Radio PA output power. 43 dBm = 20 W typical LTE base station; 30 dBm = 1 W small cell.")
        tx_gain_dbi      = st.number_input("Tx Antenna Gain (dBi)", value=15.0, step=0.5,
            help="Peak gain toward victim receiver. Use 0 dBi for isotropic worst-case bound.")
        tx_height_m      = st.number_input("Tx Height (m AGL)", value=30.0, step=5.0)

        st.markdown("**Tx-Side Losses** *(enter 0 if not applicable)*")
        ex("All losses between the PA output and the radiated signal. Each field defaults to 0 — enter only what is present in your scenario.")
        tx_cable_loss     = st.number_input("Tx Cable / Coax Loss (dB)", value=2.0, step=0.5,
            help="Transmission line loss between PA output and antenna connector. Typical: 0.5–3 dB.")
        tx_connector_loss = st.number_input("Tx Connector / Adapter Loss (dB)", value=0.0, step=0.1,
            help="Each RF connector contributes ~0.1–0.2 dB. Sum all connectors in the Tx path.")
        tx_branching_loss = st.number_input("Tx Branching / Hybrid / Combiner Loss (dB)", value=0.0, step=0.5,
            help="Diplexer, combiner, or hybrid coupler sharing one antenna. Typical: 3–4 dB for a 2-way combiner.")
        tx_filter_loss    = st.number_input("Tx Filter / Duplexer Loss (dB)", value=0.0, step=0.5,
            help="Bandpass filter or duplexer insertion loss. Used to limit OOB and spurious emissions. Typical: 0.5–2 dB.")
        tx_jumper_loss    = st.number_input("Tx Jumper / Pigtail Loss (dB)", value=0.0, step=0.1,
            help="Short flexible cable sections. Often overlooked — can add 0.2–0.5 dB.")
        tx_misc_loss      = st.number_input("Tx Other Losses (dB)", value=0.0, step=0.1,
            help="Circulators, lightning arrestors, splitters, or any other Tx-side hardware.")
        tx_total_loss = round(tx_cable_loss + tx_connector_loss + tx_branching_loss +
                              tx_filter_loss + tx_jumper_loss + tx_misc_loss, 2)
        st.metric("Total Tx Losses", f"{tx_total_loss:.2f} dB",
            help="Sum of all Tx-side losses. EIRP = Pt + Gt − this value.")

        st.markdown("---")
        st.subheader("🌐 Channel / Propagation")
        ex("FSPL is the most optimistic model — least loss, worst-case interference. Always run FSPL first. If it shows compatibility, the system is protected.")
        propagation_model = st.selectbox("Propagation Model",
            ["Free Space (FSPL)", "P.452 (Terrestrial)", "P.528 (Aeronautical)"])
        freq_mhz  = st.number_input("Frequency (MHz)", value=4300.0, step=10.0,
            help="Center frequency of the interfering emission.")
        dist_km   = st.number_input("Distance (km)", value=5.0, step=0.5)

        if propagation_model == "P.452 (Terrestrial)":
            terrain_type    = st.selectbox("Terrain / Clutter Type",
                ["open", "suburban", "urban", "dense_urban"])
            rx_height_m     = st.number_input("Rx Height (m AGL)", value=5.0, step=1.0)
            aircraft_alt_km = 0.0
        elif propagation_model == "P.528 (Aeronautical)":
            aircraft_alt_km = st.number_input("Aircraft Altitude (km)", value=3.0, step=0.5)
            rx_height_m     = 5.0
            terrain_type    = "suburban"
        else:
            rx_height_m     = 5.0
            terrain_type    = "suburban"
            aircraft_alt_km = 0.0

    with col_r:
        st.subheader("📻 Receiver (Victim)")
        ex("Use 0 dBi Rx gain for worst-case. Assumed directivity away from interferer is a common proponent tactic to challenge if not supported by a measured antenna pattern.")
        rx_gain_dbi = st.number_input("Rx Antenna Gain (dBi)", value=0.0, step=0.5,
            help="Toward interferer. 0 dBi = isotropic worst-case.")

        st.markdown("**— or auto-fill from FAA band —**")
        band_select = st.selectbox("Auto-fill from FAA band:", ["(manual)"] + list(FAA_BANDS.keys()))
        if band_select != "(manual)":
            b_sel = FAA_BANDS[band_select]
            rx_noise_floor_dbm_input = b_sel["noise_floor_dbm"]
            in_threshold_db = b_sel["in_threshold_db"]
            sf_loaded = b_sel.get("aviation_safety_factor_db", 0)
            eff_thresh = b_sel.get("effective_threshold_db", in_threshold_db)
            st.info(f"Loaded: Noise floor = {rx_noise_floor_dbm_input} dBm | I/N threshold = {in_threshold_db} dB"
                    + (f" | +{sf_loaded} dB aviation safety factor → effective {eff_thresh} dB" if sf_loaded else ""))
        else:
            rx_bw_mhz   = st.number_input("Rx Bandwidth (MHz)", value=100.0, step=10.0)
            rx_nf_db    = st.number_input("Receiver Noise Figure (dB)", value=5.0, step=0.5)
            rx_noise_floor_dbm_input = noise_floor_dbm(rx_bw_mhz * 1e6, rx_nf_db)
            in_threshold_db = st.number_input("I/N Threshold (dB)", value=-6.0, step=1.0,
                help="Protection criterion: −6 dB typical ARNS, −10 dB for GNSS/ADS-B.")
            sf_loaded = 0
            eff_thresh = in_threshold_db

        st.markdown("**Rx-Side Losses** *(enter 0 if not applicable)*")
        ex("All losses between the antenna and receiver input. Every dB here degrades sensitivity directly — Rx losses are just as damaging as the same amount of path loss.")
        rx_cable_loss     = st.number_input("Rx Cable / Coax Loss (dB)", value=0.5, step=0.5,
            help="Transmission line from antenna to LNA or receiver. Keep short in avionics — 1 dB here = 1 dB sensitivity loss.")
        rx_connector_loss = st.number_input("Rx Connector / Adapter Loss (dB)", value=0.0, step=0.1,
            help="Each connector in the Rx path adds ~0.1–0.2 dB.")
        rx_branching_loss = st.number_input("Rx Branching / Splitter Loss (dB)", value=0.0, step=0.5,
            help="If the antenna feeds multiple receivers via a splitter or diplexer. Typical: 3–4 dB for 2-way split.")
        rx_filter_loss    = st.number_input("Rx Filter / Duplexer Loss (dB)", value=0.0, step=0.5,
            help="Bandpass filter insertion loss on receive path. Rx filters protect the LNA but add loss before it — directly degrades noise figure.")
        rx_limiter_loss   = st.number_input("Rx Limiter / Protection Device Loss (dB)", value=0.0, step=0.1,
            help="PIN diode limiters or gas discharge tubes protecting the LNA. Typical insertion loss: 0.3–1.0 dB.")
        rx_misc_loss      = st.number_input("Rx Other Losses (dB)", value=0.0, step=0.1,
            help="Lightning arrestors, bias tees, long coax runs, or any other Rx-side hardware losses.")
        rx_total_loss = round(rx_cable_loss + rx_connector_loss + rx_branching_loss +
                              rx_filter_loss + rx_limiter_loss + rx_misc_loss, 2)
        st.metric("Total Rx Losses", f"{rx_total_loss:.2f} dB",
            help="Sum of all Rx-side losses. These reduce the effective signal reaching the receiver just like path loss.")

    # ── Compute ───────────────────────────────────────────────────────────────
    st.markdown("---")
    if st.button("⚡ Run Link Budget", type="primary"):
        eir = round(tx_power_dbm + tx_gain_dbi - tx_total_loss, 2)

        if propagation_model == "Free Space (FSPL)":
            pl = free_space_path_loss_db(freq_mhz, dist_km)
            model_label = "FSPL"
        elif propagation_model == "P.452 (Terrestrial)":
            pl = p452_basic_loss_db(freq_mhz, dist_km, tx_height_m, rx_height_m, terrain_type)
            model_label = f"P.452 ({terrain_type})"
        else:
            pl = p528_aero_path_loss_db(freq_mhz, dist_km, aircraft_alt_km)
            model_label = f"P.528 (alt={aircraft_alt_km} km)"

        rx_pwr = round(eir - pl + rx_gain_dbi - rx_total_loss, 2)
        i_n    = in_ratio_db(rx_pwr, rx_noise_floor_dbm_input)
        margin = protection_margin_db(i_n, in_threshold_db)
        use_eff   = sf_loaded > 0
        eff_margin = protection_margin_db(i_n, eff_thresh) if use_eff else None

        col1, col2, col3, col4, col5 = st.columns(5)
        col1.metric("EIRP", f"{eir:.2f} dBm",
            help=f"Pt({tx_power_dbm}) + Gt({tx_gain_dbi}) − Tx losses({tx_total_loss})")
        col2.metric("Path Loss", f"{pl:.2f} dB", help=model_label)
        col3.metric("Rx Power", f"{rx_pwr:.2f} dBm",
            help=f"EIRP({eir}) − PL({pl:.2f}) + Gr({rx_gain_dbi}) − Rx losses({rx_total_loss})")
        col4.metric("I/N", f"{i_n:.2f} dB")
        col5.metric("Protection Margin", f"{margin:.2f} dB",
                    delta=f"{margin:.2f} dB",
                    delta_color="normal" if margin >= 0 else "inverse")

        if use_eff:
            st.info(f"⚠️ Aviation safety factor +{sf_loaded} dB applies. Effective threshold = {eff_thresh} dB → Effective margin = **{eff_margin:.2f} dB** ({'protected' if eff_margin >= 0 else 'VIOLATED'})")

        ex("Protection Margin = threshold − I/N. Negative = threshold violated = grounds to cite harmful interference under RR No. 4.10.")

        if margin >= 10:
            ok(f"PROTECTED with {margin:.2f} dB margin.")
        elif margin >= 0:
            warn(f"MARGINALLY PROTECTED — {margin:.2f} dB margin. Consider more conservative assumptions.")
        else:
            warn(f"THRESHOLD VIOLATED by {abs(margin):.2f} dB — basis for citing harmful interference under RR No. 4.10.")

        # Loss breakdown table
        st.subheader("📋 Full Loss Breakdown")
        ex("Every gain and loss itemized. Non-zero loss rows are highlighted red; gain rows green; milestones (EIRP, Rx Power, RSL) blue.")
        loss_rows = [
            ("Transmit Power (Pt)",              f"+{tx_power_dbm:.2f} dBm",     "Input"),
            ("+ Tx Antenna Gain (Gt)",            f"+{tx_gain_dbi:.2f} dBi",      "Gain"),
            ("− Tx Cable / Coax Loss",            f"−{tx_cable_loss:.2f} dB",     "TxL" if tx_cable_loss else "zero"),
            ("− Tx Connector Loss",               f"−{tx_connector_loss:.2f} dB", "TxL" if tx_connector_loss else "zero"),
            ("− Tx Branching / Combiner Loss",    f"−{tx_branching_loss:.2f} dB", "TxL" if tx_branching_loss else "zero"),
            ("− Tx Filter / Duplexer Loss",       f"−{tx_filter_loss:.2f} dB",    "TxL" if tx_filter_loss else "zero"),
            ("− Tx Jumper / Pigtail Loss",        f"−{tx_jumper_loss:.2f} dB",    "TxL" if tx_jumper_loss else "zero"),
            ("− Tx Other Losses",                 f"−{tx_misc_loss:.2f} dB",      "TxL" if tx_misc_loss else "zero"),
            ("★ EIRP",                            f"= {eir:.2f} dBm",             "Mile"),
            (f"− Path Loss ({model_label})",      f"−{pl:.2f} dB",                "Path"),
            ("+ Rx Antenna Gain (Gr)",            f"+{rx_gain_dbi:.2f} dBi",      "Gain"),
            ("− Rx Cable / Coax Loss",            f"−{rx_cable_loss:.2f} dB",     "RxL" if rx_cable_loss else "zero"),
            ("− Rx Connector Loss",               f"−{rx_connector_loss:.2f} dB", "RxL" if rx_connector_loss else "zero"),
            ("− Rx Branching / Splitter Loss",    f"−{rx_branching_loss:.2f} dB", "RxL" if rx_branching_loss else "zero"),
            ("− Rx Filter / Duplexer Loss",       f"−{rx_filter_loss:.2f} dB",    "RxL" if rx_filter_loss else "zero"),
            ("− Rx Limiter / Protection Loss",    f"−{rx_limiter_loss:.2f} dB",   "RxL" if rx_limiter_loss else "zero"),
            ("− Rx Other Losses",                 f"−{rx_misc_loss:.2f} dB",      "RxL" if rx_misc_loss else "zero"),
            ("★ Received Power (Pr)",             f"= {rx_pwr:.2f} dBm",          "Mile"),
            ("  Noise Floor",                     f"{rx_noise_floor_dbm_input:.2f} dBm","Ref"),
            ("  I/N",                             f"{i_n:.2f} dB",                "Res"),
            ("  I/N Threshold",                   f"{in_threshold_db:.1f} dB",    "Thr"),
            ("★ Protection Margin",              f"= {margin:.2f} dB",            "Mile"),
        ]
        if use_eff:
            loss_rows.append(("★ Effective Margin (incl. +{} dB safety factor)".format(sf_loaded),
                              f"= {eff_margin:.2f} dB", "Mile"))

        df_loss = pd.DataFrame(loss_rows, columns=["Item","Value","Cat"])

        def style_loss(row):
            if row["Cat"] == "Mile":
                return ["background-color:#1a3560;font-weight:bold;color:#aaddff"]*3
            if row["Cat"] in ("TxL","RxL"):
                return ["background-color:#2a1a1a;color:#ffaaaa"]*3
            if row["Cat"] == "Gain":
                return ["background-color:#1a2a1a;color:#aaffaa"]*3
            if row["Cat"] == "Path":
                return ["background-color:#2a2a1a;color:#ffffaa"]*3
            return [""]*3

        st.dataframe(df_loss.style.apply(style_loss, axis=1),
                     use_container_width=True, hide_index=True)

        # Waterfall chart
        st.subheader("Link Budget Waterfall")
        ex("Green bars = gains, red bars = losses. Blue = Tx power start, orange = final Rx power.")
        stages = ["Tx Power","Tx Gain","Tx Losses","Path Loss","Rx Gain","Rx Losses","Rx Power"]
        deltas  = [0, tx_gain_dbi, -tx_total_loss, -pl, rx_gain_dbi, -rx_total_loss, 0]
        starts  = [tx_power_dbm]
        rw = tx_power_dbm
        for d in deltas[1:-1]:
            rw += d
            starts.append(rw)
        starts.append(rx_pwr)

        fig2, ax2 = plt.subplots(figsize=(12, 4))
        fig2.patch.set_facecolor("#0e1117"); ax2.set_facecolor("#0e1117")
        for idx in range(len(stages)):
            if idx == 0:
                ax2.bar(idx, tx_power_dbm, color="#4488ff", width=0.6)
            elif idx == len(stages)-1:
                ax2.bar(idx, rx_pwr, color="#ffaa00", width=0.6)
            else:
                d = deltas[idx]
                if d >= 0:
                    ax2.bar(idx, d, bottom=starts[idx-1], color="#44bb44", width=0.6, alpha=0.85)
                else:
                    ax2.bar(idx, -d, bottom=starts[idx-1]+d, color="#bb4444", width=0.6, alpha=0.85)
            ax2.text(idx, starts[idx]+0.5, f"{starts[idx]:.1f}",
                     ha="center", fontsize=7, color="white")
        ax2.axhline(rx_noise_floor_dbm_input, color="cyan", linestyle="--", linewidth=1.5,
                    label=f"Noise Floor ({rx_noise_floor_dbm_input:.0f} dBm)")
        ax2.axhline(rx_noise_floor_dbm_input + in_threshold_db, color="red",
                    linestyle=":", linewidth=1.5,
                    label=f"I/N Threshold ({in_threshold_db} dB)")
        ax2.set_xticks(range(len(stages)))
        ax2.set_xticklabels(stages, color="white", fontsize=9)
        ax2.set_ylabel("Power (dBm)", color="white")
        ax2.tick_params(colors="white")
        ax2.legend(fontsize=8, facecolor="#1a1a2e", labelcolor="white")
        for sp in ax2.spines.values(): sp.set_color("#444")
        plt.tight_layout()
        st.pyplot(fig2)

        # Distance sweep
        st.subheader("Path Loss vs Distance")
        ex("Required path loss = EIRP + Rx Gain − Rx Losses − noise floor − I/N threshold. Where FSPL crosses that line = minimum safe coordination distance.")
        dists = np.linspace(0.1, max(50, dist_km * 3), 200)
        pls_fspl = [free_space_path_loss_db(freq_mhz, d) for d in dists]
        pls_p452 = [p452_basic_loss_db(freq_mhz, d, tx_height_m, rx_height_m, terrain_type) for d in dists]
        req_pl = eir + rx_gain_dbi - rx_total_loss - rx_noise_floor_dbm_input - in_threshold_db
        fig3, ax3 = plt.subplots(figsize=(10, 4))
        fig3.patch.set_facecolor("#0e1117"); ax3.set_facecolor("#0e1117")
        ax3.plot(dists, pls_fspl, color="#4488ff", linewidth=2, label="FSPL")
        ax3.plot(dists, pls_p452, color="#ffaa00", linewidth=2, linestyle="--", label=f"P.452 ({terrain_type})")
        ax3.axhline(req_pl, color="red", linestyle=":", label=f"Required PL ({req_pl:.0f} dB)")
        ax3.axvline(dist_km, color="white", linestyle=":", alpha=0.5, label=f"Current ({dist_km} km)")
        ax3.set_xlabel("Distance (km)", color="white")
        ax3.set_ylabel("Path Loss (dB)", color="white")
        ax3.legend(fontsize=8, facecolor="#1a1a2e", labelcolor="white")
        ax3.tick_params(colors="white")
        for sp in ax3.spines.values(): sp.set_color("#444")
        plt.tight_layout()
        st.pyplot(fig3)

# TAB 3 — NOISE & I/N ANALYSIS
# ─────────────────────────────────────────────────────────────────────────────
elif selected_tab == "📊 Noise & I/N":
    st.title("📊 Noise Floor & I/N Analysis")
    ex("I/N quantifies the degradation of receiver sensitivity: total noise = N·(1 + I/N_linear). At I/N = −6 dB, sensitivity degrades by 0.97 dB; at 0 dB it degrades by 3 dB — unacceptable for safety-critical systems.")

    col1, col2 = st.columns(2)
    with col1:
        st.subheader("Receiver Noise Floor Calculator")
        ex("Thermal noise floor: N = kTB·NF, where k = 1.38×10⁻²³ J/K, T = 290K standard, B = noise bandwidth. The NF term captures all internal noise contributions referred to the input — any interference power above kTB·NF degrades SNR.")
        bw_mhz = st.number_input("Receiver Bandwidth (MHz)", value=100.0, min_value=0.001, step=10.0)
        nf_db = st.number_input("Receiver Noise Figure (dB)", value=5.0, step=0.5,
            help="Measure of receiver's internal noise amplification (0 dB = ideal)")
        temp_k = st.number_input("Noise Temperature (K)", value=290.0, step=10.0,
            help="Physical temperature; 290K is standard; use 255K for avionics")

        kT = 10.0 * np.log10(1.38e-23 * temp_k) + 30  # in dBm/Hz
        nf_dbm = kT + 10.0 * np.log10(bw_mhz * 1e6) + nf_db

        st.markdown("---")
        st.metric("Thermal Noise Density", f"{kT:.1f} dBm/Hz")
        st.metric("Noise Floor", f"{nf_dbm:.1f} dBm")
        ex("kT = −174 dBm/Hz at 290K is the fundamental thermal noise limit. Each decade of bandwidth adds 10 dB to the noise floor. NF is the receiver's contribution above kTB — a 1 dB NF increase directly reduces sensitivity by 1 dB.")

    with col2:
        st.subheader("I/N Threshold Selector")
        ex("GNSS threshold is −10 dB I/N per M.1477/M.1905 because GPS L1 arrives at approximately −130 dBm — only 25 dB above the noise floor of a typical receiver. Any noise floor rise materially degrades acquisition and tracking margin.")
        threshold_type = st.selectbox("System Type", [
            "Generic aeronautical ARNS (−6 dB)",
            "GNSS / GPS (−10 dB)",
            "ADS-B / SSR (−10 dB)",
            "Radio Altimeter (−6 dB)",
            "Custom",
        ])
        thresh_map = {
            "Generic aeronautical ARNS (−6 dB)": -6,
            "GNSS / GPS (−10 dB)": -10,
            "ADS-B / SSR (−10 dB)": -10,
            "Radio Altimeter (−6 dB)": -6,
        }
        if threshold_type == "Custom":
            threshold_db = st.number_input("Custom I/N Threshold (dB)", value=-6.0)
        else:
            threshold_db = thresh_map[threshold_type]
            st.metric("I/N Threshold", f"{threshold_db} dB")

        interference_dbm = st.number_input("Interference Power at Rx (dBm)", value=-100.0, step=1.0)
        i_n_calc = in_ratio_db(interference_dbm, nf_dbm)
        margin_calc = protection_margin_db(i_n_calc, threshold_db)

        st.metric("Computed I/N", f"{i_n_calc:.1f} dB")
        st.metric("Protection Margin", f"{margin_calc:.1f} dB",
                  delta=f"{margin_calc:.1f} dB",
                  delta_color="normal" if margin_calc >= 0 else "inverse")

        if margin_calc < 0:
            warn(f"Threshold violated by {abs(margin_calc):.1f} dB — this would be cited as harmful interference in a US contribution.")
        else:
            ok(f"Protected with {margin_calc:.1f} dB margin.")

    # I/N sensitivity heatmap
    st.subheader("I/N Sensitivity — Interference Power vs Bandwidth")
    ex("The contour line at your I/N threshold shows the boundary between compatible and incompatible operation. Note how wider bandwidth moves the boundary upward — a broader receiver requires proportionally stronger interference to violate the threshold.")
    bw_range = np.logspace(np.log10(0.1), np.log10(500), 40)
    int_range = np.linspace(-140, -60, 40)
    Z = np.zeros((len(int_range), len(bw_range)))
    for i, intpwr in enumerate(int_range):
        for j, bw in enumerate(bw_range):
            nf_j = kT + 10.0 * np.log10(bw * 1e6) + nf_db
            Z[i, j] = in_ratio_db(intpwr, nf_j)

    fig4, ax4 = plt.subplots(figsize=(10, 5))
    fig4.patch.set_facecolor("#0e1117")
    im = ax4.contourf(bw_range, int_range, Z, levels=20, cmap='RdYlGn_r')
    ax4.contour(bw_range, int_range, Z, levels=[threshold_db], colors='red', linewidths=2)
    ax4.set_xscale('log')
    ax4.set_xlabel("Receiver Bandwidth (MHz)", color='white')
    ax4.set_ylabel("Interference Power at Rx (dBm)", color='white')
    ax4.set_title(f"I/N (dB) — Red line = {threshold_db} dB threshold", color='white')
    ax4.tick_params(colors='white')
    for sp in ax4.spines.values():
        sp.set_color('#444')
    cbar = plt.colorbar(im, ax=ax4)
    cbar.set_label("I/N (dB)", color='white')
    cbar.ax.yaxis.set_tick_params(color='white')
    plt.setp(cbar.ax.yaxis.get_ticklabels(), color='white')
    plt.tight_layout()
    st.pyplot(fig4)

    st.markdown("---")
    # ── GNSS PROTECTION FRAMEWORK (M.1318 / M.1477 / M.1904 / M.1905) ────────
    st.header("🛰️ GNSS Protection Framework — ITU-R M.1318 / M.1477 / M.1904 / M.1905")
    ex("The M.1318/M.1477 framework is the legally grounded calculation chain for GNSS protection: receiver spec (a) minus safety margin (b) gives the maximum tolerable interference density (c) — all in dB(W/Hz) at the passive antenna terminal.")

    st.subheader("1️⃣  The c = a − b Framework (ITU-R M.1318 Annex 1)")
    ex("Note that 'a' is specified at the passive antenna terminal — not at the LNA output. This is important when comparing with interference analyses that compute received power: convert to power density using the receiver noise bandwidth.")
    st.latex(r"c = a - b")
    st.markdown("""
| Symbol | Parameter | Units | Description |
|---|---|---|---|
| **a** | Max aggregate non-RNSS interference power density | dB(W/Hz) | Specified by the RNSS receiver design — from receiver datasheet or ITU-R recommendation |
| **b** | Protection margin | dB | To ensure protection as provided by RR No. 4.10 — **6 dB aeronautical safety margin (M.1477)** |
| **c** | Tolerable interference power density at receiver | dB(W/Hz) | The level you must stay below in your interference analysis |
    """)

    st.subheader("🔢 c = a − b Calculator")
    ex("'a' comes from the RNSS receiver's minimum operational performance specification (MOPS) — typically −111.5 dB(W/Hz) for GPS/GNSS. This is not a system-level threshold; it is the worst-case receiver design point.")

    col_gnss1, col_gnss2 = st.columns(2)
    with col_gnss1:
        gnss_band = st.selectbox("GNSS Band (ITU-R M.1318):", [
            "L5 / E5 — 1164–1215 MHz",
            "L2 / E6 — 1215–1300 MHz",
            "L1 / E1 — 1559–1610 MHz",
            "L3 / future — 5010–5030 MHz",
            "Custom",
        ])
        # Pre-fill typical 'a' values per band from ITU-R M.1318
        a_defaults = {
            "L5 / E5 — 1164–1215 MHz": -111.5,
            "L2 / E6 — 1215–1300 MHz": -111.5,
            "L1 / E1 — 1559–1610 MHz": -111.5,
            "L3 / future — 5010–5030 MHz": -111.5,
            "Custom": -111.5,
        }
        param_a = st.number_input(
            "a — Max aggregate non-RNSS interference power density (dB(W/Hz))",
            value=a_defaults.get(gnss_band, -111.5), step=0.5,
            help="From ITU-R M.1318 Table / receiver spec. Typical GPS L1: −111.5 dB(W/Hz)"
        )
        interferer_bw_hz = st.number_input(
            "Interferer bandwidth (Hz) — for narrowband rule check",
            value=10e6, min_value=1.0, step=1000.0, format="%.0f",
            help="Enter the bandwidth of the interfering signal in Hz"
        )

    with col_gnss2:
        # Protection margin b — M.1477 defines 6 dB aeronautical safety margin
        margin_type = st.selectbox("b — Protection Margin Source:", [
            "M.1477 Annex 5 — Aeronautical safety margin (6 dB)",
            "M.1904 / M.1905 — GLONASS/RNSS safety margin (6 dB)",
            "M.1477 Annex 5 — Narrowband interferer ≤700 Hz (+10 dB additional)",
            "Custom margin",
        ])
        margin_map = {
            "M.1477 Annex 5 — Aeronautical safety margin (6 dB)": 6.0,
            "M.1904 / M.1905 — GLONASS/RNSS safety margin (6 dB)": 6.0,
            "M.1477 Annex 5 — Narrowband interferer ≤700 Hz (+10 dB additional)": 16.0,
            "Custom margin": None,
        }
        if margin_map[margin_type] is None:
            param_b = st.number_input("b — Custom protection margin (dB)", value=6.0, step=0.5)
        else:
            param_b = margin_map[margin_type]
            st.metric("b — Protection Margin", f"{param_b} dB")

        param_c = param_a - param_b
        st.metric("c — Tolerable interference level", f"{param_c:.1f} dB(W/Hz)")
        ex("c is defined at the passive antenna terminal, integrated over the noise bandwidth. If your analysis produces received interference power P_i (dBm), convert: c_computed = P_i(dBm) − 30 − 10·log10(B_Hz), then compare to c.")

    # Narrowband rule check — M.1477 Annex 5
    st.subheader("2️⃣  Narrowband Interferer Rule (ITU-R M.1477 Annex 5)")
    ex("The ≤700 Hz narrowband rule applies because a CW or very narrowband interferer concentrates all its energy into the receiver's phase-locked tracking loops, causing cycle slips and position errors disproportionate to its total power.")

    nb_threshold_hz = 700.0
    is_narrowband = interferer_bw_hz <= nb_threshold_hz
    col_nb1, col_nb2 = st.columns(2)
    with col_nb1:
        st.metric("Interferer Bandwidth", f"{interferer_bw_hz:.0f} Hz")
        st.metric("Narrowband Threshold", "700 Hz")
    with col_nb2:
        if is_narrowband:
            warn(f"⚠️ NARROWBAND RULE APPLIES — interferer BW ({interferer_bw_hz:.0f} Hz) ≤ 700 Hz. "
                 f"Additional +10 dB protection margin required per M.1477 Annex 5. "
                 f"Effective I/N threshold tightens to −20 dB (−10 dB standard + −10 dB additional).")
            effective_gnss_threshold = -20.0
        else:
            ok(f"Narrowband rule does NOT apply — interferer BW ({interferer_bw_hz:.0f} Hz) > 700 Hz. "
               f"Standard −10 dB I/N threshold applies.")
            effective_gnss_threshold = -10.0
        st.metric("Effective GNSS I/N Threshold", f"{effective_gnss_threshold} dB")

    # Full GNSS protection criteria table
    st.subheader("3️⃣  Complete GNSS Protection Criteria by Band")
    ex("RR No. 5.328 limits RNSS protection in 1164–1215 MHz: RNSS cannot claim protection from ARNS (DME beacons) per M.1318 Annex 1. Your interference analysis must account for DME as a pre-existing interference source in this band.")
    gnss_table = pd.DataFrame([
        ["L1 / E1 / B1", "1559–1610 MHz", "−111.5", "6 dB", "−10 dB", "−20 dB*",
         "M.1318, M.1477, M.1905", "Primary aviation GNSS; WAAS/SBAS; most protected"],
        ["L5 / E5a / B2a", "1164–1215 MHz", "−111.5", "6 dB", "−10 dB", "−20 dB*",
         "M.1318, M.1477, M.1904", "Safety-of-life signal; aviation approach procedures"],
        ["L2 / E6 / B3", "1215–1300 MHz", "−111.5", "6 dB", "−10 dB", "−20 dB*",
         "M.1318, M.1477", "Dual-freq integrity; shares with ARNS (RR 5.328 applies)"],
        ["L3 / future", "5010–5030 MHz", "−111.5", "6 dB", "−10 dB", "−20 dB*",
         "M.1318", "Under pressure from IMT-2030 studies in WP 5D"],
    ], columns=[
        "Signal", "Band (MHz)", "a (dB(W/Hz))", "b (Aero Margin)",
        "Std I/N Threshold", "Narrowband I/N*", "Key References", "Notes"
    ])
    st.dataframe(gnss_table, use_container_width=True)
    st.caption("* Narrowband threshold applies when interferer BW ≤ 700 Hz per ITU-R M.1477 Annex 5")

    # Recommendation summary box
    st.subheader("4️⃣  Recommendation Doctrine Summary")
    ex("M.1904 and M.1905 are critical because they establish the 6 dB safety margin as settled ITU-R doctrine — any proponent proposing less must argue against two explicit Recommendations, not just a methodology document.")
    col_r1, col_r2 = st.columns(2)
    with col_r1:
        st.markdown("""
**ITU-R M.1318** — *Methodology for GNSS protection*
- Defines the c = a − b framework
- Covers L1, L2, L5, L3 bands
- Establishes aggregate interference methodology
- **Cite for:** methodology basis of any GNSS interference analysis

**ITU-R M.1477 Annex 5** — *Aeronautical safety margin*
- 6 dB safety margin for GNSS safety-of-life apps
- +10 dB additional margin for narrowband (≤700 Hz) interferers
- Lists factors requiring additional margins: terrain, weather, RNSS config
- **Cite for:** justifying b = 6 dB in the c = a − b calculation
        """)
    with col_r2:
        st.markdown("""
**ITU-R M.1904** — *GLONASS spaceborne receiver protection*
- Annex 1 Table 1 Note 3: safety margin = 6 dB
- Establishes GLONASS-specific protection baseline
- **Cite for:** GLONASS L1/L2 band contributions; reinforces 6 dB doctrine

**ITU-R M.1905** — *RNSS safety applications*
- Recommends safety margin be applied for all RNSS safety-of-life interference analyses
- Note 1: Aeronautical safety margin of 6 dB
- **Cite for:** Any contribution involving aviation use of GNSS — broadest coverage
- **Key strength:** Applies to ALL RNSS systems, not just GPS or GLONASS

---
**Together:** M.1318 (method) + M.1477 (margin) + M.1904 + M.1905 (doctrine) = complete GNSS defense stack
        """)

    # Interference budget worked example
    st.subheader("5️⃣  Worked Example — GPS L1 Protection Budget")
    ex("The M.1318 budget uses power spectral density (dB(W/Hz)) rather than integrated power (dBm) — this makes the criterion independent of receiver bandwidth and directly comparable across systems with different noise bandwidths.")
    with st.expander("📐 Show GPS L1 Protection Budget Walkthrough"):
        st.markdown("""
**Scenario:** New terrestrial service proposed near 1559–1610 MHz GPS L1 band.
Is the aggregate interference tolerable?

**Step 1 — Establish 'a' from receiver spec (M.1318)**
> Maximum aggregate non-RNSS interference power density for GPS L1:
> **a = −111.5 dB(W/Hz)**
> (From ITU-R M.1318 Annex 1 Table 1)

**Step 2 — Apply aeronautical safety margin 'b' (M.1477 Annex 5 + M.1905)**
> Aeronautical safety margin = **6 dB** (M.1477 Note 1, M.1905 Recommendation 2)
> This ensures protection as required by RR No. 4.10 (M.1318 Annex 1, Step 1b)
> **b = 6 dB**

**Step 3 — Compute tolerable interference level 'c'**
> c = a − b = −111.5 − 6 = **−117.5 dB(W/Hz)**

**Step 4 — Check narrowband rule (M.1477 Annex 5, Section 4)**
> If interfering signal BW ≤ 700 Hz: apply additional 10 dB margin
> Effective tolerable level = −117.5 − 10 = **−127.5 dB(W/Hz)** for narrowband

**Step 5 — Compare your computed interference to 'c'**
> Run your link budget → convert received interference power to power density (dB(W/Hz))
> If computed interference > c → threshold exceeded → **cite harmful interference under RR 4.10**
> If computed interference ≤ c → compatible → **document margin and conditions**

**Step 6 — State the regulatory consequence**
> "The computed aggregate non-RNSS interference power density of [X] dB(W/Hz) exceeds
> the tolerable level c = −117.5 dB(W/Hz) established by ITU-R M.1318 Annex 1, with
> the 6 dB aeronautical safety margin required by ITU-R M.1477 Annex 5, M.1904, and M.1905.
> This constitutes harmful interference to a safety-of-life service under RR No. 4.10.
> The United States opposes this proposal without additional protective measures."
        """)

    # Power density converter
    st.subheader("🔢 Power → Power Density Converter")
    ex("Power density = received power − 10·log10(noise BW). GPS L1 C/A noise BW ≈ 2 MHz; L5 ≈ 20.46 MHz. Using too wide a bandwidth here artificially lowers the density and makes interference appear more benign — check the proponent's assumed bandwidth.")
    col_pd1, col_pd2 = st.columns(2)
    with col_pd1:
        pwr_dbm_conv = st.number_input("Interference power at Rx (dBm)", value=-100.0, step=1.0)
        rx_bw_hz_conv = st.number_input("Receiver noise bandwidth (Hz)", value=20.46e6,
            step=1e6, format="%.0f",
            help="GPS L1 C/A: ~1 MHz; L5: ~20.46 MHz; use receiver spec")
    pwr_dbw = pwr_dbm_conv - 30
    pwr_density_dbwHz = pwr_dbw - 10*np.log10(rx_bw_hz_conv)
    with col_pd2:
        st.metric("Power density", f"{pwr_density_dbwHz:.1f} dB(W/Hz)")
        margin_vs_c = param_c - pwr_density_dbwHz
        st.metric("Margin vs tolerable level c", f"{margin_vs_c:.1f} dB",
                  delta_color="normal" if margin_vs_c >= 0 else "inverse")
        if margin_vs_c < 0:
            warn(f"Computed interference exceeds c by {abs(margin_vs_c):.1f} dB — cite M.1318 violation in contribution.")
        else:
            ok(f"Compatible with {margin_vs_c:.1f} dB margin against M.1318 tolerable level.")

# ─────────────────────────────────────────────────────────────────────────────
# TAB 4 — PROPAGATION ANALYSIS
# ─────────────────────────────────────────────────────────────────────────────
elif selected_tab == "🌐 Propagation":
    st.title("🌐 Propagation Analysis")
    ex("P.528 accounts for atmospheric refraction, radio horizon geometry, and troposcatter — all of which increase path loss relative to FSPL for slant paths. A proponent using FSPL for an airborne scenario is using an unrealistically pessimistic interference estimate.")

    freq_mhz_p = st.slider("Frequency (MHz)", min_value=100, max_value=10000, value=4300, step=50)
    max_dist_km = st.slider("Max Distance (km)", min_value=5, max_value=500, value=100)

    col_a, col_b = st.columns(2)
    with col_a:
        aircraft_alt_km_p = st.number_input("Aircraft Altitude (km) — for P.528", value=3.0, step=0.5)
        tx_h = st.number_input("Tx Height m (P.452)", value=30.0)
        rx_h = st.number_input("Rx Height m (P.452)", value=5.0)
    with col_b:
        terrain_p = st.selectbox("Terrain (P.452)", ["open", "suburban", "urban", "dense_urban"])
        time_pct = st.slider("Time % for P.528", min_value=1, max_value=99, value=50,
            help="50% = median; 1% = exceeded only 1% of the time (worst interference case)")

    dists_p = np.linspace(0.5, max_dist_km, 300)
    fspl_vals = [free_space_path_loss_db(freq_mhz_p, d) for d in dists_p]
    p452_vals = [p452_basic_loss_db(freq_mhz_p, d, tx_h, rx_h, terrain_p) for d in dists_p]
    p528_vals = [p528_aero_path_loss_db(freq_mhz_p, d, aircraft_alt_km_p, time_pct) for d in dists_p]

    fig5, ax5 = plt.subplots(figsize=(12, 5))
    fig5.patch.set_facecolor("#0e1117")
    ax5.set_facecolor("#0e1117")
    ax5.plot(dists_p, fspl_vals, color='#4488ff', linewidth=2, label='FSPL (optimistic — worst interference case)')
    ax5.plot(dists_p, p452_vals, color='#ffaa00', linewidth=2, linestyle='--',
             label=f'P.452 simplified ({terrain_p})')
    ax5.plot(dists_p, p528_vals, color='#44ffaa', linewidth=2, linestyle='-.',
             label=f'P.528 simplified (alt={aircraft_alt_km_p} km, {time_pct}%)')
    ax5.set_xlabel("Distance (km)", color='white', fontsize=11)
    ax5.set_ylabel("Basic Transmission Loss (dB)", color='white', fontsize=11)
    ax5.set_title(f"Path Loss Comparison at {freq_mhz_p} MHz", color='white')
    ax5.legend(facecolor='#1a1a2e', labelcolor='white', fontsize=9)
    ax5.tick_params(colors='white')
    ax5.grid(color='#333', alpha=0.5)
    for sp in ax5.spines.values():
        sp.set_color('#444')
    plt.tight_layout()
    st.pyplot(fig5)

    ex("Conservative protection analysis uses the lowest plausible path loss (FSPL or P.452 open, time %=1) to maximize computed interference. If compatibility holds under these conditions, any realistic deployment will also be compatible.")

    # Atmospheric attenuation using itur P.676
    st.subheader("Atmospheric Gaseous Attenuation (ITU-R P.676 via itur)")
    ex("P.676 gaseous attenuation is negligible below ~3 GHz (<0.01 dB/km), becomes relevant above 6 GHz, and peaks at 60 GHz (O₂ resonance, ~15 dB/km). Do not cite it as a protection mechanism for L-band or C-band ground-to-air paths.")

    freq_range_ghz = np.linspace(0.1, min(freq_mhz_p / 1000 * 2, 10), 100)
    gas_atten_vals = []
    path_len_km = st.slider("Path length for P.676 (km)", 1, 100, 20)

    for f_ghz in freq_range_ghz:
        try:
            g = itu676.gaseous_attenuation_terrestrial_path(
                f_ghz, p=50, T=15, H=50, P=1013.25, d=path_len_km, mode="approx"
            )
            gas_atten_vals.append(float(g.value) if hasattr(g, 'value') else float(g))
        except Exception:
            gas_atten_vals.append(np.nan)

    fig6, ax6 = plt.subplots(figsize=(12, 4))
    fig6.patch.set_facecolor("#0e1117")
    ax6.set_facecolor("#0e1117")
    ax6.plot(freq_range_ghz * 1000, gas_atten_vals, color='#ff88aa', linewidth=2)
    ax6.axvline(freq_mhz_p, color='yellow', linestyle='--', alpha=0.7, label=f'Selected: {freq_mhz_p} MHz')
    ax6.set_xlabel("Frequency (MHz)", color='white')
    ax6.set_ylabel(f"Gaseous Attenuation (dB) over {path_len_km} km", color='white')
    ax6.set_title("ITU-R P.676 Atmospheric Attenuation (via itur)", color='white')
    ax6.legend(facecolor='#1a1a2e', labelcolor='white')
    ax6.tick_params(colors='white')
    ax6.grid(color='#333', alpha=0.5)
    for sp in ax6.spines.values():
        sp.set_color('#444')
    plt.tight_layout()
    st.pyplot(fig6)

    # Model selector guidance
    st.subheader("Which Model Should You Use?")
    ex("Model selection is a regulatory choice, not just a technical one — ITU-R contributions cite the specific model by name. A challenge to your model choice in a Working Party session can invalidate your entire analysis, so document your selection rationale explicitly.")
    guidance = pd.DataFrame([
        ["Terrestrial base station → ground receiver", "P.452", "Point-to-point interference, terrain profile needed for full implementation"],
        ["Terrestrial base station → airborne receiver", "P.528", "Specific to aeronautical scenarios; slant path + atmosphere"],
        ["Satellite downlink → ground/airborne receiver", "P.619 / P.618", "Earth-space path; REQUIRED for WP 4C satellite contributions — P.452 is WRONG here"],
        ["Quick worst-case bound", "FSPL", "Always optimistic (most interference); use to bound the problem first"],
        ["High freq (>6 GHz) atmospheric loss", "P.676", "Gaseous absorption becomes significant above ~6 GHz"],
    ], columns=["Scenario", "Model", "Notes"])
    st.table(guidance)

    # ── P.619 Satellite Slant-Path Calculator ─────────────────────────────────
    st.markdown("---")
    st.subheader("🛰️ ITU-R P.619 Satellite Slant-Path Calculator")
    ex("P.619 is the correct propagation model for WP 4C satellite downlink contributions (DC-MSS-IMT, AI 1.13). Using P.452 or FSPL alone for a satellite-to-Earth path is a fundamental methodology error — flag it in any contribution you review.")

    col_619a, col_619b = st.columns(2)
    with col_619a:
        sat_freq_mhz   = st.number_input("Satellite Frequency (MHz)", value=950.0, step=10.0,
            help="Center frequency of the satellite downlink — e.g. 925–960 MHz for AI 1.13 candidate band")
        sat_eirp_dbw   = st.number_input("Satellite EIRP (dBW)", value=20.0, step=1.0,
            help="Peak EIRP toward Earth surface at nadir. Use worst-case (nadir-pointing, max power) for interference analysis.")
        sat_alt_km     = st.number_input("Satellite Altitude (km)", value=550.0, step=50.0,
            help="Orbital altitude. LEO: 400–1200 km. MEO: 8000–20000 km. GEO: 35786 km.")
        elev_angle_deg = st.number_input("Elevation Angle (°)", value=30.0, step=5.0, min_value=0.1, max_value=90.0,
            help="Elevation angle from ground/aircraft to satellite. 90° = directly overhead. Lower = longer path through atmosphere.")
        victim_alt_km  = st.number_input("Victim Altitude (km)", value=0.0, step=0.5,
            help="0 = ground receiver. Enter aircraft cruising altitude (e.g. 10 km) for airborne victim. Aircraft see stronger downlinks at altitude.")

    with col_619b:
        # Slant range from geometry
        R_e = 6371.0  # Earth radius km
        elev_rad = np.radians(elev_angle_deg)
        sat_orbit_r = R_e + sat_alt_km
        victim_r    = R_e + victim_alt_km
        # Slant range via law of cosines: r² = R_v² + R_s² - 2·R_v·R_s·cos(θ)
        # Using elevation angle: slant_range = sqrt(R_s² - (R_v·cos(el))²) - R_v·sin(el)
        cos_elev = np.cos(elev_rad); sin_elev = np.sin(elev_rad)
        slant_km = np.sqrt(sat_orbit_r**2 - (victim_r * cos_elev)**2) - victim_r * sin_elev

        # Free-space path loss at slant range
        fspl_619 = round(20*np.log10(slant_km) + 20*np.log10(sat_freq_mhz) + 32.44, 2)

        # Atmospheric attenuation (simplified P.619 / P.676)
        try:
            atm_atten_zenith = float(itu676.gaseous_attenuation_terrestrial_path(
                sat_freq_mhz/1000, p=50, T=15, H=50, P=1013.25, d=10, mode="approx"
            ).value)
            # Scale by elevation angle (zenith path × 1/sin(elev))
            path_factor = min(1.0/max(sin_elev, 0.1), 10)  # cap at 10× zenith
            atm_atten_slant = round(atm_atten_zenith * path_factor, 2)
        except Exception:
            atm_atten_slant = round(0.010 * slant_km * 0.1, 2)  # rough fallback

        total_path_loss = round(fspl_619 + atm_atten_slant, 2)
        pfd_dbw_m2 = round(sat_eirp_dbw - 10*np.log10(4 * np.pi * (slant_km*1e3)**2), 2)
        rx_pwr_isotropic = round(sat_eirp_dbw - total_path_loss, 2)

        st.metric("Slant Range", f"{slant_km:.1f} km",
            help="Geometric distance from satellite to victim. Longer than orbital altitude due to off-nadir geometry.")
        st.metric("Free Space Path Loss (P.619)", f"{fspl_619:.2f} dB",
            help="FSPL at slant range and frequency. This is the CORRECT metric for satellite downlinks — not P.452.")
        st.metric("Atmospheric Attenuation (P.676 approx)", f"{atm_atten_slant:.2f} dB",
            help="Gaseous attenuation along slant path. Increases at lower elevation angles.")
        st.metric("Total Path Loss", f"{total_path_loss:.2f} dB",
            help="FSPL + atmospheric absorption along slant path.")
        st.metric("PFD at victim (dBW/m²)", f"{pfd_dbw_m2:.2f} dBW/m²",
            help=f"Power Flux Density at receiver. For DME protection: epfd must not exceed −121.5 dBW/m²/MHz.")
        st.metric("Rx Power (isotropic antenna)", f"{rx_pwr_isotropic:.2f} dBW",
            help="Received power for a 0 dBi isotropic antenna at the victim. Add victim antenna gain for actual received power.")

        # DME/AMS(R)S compliance checks
        st.markdown("**Compliance Checks**")
        if 925 <= sat_freq_mhz <= 960:
            ok_dme = pfd_dbw_m2 <= -121.5
            status_color = "✅" if ok_dme else "❌"
            st.markdown(f"{status_color} DME band (925–960 MHz): PFD = {pfd_dbw_m2:.1f} dBW/m²  vs  limit −121.5 dBW/m²/MHz  →  {'COMPLIANT' if ok_dme else 'EXCEEDS LIMIT — INCOMPATIBLE'}")
        elif 1475 <= sat_freq_mhz <= 1518:
            st.info(f"AMS(R)S adjacent band (1475–1518 MHz): Use ΔT/T analysis, not PFD. ΔT/T single-entry limit = 6%. Compute: T_noise_rise / T_sys where T_noise_rise = S_interference / k.")
        elif 2620 <= sat_freq_mhz <= 2690:
            st.info(f"ASR adjacent band (2620–2690 MHz): ASR I/N threshold = −10 dB. Compute I/N from received interference power vs ASR noise floor.")

    # P.619 distance sweep
    st.markdown("**PFD vs Elevation Angle**")
    ex("Lower elevation angles = longer slant path = more path loss = less PFD at victim. BUT aircraft at altitude always see the satellite at a higher effective elevation angle than ground receivers — airborne victims receive stronger downlinks.")

    elev_sweep = np.linspace(5, 90, 200)
    pfd_sweep  = []
    for el in elev_sweep:
        el_r = np.radians(el)
        s_r  = np.sqrt(sat_orbit_r**2 - (victim_r * np.cos(el_r))**2) - victim_r * np.sin(el_r)
        fl   = 20*np.log10(s_r) + 20*np.log10(sat_freq_mhz) + 32.44
        pfd_sweep.append(sat_eirp_dbw - 10*np.log10(4 * np.pi * (s_r*1e3)**2))

    fig_619, ax_619 = plt.subplots(figsize=(10, 3))
    fig_619.patch.set_facecolor("#0e1117"); ax_619.set_facecolor("#0e1117")
    ax_619.plot(elev_sweep, pfd_sweep, color="#ce93d8", linewidth=2)
    ax_619.axvline(elev_angle_deg, color="white", linestyle=":", alpha=0.7, label=f"Current: {elev_angle_deg}°")
    if 925 <= sat_freq_mhz <= 960:
        ax_619.axhline(-121.5, color="red", linestyle="--", linewidth=1.5, label="DME limit (−121.5 dBW/m²)")
    ax_619.set_xlabel("Elevation Angle (°)", color="white", fontsize=10)
    ax_619.set_ylabel("PFD (dBW/m²)", color="white", fontsize=10)
    ax_619.set_title(f"Satellite PFD vs Elevation Angle — {sat_freq_mhz:.0f} MHz, {sat_alt_km:.0f} km orbit", color="white")
    ax_619.legend(facecolor="#1a1a2e", labelcolor="white", fontsize=9)
    ax_619.tick_params(colors="white"); ax_619.grid(color="#333", alpha=0.4)
    for sp in ax_619.spines.values(): sp.set_color("#444")
    plt.tight_layout()
    st.pyplot(fig_619)

# ─────────────────────────────────────────────────────────────────────────────
# TAB 5 — MONTE CARLO
# ─────────────────────────────────────────────────────────────────────────────
elif selected_tab == "🎲 Monte Carlo":
    st.title("🎲 Monte Carlo Aggregate Interference")
    ex("Per SM.2028, violation probability >5% is the standard incompatibility criterion. The simulation draws random transmitter positions from a uniform spatial distribution within the deployment area — the resulting I/N distribution is a function of aggregate received power summed linearly across all N interferers.")

    col1, col2 = st.columns(2)
    with col1:
        st.subheader("Interferer Parameters")
        n_interferers = st.number_input("Number of Interferers (N)", value=10, min_value=1, max_value=500, step=5)
        ex("N is the most sensitive assumption in any Monte Carlo study — aggregate I/N scales approximately as 10·log10(N) at high densities. Always document the source of your N assumption and challenge unrealistically low values in others' contributions.")
        mc_tx_power = st.number_input("Tx Power per Unit (dBm)", value=43.0, step=1.0)
        mc_tx_gain = st.number_input("Tx Antenna Gain (dBi)", value=15.0, step=1.0)
        mc_freq = st.number_input("Frequency (MHz)", value=4300.0, step=10.0)
        mc_radius = st.number_input("Deployment Radius (km)", value=20.0, step=1.0,
            help="Interferers placed randomly within this radius of victim")
        mc_model = st.selectbox("Propagation Model", ["FSPL", "P.452", "P.528"])
        mc_terrain = st.selectbox("Terrain (P.452 only)", ["suburban", "open", "urban"])
        mc_alt = st.number_input("Aircraft Altitude km (P.528 only)", value=3.0, step=0.5)

    with col2:
        st.subheader("Victim Receiver")
        band_mc = st.selectbox("FAA Protected Band:", ["(manual)"] + list(FAA_BANDS.keys()))
        if band_mc != "(manual)":
            b_mc = FAA_BANDS[band_mc]
            mc_noise = b_mc["noise_floor_dbm"]
            mc_thresh = b_mc["in_threshold_db"]
            st.info(f"Noise floor: {mc_noise} dBm | I/N threshold: {mc_thresh} dB")
        else:
            mc_noise = st.number_input("Rx Noise Floor (dBm)", value=-90.0)
            mc_thresh = st.number_input("I/N Threshold (dB)", value=-6.0)

        n_trials = st.select_slider("Monte Carlo Trials", [500, 1000, 2000, 5000, 10000], value=2000)
        ex(f"More trials = smoother statistics. 2000–5000 is sufficient for most policy contributions.")

    if st.button("🎲 Run Monte Carlo Simulation", type="primary"):
        with st.spinner("Running simulation..."):
            results = monte_carlo_aggregate(
                n_interferers=int(n_interferers),
                tx_power_dbm=mc_tx_power,
                tx_gain_dbi=mc_tx_gain,
                freq_mhz=mc_freq,
                deployment_radius_km=mc_radius,
                rx_noise_floor_dbm=mc_noise,
                in_threshold_db=mc_thresh,
                n_trials=n_trials,
                model=mc_model,
                terrain_type=mc_terrain,
                aircraft_alt_km=mc_alt,
            )

        col_a, col_b, col_c, col_d = st.columns(4)
        col_a.metric("Median I/N", f"{results['in_p50_db']:.1f} dB")
        col_b.metric("95th pct I/N", f"{results['in_p95_db']:.1f} dB")
        col_c.metric("99th pct I/N", f"{results['in_p99_db']:.1f} dB")
        col_d.metric("Violation Probability", f"{results['violation_probability']*100:.1f}%",
                     delta="High Risk" if results['violation_probability'] > 0.05 else "Acceptable",
                     delta_color="inverse" if results['violation_probability'] > 0.05 else "normal")

        ex("The 99th percentile bounds the upper tail of the I/N distribution — for a 5% violation criterion, you need the exceedance probability at the threshold, not just the percentile. Read the CCDF directly at the I/N threshold to get the violation probability.")

        if results['violation_probability'] > 0.05:
            warn(f"Violation probability {results['violation_probability']*100:.1f}% exceeds 5% — this scenario poses a credible interference threat and supports a protection requirement.")
        elif results['violation_probability'] > 0.01:
            warn(f"Violation probability {results['violation_probability']*100:.1f}% is marginal — conservative assumptions needed.")
        else:
            ok(f"Violation probability only {results['violation_probability']*100:.2f}% — compatible under these assumptions.")

        # Distribution plot
        fig7, axes = plt.subplots(1, 2, figsize=(14, 5))
        fig7.patch.set_facecolor("#0e1117")

        ax7 = axes[0]
        ax7.set_facecolor("#0e1117")
        ax7.hist(results['in_values'], bins=60, color='#4488ff', alpha=0.75, density=True)
        ax7.axvline(mc_thresh, color='red', linewidth=2, linestyle='--', label=f'I/N Threshold ({mc_thresh} dB)')
        ax7.axvline(results['in_p50_db'], color='yellow', linewidth=1.5, linestyle=':', label=f'Median ({results["in_p50_db"]:.1f} dB)')
        ax7.axvline(results['in_p99_db'], color='orange', linewidth=1.5, linestyle=':', label=f'99th pct ({results["in_p99_db"]:.1f} dB)')
        ax7.set_xlabel("I/N (dB)", color='white')
        ax7.set_ylabel("Probability Density", color='white')
        ax7.set_title("I/N Distribution Across Trials", color='white')
        ax7.legend(facecolor='#1a1a2e', labelcolor='white', fontsize=8)
        ax7.tick_params(colors='white')
        for sp in ax7.spines.values():
            sp.set_color('#444')

        # CCDF
        ax8 = axes[1]
        ax8.set_facecolor("#0e1117")
        sorted_in = np.sort(results['in_values'])
        ccdf = 1.0 - np.arange(1, len(sorted_in) + 1) / len(sorted_in)
        ax8.semilogy(sorted_in, ccdf, color='#44ffaa', linewidth=2)
        ax8.axvline(mc_thresh, color='red', linewidth=2, linestyle='--', label=f'Threshold ({mc_thresh} dB)')
        ax8.axhline(0.05, color='orange', linewidth=1, linestyle=':', label='5% exceedance')
        ax8.axhline(0.01, color='yellow', linewidth=1, linestyle=':', label='1% exceedance')
        ax8.set_xlabel("I/N (dB)", color='white')
        ax8.set_ylabel("Probability of Exceedance", color='white')
        ax8.set_title("CCDF — Exceedance Probability", color='white')
        ax8.legend(facecolor='#1a1a2e', labelcolor='white', fontsize=8)
        ax8.tick_params(colors='white')
        ax8.grid(color='#333', alpha=0.4)
        for sp in ax8.spines.values():
            sp.set_color('#444')

        plt.tight_layout()
        st.pyplot(fig7)

        ex("The CCDF shows P(I/N > x). Reading off at x = I/N threshold gives the violation probability directly. The SM.2028 criterion requires this value < 5% — if a proponent's CCDF intersects the threshold above the 5% line, their scenario is incompatible.")

        # Sensitivity: vary N interferers
        st.subheader("Sensitivity: Violation Probability vs Number of Interferers")
        ex("Aggregate I/N grows roughly as 10·log10(N) at low densities and more slowly as the deployment radius fills. The inflection point identifies a maximum compatible density — propose this as a PFD density limit or coordination zone in your contribution text.")
        n_range = list(range(1, min(int(n_interferers) * 3, 100), max(1, int(n_interferers) // 5)))
        viol_probs = []
        with st.spinner("Running sensitivity sweep..."):
            for n in n_range:
                r = monte_carlo_aggregate(
                    n_interferers=n,
                    tx_power_dbm=mc_tx_power,
                    tx_gain_dbi=mc_tx_gain,
                    freq_mhz=mc_freq,
                    deployment_radius_km=mc_radius,
                    rx_noise_floor_dbm=mc_noise,
                    in_threshold_db=mc_thresh,
                    n_trials=500,
                    model=mc_model,
                    terrain_type=mc_terrain,
                    aircraft_alt_km=mc_alt,
                )
                viol_probs.append(r['violation_probability'] * 100)

        fig8, ax9 = plt.subplots(figsize=(10, 4))
        fig8.patch.set_facecolor("#0e1117")
        ax9.set_facecolor("#0e1117")
        ax9.plot(n_range, viol_probs, color='#ff8844', linewidth=2, marker='o', markersize=4)
        ax9.axhline(5, color='red', linestyle='--', label='5% threshold')
        ax9.axvline(n_interferers, color='white', linestyle=':', alpha=0.5, label=f'Selected N={n_interferers}')
        ax9.set_xlabel("Number of Interferers (N)", color='white')
        ax9.set_ylabel("Violation Probability (%)", color='white')
        ax9.set_title("Risk vs. Deployment Density", color='white')
        ax9.legend(facecolor='#1a1a2e', labelcolor='white', fontsize=9)
        ax9.tick_params(colors='white')
        ax9.grid(color='#333', alpha=0.4)
        for sp in ax9.spines.values():
            sp.set_color('#444')
        plt.tight_layout()
        st.pyplot(fig8)

# ─────────────────────────────────────────────────────────────────────────────
# TAB 6 — CONTRIBUTION SUMMARY
# ─────────────────────────────────────────────────────────────────────────────
elif selected_tab == "📋 Contribution Summary":
    st.title("📋 ITU-R Contribution Summary Generator")
    ex("A contribution carries the weight of your administration's regulatory position — vague language weakens it. Specific dB margins, named propagation models, and cited Recommendation/RR article numbers are what Working Party rapporteurs actually use to draft agreed text.")

    st.subheader("Scenario Setup")
    contrib_title = st.text_input("Analysis Title", value="Interference Assessment: New IMT Allocation vs. Radio Altimeter Band")
    submitter = st.text_input("Submitting Administration", value="United States of America")
    meeting = st.text_input("Target Meeting", value="ITU-R WP 5D Meeting, [Date]",
        help="e.g. 'WP 5D #44, Oct 2025' or 'WP 4C #38, Nov 2025' or 'WP 7C #30, Mar 2026'")
    wrc_agenda = st.text_input("WRC Agenda Item", value="AI [X.Y] — [description]")

    protected_band = st.selectbox("Protected FAA System", list(FAA_BANDS.keys()))
    b_c = FAA_BANDS[protected_band]

    st.subheader("Analysis Parameters (from your prior analysis)")
    col1, col2, col3 = st.columns(3)
    with col1:
        c_tx_power = st.number_input("Interferer EIRP (dBm)", value=58.0)
        c_freq = st.number_input("Interferer Frequency (MHz)", value=4300.0)
        c_sep = st.number_input("Frequency Separation (MHz)", value=200.0)
    with col2:
        c_dist = st.number_input("Min Separation Distance (km)", value=5.0)
        c_in = st.number_input("Computed I/N (dB)", value=-4.0)
        c_viol_prob = st.number_input("Violation Probability (%)", value=8.5)
    with col3:
        c_margin = b_c["in_threshold_db"] - c_in
        st.metric("Protection Margin", f"{c_margin:.1f} dB",
                  delta_color="normal" if c_margin >= 0 else "inverse")
        compatible = c_margin >= 0 and c_viol_prob < 5.0

    # Generate contribution text
    if st.button("📄 Generate Contribution Technical Summary", type="primary"):
        status = "COMPATIBLE" if compatible else "INCOMPATIBLE"
        recommendation = (
            "The analysis supports compatibility subject to the conditions specified below."
            if compatible else
            "The analysis demonstrates that the proposed new allocation/use presents an unacceptable "
            "risk of harmful interference to the protected aeronautical service. The United States "
            "opposes this proposal without additional protective measures."
        )

        contrib_text = f"""
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
ITU-R WORKING PARTY [5D / 5B] — TECHNICAL CONTRIBUTION
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
Document:      [Document Number TBD]
Submitted by:  {submitter}
Meeting:       {meeting}
Agenda Item:   {wrc_agenda}

TITLE: {contrib_title}

━━━━━━━━━━━━━━━━━━━━━━
1. EXECUTIVE SUMMARY
━━━━━━━━━━━━━━━━━━━━━━
{submitter} has conducted an interference analysis assessing the 
compatibility of proposed new spectrum use at {c_freq:.1f} MHz with 
the protected aeronautical service: {b_c['system']} operating in 
{b_c['f_low_mhz']}–{b_c['f_high_mhz']} MHz ({b_c['allocation']}).

PRELIMINARY FINDING: {status}
{recommendation}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
2. PROTECTION CRITERIA
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
The following protection criteria are applied, consistent with 
ITU-R Recommendation {b_c.get('rtca_doc','[applicable Rec.]')} 
and established ICAO/FAA standards:

  • Protected System:     {b_c['system']}
  • Operating Band:       {b_c['f_low_mhz']}–{b_c['f_high_mhz']} MHz
  • Allocation:           {b_c['allocation']}
  • Safety-of-Life:       {'YES — RR No. 4.10 applies' if b_c['safety_of_life'] else 'No'}
  • I/N Threshold:        {b_c['in_threshold_db']} dB
  • Receiver Noise Floor: {b_c['noise_floor_dbm']} dBm (estimated)
  • Frequency Separation: {c_sep:.1f} MHz from edge of protected band
  • Reference Standard:   {b_c['rtca_doc']}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
3. INTERFERENCE ANALYSIS METHODOLOGY
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
The analysis follows the methodology of ITU-R Recommendation 
SM.2028 (Monte Carlo simulation) and ITU-R M.1642 (IMT/ARNS 
compatibility). Propagation was assessed using:

  • ITU-R P.528  — Aeronautical mobile/radionavigation propagation
  • ITU-R P.452  — Terrestrial interference (point-to-area)
  • ITU-R P.676  — Gaseous attenuation (itur library implementation)

Scenario parameters:
  • Interferer EIRP:           {c_tx_power:.1f} dBm
  • Interferer Frequency:      {c_freq:.1f} MHz
  • Victim Receiver Location:  Airborne / ground-based aeronautical
  • Minimum Separation:        {c_dist:.1f} km

━━━━━━━━━━━━━━━━━━━━━━━━━━━━
4. KEY RESULTS
━━━━━━━━━━━━━━━━━━━━━━━━━━━━
  • Computed I/N (worst case):  {c_in:.1f} dB
  • I/N Threshold:              {b_c['in_threshold_db']} dB
  • Protection Margin:          {c_margin:.1f} dB  ({'VIOLATED' if c_margin < 0 else 'MET'})
  • Monte Carlo violation prob: {c_viol_prob:.1f}%  ({'EXCEEDS' if c_viol_prob > 5 else 'within'} 5% criterion)

{'⚠ The I/N threshold is violated under the assessed scenario.' if c_margin < 0 else '✓ The I/N threshold is nominally met under median conditions.'}
{'⚠ The 5% exceedance criterion is not met — harmful interference cannot be excluded.' if c_viol_prob > 5 else '✓ Exceedance probability is within the 5% criterion.'}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
5. PROPOSED REGULATORY ACTION
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
{'The United States proposes that the Working Party adopt a footnote or Resolution establishing the following protection measures to safeguard ' + b_c["system"] + ':' if not compatible else 'The United States notes that compatibility can be maintained provided the following conditions are observed:'}

  [a] Power flux density limits at the edge of the aeronautical 
      service area shall not exceed [XX] dB(W/m²)/MHz;
  [b] Minimum coordination distance of [XX] km from aeronautical 
      facilities shall be maintained;
  [c] Out-of-band emission limits into {b_c['f_low_mhz']}–{b_c['f_high_mhz']} MHz 
      shall not exceed [XX] dBc/MHz at {c_sep:.0f} MHz offset;
  [d] Further study is required under [specific WRC-27 agenda item].

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
6. REGULATORY BASIS
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
  • Radio Regulations No. 4.10 — Protection of safety-of-life services
  • Radio Regulations No. 5.444 — ARNS frequency coordination
  • RR Resolution 233 — Protection of RNSS systems
  • ITU-R M.1642 — Methodology for IMT/ARNS compatibility
  • ICAO Annex 10 — Aeronautical Telecommunications

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
7. COORDINATION NOTE
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
This analysis was developed in coordination with the Federal 
Aviation Administration (FAA) and submitted through the National 
Telecommunications and Information Administration (NTIA). 
ICAO has been informed and is developing complementary input 
through [ICAO Navigation Systems Panel / FMG].

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
8. CONCLUSION
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
{submitter} urges the Working Party to take the results of this 
analysis into account when developing the CPM Report method(s) 
for WRC-27 Agenda Item {wrc_agenda}.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
[End of Document]
Generated by FAA RF Interference Analysis Tool
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
"""
        st.text_area("Generated Contribution Draft", contrib_text, height=600)
        st.download_button("⬇️ Download as .txt", contrib_text,
                           file_name="ITU_R_Contribution_Draft.txt", mime="text/plain")
        ex("Before submission through NTIA: verify all numerical results with SEAMCAT or ITU-R SoftTools for P.452/P.528; confirm I/N threshold citations against current RTCA DO standards; have FAA spectrum office review for classification and policy alignment.")

    # Regulatory citation quick reference
    st.subheader("Regulatory Citation Quick Reference")
    ex("Regulatory citations transform a technical finding into a legal position — 'our analysis shows threshold violated' is an opinion; 'computed I/N exceeds the criterion in M.1318 Annex 1 Step 1b, invoking RR No. 4.10' is a treaty-level regulatory argument.")
    reg_refs = pd.DataFrame([
        ["RR No. 4.10", "No harmful interference to safety-of-life services", "Strongest lever — invoked for ALL FAA safety systems"],
        ["RR No. 1.153", "Occupied bandwidth definition — band containing 99% of total mean power (0.5% each edge)", "Cite when challenging a proponent's bandwidth claim; defines where OOB zone starts"],
        ["RR No. 5.444", "ARNS protection at 960–1215 MHz", "Use for DME/TACAN/SSR/TCAS/ADS-B bands"],
        ["RR No. 5.328", "ARNS at 108–137 MHz; RNSS cannot claim protection from ARNS in 1164–1215 MHz", "VOR/ILS protection; also limits RNSS protection claims vs DME"],
        ["RR Appendix 3 (Rev. WRC-12)", "Maximum permitted spurious power levels: 43+10·log(P) dB or 70 dBc (general); 60 dBc for space stations", "Cite when spurious emissions land in a protected band. For satellites: 60 dBc is the limit. Formula scales with transmitter power P in watts."],
        ["RR Resolution 233", "Protection of RNSS (GPS/GNSS)", "Use for all GPS/GNSS band defense"],
        ["RR Resolution 750", "IMT and safety services coexistence", "Relevant for all WP 5D IMT proposals"],
        ["ITU-R SM.1540", "Unwanted emissions in OOB domain falling into adjacent allocated bands", "Cite when OOB emissions from a new service threaten an adjacent aeronautical band. Establishes the OOB mask framework and the 250% bandwidth boundary."],
        ["ITU-R SM.1541", "Unwanted emissions in OOB domain — 23 dB attenuation at allocated band edge", "Cite the 23 dB rule: OOB mask must be ≥23 dB down at the edge of the allocated band (from RR 1.153 — 0.5% power criterion). This is the minimum mask floor."],
        ["ITU-R SM.329", "Spurious domain emission limits and measurement methodology", "Cite for spurious emission limit disputes. Companion to RR Appendix 3."],
        ["ITU-R M.1318 Annex 1", "c = a − b methodology for aggregate interference to GNSS", "Cite as the formal calculation framework for GPS L1/L2/L5 protection analysis"],
        ["ITU-R M.1477 Annex 5", "Aeronautical safety margin ≥6 dB for GNSS; +10 dB for narrowband (≤700 Hz) interferers", "Cite to justify b = 6 dB in c = a − b; invoke narrowband rule for CW/tonal interferers"],
        ["ITU-R M.1904", "GLONASS spaceborne receiver — safety margin = 6 dB", "Cite alongside M.1905 for GLONASS band contributions; reinforces 6 dB doctrine"],
        ["ITU-R M.1905", "Safety margin must be applied for RNSS safety-of-life interference analyses (6 dB aero)", "Broadest RNSS safety margin authority — applies to ALL RNSS systems"],
        ["ITU-R M.1642", "IMT→ARNS methodology", "Cite as methodology basis for non-GNSS aeronautical analysis"],
        ["ITU-R SM.2028", "Monte Carlo simulation methodology", "Cite to validate your simulation approach"],
        ["ITU-R P.528", "Aeronautical propagation model", "Model authority — cite when using P.528 curves"],
        ["ICAO Annex 10", "Aeronautical telecomm standards", "Aligns ITU-R work with ICAO civil aviation requirements"],
    ], columns=["Reference", "Subject", "When to Cite"])
    st.dataframe(reg_refs, use_container_width=True)
    ex("Citation stacking: SM.1540 + SM.1541 define the OOB framework → RR Appendix 3 defines spurious limits → RR No. 4.10 is the enforcement mechanism. Together they give you the complete chain from emission characterization to regulatory consequence.")
    ex("OOB domain boundary: 250% of occupied bandwidth each side of the channel (per SM.1540). Inside that zone = SM.1540/SM.1541 OOB rules. Outside it = RR Appendix 3 spurious rules. Challenge any proponent who conflates the two.")

    # OOB Domain Quick Calculator
    st.subheader("🧮 OOB Domain Boundary Calculator")
    ex("Instantly determine the OOB zone boundary and the 23 dB attenuation point for any transmitter. Use this to check whether a neighboring band falls inside the OOB zone (SM.1540/SM.1541 apply) or the spurious domain (RR Appendix 3 applies).")
    oob_c1, oob_c2 = st.columns(2)
    with oob_c1:
        oob_center  = st.number_input("Transmit Center Frequency (MHz)", value=1219.0, step=1.0,
            help="Center frequency of the interfering transmitter")
        oob_bw      = st.number_input("Occupied Bandwidth (MHz)", value=8.0, step=0.5,
            help="99% power bandwidth per RR No. 1.153 — not the channel plan bandwidth")
        oob_power_w = st.number_input("Peak Power (W)", value=8000.0, step=100.0,
            help="Transmitter peak power in Watts. Used to calculate RR Appendix 3 spurious limit.")
    with oob_c2:
        # OOB boundary
        half_bw      = oob_bw / 2
        oob_boundary = oob_bw * 2.5  # 250% each side
        band_low     = oob_center - half_bw
        band_high    = oob_center + half_bw
        oob_low      = oob_center - oob_bw * 2.5 - half_bw
        oob_high     = oob_center + oob_bw * 2.5 + half_bw

        # RR Appendix 3 spurious limit
        p_dbw        = 10 * np.log10(oob_power_w)
        spurious_rel = min(43 + 10 * np.log10(oob_power_w), 60)  # space station rule
        spurious_abs = p_dbw - spurious_rel

        st.metric("Allocated Band Edge (low)", f"{band_low:.1f} MHz")
        st.metric("Allocated Band Edge (high)", f"{band_high:.1f} MHz")
        st.metric("OOB Zone extends to (low)", f"{oob_low:.1f} MHz",
            help="Below this = spurious domain (RR Appendix 3). Above this = OOB domain (SM.1540/SM.1541)")
        st.metric("OOB Zone extends to (high)", f"{oob_high:.1f} MHz")
        st.metric("RR App.3 Spurious Limit (space stn)", f"{spurious_rel:.1f} dBc  ({spurious_abs:.1f} dBW)",
            help=f"= min(43+10·log({oob_power_w:.0f}W), 60) = {spurious_rel:.1f} dBc below {p_dbw:.1f} dBW carrier")

    oob_check_freq = st.number_input("Check a specific frequency (MHz) — is it in OOB or spurious domain?",
        value=1215.0, step=1.0)
    dist_from_center = abs(oob_check_freq - oob_center)
    dist_from_edge   = dist_from_center - half_bw
    if dist_from_edge <= 0:
        st.info(f"📡 {oob_check_freq} MHz is **inside the transmit channel** — this is in-band interference, not OOB.")
    elif dist_from_edge <= oob_boundary:
        pct = (dist_from_edge / (oob_bw)) * 100
        st.warning(f"⚠️ {oob_check_freq} MHz is in the **OOB domain** ({dist_from_edge:.1f} MHz from band edge = {pct:.0f}% of BN). "
                   f"**SM.1540 and SM.1541 apply.** OOB mask must be ≥23 dB down at the band edge.")
    else:
        st.error(f"🔴 {oob_check_freq} MHz is in the **spurious domain** ({dist_from_edge:.1f} MHz from band edge, beyond 250% BN boundary). "
                 f"**RR Appendix 3 applies.** Limit = {spurious_rel:.1f} dBc ({spurious_abs:.1f} dBW) for space stations.")

# ─────────────────────────────────────────────────────────────────────────────
# TAB 7 — TUTORIAL
# ─────────────────────────────────────────────────────────────────────────────
elif selected_tab == "📚 Tutorial":
    st.title("📚 How to Use This Tool")
    ex("The five-step workflow maps directly to the ITU-R study cycle timeline: identify threats at CPM study launch (4–5 years before WRC), quantify during WP meetings, draft contributions 6 weeks before each meeting, engage at the meeting, implement after WRC.")

    st.markdown("---")
    st.header("🗺️ The Big Picture — Your Workflow")
    ex("The study cycle runs approximately 4 years from WRC agenda item adoption to the next WRC. Planting conservative protection criteria early in Step 1 Recommendations is far more effective than fighting a nearly-agreed CPM method in the final study cycle year.")

    st.markdown("""
```
1. IDENTIFY  →  A new ITU-R proposal may affect a protected FAA band
2. ASSESS    →  Run link budget + Monte Carlo to quantify the threat
3. COMPARE   →  Check results against ITU-R protection criteria
4. DRAFT     →  Build the US contribution with technical findings
5. ENGAGE    →  Submit through FAA or NTIA, coordinate with ICAO/allies
```
    """)

    st.markdown("---")
    st.header("📡 Module 1 — Protected Bands")
    ex("Always start here — know exactly what you're defending before touching any analysis tool.")

    st.info("📌 **This walkthrough describes the 📡 Protected Bands tab.** Open it from the sidebar to use it interactively.")
    if st.button("➡️ Go to 📡 Protected Bands module", key="tb_lb1"):
        st.session_state["_nav_to"] = "📡 Protected Bands"
        st.rerun()

    col1, col2 = st.columns(2)
    with col1:
        st.subheader("What to do")
        st.markdown("""
**Step 1.** Open the **Spectrum Overview** chart — scan it visually to see if the new proposal's frequency sits near or inside a protected zone.

**Step 2.** Select the threatened band from the **Band Details** dropdown. Note these four numbers — you will use them in every subsequent module:
- Lower / upper frequency bounds
- I/N Threshold (dB)
- Noise floor (dBm)
- RTCA reference document

**Step 3.** Confirm the band is marked **Safety-of-Life** — if it is, RR No. 4.10 applies and your policy position is automatically stronger.
        """)
    with col2:
        st.subheader("Key questions to answer")
        st.markdown("""
| Question | Where to find it |
|---|---|
| What allocation protects this band? | Allocation field |
| What I/N threshold must I defend? | I/N Threshold field |
| What RTCA standard governs the receiver? | RTCA Doc field |
| Is it safety-of-life? | Red 🔴 label |
| How wide is the band? | BW (MHz) column |
        """)

    st.markdown("---")
    st.header("🔗 Module 2 — Link Budget")
    ex("The link budget tells you how much interference actually reaches the victim receiver — this is the single-scenario worst-case calculation.")

    st.info("📌 **This walkthrough describes the 🔗 Link Budget tab.** Open it in the sidebar to follow along interactively — the waterfall chart, distance sweep, and protection margin metric are all live and interactive there.")
    if st.button("➡️ Go to Link Budget module", key="tutorial_goto_lb"):
        st.session_state["_nav_to"] = "🔗 Link Budget"
        st.rerun()

    col1, col2 = st.columns(2)
    with col1:
        st.subheader("Step-by-step walkthrough")
        st.markdown("""
**Step 1 — Set up the interferer (left panel)**
- Enter the new system's transmit power in dBm
  - *Tip: 43 dBm = 20W typical LTE macro base station*
  - *Tip: 30 dBm = 1W small cell / low-power device*
- Enter antenna gain toward the victim receiver
  - *Use 0 dBi for worst case (isotropic toward victim)*
- Set cable loss (usually 1–3 dB)

**Step 2 — Choose propagation model**
- **FSPL** — Use first; gives you the worst-case bound (most interference)
- **P.452** — Use for ground-to-ground scenarios; pick terrain type
- **P.528** — Use when victim is airborne (aircraft, helicopter)
- *Rule: Always start with FSPL. If it passes, you're protected. If it fails, run P.528/P.452 for a more realistic picture.*

**Step 3 — Set up the victim receiver (right panel)**
- Click **"Auto-fill from FAA band"** and select your protected band
- This loads the correct noise floor and I/N threshold automatically

**Step 4 — Click Run Link Budget**
- Read the **Protection Margin** metric — this is your headline number
- Positive = protected; negative = threshold violated
        """)
    with col2:
        st.subheader("How to interpret results")
        st.markdown("""
| Protection Margin | Meaning | Action |
|---|---|---|
| > +10 dB | Well protected | Compatible; document and move on |
| 0 to +10 dB | Marginally protected | Flag for conservative assumptions; run Monte Carlo |
| < 0 dB | Threshold violated | Cite harmful interference; draft objection |

**Reading the waterfall chart:**
- Green bars = gains (help the signal)
- Red bars = losses (reduce interference reaching victim)
- The gap between Rx Power and the red dotted threshold line is your protection margin

**Reading the distance sweep:**
- The red dotted line shows the minimum path loss needed for protection
- Where FSPL curve crosses that line = minimum safe separation distance
- Cite this distance in your contribution as a coordination zone requirement

**Common mistakes to avoid:**
- Don't use FSPL as your final model for regulatory submissions — reviewers will push back
- Don't forget to account for OOBE (out-of-band emissions) — the interferer's frequency may be outside the protected band but its sidelobes may not be
- Always use 0 dBi receive antenna gain for worst-case analysis
        """)

    st.markdown("---")
    st.header("📊 Module 3 — Noise & I/N Analysis")
    ex("Use this module to verify protection criteria and build the receiver characterization section of your contribution.")

    st.info("📌 **This walkthrough describes the 📊 Noise & I/N tab.** Open it from the sidebar to use it interactively.")
    if st.button("➡️ Go to 📊 Noise & I/N module", key="tb_lb3"):
        st.session_state["_nav_to"] = "📊 Noise & I/N"
        st.rerun()

    st.markdown("""
**When to use this module:**
- When you need to calculate the noise floor for a specific receiver that isn't in the FAA band database
- When you want to show how I/N varies across a range of interference levels (for the heatmap)
- When reviewing another administration's claimed protection criteria — enter their numbers and verify

**Key formula to remember:**
```
Noise Floor (dBm) = -174 + 10·log₁₀(Bandwidth_Hz) + Noise_Figure_dB
```
- -174 dBm/Hz is thermal noise at 290K (room temperature)
- Bandwidth widens the noise floor upward (10 MHz BW = 10 dB higher floor than 1 MHz BW)
- Every dB of noise figure directly raises the floor

**I/N threshold selection guide:**
| System | Threshold | Rationale |
|---|---|---|
| GNSS / GPS | −10 dB | Satellite signal is extremely weak; any noise rise is catastrophic |
| ADS-B / Mode-S | −10 dB | Safety surveillance; false negatives unacceptable |
| Radio Altimeter | −6 dB | Safety-of-life; conservative protection required |
| VOR / ILS | −6 dB | Approach/landing guidance; high consequence |
| General ARNS | −6 dB | ITU-R M.1642 standard criterion |
    """)

    st.markdown("---")
    st.header("🌐 Module 4 — Propagation")
    ex("Use this to compare models side-by-side and select the most appropriate one before running your formal analysis.")

    st.info("📌 **This walkthrough describes the 🌐 Propagation tab.** Open it from the sidebar to use it interactively.")
    if st.button("➡️ Go to 🌐 Propagation module", key="tb_lb4"):
        st.session_state["_nav_to"] = "🌐 Propagation"
        st.rerun()

    st.markdown("""
**Model selection decision tree:**
```
Is the victim receiver airborne?
  YES → Use P.528
  NO  → Is the path over ocean or open terrain?
          YES → Use P.452 (open)
          NO  → Use P.452 (suburban or urban)
          
Always run FSPL first as a bounding calculation.
```

**The atmospheric attenuation chart (P.676 via itur):**
- Shows how much signal is absorbed by oxygen and water vapor over your chosen path length
- Above ~10 GHz, this becomes significant and actually helps protect receivers
- Below ~3 GHz, gaseous absorption is minimal — don't count on it for protection

**What the time percentage means in P.528:**
- 50% = median conditions (use for typical interference assessment)
- 1% = worst-case propagation (use for regulatory/protection studies — most conservative)
- ITU-R contributions for protection should use 1% or worst-case conditions
    """)

    st.markdown("---")
    st.header("🎲 Module 5 — Monte Carlo")
    ex("Monte Carlo is the ITU-R standard methodology for aggregate interference — it's what distinguishes a rigorous contribution from a back-of-envelope calculation.")

    st.info("📌 **This walkthrough describes the 🎲 Monte Carlo tab.** Open it from the sidebar to use it interactively.")
    if st.button("➡️ Go to 🎲 Monte Carlo module", key="tb_lb5"):
        st.session_state["_nav_to"] = "🎲 Monte Carlo"
        st.rerun()

    col1, col2 = st.columns(2)
    with col1:
        st.subheader("How to set it up")
        st.markdown("""
**Number of interferers (N):**
- This is the most influential assumption — challenge it in others' contributions
- For IMT base stations: 5–20 per km² is typical urban density
- For satellite downlinks: 1 (single beam) to thousands (LEO constellation)
- Always ask: is the proponent's assumed N realistic for actual deployment?

**Deployment radius:**
- How far from the victim receiver could interferers realistically be placed?
- For aeronautical: consider that an aircraft at altitude has line-of-sight to transmitters hundreds of km away
- For ground receivers: 20–50 km is a reasonable urban deployment zone

**Propagation model for Monte Carlo:**
- Use P.528 if victim is airborne
- Use P.452 for ground victims
- FSPL gives you the worst-case bound
        """)
    with col2:
        st.subheader("How to interpret results")
        st.markdown("""
**The four key numbers:**
| Metric | How to use it |
|---|---|
| Median I/N (50th pct) | Typical day; use for general compatibility claim |
| 95th percentile I/N | Near-worst case; use in contribution headline finding |
| 99th percentile I/N | Extreme case; use to argue for extra margin |
| Violation probability | Must be < 5% (ITU-R standard criterion) |

**Reading the CCDF chart:**
- X-axis = I/N level; Y-axis = probability of exceeding that level
- Draw a vertical line at your I/N threshold — read across to Y-axis
- That Y value is your violation probability
- If it's above 5% (0.05), the scenario is incompatible under ITU-R criteria

**The sensitivity sweep:**
- Shows how violation probability grows with deployment density
- Use this to argue for a maximum density limit in your contribution
- Example: "Compatibility is maintained only if deployment density does not exceed N units per km²"
        """)

    st.markdown("---")
    st.header("📋 Module 6 — Contribution Summary")
    ex("This module generates a structured ITU-R contribution draft — fill in your analysis numbers and it produces properly formatted regulatory text.")

    st.markdown("""
**Anatomy of an ITU-R contribution:**
| Section | Purpose | Tips |
|---|---|---|
| Executive Summary | One-paragraph finding | Lead with your conclusion — Compatible or Incompatible |
| Protection Criteria | What you're protecting | Cite I/N threshold, RTCA doc, RR article |
| Methodology | How you analyzed it | Cite P.528, SM.2028, M.1642 |
| Key Results | Your numbers | I/N value, margin, violation probability |
| Proposed Regulatory Action | What you want | Specific PFD limits, coordination distances, OOBE masks |
| Regulatory Basis | Legal grounding | RR 4.10, Resolution 233, etc. |
| Coordination Note | Who's behind you | NTIA, FAA, ICAO — weight in numbers |

**Tips for getting your contribution accepted:**
1. **Lead with safety** — open with the safety-of-life consequence of interference, not the math
2. **Be specific** — "interference may occur" is weak; "I/N threshold exceeded by 4.2 dB at 5 km separation" is strong
3. **Pre-coordinate with ICAO** — an ICAO liaison statement supporting your position changes the room dynamics entirely
4. **Propose text, not just concerns** — submissions that include proposed Radio Regulations language are far more influential than those that only raise problems
5. **Bracket aggressively early** — in early study cycles, propose conservative limits; there will be pressure to relax them later
    """)

    st.markdown("---")
    st.header("🔍 Module 7 — Contribution Analyzer (AI-Powered)")
    ex("The next module uses AI to read any WP 5D, 5B, 4C, 7B, 7C or other ITU-R contribution you paste in and instantly gives you FAA-focused policy guidance.")
    st.markdown("""
**How to use it:**
1. Go to the **🤖 Contribution Analyzer** module in the sidebar
2. Paste the text or key sections of any ITU-R contribution
3. Select the Working Party — **4C, 5D, 5B, 7B, 7C** and more are all supported
4. Fill in the metadata (document number, working party, submitting admin)
5. Click **Analyze** — the AI will assess:
   - Which FAA protected bands are at risk
   - Which WRC-27 agenda item (AI 1.7, 1.13, 1.15, 1.17, or 1.19) is implicated
   - What the submitting administration is trying to achieve
   - What the US/FAA policy position should be
   - What counter-arguments to raise
   - Which RR articles and Recommendations to cite in response
   - Whether to oppose, support, or propose amendments
    """)

# ─────────────────────────────────────────────────────────────────────────────
# TAB 8 — AI CONTRIBUTION ANALYZER
# ─────────────────────────────────────────────────────────────────────────────
elif selected_tab == "🤖 Contribution Analyzer":
    st.title("🤖 AI Contribution Analyzer")
    st.caption("Policy support for **WP 5B · WP 5D · WP 4C · WP 7B · WP 7C** — FAA aeronautical spectrum protection")
    ex("Upload or paste an ITU-R contribution from WP 5B, 5D, 4C, 7B, or 7C. The analysis follows the ITU-R methodology guidelines for that specific Working Party — not a general engineering review.")

    # Working party context callout
    wp_context = {
        "WP 5D (IMT/Mobile)":                           ("AI 1.7",       "IMT near Radio Altimeter 4.4–4.8 GHz and FAA fixed links"),
        "WP 5B (Maritime/Radiodetermination)":           ("Various",      "Maritime/radiolocation services near aeronautical bands"),
        "WP 4C (MSS / DC-MSS-IMT)":                     ("AI 1.13",      "MSS candidate bands adjacent to DME (960 MHz), AMS(R)S (1525 MHz), ASR (2700 MHz)"),
        "WP 7B (Space Radiocommunication / Lunar SRS)":  ("AI 1.15",      "Lunar SRS near ASR, radar, and ARNS 5 GHz bands"),
        "WP 7C (EESS / Space Weather Sensors)":          ("AI 1.17/1.19", "EESS passive in RA band and space weather sensors near HF/VHF"),
    }

    # API key handling
    api_key = None
    try:
        api_key = st.secrets["ANTHROPIC_API_KEY"]
    except Exception:
        pass

    if not api_key:
        st.warning("⚠️ Anthropic API key not configured. See setup instructions below.")
        with st.expander("🔧 How to add your API key to Streamlit Cloud"):
            st.markdown("""
**One-time setup — takes 2 minutes:**

1. Go to [console.anthropic.com](https://console.anthropic.com) and create a free account
2. Click **API Keys** → **Create Key** → copy it
3. Go to your app at [share.streamlit.io](https://share.streamlit.io)
4. Click the **⋮ menu** next to your app → **Settings** → **Secrets**
5. Paste this exactly (replace with your real key):
```toml
ANTHROPIC_API_KEY = "sk-ant-..."
```
6. Click **Save** — the app restarts automatically

Once configured, the analyzer will work every time you visit the app.
            """)
        st.stop()

    # ── Working Party selector (drives the analysis profile — only manual input needed) ─
    working_party = st.selectbox(
        "Working Party",
        [
            "WP 5D (IMT/Mobile)",
            "WP 5B (Maritime/Radiodetermination)",
            "WP 4C (MSS / DC-MSS-IMT)",
            "WP 7B (Space Radiocommunication / Lunar SRS)",
            "WP 7C (EESS / Space Weather Sensors)",
        ],
        help="Select the WP this document comes from — this loads the correct ITU-R methodology profile for that Working Party."
    )

    # Show WP-specific FAA context callout
    if working_party in wp_context:
        ai_ref, ai_desc = wp_context[working_party]
        st.markdown(
            f"<div style='background:#1a2a3a;border-left:4px solid #ff8844;"
            f"padding:8px 12px;border-radius:4px;margin:4px 0'>"
            f"<b style='color:#ff8844'>⚠️ WRC-27 Watch — {ai_ref}:</b> "
            f"<span style='color:#ffddaa'>{ai_desc}</span></div>",
            unsafe_allow_html=True
        )

    # All other metadata (doc number, admin, meeting date, agenda item, doc type)
    # is extracted automatically from the document text during analysis.
    # These variables are set to empty strings so downstream code that references
    # them still works — the AI will identify and populate them from the document.
    doc_number       = ""
    submitting_admin = ""
    meeting_date     = ""
    agenda_item      = ""
    doc_type         = "Auto-detected"

    st.subheader("Contribution Input")
    ex("Upload a PDF or Word document, paste text, or use Batch Upload. Document number, submitting administration, agenda item, and all other metadata are extracted automatically.")

    # ── Input method tabs ─────────────────────────────────────────────────────
    input_tab_paste, input_tab_file, input_tab_batch = st.tabs([
        "📋 Paste Text",
        "📎 Upload Single File",
        "📦 Batch Upload (multiple documents)",
    ])

    contrib_from_file  = ""
    batch_mode_active  = False
    batch_files        = []

    with input_tab_file:
        uploaded_file = st.file_uploader(
            "Upload ITU-R contribution document",
            type=["pdf", "txt", "docx", "doc"],
            help="Accepts PDF, Word (.docx), or plain text files."
        )

        if uploaded_file is not None:
            file_type = uploaded_file.name.lower().split(".")[-1]
            st.caption(f"📄 File: **{uploaded_file.name}** ({uploaded_file.size / 1024:.1f} KB)")

            with st.spinner(f"Extracting text from {uploaded_file.name}..."):
                try:
                    if file_type == "txt":
                        raw_bytes = uploaded_file.read()
                        contrib_from_file = raw_bytes.decode("utf-8", errors="replace")

                    elif file_type == "pdf":
                        pdf_bytes = uploaded_file.read()
                        extracted = []
                        try:
                            import fitz
                            doc = fitz.open(stream=pdf_bytes, filetype="pdf")
                            for page in doc: extracted.append(page.get_text())
                            doc.close()
                            contrib_from_file = "\n\n".join(extracted)
                        except ImportError:
                            pass
                        if not contrib_from_file:
                            try:
                                import PyPDF2
                                reader = PyPDF2.PdfReader(io.BytesIO(pdf_bytes))
                                for page in reader.pages:
                                    t = page.extract_text()
                                    if t: extracted.append(t)
                                contrib_from_file = "\n\n".join(extracted)
                            except Exception:
                                pass
                        if not contrib_from_file:
                            try:
                                from pdfminer.high_level import extract_text as pdfminer_extract
                                contrib_from_file = pdfminer_extract(io.BytesIO(pdf_bytes))
                            except Exception:
                                pass
                        if not contrib_from_file:
                            st.warning("⚠️ Could not extract text from this PDF. Try the Paste Text tab.")

                    elif file_type in ("docx", "doc"):
                        docx_bytes = uploaded_file.read()
                        contrib_from_file = ""
                        tc_summary_single = ""

                        def extract_track_changes(raw_bytes):
                            import zipfile as _zf
                            from xml.etree import ElementTree as _ET
                            W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
                            try:
                                buf = io.BytesIO(raw_bytes)
                                with _zf.ZipFile(buf) as z:
                                    xml_bytes = z.read("word/document.xml")
                                root = _ET.fromstring(xml_bytes)
                                xml_str = xml_bytes.decode("utf-8", errors="replace")
                                has_ins = f'{{{W}}}ins' in xml_str
                                has_del = f'{{{W}}}del' in xml_str
                                has_tc  = has_ins or has_del
                                insertions, deletions = [], []
                                authors = set()
                                for elem in root.iter(f'{{{W}}}ins'):
                                    text = "".join(t.text or "" for t in elem.iter(f'{{{W}}}t'))
                                    author = elem.get(f'{{{W}}}author', '')
                                    date   = (elem.get(f'{{{W}}}date') or '')[:10]
                                    if text.strip():
                                        insertions.append({'author': author, 'date': date, 'text': text.strip()})
                                        if author: authors.add(author)
                                for elem in root.iter(f'{{{W}}}del'):
                                    text = "".join(t.text or "" for t in elem.iter(f'{{{W}}}delText'))
                                    author = elem.get(f'{{{W}}}author', '')
                                    date   = (elem.get(f'{{{W}}}date') or '')[:10]
                                    if text.strip():
                                        deletions.append({'author': author, 'date': date, 'text': text.strip()})
                                        if author: authors.add(author)
                                clean_paras = []
                                for para in root.iter(f'{{{W}}}p'):
                                    parts = []
                                    for elem in para.iter():
                                        if elem.tag == f'{{{W}}}t':
                                            parts.append(('keep', elem.text or ''))
                                        elif elem.tag == f'{{{W}}}delText':
                                            parts.append(('del', elem.text or ''))
                                    text = "".join(t for kind, t in parts if kind == 'keep').strip()
                                    if text: clean_paras.append(text)
                                clean_text = "\n".join(clean_paras)
                                if not has_tc:
                                    summary = "TRACK CHANGES: None detected."
                                else:
                                    author_str = ", ".join(sorted(authors)) if authors else "unknown"
                                    lines = [f"TRACK CHANGES: {len(insertions)} insertions, {len(deletions)} deletions. Editors: {author_str}", "", "INSERTED:"]
                                    for idx, item in enumerate(insertions[:30], 1):
                                        a = f" [{item['author']}]" if item['author'] else ""
                                        lines.append(f"  +[{idx}]{a}: {item['text'][:250]}")
                                    lines += ["", "DELETED:"]
                                    for idx, item in enumerate(deletions[:30], 1):
                                        a = f" [{item['author']}]" if item['author'] else ""
                                        lines.append(f"  -[{idx}]{a}: {item['text'][:250]}")
                                    summary = "\n".join(lines)
                                return clean_text, summary, has_tc, len(insertions), len(deletions)
                            except Exception:
                                return None, "", False, 0, 0

                        tc_clean, tc_summary_single, has_tc, n_ins, n_del = extract_track_changes(docx_bytes)
                        if not contrib_from_file:
                            try:
                                import mammoth
                                result = mammoth.extract_raw_text(io.BytesIO(docx_bytes))
                                if result and result.value and result.value.strip():
                                    contrib_from_file = result.value.strip()
                            except Exception: pass
                        if not contrib_from_file:
                            try:
                                from docx import Document as DocxDocument
                                doc2 = DocxDocument(io.BytesIO(docx_bytes))
                                paras = [p.text.strip() for p in doc2.paragraphs if p.text.strip()]
                                for table in doc2.tables:
                                    for row in table.rows:
                                        for cell in row.cells:
                                            if cell.text.strip(): paras.append(cell.text.strip())
                                contrib_from_file = "\n".join(paras)
                            except Exception: pass
                        if not contrib_from_file and tc_clean:
                            contrib_from_file = tc_clean
                        st.session_state["tc_summary"] = tc_summary_single
                        st.session_state["tc_has_tc"]  = has_tc
                        st.session_state["tc_n_ins"]   = n_ins
                        st.session_state["tc_n_del"]   = n_del
                        if has_tc:
                            st.markdown(f"<div style='background:#1a2a1a;border-left:4px solid #44bb44;padding:8px 12px;border-radius:4px;margin:6px 0'><b style='color:#44bb44'>📝 Track Changes Detected</b> — <span style='color:#aaffaa'>{n_ins} insertions, {n_del} deletions. AI will also scan text for revision language and cite supporting phrases.</span></div>", unsafe_allow_html=True)
                        else:
                            st.caption("📄 No track changes in Word file — AI will scan document text for 'new' vs 'revision' signals and cite 1–2 supporting phrases.")
                        if not contrib_from_file:
                            st.warning("⚠️ Could not extract text. Please use the Paste Text tab.")
                except Exception as e:
                    st.error(f"Error reading file: {e}")

            if contrib_from_file:
                wc = len(contrib_from_file.split()); cc = len(contrib_from_file)
                ok(f"Extracted: {wc:,} words, {cc:,} characters")
                MAX_CHARS = 28000
                if cc > MAX_CHARS:
                    warn(f"Document truncated to {MAX_CHARS:,} characters to fit AI context window.")
                    contrib_from_file = contrib_from_file[:MAX_CHARS] + "\n\n[... truncated ...]"
                with st.expander("👁️ Preview extracted text"):
                    st.text(contrib_from_file[:800] + ("..." if len(contrib_from_file) > 800 else ""))
        else:
            st.markdown("<div style='background:#1a1a2a;border:1px dashed #444;padding:20px;border-radius:6px;text-align:center;color:#888'>📎 Drag and drop a PDF, Word, or text file here<br><small>ITU-R contributions are typically PDF — download from TIES and upload directly</small></div>", unsafe_allow_html=True)

    with input_tab_batch:
        st.markdown("**📦 Batch Document Analysis — Chunked Processing**")
        ex("Upload as many files as you like. The app processes them in chunks of 15 to stay within Streamlit Cloud's timeout. Results accumulate across runs — run chunk 1, see results, run chunk 2, all results merge automatically.")

        # ── Persistent accumulated results in session state ───────────────────
        if "batch_accumulated" not in st.session_state:
            st.session_state.batch_accumulated = []   # list of result dicts
        if "batch_processed_names" not in st.session_state:
            st.session_state.batch_processed_names = set()  # filenames already done

        CHUNK_SIZE = 15  # safe for Streamlit Cloud ~10 min timeout at ~25s/doc

        batch_files_uploaded = st.file_uploader(
            "Upload ITU-R contribution documents (any quantity)",
            type=["pdf", "txt", "docx", "doc"],
            accept_multiple_files=True,
            help="No document count limit. Files are processed in chunks of 15. Upload all documents for the meeting and run multiple times — results accumulate.",
            key="batch_uploader"
        )

        if batch_files_uploaded:
            batch_mode_active = True
            total_kb  = sum(f.size for f in batch_files_uploaded) / 1024
            n_total   = len(batch_files_uploaded)
            n_done    = len(st.session_state.batch_processed_names)
            pending   = [f for f in batch_files_uploaded
                         if f.name not in st.session_state.batch_processed_names]
            n_pending = len(pending)
            next_chunk = pending[:CHUNK_SIZE]

            # Status metrics
            mc1, mc2, mc3, mc4 = st.columns(4)
            mc1.metric("Files uploaded",    n_total)
            mc2.metric("Already processed", n_done)
            mc3.metric("Remaining",         n_pending)
            mc4.metric("Next chunk size",   len(next_chunk))

            if n_pending == 0:
                ok(f"✅ All {n_total} files have been processed. Use the filter and download below, or clear results to reprocess.")
            else:
                est_lo = len(next_chunk) * 20 // 60
                est_hi = len(next_chunk) * 35 // 60
                time_str = f"~{max(1,est_lo)}–{max(2,est_hi)} min" if est_lo != est_hi else f"~{max(1,est_lo)} min"
                _depth_label = analysis_depth if 'analysis_depth' in dir() else "Quick"
                st.info(
                    f"📋 **{n_pending} files remaining.** "
                    f"Next run will process **{len(next_chunk)} files** ({time_str} at "
                    f"{'Quick' if _depth_label.startswith('Quick') else 'Standard'} depth). "
                    + (f"**{n_done} already processed** — results preserved below." if n_done else "")
                )
                with st.expander(f"Next chunk — {len(next_chunk)} files"):
                    for f in next_chunk:
                        st.caption(f"  📄 {f.name}  ({f.size/1024:.0f} KB)")
                if n_pending > CHUNK_SIZE:
                    with st.expander(f"Remaining after this chunk — {n_pending - CHUNK_SIZE} files"):
                        for f in pending[CHUNK_SIZE:]:
                            st.caption(f"  📄 {f.name}  ({f.size/1024:.0f} KB)")

            # Clear button
            if st.session_state.batch_accumulated:
                if st.button("🗑️ Clear all accumulated results and start fresh",
                             key="batch_clear", type="secondary"):
                    st.session_state.batch_accumulated   = []
                    st.session_state.batch_processed_names = set()
                    st.rerun()

            batch_files = next_chunk  # only next chunk goes to the run button

        else:
            # Show accumulated results from prior runs even without new upload
            if st.session_state.batch_accumulated:
                ok(f"📂 {len(st.session_state.batch_accumulated)} results from previous run(s) available below. Upload more files to continue, or use the download buttons.")
            else:
                st.markdown(
                    "<div style='background:#1a1a2a;border:1px dashed #444;padding:20px;"
                    "border-radius:6px;text-align:center;color:#888'>"
                    "📎 Upload your meeting documents here — no limit on quantity<br>"
                    "<small>They will be processed in chunks of 15 automatically</small>"
                    "</div>",
                    unsafe_allow_html=True
                )

    with input_tab_paste:
        contrib_pasted = st.text_area(
            "Paste contribution text here:",
            height=280,
            placeholder="""Paste the ITU-R contribution text here. Include:
- Executive summary / introduction
- Technical analysis sections (frequencies, propagation, results)
- Proposed regulatory text or amendments
- Conclusions and proposals"""
        )

    # ── Resolve single-document input ─────────────────────────────────────────
    contrib_input = contrib_from_file.strip() if contrib_from_file.strip() else (
        contrib_pasted.strip() if 'contrib_pasted' in dir() else ""
    )
    if contrib_from_file.strip() and 'contrib_pasted' in dir() and contrib_pasted.strip():
        st.info("📄 Both file and pasted text provided — file text will be used.")
        contrib_input = contrib_from_file.strip()
    if not batch_mode_active:
        if contrib_input:
            st.caption(f"✅ Ready — {len(contrib_input.split()):,} words from {'uploaded file' if contrib_from_file.strip() else 'pasted text'}")
        else:
            st.caption("⬆️ Upload a file, paste text, or use Batch Upload above")

    # ── Interference Classification Reference Card ────────────────────────────
    with st.expander("📋 ITU-R Interference Classification Reference — Applied in Every Analysis"):
        ex("The AI uses this exact framework from your slide to characterize interference in every contribution it analyzes.")

        st.markdown("### Unwanted Emissions — Source Characterization")
        col_r1, col_r2 = st.columns(2)
        with col_r1:
            st.markdown("""
<div style='background:#1a2a3a;border-left:4px solid #4488ff;padding:10px 14px;border-radius:4px;margin-bottom:8px'>
<b style='color:#4488ff'>SPURIOUS EMISSIONS</b><br>
<span style='color:#cce;font-size:0.88em'>
Frequencies just outside the necessary bandwidth whose level <b>may be reduced without 
affecting the corresponding transmission of information.</b><br><br>
Includes: harmonics, parasitic emissions, intermodulation products, frequency conversion products.<br>
<b>Explicitly EXCLUDES out-of-band emissions.</b><br><br>
Regulatory handle: RR Appendix 3 spurious emission limits
</span>
</div>""", unsafe_allow_html=True)
        with col_r2:
            st.markdown("""
<div style='background:#1a2a3a;border-left:4px solid #ff8844;padding:10px 14px;border-radius:4px;margin-bottom:8px'>
<b style='color:#ff8844'>OUT-OF-BAND (OOB) EMISSIONS</b><br>
<span style='color:#cce;font-size:0.88em'>
Frequencies <b>immediately outside the necessary bandwidth</b> which result from 
<b>the modulation process itself</b>, but excluding spurious emissions.<br><br>
Key distinction: OOB = caused by modulation. Spurious = hardware non-idealities.<br><br>
Regulatory handle: Emission designator masks and guard band requirements
</span>
</div>""", unsafe_allow_html=True)

        st.markdown("### Interference Classification — RR Article 1.166")
        col_i1, col_i2, col_i3 = st.columns(3)
        with col_i1:
            st.markdown("""
<div style='background:#3a1a1a;border-left:4px solid #ff4444;padding:10px 14px;border-radius:4px'>
<b style='color:#ff6666'>🔴 HARMFUL INTERFERENCE</b><br>
<span style='color:#aaa;font-size:0.82em'>RR 1.169</span><br><br>
<span style='color:#ffaaaa;font-size:0.88em'>
Interference which <b>endangers the functioning</b> of a radionavigation service or 
other safety services, OR <b>seriously degrades, obstructs, or repeatedly interrupts</b> 
a radiocommunication service operating in accordance with the RR.<br><br>
<b>→ Triggers RR No. 4.10<br>→ US must oppose<br>→ Demand cessation</b>
</span>
</div>""", unsafe_allow_html=True)
        with col_i2:
            st.markdown("""
<div style='background:#1a3a1a;border-left:4px solid #44bb44;padding:10px 14px;border-radius:4px'>
<b style='color:#66dd66'>🟢 PERMISSIBLE INTERFERENCE</b><br>
<span style='color:#aaa;font-size:0.82em'>RR 1.167</span><br><br>
<span style='color:#aaffaa;font-size:0.88em'>
Observed or predicted interference which <b>complies with quantitative interference 
and sharing criteria</b> in the Radio Regulations or in ITU-R Recommendations or 
in special agreements.<br><br>
<b>→ Acceptable if within criteria<br>→ Used in frequency coordination<br>→ Monitor the criteria</b>
</span>
</div>""", unsafe_allow_html=True)
        with col_i3:
            st.markdown("""
<div style='background:#2a2a1a;border-left:4px solid #ffbb44;padding:10px 14px;border-radius:4px'>
<b style='color:#ffdd66'>🟡 ACCEPTED INTERFERENCE</b><br>
<span style='color:#aaa;font-size:0.82em'>RR 1.168</span><br><br>
<span style='color:#ffffaa;font-size:0.88em'>
Interference at a <b>higher level than permissible</b>, which has been <b>agreed upon 
between two or more administrations</b> without prejudice to other administrations.<br><br>
<b>→ Requires explicit bilateral agreement<br>→ Cannot bind other admins<br>→ Verify scope carefully</b>
</span>
</div>""", unsafe_allow_html=True)

        st.caption("Both permissible interference and accepted interference are used in the coordination of frequency assignments between administrations (RR 1.167/1.168).")

    st.markdown("---")

    # Optional context
    with st.expander("➕ Additional Context (optional but helpful)"):
        user_concern = st.text_area("Specific FAA concern you want addressed:",
            placeholder="e.g., We are concerned this proposal could affect radio altimeter operations during CAT III approaches near 5G deployments at major airports.",
            height=100)
        prior_us_position = st.text_area("Prior US position on this topic (if known):",
            placeholder="e.g., The US has previously opposed new primary allocations in the 4200-4400 MHz band and submitted 5D/456 opposing similar proposals.",
            height=100)

    analysis_depth = st.selectbox("Analysis Depth", [
        "Quick assessment (key risks + recommended US position)",
        "Standard analysis (full policy brief with citations)",
        "Deep dive (comprehensive brief + draft response contribution outline)",
    ])

    # ── Analyze buttons ───────────────────────────────────────────────────────
    btn_col1, btn_col2 = st.columns([1, 1])
    with btn_col1:
        single_btn = st.button(
            "🔍 Analyze Contribution",
            type="primary",
            disabled=not contrib_input.strip(),
            use_container_width=True,
        )
    with btn_col2:
        batch_btn = st.button(
            f"📦 Run Batch Analysis ({len(batch_files)} files)" if batch_files else "📦 Batch Analysis",
            type="secondary",
            disabled=not batch_files,
            use_container_width=True,
        )

    # ══════════════════════════════════════════════════════════════════════════
    # BATCH ANALYSIS ENGINE
    # ══════════════════════════════════════════════════════════════════════════
    if batch_btn and batch_files:

        def extract_text_from_file(uploaded_f):
            """Extract text from any uploaded file type. Returns (text, tc_summary)."""
            ft = uploaded_f.name.lower().split(".")[-1]
            text = ""
            tc_sum = ""
            try:
                raw = uploaded_f.read()
                if ft == "txt":
                    text = raw.decode("utf-8", errors="replace")
                elif ft == "pdf":
                    extracted = []
                    try:
                        import fitz
                        doc = fitz.open(stream=raw, filetype="pdf")
                        for page in doc: extracted.append(page.get_text())
                        doc.close()
                        text = "\n\n".join(extracted)
                    except Exception: pass
                    if not text:
                        try:
                            import PyPDF2
                            reader = PyPDF2.PdfReader(io.BytesIO(raw))
                            for page in reader.pages:
                                t = page.extract_text()
                                if t: extracted.append(t)
                            text = "\n\n".join(extracted)
                        except Exception: pass
                elif ft in ("docx", "doc"):
                    import zipfile as _zf2
                    from xml.etree import ElementTree as _ET2
                    W2 = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
                    try:
                        buf2 = io.BytesIO(raw)
                        with _zf2.ZipFile(buf2) as z2:
                            xml2 = z2.read("word/document.xml")
                        root2 = _ET2.fromstring(xml2)
                        xml_str2 = xml2.decode("utf-8", errors="replace")
                        has_tc2 = f'{{{W2}}}ins' in xml_str2 or f'{{{W2}}}del' in xml_str2
                        ins2, del2 = [], []
                        for elem in root2.iter(f'{{{W2}}}ins'):
                            t = "".join(x.text or "" for x in elem.iter(f'{{{W2}}}t'))
                            if t.strip(): ins2.append(t.strip())
                        for elem in root2.iter(f'{{{W2}}}del'):
                            t = "".join(x.text or "" for x in elem.iter(f'{{{W2}}}delText'))
                            if t.strip(): del2.append(t.strip())
                        clean2 = []
                        for para in root2.iter(f'{{{W2}}}p'):
                            parts2 = []
                            for elem in para.iter():
                                if elem.tag == f'{{{W2}}}t':
                                    parts2.append(elem.text or '')
                            t2 = "".join(parts2).strip()
                            if t2: clean2.append(t2)
                        text = "\n".join(clean2)
                        if has_tc2:
                            tc_sum = f"TRACK CHANGES: {len(ins2)} insertions, {len(del2)} deletions."
                        else:
                            tc_sum = "TRACK CHANGES: None detected."
                    except Exception: pass
                    if not text:
                        try:
                            import mammoth
                            r2 = mammoth.extract_raw_text(io.BytesIO(raw))
                            text = r2.value.strip() if r2.value else ""
                        except Exception: pass
            except Exception:
                pass
            MAX = 22000
            if len(text) > MAX:
                text = text[:MAX] + "\n[truncated]"
            return text.strip(), tc_sum

        def run_single_analysis(text, tc_summary_text, fname, client_obj, sys_prompt, analysis_qs, depth_inst):
            """Run one API call for a single document. Returns analysis text."""
            um = f"""Analyze this ITU-R contribution. Filename: {fname}

METADATA — EXTRACT FROM DOCUMENT TEXT:
Identify and state in Section A: document number, submitting administration, meeting/date, WRC-27 agenda item, document type. Extract these from the document itself.

{f"TRACK CHANGES:{chr(10)}{tc_summary_text}{chr(10)}" if tc_summary_text else ""}
CONTRIBUTION TEXT:
{text}

{analysis_qs}

{depth_inst.replace('OUTPUT FORMAT: ','').replace('CONCISE — ','Concise — ')}

{'''═══════════════════════════════════════════════════════════════════
REQUIRED OUTPUT STRUCTURE
═══════════════════════════════════════════════════════════════════

## ⚡ FREQUENCY RELEVANCE SUMMARY
Compact table with actual MHz/GHz numbers:
| Proposed Frequency | FAA Band | FAA System | Gap/Overlap | Relationship | Study Type |
|---|---|---|---|---|---|

**REVIEW VERDICT: [REQUIRES HUMAN REVIEW / LIKELY NOT RELEVANT / FLAG FOR CLARIFICATION]**
One sentence.

## A) Document Overview + Status
- Title, source/administration, date — extracted from the document
- WRC-27 AI reference (if stated)
- 📋 STATUS: [NEW DOCUMENT / REVISION / UNCLEAR] — [confidence]
- Evidence: cite 1–2 verbatim phrases from document

## B) Key FAA Concerns (3 bullets max for batch mode)
## C) Recommended Actions (top 2 only)
## D) Draft U.S. Response (PATH 3 only — omit for US contributions or not-relevant documents)
One paragraph, 100–150 words. Ready-to-use US floor intervention citing specific RR articles and ITU-R Recommendations. State the US position and key technical objection in ITU-R meeting language.'''}"""
            response = client_obj.messages.create(
                model="claude-sonnet-4-5",   # Sonnet: 3× faster than Opus, higher rate limits — optimal for batch triage
                max_tokens=3500,
                messages=[{"role": "user", "content": um}],
                system=sys_prompt,
            )
            return response.content[0].text

        # Build shared system prompt
        wp_profile_key_b = WP_PROFILE_MAP.get(working_party)
        wp_profile_b     = WP_ANALYSIS_PROFILES.get(wp_profile_key_b) if wp_profile_key_b else None
        wp_framework_b   = f"Working Party: {working_party}. Apply appropriate analysis framework." if not wp_profile_b else f"Working Party: {working_party} — {wp_profile_b['primary_threat']}"
        faa_bands_b      = "\n".join([f"- {n}: {b['f_low_mhz']}–{b['f_high_mhz']} MHz ({b['allocation']}), I/N {b['in_threshold_db']} dB" for n,b in FAA_BANDS.items()])
        depth_b = {"Quick assessment (key risks + recommended US position)": "OUTPUT FORMAT: Concise — 2–3 bullets per section.",
                   "Standard analysis (full policy brief with citations)": "OUTPUT FORMAT: Standard structured analysis.",
                   "Deep dive (comprehensive brief + draft response contribution outline)": "OUTPUT FORMAT: Full detail."}[analysis_depth]
        sys_prompt_batch = f"You are an expert RF engineer supporting FAA/NTIA in ITU-R proceedings.\n{wp_framework_b}\nFAA PROTECTED BANDS:\n{faa_bands_b}\nACCURACY RULE: Never fabricate citations or frequency values. If uncertain, say 'Cannot confirm.'"

        client_b = anthropic.Anthropic(api_key=api_key)

        st.markdown("---")
        st.subheader("📊 Batch Analysis Progress")

        # Progress bar and status
        progress_bar = st.progress(0)
        status_text  = st.empty()

        batch_results = []  # list of dicts

        for idx, f in enumerate(batch_files):
            fname = f.name
            status_text.text(f"Analyzing {idx+1}/{len(batch_files)}: {fname}...")
            progress_bar.progress((idx) / len(batch_files))

            text_b, tc_b = extract_text_from_file(f)
            if not text_b:
                batch_results.append({"file": fname, "text": "", "tc": tc_b, "analysis": f"⚠️ Could not extract text from {fname}", "error": True})
                continue

            analysis_questions_b = analysis_questions if 'analysis_questions' in dir() else ""
            try:
                result_b = run_single_analysis(text_b, tc_b, fname, client_b, sys_prompt_batch, analysis_questions_b, depth_b)
                batch_results.append({"file": fname, "text": text_b, "tc": tc_b, "analysis": result_b, "error": False})
                # Write to Neo4j
                _neo4j_b = _neo4j_driver()
                if _neo4j_b:
                    try:
                        _neo4j_write_analysis(_neo4j_b, result_b,
                            {"doc_number": fname, "working_party": working_party,
                             "tc_summary": tc_b}, text_b)
                        _neo4j_b.close()
                    except Exception: pass
            except Exception as e:
                batch_results.append({"file": fname, "text": text_b, "tc": tc_b, "analysis": f"❌ API error: {e}", "error": True})

        progress_bar.progress(1.0)
        chunk_done = len([r for r in batch_results if not r["error"]])
        chunk_err  = len([r for r in batch_results if r["error"]])
        status_text.text(
            f"✅ Chunk complete — {chunk_done} analyzed"
            + (f", {chunk_err} error(s)" if chunk_err else "")
        )

        # ── Accumulate results into session state ─────────────────────────────
        for r in batch_results:
            if r["file"] not in st.session_state.batch_processed_names:
                st.session_state.batch_accumulated.append(r)
                st.session_state.batch_processed_names.add(r["file"])

        # Use the full accumulated set for all downstream display
        batch_results = st.session_state.batch_accumulated

        n_acc = len(batch_results)
        n_pending_after = len([
            f for f in (batch_files_uploaded or [])
            if f.name not in st.session_state.batch_processed_names
        ])

        acc_col1, acc_col2 = st.columns(2)
        with acc_col1:
            ok(f"📂 {n_acc} total results accumulated across all runs.")
        with acc_col2:
            if n_pending_after > 0:
                st.info(f"⏭️ {n_pending_after} files still pending — click **Run Batch Analysis** again to process the next chunk.")
            else:
                ok("✅ All uploaded files processed.")

        # ── Triage summary table ──────────────────────────────────────────────
        st.markdown("---")
        st.subheader("🗂️ Triage Summary")
        ex("Filter by Agenda Item or FAA system to quickly isolate documents that matter to a specific concern. The triage table and filtered results update instantly.")

        import re as _re_b

        # ── Known FAA entity aliases for matching ──────────────────────────────
        FAA_ENTITY_PATTERNS = {
            "Radio Altimeter (RA)":     r"radio alt(?:imeter)?|RA\b|WAICS|4[,.]?2.+4[,.]?4",
            "DME / TACAN":              r"DME|TACAN|distance measur|960.+1215|1215.+960",
            "GPS / GNSS L1":            r"GPS L1|GNSS L1|1559.+1610|1575",
            "GNSS L5 / ARNS":           r"GNSS L5|GPS L5|L5\b|1164.+1215",
            "ADS-B / Mode-S":           r"ADS-?B|Mode-?S|1090",
            "ASR (Airport Radar)":      r"ASR\b|airport surv|2700.+2900|2900.+2700",
            "ARSR (En-Route Radar)":    r"ARSR\b|en.?route.+radar|air route surv",
            "ILS / VOR":                r"ILS\b|VOR\b|localizer|glide slope|108.+118",
            "L-band AMS(R)S":           r"AMS\(R\)S|L.band.+sat|1525.+1559|1559.+1525",
            "MLS":                      r"\bMLS\b|microwave landing",
            "ARNS 5 GHz":               r"ARNS.+5|5[,.]?0.+5[,.]?15|5350.+5470",
            "Weather Radar":            r"weather radar|airborne.+radar|9[,.]?0.+9[,.]?5",
        }

        # ── Known Agenda Item patterns ─────────────────────────────────────────
        AI_PATTERNS = {
            "AI 1.7":  r"AI\s*1\.7|agenda item\s*1\.7|1\.7\b.*IMT|IMT.*4[,.]?4|4[,.]?8.*GHz",
            "AI 1.13": r"AI\s*1\.13|agenda item\s*1\.13|1\.13\b|MSS.+694|DC.MSS|925.+960|1475.+1518|2620.+2690",
            "AI 1.15": r"AI\s*1\.15|agenda item\s*1\.15|1\.15\b|lunar|SRS.*space.+space",
            "AI 1.17": r"AI\s*1\.17|agenda item\s*1\.17|1\.17\b|space weather|EESS.*passive.*27|27.*28.*GHz",
            "AI 1.19": r"AI\s*1\.19|agenda item\s*1\.19|1\.19\b|EESS.*4[,.]?2|4[,.]?2.*EESS|passive.*4[,.]?4",
        }

        triage_rows = []
        for res in batch_results:
            if res["error"]:
                triage_rows.append({
                    "File": res["file"], "Verdict": "ERROR", "Doc Status": "—",
                    "Proposed Freq": "—", "Agenda Item": "—", "FAA Systems": "—",
                    "_analysis": res["analysis"], "_error": True,
                })
                continue

            text = res["analysis"]

            # Verdict
            vm = _re_b.search(r'REVIEW VERDICT[:\s]+([A-Z ]+(?:HUMAN REVIEW|NOT RELEVANT|CLARIFICATION))', text)
            verdict = vm.group(1).strip() if vm else "SEE ANALYSIS"

            # Doc status
            sm = _re_b.search(r'STATUS[:\s]+(NEW DOCUMENT|REVISION|UNCLEAR)', text, _re_b.IGNORECASE)
            status = sm.group(1).upper() if sm else "—"

            # Proposed frequency — extract all freq mentions from table rows
            freq_matches = _re_b.findall(
                r'\|\s*(\d[\d,\s]*(?:\.\d+)?(?:–|-)\d[\d,\s]*(?:\.\d+)?\s*(?:MHz|GHz))[^|]*\|',
                text, _re_b.IGNORECASE
            )
            freq_str = "; ".join(dict.fromkeys(f.strip() for f in freq_matches[:3])) or "—"

            # Agenda item — match from analysis text
            ai_hits = []
            for ai_label, pattern in AI_PATTERNS.items():
                if _re_b.search(pattern, text, _re_b.IGNORECASE):
                    ai_hits.append(ai_label)
            ai_str = ", ".join(ai_hits) if ai_hits else "—"

            # FAA entities — scan for mentions
            faa_hits = []
            for entity, pattern in FAA_ENTITY_PATTERNS.items():
                if _re_b.search(pattern, text, _re_b.IGNORECASE):
                    faa_hits.append(entity)
            faa_str = ", ".join(faa_hits) if faa_hits else "—"

            triage_rows.append({
                "File":         res["file"],
                "Verdict":      verdict,
                "Doc Status":   status,
                "Proposed Freq":freq_str[:50],
                "Agenda Item":  ai_str,
                "FAA Systems":  faa_str,
                "_analysis":    text,
                "_error":       False,
            })

        # ── Filter controls ────────────────────────────────────────────────────
        all_ai  = sorted({r["Agenda Item"] for r in triage_rows if r["Agenda Item"] not in ("—","")})
        all_faa = sorted({
            ent.strip()
            for r in triage_rows
            for ent in r["FAA Systems"].split(",")
            if ent.strip() not in ("—","")
        })

        fcol1, fcol2, fcol3 = st.columns([1, 2, 2])
        with fcol1:
            verdict_filter = st.multiselect(
                "Filter by Verdict",
                ["REQUIRES HUMAN REVIEW", "FLAG FOR CLARIFICATION", "LIKELY NOT RELEVANT", "ERROR"],
                default=[],
                key="batch_verdict_filter",
            )
        with fcol2:
            ai_filter = st.multiselect(
                "Filter by Agenda Item",
                options=all_ai if all_ai else ["— none detected —"],
                default=[],
                key="batch_ai_filter",
            )
        with fcol3:
            faa_filter = st.multiselect(
                "Filter by FAA System",
                options=all_faa if all_faa else ["— none detected —"],
                default=[],
                key="batch_faa_filter",
            )

        # Apply filters
        def row_matches(r):
            if verdict_filter and not any(v in r["Verdict"] for v in verdict_filter):
                return False
            if ai_filter and not any(ai in r["Agenda Item"] for ai in ai_filter):
                return False
            if faa_filter and not any(faa in r["FAA Systems"] for faa in faa_filter):
                return False
            return True

        filtered = [r for r in triage_rows if row_matches(r)]
        n_total    = len(triage_rows)
        n_filtered = len(filtered)
        n_review   = sum(1 for r in filtered if "HUMAN REVIEW" in r["Verdict"])

        # Filter stats
        sf1, sf2, sf3, sf4 = st.columns(4)
        sf1.metric("Total documents",      n_total)
        sf2.metric("Matching filters",     n_filtered)
        sf3.metric("Requires human review", n_review,
                   delta=f"{n_review}/{n_filtered}" if n_filtered else None)
        sf4.metric("Not relevant",
                   sum(1 for r in filtered if "NOT RELEVANT" in r["Verdict"]))

        # ── Triage table (filtered) ────────────────────────────────────────────
        display_cols = ["File","Verdict","Doc Status","Proposed Freq","Agenda Item","FAA Systems"]
        triage_df = pd.DataFrame([{k: r[k] for k in display_cols} for r in filtered])

        def color_verdict(val):
            v = str(val)
            if "HUMAN REVIEW" in v:   return "background-color:#3a1a1a;color:#ff8888;font-weight:bold"
            if "NOT RELEVANT"  in v:  return "background-color:#1a3a1a;color:#88ff88"
            if "CLARIFICATION" in v:  return "background-color:#3a3a1a;color:#ffff88"
            if v == "ERROR":          return "background-color:#2a1a2a;color:#ff88ff"
            return ""

        if not triage_df.empty:
            st.dataframe(
                triage_df.style.map(color_verdict, subset=["Verdict"]),
                use_container_width=True, hide_index=True
            )
        else:
            st.info("No documents match the selected filters.")

        # ── Individual results (filtered) ─────────────────────────────────────
        st.markdown("---")
        active_filters = []
        if verdict_filter: active_filters.append(f"Verdict: {', '.join(verdict_filter)}")
        if ai_filter:      active_filters.append(f"AI: {', '.join(ai_filter)}")
        if faa_filter:     active_filters.append(f"FAA system: {', '.join(faa_filter)}")
        filter_desc = f" — Filtered: {', '.join(active_filters)}" if active_filters else f" — All {n_total} documents"
        st.subheader(f"📋 Full Analysis ({n_filtered} document{'s' if n_filtered != 1 else ''}){filter_desc}")

        for res in filtered:
            icon = "🔴" if "HUMAN REVIEW" in res["Verdict"] \
                   else ("🟢" if "NOT RELEVANT" in res["Verdict"] \
                   else ("⚠️" if "CLARIFICATION" in res["Verdict"] else "⚪"))
            ai_tag  = f" [{res['Agenda Item']}]"  if res["Agenda Item"]  != "—" else ""
            faa_tag = f" [{res['FAA Systems'][:40]}]" if res["FAA Systems"] != "—" else ""
            label   = f"{icon} {res['File']}{ai_tag}{faa_tag}"
            with st.expander(label):
                st.markdown(res["_analysis"])


        # ── Downloads: Word + Excel ────────────────────────────────────────────
        st.markdown("---")
        n_label = f"{n_filtered} document{'s' if n_filtered != 1 else ''}"
        filter_note = f" (filtered: {', '.join(active_filters)})" if active_filters else f" — all {n_total} docs"
        st.subheader(f"📥 Download Reports — {n_label}{filter_note}")

        try:
            from datetime import date as _bdate2
            import re as _brc2
            def _bc2(s, n=20): return _brc2.sub(r'[^A-Za-z0-9_-]','_',str(s or ''))[:n].strip('_')

            filtered_triage = [{k: r[k] for k in display_cols} for r in filtered]
            filtered_batch  = [{"file": r["File"], "analysis": r["_analysis"],
                                 "tc": "", "error": r["_error"]} for r in filtered]

            filter_tag = ""
            if ai_filter:  filter_tag += "_" + "_".join(_bc2(a,8) for a in ai_filter)
            if faa_filter: filter_tag += "_" + "_".join(_bc2(f,8) for f in faa_filter)
            base_name = f"FAA_Batch_{_bc2(working_party,10)}_{n_filtered}docs{filter_tag}_{_bdate2.today()}"

            # Generate and persist to session state
            try:
                st.session_state["batch_docx_bytes"]    = _make_batch_docx(filtered_batch, filtered_triage, working_party, analysis_depth)
                st.session_state["batch_docx_filename"] = f"{base_name}.docx"
                st.session_state["batch_docx_label"]    = f"📄 Word — Full Detailed Report ({n_label}) (.docx)"
            except Exception as _we:
                st.session_state.pop("batch_docx_bytes", None)
                st.error(f"❌ Word error: {_we}")

            try:
                _xlsx_rows = [
                    _extract_analysis_fields(
                        r.get("_analysis",""),
                        {"doc_number": r.get("File",""), "working_party": working_party,
                         "submitting_admin": r.get("Source / Admin",""),
                         "agenda_item": r.get("Agenda Item",""), "analysis_depth": analysis_depth}
                    )
                    for r in filtered
                ]
                st.session_state["batch_xlsx_bytes"]    = _make_summary_xlsx(_xlsx_rows)
                st.session_state["batch_xlsx_filename"] = f"{base_name}_summary.xlsx"
                st.session_state["batch_xlsx_label"]    = f"📊 Excel — Summary Table ({n_label}) (.xlsx)"
            except Exception as _xe:
                st.session_state.pop("batch_xlsx_bytes", None)
                st.error(f"❌ Excel error: {_xe}")

        except Exception as be:
            st.error(f"❌ Report generation error: {be}")
            import traceback; st.code(traceback.format_exc())

    # ── Persistent batch download buttons — survive reruns and filter changes ─
    if st.session_state.get("batch_docx_bytes") or st.session_state.get("batch_xlsx_bytes"):
        st.markdown("---")
        st.subheader("📥 Downloads — Last Generated Report")
        st.caption("These buttons stay available until you clear results or reset the app.")
        _bdl1, _bdl2 = st.columns(2)
        with _bdl1:
            if st.session_state.get("batch_docx_bytes"):
                st.download_button(
                    label=st.session_state.get("batch_docx_label","📄 Word Report (.docx)"),
                    data=st.session_state["batch_docx_bytes"],
                    file_name=st.session_state.get("batch_docx_filename","FAA_Batch.docx"),
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    type="primary",
                    use_container_width=True,
                    key="persist_batch_docx",
                )
        with _bdl2:
            if st.session_state.get("batch_xlsx_bytes"):
                st.download_button(
                    label=st.session_state.get("batch_xlsx_label","📊 Excel Summary (.xlsx)"),
                    data=st.session_state["batch_xlsx_bytes"],
                    file_name=st.session_state.get("batch_xlsx_filename","FAA_Batch_Summary.xlsx"),
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    use_container_width=True,
                    key="persist_batch_xlsx",
                )
        if st.button("🗑️ Clear batch downloads", key="clear_batch_dl", type="secondary"):
            for k in ("batch_docx_bytes","batch_docx_filename","batch_docx_label",
                      "batch_xlsx_bytes","batch_xlsx_filename","batch_xlsx_label"):
                st.session_state.pop(k, None)
            st.rerun()

    # ── Show accumulated results between runs (without re-running analysis) ──
    elif (not batch_btn) and st.session_state.get("batch_accumulated") and batch_mode_active:
        batch_results = st.session_state.batch_accumulated
        st.markdown("---")
        st.subheader(f"🗂️ Accumulated Results — {len(batch_results)} document(s)")
        ex("Results from previous batch run(s). Upload more files and click Run to continue processing. All results persist until you click Clear.")

        import re as _re_b  # needed for triage extraction below
        # Fall through to triage/filter display — reuse the same code path
        # by setting batch_results and proceeding normally

    if single_btn and contrib_input.strip():

        # ── Resolve WP profile ──────────────────────────────────────────────
        wp_profile_key = WP_PROFILE_MAP.get(working_party)
        wp_profile = WP_ANALYSIS_PROFILES.get(wp_profile_key) if wp_profile_key else None

        depth_instruction = {
            "Quick assessment (key risks + recommended US position)":
                "OUTPUT FORMAT: Use the structured format below but be CONCISE — 2–4 bullet points per section. Focus on the highest-priority risks and most actionable steps.",
            "Standard analysis (full policy brief with citations)":
                "OUTPUT FORMAT: Complete all sections of the structured format below with full technical detail and regulatory citations.",
            "Deep dive (comprehensive brief + draft response contribution outline)":
                "OUTPUT FORMAT: Complete all sections in full detail, PLUS provide Section F: Draft Response — a full outline of a US counter-contribution including proposed regulatory text, specific parameter changes, and floor intervention language.",
        }[analysis_depth]

        faa_bands_summary = "\n".join([
            f"- {name}: {b['f_low_mhz']}–{b['f_high_mhz']} MHz ({b['allocation']}), "
            f"I/N threshold {b['in_threshold_db']} dB, "
            f"Aviation safety factor {b.get('aviation_safety_factor_db',0)} dB, "
            f"Effective threshold {b.get('effective_threshold_db', b['in_threshold_db'])} dB, "
            f"Safety Service (RR 1.59): {b.get('rr_1_59', True)}"
            for name, b in FAA_BANDS.items()
        ])

        system_protection_table = """
FAA SYSTEM PROTECTION LEVELS:
| System | Protection Level |
|--------|-----------------|
| ARSR | I/N = −6 dB |
| ASR | I/N = −10 dB |
| RNSS Feeder Links | ΔT/T = 6% |
| L1 SBAS Type 1 | I < −146.5 dBW/MHz; I/N ≈ −5 dB + 6 dB safety margin |
| L2 SBAS Ground Reference Rx | I < −147.5 dBW/MHz; I/N ≈ −6 dB |
| L-band AMS(R)S | ΔT/T = 20% aggregate = 6% single-entry |
| DME | epfd ≤ −121.5 dBW/m²/MHz in any 1 MHz band |
| MLS | pfd ≤ −124.5 dBW/m² in 150 kHz band |
AVIATION SAFETY FACTOR: +6 dB for precision approach/landing applications."""

        # ── Build WP-specific analytical framework block ────────────────────
        if wp_profile:
            wp_framework = f"""
═══════════════════════════════════════════════════════════════════
WORKING PARTY SPECIFIC FRAMEWORK — {wp_profile['label']}
THIS OVERRIDES GENERIC ASSUMPTIONS — APPLY THIS FRAMEWORK PRECISELY
═══════════════════════════════════════════════════════════════════

INTERFERER TYPE: {wp_profile['interferer_type']}
VICTIM TYPE: {wp_profile['victim_type']}
PRIMARY THREAT: {wp_profile['primary_threat']}

CORRECT PROPAGATION MODELS FOR THIS WP:
{chr(10).join(f'  - {m}' for m in wp_profile['propagation_models'])}

CORRECT INTERFERENCE METRICS FOR THIS WP:
{chr(10).join(f'  - {m}' for m in wp_profile['interference_metrics'])}

KEY ITU-R RECOMMENDATIONS FOR THIS WP:
{chr(10).join(f'  - {r}' for r in wp_profile['key_recommendations'])}

PROTECTION CRITERIA: {wp_profile['protection_criteria']}

AGGREGATE INTERFERENCE METHOD: {wp_profile['aggregate_method']}

SPECIFIC CHECKS TO PERFORM ON THIS CONTRIBUTION:
{chr(10).join(f'  {i+1}. {c}' for i,c in enumerate(wp_profile['specific_checks']))}

COMMON PROPONENT TACTICS TO WATCH FOR AND CHALLENGE:
{chr(10).join(f'  - {t}' for t in wp_profile['common_proponent_tactics'])}

POLICY LEVERS AVAILABLE FOR THIS WP:
{chr(10).join(f'  - {p}' for p in wp_profile['policy_levers'])}

WRC-27 ITEMS: {', '.join(wp_profile['wrc27_items']) if wp_profile['wrc27_items'] else 'None directly — check for strategic linkage'}
═══════════════════════════════════════════════════════════════════
"""
        else:
            wp_framework = f"""
NOTE: No specific WP profile found for '{working_party}'. Apply general FAA interference 
analysis framework. Identify which WRC-27 agenda item (if any) this contribution relates to.
"""

        # ── Track-changes handling flag ──────────────────────────────────────
        track_changes_note = """
TRACK CHANGES AND DOCUMENT STATUS DETECTION:

Step 1 — TECHNICAL SIGNAL: Check whether track changes data is provided above.
  Track changes present → strong evidence this is a REVISION of a previous document.
  No track changes → document may still be a revision (authors sometimes accept changes before submitting).

Step 2 — LINGUISTIC SIGNAL: Search the contribution text carefully for phrases that indicate status.
  REVISION signals (look for these verbatim phrases or close variants):
    "this document updates / revises / replaces / supersedes"
    "in response to / following the discussion / based on comments received"
    "Rev. [N]" or "Revision [N]" anywhere in the document
    "the Working Party agreed at its previous meeting / at [city]"
    "as agreed in [meeting name or document reference]"
    References to a prior document number (e.g. "5D/123-E", "4C/456", "7B/78")
    "following [administration]'s contribution", "in liaison with"
    "the previous version of this document"
  NEW DOCUMENT signals:
    "hereby proposes", "first submission", "new proposal", "introduces"
    No reference to any prior document number
    No mention of a previous meeting discussion on this specific topic

Step 3 — CITE EVIDENCE: Quote 1–2 actual phrases from the document verbatim to support your conclusion.
  Use quotation marks. Only cite text you can find IN the provided document.
  If no clear signal exists: state "No revision language found — document status unclear."

Step 4 — DETERMINE STATUS with confidence level:
  NEW DOCUMENT (High confidence): No track changes + no revision language found
  NEW DOCUMENT (Low confidence): No track changes, no revision language, but document may reference prior work
  REVISION / UPDATE (High confidence): Track changes present AND revision language found + cited
  REVISION / UPDATE (Medium confidence): Track changes present but no explicit revision language
  REVISION / UPDATE (Medium confidence): No track changes but revision language found + cited
  UNCLEAR: Neither signal present — state this explicitly
"""



        system_prompt = f"""You are an expert RF engineer with a strong background in IMT and aviation/FAA spectrum protection, supporting the FAA and NTIA in ITU-R proceedings.

You are analyzing a contribution from: {working_party}

{wp_framework}

FAA PROTECTED FREQUENCY LIST — cross-check all contributions against these:
{faa_bands_summary}

{system_protection_table}

{track_changes_note}

WRC-27 ITEMS THREATENING FAA BANDS: AI 1.7 (WP 5D, IMT near RA 4.4–4.8 GHz), AI 1.13 (WP 4C, MSS near DME/AMS(R)S/ASR), AI 1.15 (WP 7B, lunar SRS), AI 1.17/1.19 (WP 7C, EESS passive).

═══════════════════════════════════════════════════════════════════
MANDATORY REVIEW CHECKLIST — APPLY TO EVERY CONTRIBUTION
SCOPE RULE: Only flag deviations from the methodology REQUIRED by the applicable
ITU-R Recommendations for this Working Party. Do NOT critique general engineering
design choices, system architecture, or implementation decisions that are outside
the ITU-R mandate. Every finding must cite the specific Recommendation or RR
article that requires what the contribution is missing or violating.
═══════════════════════════════════════════════════════════════════

ACCURACY RULE — NON-NEGOTIABLE:
If you cannot find or verify a specific fact, frequency, document number, study result,
or regulatory citation IN THE CONTRIBUTION TEXT PROVIDED, state "Cannot confirm from
document — requires manual verification." NEVER invent frequencies, dB values, document
references, or regulatory conclusions. It is better to flag a gap than to fabricate.

COMPATIBILITY vs SHARING DISTINCTION (use correct term in output):
- SHARING STUDY: Proposed service and FAA service are IN THE SAME BAND — they share spectrum.
- COMPATIBILITY STUDY: Proposed service is in an ADJACENT or NEARBY BAND — ITU-R requires
  the proposer to demonstrate the FAA incumbent is protected.

═══════════════════════════════════════════════════════════════════
DOCUMENT ROUTING — APPLY BEFORE ANY OTHER ANALYSIS
Three paths. Determine which path applies FIRST, then stop at that level.
═══════════════════════════════════════════════════════════════════

PATH 1 — US CONTRIBUTION (submitting administration is USA / United States / NTIA / FCC):
  → SUMMARIZE ONLY. Do not critique, challenge, or flag concerns.
  → Produce only: document title, submitting admin, WRC-27 AI reference, and a
    3–5 sentence neutral summary of what the US is proposing or supporting.
  → End with: "US contribution — summary only per review policy."
  → Do NOT proceed to Sections B–F.

PATH 2 — FOREIGN CONTRIBUTION, NO FAA IMPACT:
  → The proposed frequencies/bands are NOT in-band, adjacent, or otherwise threatening
    to any FAA protected band. Clearly state the gap in MHz/GHz.
  → Produce only: the Frequency Relevance Summary table (with the gap), a one-sentence
    explanation of why this does not affect FAA interests, and the REVIEW VERDICT:
    "LIKELY NOT RELEVANT — [reason]."
  → End with: "No FAA band affected — no further analysis required."
  → Do NOT proceed to Sections B–F.

PATH 3 — FOREIGN CONTRIBUTION WITH FAA IMPACT (or uncertain relevance):
  → Proceed with the full analysis below (Sections A–F).
  → Apply all methodology compliance checks for the applicable Working Party.
  → Every finding must cite the specific ITU-R Recommendation that requires it.

To determine the path:
  - Check the submitting administration field first. Any US-origin contribution → PATH 1.
  - If foreign: run the frequency screen. No FAA band in-band or adjacent → PATH 2.
  - If foreign AND frequencies threaten FAA bands → PATH 3.
  - If submitting administration is unclear: state "Administration not identified — treating
    as PATH 3 pending clarification."

1. RELEVANCE SCREEN — FREQUENCY-FIRST (PATH 3 only)
   Step 1: Extract ALL frequencies/bands from the document — exact low MHz, high MHz, and bandwidth.
           State what the document actually says, verbatim if possible.
   Step 2: For each proposed band, identify the closest FAA protected band and calculate:
           - If proposed band overlaps FAA band: OVERLAP = min(prop_high, faa_high) − max(prop_low, faa_low) MHz
           - If proposed band is below FAA band:  GAP = FAA_low − prop_high  MHz
           - If proposed band is above FAA band:  GAP = prop_low − FAA_high  MHz
           State the arithmetic result. Never write "adjacent" or "nearby" without the number.
   Step 3: State RELEVANCE VERDICT with the actual calculated overlap/gap.
           Identify WRC-27 AI and study type (SHARING if overlap > 0, COMPATIBILITY if gap = 0 or small).

2. AVIATION/FAA IMPACT ASSESSMENT
   Only raise concerns that are grounded in a specific ITU-R requirement or RR provision.
   For each concern, state the applicable rule:
   - OOB emissions → cite SM.1540/SM.1541 (250% BN boundary, 23 dB mask rule, RR 1.153)
   - Spurious emissions → cite RR Appendix 3 (43+10·log(P), 60/70 dBc limits)
   - Aggregate interference → cite SM.2028 (required methodology for aggregate studies)
   - I/N exceedance → cite the specific protection level (ARSR −6 dB, ASR −10 dB, etc.)
   - Aviation safety factor omission → cite M.1477 Annex 5 (+6 dB for precision approach)
   - Blocking/desensitization omission → cite M.1642 §X or the applicable WP methodology doc
   Do NOT raise concerns about parameters or scenarios that the applicable Recommendation
   does not specifically require to be addressed.

3. METHODOLOGY COMPLIANCE — per applicable ITU-R Recommendations only
   Check ONLY whether the study follows the methodology mandated for this WP:

   WP 5D (IMT/ARNS): Required methodology is ITU-R M.1642. Check:
   - Propagation: M.1642 requires P.452 (terrestrial) or P.528 (airborne victim)
   - Monte Carlo per SM.2028 if aggregate interference is claimed
   - Protection criteria per M.1477 (I/N thresholds + 6 dB safety factor)
   - Time percentage: 1% required for worst-case per SM.2028 — NOT 50%

   WP 4C (MSS/satellite): Required methodology is P.619 + S.1586 (epfd). Check:
   - Propagation MUST be P.619 — P.452 is terrestrial-only and WRONG for satellite
   - Metric MUST be epfd (constellation aggregate) — single-satellite pfd is insufficient
   - ΔT/T for AMS(R)S — single-entry ≤6% per system protection table
   - epfd for DME — ≤−121.5 dBW/m²/MHz per system protection table

   WP 7B (Space Research/Lunar): No established ITU-R methodology exists for Earth-Moon
   to terrestrial ARNS. If the contribution proposes one, assess whether it is grounded
   in existing ITU-R Recommendations. If no methodology is cited, flag the gap.

   WP 7C (EESS passive): No interference methodology applies — passive sensors do not
   transmit. Assess only allocation policy implications per the Radio Regulations.

   WP 5B (Maritime/Radiodetermination): Required methodology is M.1849 for radar,
   P.452/P.528 as applicable. Check protection criteria for co-channel ARNS systems.

   Do NOT flag methodology choices that the applicable Recommendation does not prohibit
   or that are within the discretion of the submitting administration.

4. REGULATORY AND PROCEDURAL ISSUES
   Only cite RR provisions that are directly applicable to the band and service:
   - RR No. 4.10: harmful interference to safety-of-life — cite only if threshold is exceeded
   - RR 5.444: ARNS protection at 960–1215 MHz — cite only for that band
   - RR Appendix 3: spurious limits — cite only if spurious products are identified
   - SM.1540/SM.1541: OOB domain — cite only if OOB boundary analysis is missing or wrong
   Do NOT cite a regulation unless it is specifically applicable to the scenario in question.

INTERFERENCE CLASSIFICATION (apply only if the study addresses the FAA band):
- Harmful (RR 1.169) → triggers RR 4.10 → cite only if I/N threshold is exceeded
- Permissible (RR 1.167) → within agreed criteria
- Accepted (RR 1.168) → bilateral agreement only

REGULATORY TOOLKIT (cite only what is applicable to this WP and scenario):
- RR No. 4.10, 1.59, 5.444, Appendix 3
- SM.1540/SM.1541 (OOB domain); SM.2028 (Monte Carlo); SM.329 (spurious measurement)
- M.1642 (WP 5D/5B terrestrial IMT→ARNS); M.1477/M.1318/M.1904/M.1905 (GNSS)
- P.619 + S.1586 (WP 4C satellite); P.452/P.528 (terrestrial/aeronautical propagation)
- RTCA DO-155 (RA), DO-235B (GNSS), DO-260B (ADS-B), DO-189 (DME) — for victim parameters

TONE: Technically rigorous, grounded only in applicable ITU-R requirements.
Flag what the guidelines require. Do not expand the scope beyond the WP mandate.

{depth_instruction}"""

        # ── WP-specific analysis questions ──────────────────────────────────
        if wp_profile_key == "WP 7C (EESS / Space Weather Sensors)":
            # Passive sensors — allocation policy analysis, not interference analysis
            analysis_questions = """
ANALYSIS STRUCTURE FOR THIS WP 7C CONTRIBUTION:

⚠️ IMPORTANT: EESS passive sensors DO NOT TRANSMIT. Do NOT perform interference analysis 
as if this is a transmitting service. The FAA concern is ALLOCATION POLICY, not interference.

1. ALLOCATION POLICY ANALYSIS
   a) What allocation status is proposed — secondary or co-primary?
   b) Which specific band(s) are affected? Does any overlap 4.2–4.4 GHz (RA band)?
   c) Does a co-primary EESS allocation in 4.2–4.4 GHz weaken FAA's exclusive ARNS status for AI 1.7?
   d) Does the allocation create any coordination or notification obligations on FAA transmitters?
   e) Is this a stepping stone for future active allocation in the band?

2. STRATEGIC ASSESSMENT — WRC-27 AI LINKAGE
   Assess how this contribution affects FAA's position on AI 1.7, AI 1.13, or other active WRC-27 AIs.

3. SUBMITTER'S OBJECTIVE — What allocation outcome is the proponent seeking?

4. RECOMMENDED US POSITION — What FAA/NTIA should argue (allocation table language)

5. REGULATORY CITATIONS — Relevant RR provisions (allocation table, footnotes)

6. DRAFT RESPONSE LANGUAGE — Proposed US intervention text"""

        elif wp_profile_key == "WP 7B (Space Radiocommunication / Lunar SRS)":
            # Novel use case — methodology gap argument
            analysis_questions = """
ANALYSIS STRUCTURE FOR THIS WP 7B CONTRIBUTION:

⚠️ IMPORTANT: Lunar SRS is a NOVEL USE CASE. There is NO established ITU-R methodology
for assessing Earth-Moon SRS interference to terrestrial aeronautical systems. 
The absence of methodology is itself a FAA policy argument — allocation before methodology = wrong order.

1. METHODOLOGY ASSESSMENT
   a) Does the contribution propose a coordination methodology for lunar SRS vs terrestrial ARNS? 
      If NOT: this is the primary objection — demand methodology before allocation.
   b) If a methodology is proposed: is it valid? Does it use appropriate Earth-Moon propagation geometry?
   c) Is EIRP from lunar surface transmitters quantified? What is the pfd at Earth surface?

2. FAA BAND IMPACT ASSESSMENT
   Check interference/coordination zones for: 2700–2900 MHz (ASR), 3600–4200 MHz, 
   5350–5470 MHz (ARNS 5 GHz), 7190–7235 MHz, 8450–8500 MHz.

3. SUBMITTER'S OBJECTIVE — What SRS allocation is being sought?

4. RECOMMENDED US POSITION — Methodology first, allocation second. Demand coordination mechanism.

5. REGULATORY CITATIONS

6. DRAFT RESPONSE LANGUAGE — Propose that WP 7B establish methodology BEFORE finalizing allocation"""

        elif wp_profile_key == "WP 4C (MSS / DC-MSS-IMT)":
            analysis_questions = """
ANALYSIS STRUCTURE FOR THIS WP 4C CONTRIBUTION:
SCOPE: Assess compliance with ITU-R P.619, S.1586, SM.2028, and the WP 4C methodology
for satellite-to-Earth interference. Do not critique satellite system design choices
beyond what these Recommendations require.

1. METHODOLOGY COMPLIANCE (per P.619 / S.1586 / SM.2028)
   a) Propagation: Is P.619 used? P.452 is terrestrial-only — wrong for satellite downlinks. Cite P.619.
   b) Metric: Is aggregate epfd calculated per S.1586? Single-satellite pfd alone is non-compliant with S.1586.
   c) For 925–960 MHz: does epfd comply with the −121.5 dBW/m²/MHz DME limit? Cite RR 5.444.
   d) For 1475–1518 MHz: does ΔT/T comply with 6% single-entry per the AMS(R)S protection level?
   e) For 2620–2690 MHz: does I/N comply with −10 dB per the ASR protection level?
   f) Is aggregate from all visible satellites computed? SM.2028 requires this for aggregate studies.

2. INTERFERENCE CLASSIFICATION — per RR 1.166–1.169 using correct WP 4C metrics (epfd/ΔT/T)

3. THREAT ASSESSMENT — which candidate band(s) pose risk, with the actual numbers

4. SUBMITTER'S OBJECTIVE

5. RECOMMENDED US POSITION — band-by-band

6. COUNTER-ARGUMENTS — grounded in P.619, S.1586, SM.2028, RR 5.444, RR 4.10

7. REGULATORY CITATIONS

8. COALITION STRATEGY

9. REQUIRED ANALYSIS — what FAA needs to compute

10. DRAFT RESPONSE LANGUAGE"""

        else:
            # Generic terrestrial WP (5D, 5B, 5A, 4A, etc.)
            analysis_questions = """
ANALYSIS STRUCTURE:
SCOPE: Assess compliance with the methodology required by the applicable ITU-R Recommendations
for this Working Party (M.1642 for WP 5D/5B, P.452/P.528 for propagation, SM.2028 for
aggregate studies). Every finding must cite the specific Recommendation that requires it.

1. METHODOLOGY COMPLIANCE
   a) Propagation: P.452 (terrestrial) or P.528 (airborne victim) as required by M.1642?
   b) OOB emissions: does the study address the 250% BN boundary per SM.1540/SM.1541?
      Only flag if OOB products land in a FAA protected band.
   c) Spurious: does it comply with RR Appendix 3? Only flag if spurious lands in FAA band.
   d) Aviation safety factor: is the +6 dB applied where M.1477 requires it?
   e) Aggregate: does SM.2028 apply? Flag only if aggregate study is claimed but methodology absent.
   f) I/N thresholds: are the correct protection levels used per the FAA system protection table?

2. INTERFERENCE CLASSIFICATION — per RR 1.166–1.169 with actual numbers from study

3. THREAT ASSESSMENT — FAA bands at risk with specific dB values from the document

4. SUBMITTER'S OBJECTIVE

5. RECOMMENDED US POSITION — Oppose / Support / Propose amendments

6. COUNTER-ARGUMENTS — grounded in M.1642, P.452/P.528, SM.2028, RR 4.10 as applicable

7. REGULATORY CITATIONS

8. COALITION STRATEGY

9. REQUIRED ANALYSIS

10. DRAFT RESPONSE LANGUAGE"""

        system_prompt = f"""You are a senior RF spectrum policy advisor supporting the FAA and NTIA in ITU-R proceedings.

You are analyzing a contribution from: {working_party}

{wp_framework}

FAA PROTECTED BANDS (defend all of these):
{faa_bands_summary}

{system_protection_table}

WRC-27 ITEMS THREATENING FAA BANDS: AI 1.7 (WP 5D, IMT near RA), AI 1.13 (WP 4C, MSS near DME/AMS(R)S/ASR), AI 1.15 (WP 7B, lunar SRS), AI 1.17/1.19 (WP 7C, EESS passive).

REGULATORY TOOLKIT:
- RR No. 4.10: No harmful interference to safety-of-life services
- RR No. 1.59: Safety service definition
- RR No. 5.444: ARNS protection 960–1215 MHz
- RR Appendix 3: Spurious limits (43+10·log(P) or 60/70 dBc)
- ITU-R SM.1540/SM.1541: OOB domain (250% BN boundary; 23 dB mask rule)
- ITU-R M.1318/M.1477/M.1904/M.1905: GNSS protection methodology + 6 dB safety margin
- ITU-R M.1642: IMT→ARNS methodology (terrestrial only — do not apply to satellite contributions)
- ITU-R P.619: Earth-space propagation (required for WP 4C satellite analysis)
- ITU-R SM.2028: Monte Carlo aggregate interference
- ITU-R S.1586: epfd methodology (required for WP 4C/4A satellite downlink analysis)

INTERFERENCE CLASSIFICATION (apply to every analysis):
- Harmful (RR 1.169) → triggers RR 4.10 → oppose
- Permissible (RR 1.167) → within criteria → monitor
- Accepted (RR 1.168) → bilateral agreement → ensure it doesn't bind US
Emission types: Spurious vs OOB (see SM.1540/SM.1541 for OOB domain)
Mechanisms: In-band, OOB coupling, blocking, intermodulation, spurious response

{depth_instruction}

Use clear headers. Plain language. Flag NTIA/ICAO escalation needs."""

        # Pull track change data from session state if available
        tc_summary_for_prompt = st.session_state.get("tc_summary", "")

        user_message = f"""Analyze this ITU-R contribution using the mandatory review checklist and produce the structured output below.

DOCUMENT METADATA — EXTRACT FROM THE DOCUMENT TEXT:
All metadata fields (document number, submitting administration, meeting/date, agenda item, document type) must be identified from the contribution text itself and stated clearly in Section A. Do not leave them blank.
- Working Party: {working_party}
- All other fields: extract from the document text and state them in Section A of the output.

{f"TRACK CHANGES (extracted from uploaded Word document — use for Section B):{chr(10)}{tc_summary_for_prompt}{chr(10)}" if tc_summary_for_prompt else "TRACK CHANGES: No Word document uploaded — if pasting text, note whether this is a revision."}

CONTRIBUTION TEXT (full document):
{contrib_input}

{f"SPECIFIC FAA CONCERN TO PRIORITIZE: {user_concern}" if user_concern else ""}
{f"PRIOR US POSITION: {prior_us_position}" if prior_us_position else ""}

{analysis_questions}

═══════════════════════════════════════════════════════════════════
REQUIRED OUTPUT STRUCTURE
Apply the correct path determined above.
═══════════════════════════════════════════════════════════════════

━━━ PATH 1 — US CONTRIBUTION (use this format only) ━━━

## Document Overview
- Title (if found in document)
- Submitting Administration: USA / United States / NTIA / FCC
- WRC-27 Agenda Item: [number if stated]
- Summary: [3–5 neutral sentences describing what the US is proposing or supporting]

*US contribution — summary only per review policy.*

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

━━━ PATH 2 — FOREIGN, NO FAA IMPACT (use this format only) ━━━

## ⚡ FREQUENCY RELEVANCE SUMMARY
Show the proposed band AND the nearest FAA band with exact numbers — always both columns populated.

| Proposed Band (from doc) | FAA Band (from protected list) | FAA System | Gap (MHz) | Verdict |
|---|---|---|---|---|
| [low MHz]–[high MHz] ([BW] MHz) | [low MHz]–[high MHz] | [system name] | [exact gap in MHz — calculated, not estimated] | NOT RELEVANT |

Rule for gap calculation: Gap = FAA_low − Proposed_high  (or Proposed_low − FAA_high if above).
If proposed band is above the FAA band: Gap = Proposed_low_MHz − FAA_high_MHz.
If proposed band is below the FAA band: Gap = FAA_low_MHz − Proposed_high_MHz.
State the number explicitly, e.g. "Gap: 340 MHz" — never write "significant gap" without the number.

**REVIEW VERDICT: LIKELY NOT RELEVANT — [one sentence with the actual gap number]**
Example: "Proposed 27.5–28 GHz is 23,100 MHz above the nearest FAA band (Radio Altimeter 4.2–4.4 GHz). No in-band or adjacent overlap."

- [Confirm no overlap and state the gap]
- [Note indirect risks only if a specific harmonic or OOB path exists — state the math, e.g. "2nd harmonic of 960 MHz = 1920 MHz, not in FAA band"]

*No FAA band affected — no further analysis required.*

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

━━━ PATH 3 — FOREIGN WITH FAA IMPACT (full analysis) ━━━

## ⚡ FREQUENCY RELEVANCE SUMMARY  ← ALWAYS FIRST, ALWAYS EXPLICIT NUMBERS

One row per proposed band / FAA band pair. Both columns MUST contain exact MHz or GHz numbers — never abbreviations or names alone.

| Proposed Band | Proposed BW | FAA Band | FAA System | Overlap or Gap | Relationship | Study Type |
|---|---|---|---|---|---|---|
| [low]–[high] MHz/GHz | [BW] MHz | [low]–[high] MHz/GHz | [system name] | OVERLAP: [X] MHz  –or–  GAP: [X] MHz | IN-BAND / ADJACENT / NEARBY | SHARING / COMPATIBILITY |

Overlap/Gap calculation rules (show the arithmetic):
- IN-BAND overlap: state "OVERLAP: [min(prop_high, faa_high) − max(prop_low, faa_low)] MHz"
  Example: Proposed 1085–1095 MHz overlaps ADS-B 1085–1095 MHz → OVERLAP: 10 MHz (full band)
- ADJACENT gap: state "GAP: [FAA_low − Prop_high] MHz" or "GAP: [Prop_low − FAA_high] MHz"
  Example: Proposed 4400–4800 MHz, FAA 4200–4400 MHz → GAP: 0 MHz (immediately adjacent, touching)
  Example: Proposed 925–960 MHz, FAA 960–1215 MHz → GAP: 0 MHz (touching at 960 MHz)
- NEARBY: state "GAP: [X] MHz — OOB/harmonic risk possible"
  Example: Proposed 3700–3980 MHz, FAA 4200–4400 MHz → GAP: 220 MHz (OOB risk)

If the document does not state exact frequencies: write "Exact band not stated in document — [what was stated]. Flag for clarification."

**REVIEW VERDICT: [REQUIRES HUMAN REVIEW / FLAG FOR CLARIFICATION]**
One sentence including the actual overlap or gap number.
Example: "REQUIRES HUMAN REVIEW — proposed 4.4–4.8 GHz is 0 MHz from Radio Altimeter 4.2–4.4 GHz (immediately adjacent, touching at 4.4 GHz)."

---

## A) Document Overview + Status
- Title, source/administration, date (only if in document — do not invent)
- WRC-27 Agenda Item (if stated)
- 📋 STATUS: [NEW DOCUMENT / REVISION / UNCLEAR] — [confidence]
- Evidence: cite 1–2 verbatim phrases supporting the status determination

## B) Track Changes Summary
- If track changes detected: bullet list of consequential changes and FAA impact
- If new document: "New document — full text analyzed"

## C) Relevance Screen
Two parts:
1. **Proposed bands:** List every frequency band proposed in the document with exact MHz/GHz numbers and bandwidth.
2. **FAA band impacts:** For each proposed band, state the FAA band it overlaps or is adjacent to, the exact overlap in MHz (or gap if no overlap), and whether it requires a sharing study (co-band) or compatibility study (adjacent). Use numbers throughout — no prose substitutions for numerical values.

## D) FAA Impact Findings
For each concern — both the proposed band AND the FAA band must appear as exact numbers:

**Issue [N]:** [one-line description]
- Proposed band: [low]–[high] MHz/GHz ([BW] MHz)
- FAA band: [low]–[high] MHz/GHz — [FAA system name]
- Overlap / Gap: [OVERLAP: X MHz]  or  [GAP: X MHz]
- Mechanism: [in-band / OOB (state the OOB boundary) / blocking / spurious / intermod]
- Required by: [specific Recommendation/RR article]
- Severity: [Low / Medium / High] | Confidence: [Low / Medium / High]
- Mitigation: [specific fix with numbers where possible]
- ⚠️ UNVERIFIED: [anything that cannot be confirmed from the document]

## E) Methodology Compliance
**Compliant:** [what the study does correctly per the applicable Recommendation]
**Non-compliant:** [deviations — cite the specific Recommendation for each]
**Missing:** [what the applicable Recommendation requires but is absent — cite it]

## F) Recommended Actions
Numbered list — actionable, specific, with owner and priority:
1. **Action:** | **Target:** | **Priority:** Immediate / Before next meeting / Long-term | **Owner:** FAA / NTIA / US delegation

## G) Draft U.S. Response
Write a concise, ready-to-use US floor intervention or contribution response for this document. This should be suitable for delivery at the WP session or as a formal written comment. Include:

**Intervention Header:** US Position on [Document Title / Doc Number]

**Opening statement** (1–2 sentences): State the US concern and the applicable WRC-27 agenda item.

**Technical objections** (bullet points): Cite specific ITU-R Recommendations and Radio Regulations that are implicated. Reference actual frequencies and FAA systems at risk with the calculated overlap/gap.

**Proposed language or amendments** (if applicable): Suggest specific text changes, conditions, or requirements the US would need to see before supporting the proposal. Frame in ITU-R regulatory language.

**Closing**: State the US position (support/oppose/condition) and any coalition language.

Keep the draft to 200–300 words. Use ITU-R meeting language. Cite specific Recommendations and RR articles. This draft should be actionable — something a delegation representative could actually use.

If this is PATH 1 (US contribution) or PATH 2 (not relevant) — omit this section entirely."""

        with st.spinner("Analyzing contribution... this takes 15–30 seconds for deep analysis."):
            try:
                client = anthropic.Anthropic(api_key=api_key)
                response = client.messages.create(
                    model="claude-opus-4-5",
                    max_tokens=6000,
                    messages=[{"role": "user", "content": user_message}],
                    system=system_prompt,
                )
                analysis_text = response.content[0].text

                # Store for interactive Q&A
                st.session_state["qa_contrib_text"]    = contrib_input
                st.session_state["qa_analysis_text"]   = analysis_text
                st.session_state["qa_doc_meta"]        = docx_meta
                st.session_state["qa_history"]         = []   # reset on new analysis
                st.session_state["qa_active_doc"]      = doc_number or "this document"

                # ── Write to Neo4j (if configured) ────────────────────────────
                neo4j_driver = _neo4j_driver()
                if neo4j_driver:
                    try:
                        tc_meta = {**docx_meta, "tc_summary": st.session_state.get("tc_summary","")}
                        nc, rc = _neo4j_write_analysis(neo4j_driver, analysis_text, tc_meta, contrib_input)
                        neo4j_driver.close()
                        st.caption(f"🔗 Neo4j: +{nc} nodes, +{rc} relationships written to knowledge graph")
                    except Exception as neo_err:
                        st.caption(f"⚠️ Neo4j write skipped: {neo_err}")

                st.success("✅ Analysis complete")
                st.markdown("---")

                # Interference characterization quick-badge at top
                st.subheader("⚡ Interference Classification (ITU-RR Taxonomy)")
                ex("The AI has characterized the interference using the formal ITU Radio Regulations definitions from RR Articles 1.166–1.169. This framing is what Working Party rapporteurs and experienced delegations use to assess the regulatory weight of a contribution.")

                ic1, ic2, ic3 = st.columns(3)
                with ic1:
                    st.markdown("""
<div style='background:#1a2a3a;border-left:4px solid #4488ff;padding:10px;border-radius:4px'>
<b style='color:#4488ff'>📡 Emission Type</b><br>
<span style='color:#aaddff;font-size:0.85em'>
Spurious (RR 1.145) — hardware non-linearity<br>
Out-of-Band / OOB — from modulation process<br>
In-band — co-frequency operation
</span></div>""", unsafe_allow_html=True)
                with ic2:
                    st.markdown("""
<div style='background:#1a3a1a;border-left:4px solid #44bb44;padding:10px;border-radius:4px'>
<b style='color:#44bb44'>⚖️ RR Classification</b><br>
<span style='color:#aaffaa;font-size:0.85em'>
Harmful (RR 1.169) → triggers RR 4.10<br>
Permissible (RR 1.167) → within criteria<br>
Accepted (RR 1.168) → bilateral agreement
</span></div>""", unsafe_allow_html=True)
                with ic3:
                    st.markdown("""
<div style='background:#3a1a1a;border-left:4px solid #ff4444;padding:10px;border-radius:4px'>
<b style='color:#ff4444'>🔧 Mechanism</b><br>
<span style='color:#ffaaaa;font-size:0.85em'>
In-band / Adjacent band coupling<br>
Receiver blocking / desensitization<br>
Intermodulation / spurious response
</span></div>""", unsafe_allow_html=True)
                st.caption("See full definitions in the 📖 Glossary module under 'ITU-R & Regulatory' category.")
                st.markdown("---")
                st.subheader("Policy Guidance")
                st.markdown(analysis_text)

                # ── Word Document Download ────────────────────────────────────
                st.markdown("---")
                st.subheader("📥 Download Analysis Report")

                def _clean(s, maxlen=20):
                    import re as _rc
                    return _rc.sub(r'[^A-Za-z0-9_-]', '_', str(s or ""))[:maxlen].strip('_')

                # Extract metadata automatically from the analysis text
                _auto = _extract_analysis_fields(analysis_text, {"working_party": working_party})
                auto_doc_num  = _auto.get("Document No.","") or doc_number
                auto_admin    = _auto.get("Source / Admin","") or submitting_admin
                auto_ai       = _auto.get("Agenda Item(s)","") or agenda_item

                from datetime import date as _dl
                safe_name = "_".join(filter(None, [
                    _clean(auto_doc_num,  20),
                    _clean(auto_admin,    15),
                    _clean(working_party, 10),
                    str(_dl.today()),
                ]))
                if not safe_name.strip("_"):
                    safe_name = f"analysis_{_dl.today()}"

                docx_meta = {
                    "doc_number":      auto_doc_num,
                    "working_party":   working_party,
                    "submitting_admin":auto_admin,
                    "meeting_date":    meeting_date,
                    "agenda_item":     auto_ai,
                    "doc_type":        doc_type,
                    "analysis_depth":  analysis_depth,
                }

                # Generate and persist bytes in session state so buttons survive page reruns
                try:
                    st.session_state["single_docx_bytes"]    = _make_analysis_docx(analysis_text, docx_meta)
                    st.session_state["single_docx_filename"] = f"FAA_Analysis_{safe_name}.docx"
                except Exception as _de:
                    st.session_state.pop("single_docx_bytes", None)
                    st.warning(f"⚠️ Word generation failed: {_de}")

                try:
                    _xlsx_row = _extract_analysis_fields(analysis_text, docx_meta)
                    st.session_state["single_xlsx_bytes"]    = _make_summary_xlsx([_xlsx_row])
                    st.session_state["single_xlsx_filename"] = f"FAA_Summary_{safe_name}.xlsx"
                except Exception as _xe:
                    st.session_state.pop("single_xlsx_bytes", None)
                    st.warning(f"⚠️ Excel generation failed: {_xe}")

            except anthropic.AuthenticationError:
                st.error("❌ Invalid API key. Check your Streamlit secrets configuration.")
            except anthropic.RateLimitError:
                st.error("❌ API rate limit reached. Wait a moment and try again.")
            except Exception as e:
                st.error(f"❌ Error: {str(e)}")

    # ── Persistent single-analysis download buttons ───────────────────────────
    # Always rendered when bytes exist — survives reruns, filter changes, Q&A interactions
    if st.session_state.get("single_docx_bytes") or st.session_state.get("single_xlsx_bytes"):
        st.markdown("---")
        st.subheader("📥 Download Last Analysis")
        _pdl1, _pdl2 = st.columns(2)
        with _pdl1:
            if st.session_state.get("single_docx_bytes"):
                st.download_button(
                    label="📄 Word — Full Analysis Report (.docx)",
                    data=st.session_state["single_docx_bytes"],
                    file_name=st.session_state.get("single_docx_filename", "FAA_Analysis.docx"),
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    type="primary",
                    use_container_width=True,
                    key="persist_docx_dl",
                )
        with _pdl2:
            if st.session_state.get("single_xlsx_bytes"):
                st.download_button(
                    label="📊 Excel — Summary Table (.xlsx)",
                    data=st.session_state["single_xlsx_bytes"],
                    file_name=st.session_state.get("single_xlsx_filename", "FAA_Summary.xlsx"),
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    use_container_width=True,
                    key="persist_xlsx_dl",
                )
        if st.button("🗑️ Clear downloads", key="clear_single_dl", type="secondary"):
            for k in ("single_docx_bytes","single_docx_filename",
                      "single_xlsx_bytes","single_xlsx_filename"):
                st.session_state.pop(k, None)
            st.rerun()

    elif not contrib_input.strip():
        st.info("👆 Paste a contribution above and click Analyze to get policy guidance.")

    # ── Interactive Q&A — persists after analysis, survives page interactions ─
    if st.session_state.get("qa_analysis_text") and st.session_state.get("qa_contrib_text"):
        st.markdown("---")
        st.subheader("💬 Ask About This Analysis")
        ex(f"Question any finding in the analysis for **{st.session_state.get('qa_active_doc','this document')}**. "
           "The AI will answer from the document and analysis, citing specific passages.")

        # Initialize chat history
        if "qa_history" not in st.session_state:
            st.session_state.qa_history = []

        # Show existing conversation
        for msg in st.session_state.qa_history:
            with st.chat_message(msg["role"],
                                 avatar="🧑‍💼" if msg["role"] == "user" else "🤖"):
                st.markdown(msg["content"])

        # Suggested questions
        if not st.session_state.qa_history:
            st.markdown("**Suggested questions:**")
            sq_cols = st.columns(3)
            suggestions = [
                "Where in the document does it mention this frequency?",
                "Why is this flagged as requiring human review?",
                "What evidence supports the severity rating?",
                "Where is the propagation model stated?",
                "What exact text was cited for the revision status?",
                "Which part of the document mentions the FAA band?",
            ]
            for i, sq in enumerate(suggestions):
                with sq_cols[i % 3]:
                    if st.button(sq, key=f"sq_{i}", use_container_width=True):
                        st.session_state.qa_history.append({"role": "user", "content": sq})
                        st.rerun()

        # Chat input
        user_q = st.chat_input(
            f"Ask about the analysis of {st.session_state.get('qa_active_doc','this document')}…",
            key="qa_input"
        )

        if user_q:
            st.session_state.qa_history.append({"role": "user", "content": user_q})

            qa_system = f"""You are reviewing a completed FAA spectrum interference analysis of an ITU-R contribution.
You have access to:
1. THE ORIGINAL CONTRIBUTION TEXT — the source document that was analyzed
2. THE COMPLETED ANALYSIS — the findings, verdicts, and recommendations already produced

Your job is to answer the user's question about a specific finding or conclusion.

RULES:
- Always cite the exact passage from the ORIGINAL DOCUMENT that supports or contradicts the finding.
  Use "quote marks" around direct quotes from the document.
- If the finding is derived from absence (i.e. something was NOT in the document), say so explicitly.
- If you cannot find the relevant passage in the provided text, say "Cannot locate in the provided document text — the original document may have more context."
- Be concise but precise. One to three short paragraphs maximum.
- Do not re-summarize the entire analysis — answer only the specific question asked.

DOCUMENT METADATA:
{st.session_state.get('qa_doc_meta', {})}

ORIGINAL CONTRIBUTION TEXT:
{st.session_state['qa_contrib_text'][:15000]}

COMPLETED ANALYSIS:
{st.session_state['qa_analysis_text'][:8000]}"""

            # Build messages from history
            qa_messages = [
                {"role": m["role"], "content": m["content"]}
                for m in st.session_state.qa_history
            ]

            try:
                with st.spinner("Checking the document…"):
                    qa_client = anthropic.Anthropic(api_key=api_key)
                    qa_resp = qa_client.messages.create(
                        model="claude-sonnet-4-5",
                        max_tokens=1000,
                        system=qa_system,
                        messages=qa_messages,
                    )
                qa_answer = qa_resp.content[0].text
                st.session_state.qa_history.append({"role": "assistant", "content": qa_answer})
                st.rerun()
            except Exception as qa_err:
                st.error(f"❌ {qa_err}")

        # Clear conversation button
        if st.session_state.qa_history:
            if st.button("🗑️ Clear conversation", key="qa_clear", type="secondary"):
                st.session_state.qa_history = []
                st.rerun()

    # ── Neo4j cross-document query panel ──────────────────────────────────────
    neo4j_drv = _neo4j_driver()
    if neo4j_drv:
        neo4j_drv.close()
        st.markdown("---")
        st.subheader("🔗 Knowledge Graph — Cross-Document Query")
        ex("Ask questions across ALL analyzed documents in the graph — not just this session. The AI converts your question to Cypher, queries Neo4j, and summarizes the results.")

        if "neo4j_history" not in st.session_state:
            st.session_state.neo4j_history = []

        # Show conversation
        for msg in st.session_state.neo4j_history:
            with st.chat_message(msg["role"], avatar="🗄️" if msg["role"] == "assistant" else "🧑‍💼"):
                st.markdown(msg["content"])
                if msg.get("cypher"):
                    with st.expander("🔍 Cypher query used"):
                        st.code(msg["cypher"], language="cypher")

        # Suggested cross-document questions
        if not st.session_state.neo4j_history:
            st.markdown("**Example cross-document queries:**")
            neo_cols = st.columns(2)
            neo_suggestions = [
                "Which documents from China affect the DME band?",
                "How many HIGH severity findings are there across all documents?",
                "Which administrations have submitted revisions (not new docs)?",
                "List all documents where US stance is Oppose",
                "Which documents use the P.452 propagation model?",
                "Show all documents that affect the Radio Altimeter band",
                "Which documents have tracked changes?",
                "List documents related to AI 1.13",
            ]
            for i, sq in enumerate(neo_suggestions):
                with neo_cols[i % 2]:
                    if st.button(sq, key=f"neo_sq_{i}", use_container_width=True):
                        st.session_state.neo4j_history.append({"role": "user", "content": sq})
                        st.rerun()

        neo_q = st.chat_input("Ask anything about the full document corpus…", key="neo4j_input")

        if neo_q:
            st.session_state.neo4j_history.append({"role": "user", "content": neo_q})
            try:
                with st.spinner("Querying knowledge graph…"):
                    drv2 = _neo4j_driver()
                    answer, cypher_used, raw_records = _neo4j_nl_query(drv2, neo_q, api_key)
                    drv2.close()
                st.session_state.neo4j_history.append({
                    "role": "assistant", "content": answer, "cypher": cypher_used
                })
                # Show raw records toggle
                if raw_records:
                    st.session_state[f"neo4j_raw_{len(st.session_state.neo4j_history)}"] = raw_records
                st.rerun()
            except Exception as neo_err:
                st.error(f"❌ Knowledge graph error: {neo_err}")

        if st.session_state.neo4j_history:
            if st.button("🗑️ Clear graph queries", key="neo4j_clear", type="secondary"):
                st.session_state.neo4j_history = []
                st.rerun()

    # ── Batch document Q&A ─────────────────────────────────────────────────────
    if st.session_state.get("batch_accumulated"):
        acc = st.session_state.batch_accumulated
        valid = [r for r in acc if not r.get("error") and r.get("analysis")]
        if valid:
            st.markdown("---")
            st.subheader("💬 Ask About a Batch Document")
            ex("Select any processed document and ask follow-up questions about its analysis findings.")

            doc_names  = [r["file"] for r in valid]
            selected_doc_name = st.selectbox(
                "Select document to interrogate:",
                doc_names,
                key="qa_batch_select"
            )
            selected_doc = next(r for r in valid if r["file"] == selected_doc_name)

            if "qa_batch_history" not in st.session_state:
                st.session_state.qa_batch_history = {}
            if selected_doc_name not in st.session_state.qa_batch_history:
                st.session_state.qa_batch_history[selected_doc_name] = []

            # Show history for this document
            for msg in st.session_state.qa_batch_history[selected_doc_name]:
                with st.chat_message(msg["role"],
                                     avatar="🧑‍💼" if msg["role"] == "user" else "🤖"):
                    st.markdown(msg["content"])

            batch_q = st.chat_input(
                f"Ask about {selected_doc_name}…",
                key="qa_batch_input"
            )

            if batch_q:
                st.session_state.qa_batch_history[selected_doc_name].append(
                    {"role": "user", "content": batch_q}
                )

                bqa_system = f"""You are reviewing a completed FAA spectrum interference analysis of the ITU-R contribution: {selected_doc_name}

You have:
1. THE COMPLETED ANALYSIS of this document
2. THE ORIGINAL DOCUMENT TEXT (if available from extraction)

RULES:
- Cite exact passages from the document using "quote marks" when available.
- If a finding is based on absence of information, say so explicitly.
- If you cannot locate the relevant text, say "Cannot locate in provided text."
- Concise — one to three paragraphs. Answer only what is asked.

COMPLETED ANALYSIS:
{selected_doc.get('analysis','')[:8000]}

ORIGINAL DOCUMENT TEXT (extracted):
{selected_doc.get('text','')[:10000] if selected_doc.get('text') else '(Not available — document text was not stored in this session.)'}"""

                bqa_messages = [
                    {"role": m["role"], "content": m["content"]}
                    for m in st.session_state.qa_batch_history[selected_doc_name]
                ]

                try:
                    with st.spinner(f"Checking {selected_doc_name}…"):
                        bqa_client = anthropic.Anthropic(api_key=api_key)
                        bqa_resp = bqa_client.messages.create(
                            model="claude-sonnet-4-5",
                            max_tokens=1000,
                            system=bqa_system,
                            messages=bqa_messages,
                        )
                    bqa_answer = bqa_resp.content[0].text
                    st.session_state.qa_batch_history[selected_doc_name].append(
                        {"role": "assistant", "content": bqa_answer}
                    )
                    st.rerun()
                except Exception as bqa_err:
                    st.error(f"❌ {bqa_err}")

            if st.session_state.qa_batch_history.get(selected_doc_name):
                bcol1, bcol2 = st.columns([1, 4])
                with bcol1:
                    if st.button("🗑️ Clear", key="qa_batch_clear", type="secondary"):
                        st.session_state.qa_batch_history[selected_doc_name] = []
                        st.rerun()

    # Example contributions to try
    with st.expander("📋 Example contributions to test with"):
        st.markdown("""
**Example 1 — IMT identification near radio altimeter band (WP 5D)**
> *"This contribution proposes to identify the frequency range 4800–4990 MHz for IMT under WRC-27 Agenda Item 1.2. 
Sharing studies indicate that IMT base stations operating with EIRP not exceeding 58 dBm and using downtilted antennas 
can coexist with existing services. The submitting administrations request that the Working Party develop a draft 
CPM Report method supporting IMT identification in this band."*

**Example 2 — VDES proposal near aeronautical band (WP 5B)**
> *"This document proposes expansion of the VHF Data Exchange System (VDES) to include additional channels 
in the 156–174 MHz band to support maritime e-navigation. The proposal includes shore-based transmitters 
with power levels up to 50W ERP. Compatibility with aeronautical VHF communications has not been fully assessed."*

**Example 3 — LEO satellite downlinks (WP 4A / SG 5)**
> *"The proposed non-GSO satellite constellation will operate downlinks in the 1525–1559 MHz band with 
aggregate EPFD levels approaching the limits specified in RR Appendix 7. Individual satellite EIRP is 
limited to 15 dBW per beam with a minimum elevation angle of 10 degrees."*
        """)

# ─────────────────────────────────────────────────────────────────────────────
# TAB 9 — RF TRAINING
# ─────────────────────────────────────────────────────────────────────────────
elif selected_tab == "🎓 RF Training":
    st.title("🎓 RF Fundamentals Training")
    st.markdown("*Ground-up refresher designed for engineers returning to RF after time away — focused on what you need for ITU-R policy support.*")
    ex("Work through these lessons in order. Each builds on the last. Every concept has a worked example and a self-check.")

    lesson = st.selectbox("Select Lesson:", [
        "Lesson 1 — The Decibel: Your Most Important Tool",
        "Lesson 2 — Frequency, Wavelength & The EM Spectrum",
        "Lesson 3 — Transmit Power, EIRP & Antenna Gain",
        "Lesson 4 — Free Space Path Loss",
        "Lesson 5 — The Link Budget: Putting It All Together",
        "Lesson 6 — Noise, Sensitivity & the Noise Floor",
        "Lesson 7 — Interference: I/N, C/I & Protection Criteria",
        "Lesson 8 — Propagation Models: P.452, P.528, P.676",
        "Lesson 9 — Monte Carlo & Aggregate Interference",
        "Lesson 10 — From RF Math to ITU-R Policy",
        "── WP-Specific Case Studies ──",
        "Lesson 11 — WP 5D: IMT vs Radio Altimeter (AI 1.7)",
        "Lesson 12 — WP 5B: Maritime Radar vs ARSR/ASR",
        "Lesson 13 — WP 4C: MSS Satellite Downlinks vs DME/AMS(R)S (AI 1.13)",
        "Lesson 14 — WP 7B & 7C: Lunar SRS and EESS Passive (AI 1.15, 1.17, 1.19)",
    ])

    st.markdown("---")

    # ── LESSON 1 ──────────────────────────────────────────────────────────────
    if lesson.startswith("Lesson 1"):
        st.header("📐 Lesson 1 — The Decibel: Your Most Important Tool")
        ex("The dB scale is log base 10 of a power ratio multiplied by 10. Because the Friis equation is multiplicative in linear scale, it becomes additive in dB — this is why every RF link analysis is a simple arithmetic chain of dB values.")

        st.subheader("Why Decibels?")
        st.markdown("""
RF signals span an enormous range — a transmitter might produce 100 watts while a GPS receiver needs just 
0.000000000000001 watts (10⁻¹⁵ W). Writing those numbers linearly is unworkable. Decibels compress 
this range into manageable numbers by using logarithms.

**The core formula:**
""")
        st.latex(r"\text{Value (dB)} = 10 \cdot \log_{10}\left(\frac{P_2}{P_1}\right)")
        st.markdown("""
This compares two power levels. When you say a signal is "30 dB stronger," you mean it is **1000× more powerful**.

**The three numbers you must memorize:**

| dB | Power Ratio | Memory Hook |
|---|---|---|
| **+3 dB** | 2× (double) | "3 dB = double the power" |
| **+10 dB** | 10× | "10 dB = ten times the power" |
| **−3 dB** | 0.5× (half) | "−3 dB = half the power" |
| **−10 dB** | 0.1× | "−10 dB = one tenth the power" |
| **+30 dB** | 1000× | Combine: 10 × 10 × 10 |
| **−60 dB** | 0.000001× | Combine: −10 − 10 − 10 − 10 − 10 − 10 |

**In RF we use two reference points:**
- **dBm** = decibels relative to 1 milliwatt. Used for signal power.
- **dBi** = decibels relative to an isotropic antenna. Used for antenna gain.
- **dBW** = decibels relative to 1 watt. Used in satellite/high-power work (0 dBW = 30 dBm).
        """)

        st.subheader("Converting Between Watts and dBm")
        st.latex(r"P_{\text{dBm}} = 10 \cdot \log_{10}(P_{\text{mW}}) = 10 \cdot \log_{10}(P_{\text{W}} \times 1000)")
        st.latex(r"P_{\text{W}} = \frac{10^{P_{\text{dBm}}/10}}{1000} = 10^{(P_{\text{dBm}} - 30)/10}")
        st.markdown("""
> **Quick rule:** dBW = dBm − 30 &nbsp;&nbsp;|&nbsp;&nbsp; dBm = dBW + 30 &nbsp;&nbsp;|&nbsp;&nbsp; 0 dBW = 30 dBm = 1 W
        """)

        st.markdown("**Common conversions to memorize:**")
        conv_data = {
            "Power (Watts)": ["0.001 W (1 mW)", "0.01 W (10 mW)", "0.1 W (100 mW)", "1 W", "10 W", "20 W", "100 W"],
            "Power (dBm)": ["0 dBm", "10 dBm", "20 dBm", "30 dBm", "40 dBm", "43 dBm", "50 dBm"],
            "Typical Use": ["Reference level", "Low-power IoT device", "Wi-Fi access point", "Small radio", "Mobile phone max", "LTE base station typical", "FM broadcast transmitter"],
        }
        st.table(pd.DataFrame(conv_data))

        st.subheader("🔢 Interactive dB Calculator")
        ex("Note that dBm is an absolute power level (referenced to 1 mW), while dB is a dimensionless ratio. You can add dB values to dBm, but you cannot add two dBm values — that would require linear addition of milliwatts, then convert back.")
        col1, col2 = st.columns(2)
        with col1:
            p_watts = st.number_input("Power (Watts)", value=1.0, min_value=0.000001, format="%.6f")
            p_dbm_calc = 10 * np.log10(p_watts * 1000)
            st.metric("In dBm", f"{p_dbm_calc:.2f} dBm")
            st.metric("In dBW", f"{p_dbm_calc - 30:.2f} dBW")
        with col2:
            p_dbm_in = st.number_input("Power (dBm)", value=30.0, step=1.0)
            p_watts_calc = 10 ** (p_dbm_in / 10) / 1000
            st.metric("In Watts", f"{p_watts_calc:.6f} W")
            st.metric("In milliwatts", f"{p_watts_calc*1000:.4f} mW")

        st.subheader("The Golden Rule of dB Arithmetic")
        st.markdown("""
In dB, **multiplication becomes addition and division becomes subtraction.**
This is why link budgets work — instead of multiplying and dividing many large/small numbers,
you just add and subtract dB values.

**Example:**
> A transmitter puts out 43 dBm. It goes through a 3 dB cable loss, then a 15 dBi antenna.
> What is the EIRP?
>
> EIRP = 43 − 3 + 15 = **55 dBm** ✓
>
> (In linear: 20W × 0.5 × 31.6 = 316W = 55 dBm. Same answer, harder math.)
        """)

        st.subheader("✅ Self-Check")
        with st.expander("Try these — click to reveal answers"):
            st.markdown("""
**Q1:** A signal drops by 20 dB. By what factor did the power decrease?
> **Answer:** 100× decrease (−10 dB = ÷10, twice = ÷100)

**Q2:** A GPS satellite transmits at 27 dBW. What is that in dBm? In Watts?
> **Answer:** 27 dBW = 57 dBm = 500 Watts

**Q3:** You have 43 dBm Tx power, 2 dB cable loss, 17 dBi antenna gain. What is the EIRP?
> **Answer:** 43 − 2 + 17 = **58 dBm** (631 Watts equivalent)

**Q4:** A receiver needs −90 dBm to work. The received signal is −105 dBm. By how many dB is it below threshold?
> **Answer:** −105 − (−90) = **−15 dB** below threshold (about 30× too weak)
            """)

    # ── LESSON 2 ──────────────────────────────────────────────────────────────
    elif lesson.startswith("Lesson 2"):
        st.header("📻 Lesson 2 — Frequency, Wavelength & The EM Spectrum")
        ex("The free-space path loss equation has a 20·log10(f) term — doubling frequency adds 6 dB of path loss at fixed distance. This frequency dependence is why UHF/L-band systems like GPS require extremely sensitive receivers to overcome orbital path losses exceeding 180 dB.")

        st.subheader("The Fundamental Relationship")
        st.latex(r"c = f \cdot \lambda \quad \Rightarrow \quad \lambda = \frac{c}{f}")
        st.markdown("""
Where:
- **c** = speed of light = 3 × 10⁸ m/s
- **f** = frequency in Hz
- **λ** = wavelength in meters

**Quick rule:** λ(meters) ≈ 300 / f(MHz)
        """)

        st.subheader("🔢 Frequency ↔ Wavelength Calculator")
        col1, col2 = st.columns(2)
        with col1:
            f_mhz_l2 = st.number_input("Frequency (MHz)", value=1090.0, step=10.0)
            lam = 300 / f_mhz_l2
            st.metric("Wavelength", f"{lam*100:.2f} cm")
            st.metric("Half-wavelength (antenna)", f"{lam/2*100:.2f} cm")
        with col2:
            lam_cm = st.number_input("Wavelength (cm)", value=27.5, step=1.0)
            f_from_lam = 300 / (lam_cm / 100)
            st.metric("Frequency", f"{f_from_lam:.1f} MHz")

        st.subheader("The Aeronautical Spectrum — Bands You Work With")
        ex("The key parameter is the ratio of wavelength to obstacle size: wavelength < obstacle → shadowing/reflection dominate; wavelength ≈ obstacle → diffraction and scattering matter. At C-band (λ ≈ 7 cm), rain drops scatter efficiently — hence airborne weather radar operates here.")
        band_data = pd.DataFrame([
            ["VHF", "30–300 MHz", "~1–10 m", "Line-of-sight + some diffraction", "VOR, ILS, VHF comms"],
            ["UHF", "300 MHz–3 GHz", "~10 cm–1 m", "Line-of-sight dominant", "DME, TACAN, ADS-B, GPS L1"],
            ["L-band", "1–2 GHz", "~15–30 cm", "Good penetration, long range", "GPS, GLONASS, DME, Mode-S"],
            ["C-band", "4–8 GHz", "~4–7 cm", "Line-of-sight, rain starts to matter", "Radio altimeter, en-route radar"],
            ["X-band", "8–12 GHz", "~2.5–4 cm", "Rain attenuation significant", "Weather radar, surface radar"],
        ], columns=["Band", "Frequency Range", "Wavelength", "Propagation Character", "FAA Systems"])
        st.table(band_data)

        st.subheader("Why Frequency Matters for Interference")
        st.markdown("""
**1. Path Loss scales with frequency** — higher frequency = more free-space path loss per km
(we'll calculate this in Lesson 4). This can work in your favor — a 5 GHz interferer loses
more power over distance than a 500 MHz one.

**2. Antenna size scales with wavelength** — a half-wave dipole at 1090 MHz (ADS-B) is
~13.7 cm. At 121.5 MHz (VHF emergency) it's ~1.23 m. This affects what fits on an aircraft.

**3. Atmospheric absorption** — above ~10 GHz, oxygen and water vapor absorb RF energy.
At 60 GHz, absorption is so high signals can barely travel a few km (used in 5G mmWave
for this reason — natural isolation).

**4. Rain attenuation** — above ~3 GHz, raindrops are comparable to the wavelength
and scatter energy. Airborne weather radar at 9 GHz exploits this to detect storms.
        """)

        st.subheader("✅ Self-Check")
        with st.expander("Click to reveal answers"):
            st.markdown("""
**Q1:** ADS-B operates at 1090 MHz. What is its wavelength?
> **Answer:** λ = 300/1090 = **0.275 m = 27.5 cm**

**Q2:** GPS L1 is at 1575.42 MHz. GPS L5 is at 1176.45 MHz. Which has a longer wavelength?
> **Answer:** L5 (lower frequency = longer wavelength). L5: 25.5 cm vs L1: 19.0 cm

**Q3:** Why does the radio altimeter band (4200–4400 MHz) require more careful protection from nearby 5G (3700–3980 MHz) than, say, VHF comms at 121 MHz?
> **Answer:** The 5G band is only ~220 MHz away from the radio altimeter band. At these frequencies, receiver front-ends can be overloaded by strong adjacent signals even if the 5G signal is not technically "in-band." At VHF, the fractional frequency separation would be enormous.
            """)

    # ── LESSON 3 ──────────────────────────────────────────────────────────────
    elif lesson.startswith("Lesson 3"):
        st.header("📡 Lesson 3 — Transmit Power, EIRP & Antenna Gain")
        ex("EIRP is the product of transmit power and antenna gain in the direction of the victim, less cable losses. It is invariant to whether gain is achieved by increasing power or focusing the antenna — both create identical interference at a distant receiver.")

        st.subheader("Transmit Power")
        st.markdown("""
Transmit power (Pt) is the RF power delivered to the antenna input. It does NOT account for the antenna's directionality.

Common values in ITU-R work:
| System | Typical Tx Power |
|---|---|
| LTE macro base station | 43–46 dBm (20–40W) |
| 5G NR base station | 46–53 dBm (40–200W) |
| Mobile handset | 23–30 dBm (200mW–1W) |
| VOR ground station | 50 dBm (100W) |
| GPS satellite | 27 dBW = 57 dBm (500W) |
| ADS-B transponder | 18–21 dBW = 48–51 dBm |
        """)

        st.subheader("Antenna Gain")
        st.markdown("""
An antenna does not amplify power — it **redirects** it. Gain (G) describes how much more power 
is concentrated in the main beam versus an isotropic (perfectly omnidirectional) antenna.

- **0 dBi** — isotropic; radiates equally in all directions (theoretical)
- **2 dBi** — simple dipole; slight directivity
- **15 dBi** — typical sector antenna on a cell tower
- **30+ dBi** — high-gain dish or phased array

**Key insight for interference analysis:** When assessing interference, use the antenna gain 
**in the direction of the victim receiver** — not the peak gain. For worst-case analysis, 
assume 0 dBi (isotropic) toward the victim unless you have a specific antenna pattern.
        """)

        st.subheader("EIRP — Effective Isotropic Radiated Power")
        st.latex(r"\text{EIRP (dBm)} = P_t \text{(dBm)} + G_t \text{(dBi)} - L_{cable} \text{(dB)}")
        st.markdown("""
EIRP represents the power that would need to be fed into an isotropic antenna to produce 
the same signal strength in the direction of maximum radiation.

**It is the standard metric used in:**
- ITU-R Radio Regulations (PFD and EIRP limits)
- Coordination agreements between administrations
- Interference analysis inputs
        """)

        st.subheader("🔢 Interactive EIRP Calculator")
        col1, col2, col3 = st.columns(3)
        with col1:
            pt = st.number_input("Tx Power (dBm)", value=43.0, step=1.0)
        with col2:
            gt = st.number_input("Antenna Gain (dBi)", value=15.0, step=1.0)
        with col3:
            lc = st.number_input("Cable Loss (dB)", value=2.0, step=0.5)
        eirp_val = pt + gt - lc
        st.metric("EIRP", f"{eirp_val:.1f} dBm = {10**(eirp_val/10)/1000:.1f} W equivalent")
        ex("In ITU-R Radio Regulations, EIRP limits are specified per carrier, per antenna beam, or as aggregate — always check which definition applies. A proponent quoting per-carrier EIRP for a MIMO system may be understating the effective aggregate EIRP toward your victim.")

        st.subheader("Power Flux Density (PFD)")
        st.markdown("""
PFD describes how much power arrives per unit area at a given distance. 
ITU-R uses PFD limits in the Radio Regulations to protect Earth-based receivers from satellites.
PFD depends only on EIRP and distance — it is **frequency-independent**.
        """)
        st.latex(r"\text{PFD (dBW/m}^2) = \text{EIRP (dBW)} - 10\log_{10}(4\pi d^2)")
        st.markdown("> where *d* is in metres. To use km: substitute $d_m = d_{km} \\times 1000$")

        d_km_l3 = st.slider("Distance (km)", 1, 500, 10)
        eirp_dbw = eirp_val - 30
        pfd_val = eirp_dbw - 10 * np.log10(4 * np.pi * (d_km_l3 * 1000) ** 2)
        st.metric("PFD at distance", f"{pfd_val:.1f} dBW/m²")
        ex("PFD (W/m²) is power per unit area at the victim, independent of the victim's antenna aperture. RR Appendix 7 PFD limits are specified by elevation angle — the lower the elevation, the tighter the limit, because low-elevation satellite paths have longer atmospheric path lengths.")

        st.subheader("✅ Self-Check")
        with st.expander("Click to reveal answers"):
            st.markdown("""
**Q1:** A 5G base station has Tx power 46 dBm, antenna gain 18 dBi, cable loss 1 dB. What is the EIRP?
> **Answer:** 46 + 18 − 1 = **63 dBm** (2000W equivalent — this is why 5G can interfere with things far away)

**Q2:** For worst-case interference analysis, should you use the base station's peak antenna gain or 0 dBi toward the victim?
> **Answer:** Use the gain **toward the victim** — which may be a side lobe. For a worst-case bound, use 0 dBi. In a realistic analysis, model the actual antenna pattern.

**Q3:** Why is EIRP used instead of just transmit power in regulatory limits?
> **Answer:** EIRP captures both the transmitter power AND the antenna focusing effect. Two transmitters with the same Tx power but different antennas can have very different interference impacts — EIRP accounts for this in a single number.
            """)

    # ── LESSON 4 ──────────────────────────────────────────────────────────────
    elif lesson.startswith("Lesson 4"):
        st.header("📉 Lesson 4 — Free Space Path Loss")
        ex("FSPL arises from the inverse square law: power density decreases as 1/4πr². The 20·log10(d) + 20·log10(f) form captures both the geometric spreading and the frequency dependence of effective antenna aperture (Ae = λ²G/4π).")

        st.subheader("The Formula")
        st.latex(r"FSPL \text{ (dB)} = 20\log_{10}(d_{km}) + 20\log_{10}(f_{MHz}) + 32.44")
        st.markdown("""
This is the **Friis free-space path loss** equation. It tells you how much signal is lost 
between a transmitter and receiver separated by distance d, at frequency f.

**Key observations:**
- **Distance doubles → path loss increases by 6 dB** (power drops to ¼)
- **Frequency doubles → path loss increases by 6 dB**
- Higher frequency signals lose more power over the same distance
- This is WHY 5G mmWave (26 GHz) has short range and GPS (1.5 GHz) needs very sensitive receivers
        """)

        st.subheader("🔢 Interactive FSPL Calculator")
        col1, col2 = st.columns(2)
        with col1:
            f_l4 = st.number_input("Frequency (MHz)", value=1090.0, step=10.0, key="l4f")
            d_l4 = st.number_input("Distance (km)", value=10.0, step=1.0, key="l4d")
        fspl_val = 20*np.log10(d_l4) + 20*np.log10(f_l4) + 32.44
        with col2:
            st.metric("Free Space Path Loss", f"{fspl_val:.1f} dB")
            st.metric("Signal reduced by factor of", f"{10**(fspl_val/10):.2e}")
        ex("FSPL is the dominant loss term for most aeronautical links — it typically ranges from 100 dB (VHF, short range) to 190 dB (GPS orbital distance). All other loss mechanisms (atmospheric, terrain, clutter) are additive corrections to this baseline.")

        st.subheader("FSPL vs Distance — Interactive Chart")
        freqs_to_plot = st.multiselect("Select frequencies to compare:",
            [121.5, 328.6, 1090, 1575, 4300, 9375],
            default=[1090, 1575, 4300],
            format_func=lambda x: f"{x} MHz")

        if freqs_to_plot:
            d_range = np.linspace(0.1, 100, 300)
            fig_l4, ax_l4 = plt.subplots(figsize=(10, 5))
            fig_l4.patch.set_facecolor("#0e1117")
            ax_l4.set_facecolor("#0e1117")
            colors = plt.cm.plasma(np.linspace(0.1, 0.9, len(freqs_to_plot)))
            for freq, col in zip(freqs_to_plot, colors):
                fspl_curve = 20*np.log10(d_range) + 20*np.log10(freq) + 32.44
                ax_l4.plot(d_range, fspl_curve, color=col, linewidth=2, label=f"{freq} MHz")
            ax_l4.set_xlabel("Distance (km)", color='white')
            ax_l4.set_ylabel("Path Loss (dB)", color='white')
            ax_l4.set_title("Free Space Path Loss vs Distance", color='white')
            ax_l4.legend(facecolor='#1a1a2e', labelcolor='white')
            ax_l4.tick_params(colors='white')
            ax_l4.grid(color='#333', alpha=0.5)
            for sp in ax_l4.spines.values(): sp.set_color('#444')
            plt.tight_layout()
            st.pyplot(fig_l4)

        st.subheader("Why FSPL is Just the Starting Point")
        st.markdown("""
FSPL assumes a perfect vacuum with no obstructions. Real paths add loss from:

| Additional Loss Mechanism | Typical Extra Loss | When It Applies |
|---|---|---|
| Atmospheric gases (O₂, H₂O) | 0.01–10 dB/km | Above ~1 GHz, long paths |
| Rain attenuation | 0.01–10 dB/km | Above ~3 GHz |
| Building/terrain diffraction | 10–30 dB | Ground paths, urban |
| Vegetation/foliage | 5–20 dB | Low-altitude paths |
| Multipath fading | ±20 dB | Reflective environments |

**For interference analysis:** FSPL gives you the **optimistic** interference estimate 
(most interference at the victim). If even FSPL shows compatibility, you're safe.
If FSPL fails, use P.452 or P.528 for a more realistic (higher loss) calculation.
        """)

        st.subheader("✅ Self-Check")
        with st.expander("Click to reveal answers"):
            st.markdown("""
**Q1:** Calculate FSPL at 4300 MHz, 5 km.
> **Answer:** 20·log(5) + 20·log(4300) + 32.44 = 13.98 + 72.67 + 32.44 = **119.1 dB**

**Q2:** If you double the distance, how much does FSPL increase?
> **Answer:** 20·log(2) = **6 dB** increase

**Q3:** GPS signal arrives at ~−130 dBm at a receiver. GPS satellites transmit at ~57 dBm EIRP. 
What is the approximate path loss over the ~20,200 km orbital distance at 1575 MHz?
> **Answer:** FSPL = 20·log(20200) + 20·log(1575) + 32.44 ≈ 86.1 + 63.9 + 32.4 = **182.4 dB**. 
This is why GPS receivers are so sensitive — and why any interference is so damaging.
            """)

    # ── LESSON 5 ──────────────────────────────────────────────────────────────
    elif lesson.startswith("Lesson 5"):
        st.header("🔗 Lesson 5 — The Link Budget: Putting It All Together")
        ex("The Friis transmission equation: Pr(dBm) = Pt + Gt − Lcable − FSPL + Gr. For interference analysis, Pr is the interference power at the victim input — compare it to the noise floor (N = kTBF) using I/N = Pr − N to assess compatibility.")

        st.subheader("The Friis Transmission Equation")
        st.latex(r"P_r = P_t + G_t - L_{cable} - FSPL + G_r \quad \text{(all in dB)}")
        st.markdown("""
| Term | Symbol | Typical Range | Description |
|---|---|---|---|
| Received power | Pr | −60 to −130 dBm | What arrives at receiver input |
| Transmit power | Pt | 20 to 53 dBm | PA output power |
| Tx antenna gain | Gt | 0 to 30 dBi | Toward receiver |
| Cable/feeder loss | Lcable | 0.5 to 5 dB | Physical losses in hardware |
| Path loss | FSPL | 60 to 200 dB | Distance + frequency dependent |
| Rx antenna gain | Gr | 0 to 30 dBi | Toward transmitter |
        """)

        st.subheader("🔢 Build a Link Budget Step by Step")
        ex("Sensitivity analysis: vary one parameter at a time to identify which term most strongly drives the I/N result. If path loss dominates, argue for a more realistic propagation model. If EIRP dominates, argue for a PFD or EIRP limit in the regulatory text.")

        col1, col2 = st.columns(2)
        with col1:
            lb_pt = st.number_input("① Tx Power (dBm)", value=43.0, step=1.0, key="lb1")
            lb_gt = st.number_input("② Tx Antenna Gain (dBi)", value=15.0, step=1.0, key="lb2")
            lb_lc = st.number_input("③ Cable Loss (dB)", value=2.0, step=0.5, key="lb3")
            lb_f  = st.number_input("④ Frequency (MHz)", value=4300.0, step=100.0, key="lb4")
            lb_d  = st.number_input("⑤ Distance (km)", value=10.0, step=1.0, key="lb5")
            lb_gr = st.number_input("⑥ Rx Antenna Gain (dBi)", value=0.0, step=1.0, key="lb6")

        lb_eirp = lb_pt + lb_gt - lb_lc
        lb_fspl = 20*np.log10(lb_d) + 20*np.log10(lb_f) + 32.44
        lb_pr = lb_eirp - lb_fspl + lb_gr

        with col2:
            st.markdown("**Running Total:**")
            st.markdown(f"① Tx Power: **{lb_pt} dBm**")
            st.markdown(f"② + Tx Gain: **+{lb_gt} dBi** → {lb_pt+lb_gt:.1f} dBm")
            st.markdown(f"③ − Cable Loss: **−{lb_lc} dB** → {lb_eirp:.1f} dBm (EIRP)")
            st.markdown(f"④⑤ − Path Loss ({lb_f:.0f}MHz, {lb_d}km): **−{lb_fspl:.1f} dB** → {lb_eirp-lb_fspl:.1f} dBm")
            st.markdown(f"⑥ + Rx Gain: **+{lb_gr} dBi** → **{lb_pr:.1f} dBm** ← Received Power")
            st.metric("Received Power", f"{lb_pr:.1f} dBm")

        st.subheader("The Interference Link Budget")
        st.markdown("""
When you're assessing interference (not a wanted signal), the link budget is the same —
but instead of comparing received power to a minimum detectable signal,
you compare it to the **noise floor** using the I/N criterion:
        """)
        st.latex(r"\frac{I}{N} \text{ (dB)} = P_{interference} \text{ (dBm)} - \text{Noise Floor (dBm)}")
        st.markdown("""
**Protection is maintained when:** I/N < threshold (typically −6 dB or −10 dB)

**Protection margin** = threshold − I/N  
- Positive margin → protected  
- Negative margin → interference threshold exceeded → you have grounds to object
        """)

        lb_nf_bw = st.number_input("Rx Bandwidth (MHz)", value=100.0, step=10.0)
        lb_nf_nf = st.number_input("Rx Noise Figure (dB)", value=5.0, step=0.5)
        lb_thresh = st.number_input("I/N Threshold (dB)", value=-6.0, step=1.0)
        lb_noise = -174 + 10*np.log10(lb_nf_bw*1e6) + lb_nf_nf
        lb_in = lb_pr - lb_noise
        lb_margin = lb_thresh - lb_in

        col_a, col_b, col_c = st.columns(3)
        col_a.metric("Noise Floor", f"{lb_noise:.1f} dBm")
        col_b.metric("I/N", f"{lb_in:.1f} dB")
        col_c.metric("Protection Margin", f"{lb_margin:.1f} dB",
                     delta_color="normal" if lb_margin >= 0 else "inverse")

        if lb_margin < 0:
            warn(f"Threshold violated by {abs(lb_margin):.1f} dB — at {lb_d} km this interferer exceeds the protection criterion.")
        else:
            ok(f"Protected with {lb_margin:.1f} dB margin at {lb_d} km separation.")

        st.subheader("✅ Self-Check")
        with st.expander("Click to reveal answers"):
            st.markdown("""
**Q1:** A 5G base station (EIRP 58 dBm) is 5 km from a radio altimeter receiver (4300 MHz, noise floor −89 dBm, I/N threshold −6 dB). Is there an interference problem?
> FSPL = 20·log(5) + 20·log(4300) + 32.44 = 14.0 + 72.7 + 32.4 = **119.1 dB**
> Received power = 58 − 119.1 + 0 = **−61.1 dBm**
> I/N = −61.1 − (−89) = **+27.9 dB**
> Threshold = −6 dB. Margin = −6 − 27.9 = **−33.9 dB** ← Severely violated!
> This is exactly the 5G/rad-alt interference problem that grounded aircraft in 2022.

**Q2:** What separation distance would be needed for I/N ≤ −6 dB?
> Need: received power ≤ noise floor + threshold = −89 + (−6) = −95 dBm
> Need FSPL ≥ 58 − (−95) = 153 dB
> 153 = 20·log(d) + 72.7 + 32.4 → 20·log(d) = 47.9 → d = 10^(47.9/20) = **248 km!**
            """)

    # ── LESSON 6 ──────────────────────────────────────────────────────────────
    elif lesson.startswith("Lesson 6"):
        st.header("📊 Lesson 6 — Noise, Sensitivity & the Noise Floor")
        ex("The receiver noise floor N = kTB·NF is a physical lower bound set by thermodynamics. Interference raises the effective noise floor to N·(1 + I/N_linear), degrading SNR for the desired signal by the same amount — this is why I/N is the correct metric.")

        st.subheader("Thermal Noise — The Unavoidable Baseline")
        st.markdown("""
All electronic components generate noise due to random thermal motion of electrons.
This sets the absolute minimum noise any real receiver will experience.
        """)
        st.latex(r"N_{thermal} = k \cdot T \cdot B")
        st.latex(r"N_{thermal} \text{ (dBm)} = -174 + 10\log_{10}(B_{Hz}) \quad \text{at 290K}")
        st.markdown("""
Where:
- **k** = Boltzmann's constant = 1.38 × 10⁻²³ J/K
- **T** = temperature in Kelvin (290K = room temperature standard)
- **B** = bandwidth in Hz
- **−174 dBm/Hz** is the thermal noise spectral density at 290K — memorize this number
        """)

        st.subheader("🔢 Noise Floor Calculator")
        col1, col2 = st.columns(2)
        with col1:
            bw_l6 = st.number_input("Receiver Bandwidth (MHz)", value=100.0, min_value=0.001, step=10.0)
            nf_l6 = st.number_input("Noise Figure (dB)", value=5.0, step=0.5,
                help="Real receivers amplify the noise — NF captures how much")
            temp_l6 = st.number_input("Temperature (K)", value=290.0, step=10.0)
        kT_l6 = 10*np.log10(1.38e-23 * temp_l6) + 30
        nf_total = kT_l6 + 10*np.log10(bw_l6*1e6) + nf_l6
        with col2:
            st.metric("kT noise density", f"{kT_l6:.1f} dBm/Hz")
            st.metric("Noise floor (kTB)", f"{kT_l6 + 10*np.log10(bw_l6*1e6):.1f} dBm")
            st.metric("Noise floor (kTBNF)", f"{nf_total:.1f} dBm")
        ex("Noise floor = −174 + 10·log10(B) + NF. A GPS L1 receiver (2 MHz BW, 2 dB NF) has a noise floor of −174 + 63 + 2 = −109 dBm. A radio altimeter (200 MHz BW, 5 dB NF) has −174 + 83 + 5 = −86 dBm — the wider bandwidth raises the floor by 23 dB.")

        st.subheader("Noise Figure")
        st.markdown("""
No real receiver is perfect — its internal components (LNA, mixer, filters) add noise.
The **noise figure (NF)** measures how much worse the real receiver is compared to an ideal one.

| Receiver Type | Typical NF | Why |
|---|---|---|
| GPS/GNSS LNA | 1–2 dB | Extremely sensitive; expensive low-noise design |
| Radio Altimeter | 4–6 dB | Aviation grade; well-designed |
| ADS-B receiver | 4–8 dB | Varies by implementation |
| General avionic | 5–10 dB | Depends on age and design |
| Consumer Wi-Fi | 8–12 dB | Cost-optimized |

A 1 dB increase in noise figure directly raises the noise floor by 1 dB —
and reduces sensitivity by 1 dB. This is why avionics engineers fight hard for low NF.
        """)

        st.subheader("Sensitivity vs Selectivity")
        st.markdown("""
**Sensitivity** — minimum signal the receiver can detect above the noise floor.
(Usually defined as SNR = some threshold, e.g., 10 dB above noise floor)

**Selectivity** — how well the receiver rejects signals on adjacent or nearby frequencies.
Determined by the filter characteristics (bandwidth, roll-off, out-of-band rejection).

**Why this matters for interference:**
- A strong out-of-band signal can **overload** (block/desensitize) the receiver front-end 
  even if the filter should reject it — the amplifier saturates before the filter acts.
- This was the exact mechanism in the 5G/radio altimeter problem:
  5G at 3.8 GHz was outside the 4.2–4.4 GHz altimeter band, but the 5G signal was
  strong enough to drive the altimeter's LNA into compression, raising its noise floor
  and killing sensitivity across the whole band.
- **In your contributions:** distinguish between **in-band interference** (signal inside 
  the protected band) and **receiver blocking/desensitization** (out-of-band signal 
  overloading the front end). Both are harmful; blocking is often harder to defend against.
        """)

        st.subheader("✅ Self-Check")
        with st.expander("Click to reveal answers"):
            st.markdown("""
**Q1:** Calculate the noise floor for a radio altimeter with 200 MHz bandwidth and 5 dB noise figure.
> −174 + 10·log(200×10⁶) + 5 = −174 + 83.0 + 5 = **−86 dBm**

**Q2:** If you narrow the bandwidth from 200 MHz to 20 MHz, what happens to the noise floor?
> 10·log(20M/200M) = −10 dB. Noise floor drops to **−96 dBm**. 
> Narrower bandwidth = lower noise floor = better sensitivity. But you also reject more of the desired signal.

**Q3:** What is the difference between a receiver being "jammed" and being "desensitized"?
> **Jamming** = an in-band interferer that directly competes with the desired signal.
> **Desensitization/blocking** = a strong out-of-band signal overloads the front-end amplifier, 
> raising the effective noise floor and reducing sensitivity to the desired signal. 
> The interferer doesn't need to be in-band to cause this — it just needs to be strong enough.
            """)

    # ── LESSON 7 ──────────────────────────────────────────────────────────────
    elif lesson.startswith("Lesson 7"):
        st.header("⚡ Lesson 7 — Interference: I/N, C/I & Protection Criteria")
        ex("I/N = I(dBm) − N(dBm). The noise floor rise ΔN = 10·log10(1 + 10^(I/N/10)). For I/N = −6 dB: ΔN = 0.97 dB. For I/N = 0 dB: ΔN = 3 dB. The ITU-R thresholds (−6 dB, −10 dB) correspond to ΔN < 1 dB for safety-critical systems.")

        st.subheader("The Three Key Ratios")
        col1, col2, col3 = st.columns(3)
        with col1:
            st.markdown("**I/N — Interference to Noise**")
            st.latex(r"\frac{I}{N} \text{(dB)} = I_{\text{dBm}} - N_{\text{floor dBm}}")
            st.markdown("Used in: ITU-R protection criteria, most aeronautical compatibility studies")
        with col2:
            st.markdown("**C/I — Carrier to Interference**")
            st.latex(r"\frac{C}{I} \text{(dB)} = C_{\text{dBm}} - I_{\text{dBm}}")
            st.markdown("Used in: Communications system design, co-channel interference assessment")
        with col3:
            st.markdown("**C/N — Carrier to Noise**")
            st.latex(r"\frac{C}{N} \text{(dB)} = C_{\text{dBm}} - N_{\text{floor dBm}}")
            st.markdown("Used in: Link quality assessment, minimum SNR for operation")

        st.subheader("Why I/N and Not Just 'Is the Interference Below the Noise'?")
        st.markdown("""
The noise floor is not a hard wall — it's a statistical level. Adding an interferer 
**raises the effective noise floor** by:

| I/N | Noise floor increase | Effect |
|---|---|---|
| −10 dB | +0.41 dB | Barely noticeable — standard GNSS criterion |
| −6 dB | +0.97 dB | ~1 dB degradation — standard ARNS criterion |
| 0 dB | +3.0 dB | Significant degradation |
| +10 dB | +10.4 dB | Receiver largely unusable |

Setting I/N ≤ −6 dB means the interference raises the noise floor by less than 1 dB —
a small, tolerable degradation for safety-of-life systems.

Setting I/N ≤ −10 dB for GNSS is more conservative because GPS signals arrive at 
only ~−130 dBm — any noise floor increase materially degrades position accuracy.
        """)

        st.subheader("🔢 Interactive I/N Explorer")
        col1, col2 = st.columns(2)
        with col1:
            i_dbm = st.number_input("Interference Power (dBm)", value=-100.0, step=1.0)
            n_dbm = st.number_input("Noise Floor (dBm)", value=-89.0, step=1.0)
            thresh_l7 = st.number_input("I/N Threshold (dB)", value=-6.0, step=1.0)
        i_n_l7 = i_dbm - n_dbm
        margin_l7 = thresh_l7 - i_n_l7
        noise_rise = 10*np.log10(1 + 10**(i_n_l7/10))
        with col2:
            st.metric("I/N", f"{i_n_l7:.1f} dB")
            st.metric("Protection Margin", f"{margin_l7:.1f} dB",
                      delta_color="normal" if margin_l7 >= 0 else "inverse")
            st.metric("Effective Noise Floor Rise", f"+{noise_rise:.2f} dB")

        st.subheader("Protection Criteria by System")
        ex("RTCA Minimum Operational Performance Standards (MOPS) define receiver sensitivity and selectivity under interference. The I/N thresholds are derived from MOPS by determining the interference level that degrades receiver output (accuracy, availability, integrity) beyond acceptable limits.")
        prot_data = pd.DataFrame([
            ["GPS/GNSS L1/L5", "−10 dB", "DO-235B / DO-253", "Satellite signal extremely weak (~−130 dBm)"],
            ["ADS-B / Mode-S", "−10 dB", "DO-260B", "False target / missed detection consequences"],
            ["Radio Altimeter", "−6 dB", "DO-155 / ETSO-C87", "CAT III landing; 1 dB degradation max"],
            ["ILS Localizer/GS", "−6 dB", "DO-148", "Precision approach; safety-of-life"],
            ["VOR", "−6 dB", "DO-196", "En-route navigation"],
            ["DME / TACAN", "−6 dB", "DO-189", "Distance measuring, co-located with ILS"],
            ["General ARNS", "−6 dB", "ITU-R M.1642", "Default ITU-R criterion for aeronautical"],
        ], columns=["System", "I/N Threshold", "Reference", "Rationale"])
        st.table(prot_data)

        st.subheader("✅ Self-Check")
        with st.expander("Click to reveal answers"):
            st.markdown("""
**Q1:** An interference study shows I/N = −8 dB at a GNSS receiver. Is this acceptable?
> **No** — the GPS threshold is −10 dB. At −8 dB, the threshold is violated by 2 dB. 
> The US should oppose unless mitigation measures are applied.

**Q2:** The same study shows I/N = −8 dB at a radio altimeter. Is this acceptable?
> **Yes** — the radio altimeter threshold is −6 dB. At −8 dB, there is 2 dB of margin.
> Compatible, but flag it as marginal and watch for more aggressive scenarios.

**Q3:** A contribution claims "interference is 20 dB below the noise floor, so it's negligible." 
Is that a strong argument?
> **Yes, actually** — I/N = −20 dB is well within any protection criterion. However, you should 
> verify the assumptions: What noise floor did they use? What propagation model? What deployment 
> density? A −20 dB result with aggressive assumptions may be −6 dB under realistic ones.
            """)

    # ── LESSON 8 ──────────────────────────────────────────────────────────────
    elif lesson.startswith("Lesson 8"):
        st.header("🌐 Lesson 8 — Propagation Models: P.452, P.528, P.676")
        ex("Path loss model choice can swing computed I/N by 10–30 dB depending on terrain and scenario. This is the single parameter most worth challenging in a proponent's study — an unjustified urban clutter correction or inappropriate P.452 time percentage can flip the compatibility finding entirely.")

        st.subheader("Why FSPL Is Just the Starting Point")
        st.markdown("""
Free Space Path Loss assumes a perfect vacuum. Real paths include:
- **Atmospheric gases** — oxygen and water vapor absorb RF above ~1 GHz
- **Terrain** — hills cause diffraction (some loss), but can also provide shielding
- **Clutter** — buildings, trees add loss in urban/suburban areas
- **Multipath** — reflections can add constructively or destructively

Different ITU-R models capture these effects for different scenarios.
        """)

        st.subheader("ITU-R P.452 — Terrestrial Point-to-Point")
        st.markdown("""
**Use when:** Both transmitter and receiver are on or near the ground.

**What it models:**
- Line-of-sight propagation with atmospheric refraction
- Diffraction over terrain obstacles
- Tropospheric scatter (important for long paths)
- Ducting and anomalous propagation

**Key input:** Terrain profile between Tx and Rx — the full P.452 requires actual terrain 
elevation data along the path. Our simplified version uses empirical clutter corrections.

**Time percentage:** P.452 calculates loss exceeded for a given percentage of time.
- 50% = median (use for general studies)
- 1% = nearly worst-case propagation (use for protection studies — less path loss = more interference)
- Always use 1% for regulatory protection work

**Terrain/clutter corrections (simplified):**
| Environment | Additional Loss vs FSPL |
|---|---|
| Open (flat, no obstacles) | ~0 dB |
| Suburban | ~8 dB |
| Urban | ~18 dB |
| Dense urban | ~26 dB |
        """)

        st.subheader("ITU-R P.528 — Aeronautical Propagation")
        st.markdown("""
**Use when:** The victim receiver is airborne (aircraft, helicopter).

**What it models:**
- Slant-path geometry (ground-to-air or air-to-air)
- Atmospheric refraction at low elevation angles
- Tropospheric scatter
- Radio horizon effects

**Why it's different from P.452:**
Aircraft receivers have line-of-sight to transmitters over much longer distances.
An aircraft at 10,000 ft altitude can "see" a ground transmitter 130+ km away.
This dramatically increases the potential number of interferers — and is why 
airborne receivers are especially vulnerable to aggregate interference.

**Key parameters:**
- Aircraft altitude (higher = more line-of-sight, more interferers visible)
- Ground distance (horizontal separation)
- Time percentage (use 1% for protection studies)
        """)

        st.subheader("ITU-R P.676 — Atmospheric Gaseous Attenuation")
        st.markdown("""
**Use when:** Assessing attenuation on paths above ~1 GHz where atmospheric absorption matters.

**What it models:**
- Oxygen absorption (peaks at 60 GHz — used by 5G mmWave for natural isolation)
- Water vapor absorption (peaks at 22 GHz)
- Path length through the atmosphere

**Practical relevance for FAA work:**
- Below 3 GHz: gaseous absorption is minimal (<0.1 dB/km), not a useful protection mechanism
- Above 6 GHz: starts to matter for longer paths
- Above 10 GHz: significant; helps protect airborne weather radar (9–10 GHz)
- The itur library implements P.676 directly — you can run it live in the Propagation module

**Important:** Do NOT claim gaseous attenuation as a protection mechanism for ground-to-air 
paths below 6 GHz in ITU-R submissions — reviewers will challenge it as negligible.
        """)

        st.subheader("Model Selection Decision Tree")
        ex("The ITU-R rapporteur and experienced delegations will immediately flag model misapplication — e.g., using P.452 for an airborne receiver, using 50% time percentage for a protection study, or applying urban clutter corrections in open terrain near airports. These are grounds to request reanalysis.")
        st.markdown("""
```
Is the victim airborne?
├── YES → Use ITU-R P.528
│         (slant path, altitude matters)
└── NO → Is the path over open ocean or flat terrain?
         ├── YES → Use P.452 (open, time%=1)
         └── NO  → Use P.452 (suburban/urban, time%=1)

Always run FSPL first:
├── If FSPL shows compatibility → you're protected (FSPL is most optimistic)
└── If FSPL shows violation → run P.452/P.528 for realistic assessment
    ├── If realistic model still shows violation → cite harmful interference
    └── If realistic model shows compliance → document the margin
```
        """)

        st.subheader("✅ Self-Check")
        with st.expander("Click to reveal answers"):
            st.markdown("""
**Q1:** You're analyzing whether a new 5G base station could interfere with aircraft radio altimeters during approach. Which propagation model do you use?
> **P.528** — the victim is airborne, so slant-path aeronautical propagation applies.

**Q2:** A contribution uses FSPL and shows I/N = −8 dB. Should you be concerned?
> **Yes** — FSPL is the most optimistic model (least path loss = most interference at victim). 
> If FSPL only shows −8 dB margin, more realistic models (P.452/P.528) will show even less margin. 
> The scenario is close to the protection threshold and warrants further study.

**Q3:** A submitting administration uses P.452 with urban clutter correction (+18 dB) to show compatibility. What should you check?
> Check whether the deployment scenario actually justifies "urban" clutter. If the interferers 
> could be deployed in open or suburban areas near airports, the 18 dB clutter correction is 
> not valid and the real interference could be 10–18 dB higher than claimed.
            """)

    # ── LESSON 9 ──────────────────────────────────────────────────────────────
    elif lesson.startswith("Lesson 9"):
        st.header("🎲 Lesson 9 — Monte Carlo & Aggregate Interference")
        ex("Single-entry analysis captures the worst individual interferer but misses aggregate effects. For a deployment of N interferers, aggregate I/N ≈ single-entry I/N + 10·log10(N) at low N — 10 interferers add ~10 dB to the single-entry result, potentially flipping a compliant scenario into a violation.")

        st.subheader("Why Single-Scenario Analysis Isn't Enough")
        st.markdown("""
A link budget with one interferer at one distance answers: "What if the worst case happens?"
But in real deployments:
- Thousands of base stations exist, each at different distances
- Not all are transmitting at maximum power all the time
- The combined (aggregate) effect of many interferers can exceed the threshold 
  even if each individual one is below it

**Aggregate interference** is the sum of all interferers' contributions at the victim receiver.
This is what the ITU-R actually protects against, and it's what Monte Carlo quantifies.
        """)

        st.subheader("How Monte Carlo Works")
        st.markdown("""
**One trial:**
1. Randomly place N interferers within the deployment area
2. Calculate each interferer's received power at the victim (using path loss model)
3. Sum all N interference powers linearly (in milliwatts, not dB)
4. Convert back to dBm → compare to noise floor → compute I/N
5. Record whether I/N > threshold (a "violation")

**Many trials (typically 2,000–10,000):**
- Repeat steps 1–5 thousands of times
- Each trial uses different random positions
- Result: a statistical distribution of I/N values

**Key output:**
- **Violation probability** = fraction of trials where I/N > threshold
- **ITU-R criterion:** violation probability must be < 5% (or 1% for stringent cases)
        """)

        st.subheader("🔢 Mini Monte Carlo — Watch It Work")
        ex("Each trial places N interferers uniformly in area (positions drawn from P(r) ∝ r, i.e., r = √(U)·R_max for uniform areal density), computes path loss per interferer, sums interference powers linearly in mW, then converts aggregate to dBm for the I/N comparison.")

        col1, col2 = st.columns(2)
        with col1:
            mc_n = st.slider("Number of interferers (N)", 1, 50, 10)
            mc_eirp = st.number_input("Each interferer EIRP (dBm)", value=58.0, step=1.0)
            mc_f_l9 = st.number_input("Frequency (MHz)", value=4300.0, step=100.0)
            mc_r = st.number_input("Deployment radius (km)", value=20.0, step=1.0)
            mc_noise_l9 = st.number_input("Noise floor (dBm)", value=-89.0, step=1.0)
            mc_thresh_l9 = st.number_input("I/N threshold (dB)", value=-6.0, step=1.0)
            n_trials_l9 = st.select_slider("Trials", [500, 1000, 2000], value=1000)

        if st.button("▶ Run Mini Monte Carlo", key="mc_l9"):
            trials_in = []
            for _ in range(n_trials_l9):
                r = np.sqrt(np.random.uniform(0.1**2, mc_r**2, mc_n))
                pl = 20*np.log10(np.maximum(r, 0.01)) + 20*np.log10(mc_f_l9) + 32.44
                rx_pw = mc_eirp - pl
                agg_mw = np.sum(10**(rx_pw/10))
                agg_dbm = 10*np.log10(max(agg_mw, 1e-30))
                trials_in.append(agg_dbm - mc_noise_l9)

            trials_in = np.array(trials_in)
            vp = np.mean(trials_in > mc_thresh_l9) * 100

            with col2:
                st.metric("Median I/N", f"{np.percentile(trials_in,50):.1f} dB")
                st.metric("95th pct I/N", f"{np.percentile(trials_in,95):.1f} dB")
                st.metric("Violation Probability", f"{vp:.1f}%",
                          delta="High Risk" if vp > 5 else "OK",
                          delta_color="inverse" if vp > 5 else "normal")

            fig_l9, ax_l9 = plt.subplots(figsize=(8, 4))
            fig_l9.patch.set_facecolor("#0e1117")
            ax_l9.set_facecolor("#0e1117")
            ax_l9.hist(trials_in, bins=50, color='#4488ff', alpha=0.8, density=True)
            ax_l9.axvline(mc_thresh_l9, color='red', lw=2, linestyle='--', label=f'Threshold ({mc_thresh_l9} dB)')
            ax_l9.axvline(np.percentile(trials_in,95), color='orange', lw=1.5, linestyle=':', label='95th pct')
            ax_l9.set_xlabel("I/N (dB)", color='white')
            ax_l9.set_ylabel("Density", color='white')
            ax_l9.set_title("Monte Carlo I/N Distribution", color='white')
            ax_l9.legend(facecolor='#1a1a2e', labelcolor='white', fontsize=8)
            ax_l9.tick_params(colors='white')
            for sp in ax_l9.spines.values(): sp.set_color('#444')
            plt.tight_layout()
            st.pyplot(fig_l9)

        st.subheader("How to Challenge a Monte Carlo Study")
        st.markdown("""
When reviewing another administration's Monte Carlo contribution, examine these assumptions:

| Assumption | What to look for | Your challenge |
|---|---|---|
| **N (number of interferers)** | Is the deployment density realistic? | Cite actual deployment data; argue for higher N |
| **Deployment area** | Did they exclude areas near airports? | Challenge exclusion zones as unrealistic |
| **Propagation model** | Did they use urban clutter inappropriately? | Argue for open/suburban for areas near airports |
| **Time percentage** | Did they use 50% instead of 1%? | Cite ITU-R SM.2028 — use 1% for protection studies |
| **Tx power distribution** | Did they assume all stations at max power? | If not, challenge as non-conservative |
| **Antenna pattern** | Did they model realistic azimuth patterns? | Check whether downtilt assumptions are valid near airports |
        """)

        st.subheader("✅ Self-Check")
        with st.expander("Click to reveal answers"):
            st.markdown("""
**Q1:** A Monte Carlo study shows violation probability of 3.2%. Is this compatible?
> **Yes** — below the ITU-R 5% criterion. However, check whether conservative assumptions were used. 
> If they used 50% time percentage, the real 1% result could exceed 5%.

**Q2:** Why do we sum interference powers linearly and then convert to dB — why not just add dBm values?
> **Because dB values represent logarithms — you can't add logarithms to get the sum.**
> 10 dBm + 10 dBm ≠ 20 dBm. In linear: 10mW + 10mW = 20mW = 13 dBm.
> Always convert to mW, sum linearly, then convert back to dBm.

**Q3:** What does the CCDF curve in the Monte Carlo module tell you?
> The CCDF (Complementary CDF) shows the probability that I/N *exceeds* a given level.
> Reading it at your I/N threshold directly gives you the violation probability — 
> the key number for the ITU-R compatibility criterion.
            """)

    # ── LESSON 10 ──────────────────────────────────────────────────────────────
    elif lesson.startswith("Lesson 10"):
        st.header("🏛️ Lesson 10 — From RF Math to ITU-R Policy")
        ex("The CPM Report method that gains traction at WRC is the one with the most administrations behind it, the most rigorous technical basis, and the clearest proposed regulatory text. Your job is to build all three simultaneously — analysis alone without proposed text is advisory, not decisive.")

        st.subheader("The Translation Problem")
        st.markdown("""
RF engineers speak in dB, dBm, I/N, path loss, and propagation models.
Policy officials speak in regulatory text, agenda items, resolutions, and geopolitical interests.

**Your job** is to be fluent in both. A technically perfect analysis that can't be 
communicated as policy language will not change the outcome of a Working Party session.
        """)

        st.subheader("Mapping RF Results to Policy Language")
        ex("The translation from dB margins to regulatory language is not mechanical — you must also characterize the consequence of interference (what flight operation fails, what safety margin is lost) to justify the regulatory weight of your proposed protection measure.")
        mapping = pd.DataFrame([
            ["I/N threshold violated", "Harmful interference cannot be excluded", "Grounds to oppose"],
            ["I/N margin < 3 dB", "Compatibility is marginal; further study required", "Request additional studies"],
            ["I/N margin > 10 dB", "Compatible under assessed conditions", "Support with conditions documented"],
            ["Violation probability > 5%", "The proposed use poses unacceptable interference risk", "Oppose or require mitigation"],
            ["Violation probability < 1%", "Compatible under conservative assumptions", "Support with caveats"],
            ["FSPL violates, P.528 complies", "Compatibility depends on deployment restrictions", "Propose coordination distance limits"],
            ["Aggregate > single-source", "Deployment density must be limited", "Propose density or PFD limits in RR footnote"],
            ["Receiver blocking identified", "Out-of-band OOBE limits are required", "Propose emission mask in RR or Recommendation"],
        ], columns=["RF Finding", "Policy Translation", "US Position"])
        st.table(mapping)

        st.subheader("The Regulatory Toolkit")
        st.markdown("""
Once you have a finding, you need to know what regulatory mechanism to request.
These are listed from most protective to least protective:

**1. Primary allocation protection (strongest)**
— Keep the protected band free of new allocations. Cite RR No. 4.10.

**2. Radio Regulations footnote**
— Add a footnote to the frequency table requiring coordination with aeronautical services. Negotiated at WRC.

**3. WRC Resolution**
— A formal decision requiring future studies or imposing conditions. Less binding than the RR table.

**4. ITU-R Recommendation**
— Technical guidance document with protection criteria. Not legally binding but heavily cited.

**5. Coordination zone / distance (weakest)**
— Informal agreement to coordinate around airports. Relies on national implementation.

**Rule:** Always ask for the strongest mechanism you can justify technically.
Start with footnote/Resolution language; let others negotiate you down. 
Never start by proposing weaker protection than you actually need.
        """)

        st.subheader("Writing the US Position — A Framework")
        st.markdown("""
When preparing for a Working Party meeting, your position paper should follow this logic:

```
1. IDENTIFY the threat
   "Document 5D/[X] proposes [Y] which could affect [Z FAA system] 
    operating in [band] under [allocation]."

2. QUANTIFY the risk  
   "Our analysis using ITU-R P.528 and SM.2028 shows I/N = [X] dB,
    exceeding the [−6/−10] dB threshold by [Y] dB under [scenario].
    Monte Carlo results indicate [Z]% violation probability."

3. STATE the consequence
   "This would degrade [system] performance, affecting [safety 
    consequence — e.g., CAT III approach capability, GPS accuracy]."

4. INVOKE the legal basis
   "Under RR No. 4.10, this constitutes harmful interference to a 
    safety-of-life service. Resolution 233 requires that new 
    allocations not degrade RNSS performance."

5. PROPOSE the solution
   "The United States proposes [specific text — PFD limit / 
    coordination distance / OOBE mask / footnote language]."

6. BUILD the coalition
   "We note that ICAO, [Admin A], and [Admin B] share these 
    concerns and have submitted aligned contributions."
```
        """)

        st.subheader("Red Flags in Others' Contributions")
        ex("Proponents have strong incentives to reach a compatibility finding — check every assumption that reduces computed interference: time percentage, clutter correction, deployment exclusion zones, antenna downtilt, frequency separation, and receiver bandwidths used for power density normalization.")
        st.markdown("""
| Red Flag | What It Means | Your Response |
|---|---|---|
| Uses 50% time percentage | Should be 1% for protection studies | Cite SM.2028; request reanalysis at 1% |
| Uses urban clutter near airports | Airports are open terrain | Challenge clutter classification |
| Excludes area within X km of airports | No regulatory basis for exclusion | Demand coordination distance be justified |
| Uses receiver characteristics worse than RTCA standard | Understates protection need | Cite RTCA DO-xxx; use correct receiver parameters |
| Only analyzes single interferer | Misses aggregate effect | Request Monte Carlo analysis per SM.2028 |
| Claims "no co-frequency operation" | Ignores OOBE/blocking | Raise receiver blocking/desensitization mechanism |
| Analysis only at median conditions | Should use worst-case | Request 99th percentile analysis |
        """)

        st.subheader("Your 30-Day Readiness Plan")
        ex("Follow this to get fully confident before your first Working Party meeting.")
        st.markdown("""
**Week 1 — Foundation**
- Complete Lessons 1–5 in this module
- Run the Link Budget module for the Radio Altimeter band (4200–4400 MHz) vs 5G at 4300 MHz
- Read the FAA/NTIA C-band technical report (publicly available) — it's a complete worked example

**Week 2 — Analysis**
- Complete Lessons 6–9
- Run Monte Carlo for Radio Altimeter, varying N from 5 to 50
- Read ITU-R M.1642 (free at itu.int) — 20 pages, the foundation of all IMT/ARNS studies

**Week 3 — Policy**
- Complete Lesson 10
- Use the Contribution Analyzer with the three example contributions
- Read two real US contributions from a past WP 5D meeting (ask FAA or NTIA for access)

**Week 4 — Integration**
- Draft a mock US contribution using the Contribution Summary module
- Present your analysis to a colleague and defend your assumptions
- Review the WRC-27 agenda items list and identify which ones affect FAA bands
        """)

        ok("You've completed the RF Training curriculum. You now have the fundamentals to run interference analyses, interpret results, and translate findings into ITU-R policy language.")

    # ── LESSON 11 — WP 5D ────────────────────────────────────────────────────
    elif lesson.startswith("Lesson 11"):
        st.header("📡 Lesson 11 — WP 5D: IMT vs Radio Altimeter (AI 1.7)")
        ex("WRC-27 Agenda Item 1.7 is the highest-priority FAA concern: IMT identification at 4.4–4.8 GHz sits immediately adjacent to the Radio Altimeter band at 4.2–4.4 GHz. This lesson walks through the complete interference analysis as WP 5D sees it.")

        st.subheader("The Threat in Numbers")
        st.markdown("""
| Parameter | Value | Significance |
|---|---|---|
| IMT proposed band | 4,400–4,800 MHz | WRC-27 AI 1.7 candidate |
| Radio Altimeter band | 4,200–4,400 MHz | ARNS safety-of-life |
| Gap at band edge | **0 MHz** (touching at 4,400 MHz) | No guard band whatsoever |
| RA I/N threshold | −6 dB + 6 dB safety factor = **−12 dB** | ITU-R M.1477 Annex 5 |
| RA noise floor | −90 dBm (200 MHz BW, 5 dB NF) | Derived from DO-155 |
| Max interference allowed | −90 + (−12) = **−102 dBm** | At RA receiver input |

The Radio Altimeter is a FMCW ranging radar — it transmits at 4.2–4.4 GHz and measures the return time.
Its receiver front-end is **broadband** — it cannot filter out an adjacent 4.4 GHz IMT signal.
This creates a **blocking and desensitization** risk, not just in-band interference.
        """)

        st.subheader("Step 1 — Identify the Interference Mechanism")
        st.markdown("""
WP 5D's methodology (ITU-R M.1642) requires two separate analyses:

**Path A — OOB emissions from IMT landing in RA band:**
- ITU-R SM.1541 requires IMT OOB emissions to be ≥23 dB below the carrier at the 250% bandwidth boundary
- The 250% BN boundary of 4,400–4,800 MHz (400 MHz BW) = 4,400 − (400 × 1.5) = **3,800 MHz** (lower boundary)
- The upper OOB boundary is at 4,400 + (400 × 1.5) = **5,000 MHz**
- The RA band (4,200–4,400 MHz) is **inside** the lower OOB domain of the IMT band
- SM.1541 mask level at 4,200 MHz: OOB emissions must be ≤ carrier − 23 dB

**Path B — IMT blocking/desensitization of RA receiver front-end:**
- A strong IMT signal at 4,400 MHz can saturate the RA LNA even if spectrally out of band
- DO-155 defines the RA blocking threshold — typically 0 dBm at antenna port
- 5G NR base station EIRP can be 60–70 dBm — FSPL at 1 km is only ~125 dB → received = −65 dBm >> blocking threshold
        """)

        st.subheader("Step 2 — Run the Link Budget")
        st.markdown("""
**Worked example — worst case airborne scenario (M.1642 §4):**

```
IMT base station EIRP:           +63 dBm  (46 dBm Tx + 17 dBi antenna)
FSPL at 1 km, 4,400 MHz:        −124 dB   (20·log10(1) + 20·log10(4400) + 32.44)
Aircraft altitude gain bonus:     +6 dB    (aircraft sees stronger signal than ground)
RA receive antenna gain:           0 dBi   (worst case)
────────────────────────────────────────────
Received interference at RA:    = 63 − 124 + 6 + 0 = −55 dBm

RA noise floor:                  −90 dBm
I/N ratio:                    −55 − (−90) = +35 dB
I/N threshold (M.1477 + 6 dB safety): −12 dB
Protection margin:              −12 − 35 = −47 dB  ← VIOLATED by 47 dB
```
At 1 km the IMT station would be completely incompatible. The question WP 5D must answer: **at what minimum distance is the RA protected?**
        """)

        st.subheader("Step 3 — The SM.1541 OOB Mask Test")
        st.latex(r"\text{OOB limit at } f_{\text{edge}} = P_{\text{carrier}} - 23 \text{ dB (SM.1541, 250\% BN boundary)}")
        st.markdown("""
For a 400 MHz IMT carrier at 46 dBm transmit power:
- OOB limit at boundary = 46 − 23 = **23 dBm** EIRP in the OOB domain
- The RA band sits 200 MHz below the IMT lower edge — inside the OOB domain
- A country arguing the OOB mask is met must show that emissions at 4,200–4,400 MHz are ≤23 dBm EIRP

**What FAA challenges:** Most contributions use average OOB power, not peak. For a 5G NR signal with high PAPR, peak OOB can be 10–15 dB above the average — SM.1541 specifies the average but RA receivers respond to peaks.
        """)

        st.subheader("Step 4 — US Policy Position and Key Arguments")
        st.markdown("""
**US position:** Oppose IMT identification at 4.4–4.8 GHz without:
1. OOB emission limits compliant with SM.1541 for the RA band
2. Coordination zones derived from worst-case airborne M.1642 analysis
3. Monte Carlo aggregate interference per SM.2028 (not just single-entry)
4. Recognition that aviation safety factor (+6 dB) applies per M.1477

**Key counter-arguments when proponents say "OOB is within SM.1541":**
- Ask: is the claimed OOB level per carrier or aggregate for a MIMO system?
- Ask: is it average power or peak power (relevant for FMCW RA receiver)?
- Ask: does the analysis include aircraft at altitude (not just ground receivers)?
- Cite: M.1642 §4 requires airborne victim — if absent, the study is methodologically incomplete

**Floor language template:**
> *"The delegation of the United States notes that the analysis in [Doc X] does not include an airborne victim scenario as required by ITU-R M.1642. The United States requests that the Working Party postpone any decisions on CPM text pending a methodology-complete study addressing the Radio Altimeter protection criterion of I/N ≤ −12 dB (including the aviation safety factor per M.1477)."*
        """)

        with st.expander("✅ Self-Check — WP 5D"):
            st.markdown("""
**Q1:** Why is the aviation safety factor +6 dB applied on top of the I/N threshold for Radio Altimeters?
> **Answer:** M.1477 Annex 5 requires an additional 6 dB margin for safety-of-life precision approach systems to account for real-world variability, model uncertainty, and the consequence of failure during CAT III approaches.

**Q2:** A proponent shows their IMT OOB power at 4,200 MHz is −80 dBm EIRP. The RA maximum interference limit is −102 dBm at the receiver. Does this pass?
> **Answer:** −80 dBm EIRP is the transmitted power — it must then be attenuated by path loss before reaching the RA. At 1 km FSPL is ~124 dB, so received = −80 − 124 = −204 dBm — easily below −102 dBm at that distance. The analysis must show at what distance the limit is met, and that distance becomes the coordination zone.

**Q3:** What is the SM.1541 OOB emission boundary for a 100 MHz wide IMT carrier?
> **Answer:** 250% of 100 MHz BW = 250 MHz from carrier edge. For a 4,400–4,500 MHz carrier, the lower boundary is 4,500 − 250 = **4,250 MHz** — inside the RA band.
            """)

    # ── LESSON 12 — WP 5B ────────────────────────────────────────────────────
    elif lesson.startswith("Lesson 12"):
        st.header("🚢 Lesson 12 — WP 5B: Maritime Radar vs ARSR/ASR")
        ex("WP 5B governs maritime and radiodetermination services. The key FAA concern is co-channel and adjacent-band interference between ship-borne and coastal radiolocation systems and ATC radar (ASR/ARSR) in the 2,700–2,900 MHz band.")

        st.subheader("The Interference Scenario")
        st.markdown("""
**Victim system:** ATC En-Route Surveillance Radar (ARSR/ASR)
- Band: 2,700–2,900 MHz (200 MHz, fully shared with marine radar)
- Protection criterion: ARSR I/N ≤ −6 dB; ASR I/N ≤ −10 dB
- Receiver noise floor: −100 dBm (typical ASR bandwidth ~1 MHz)
- Maximum interference: −100 + (−10) = **−110 dBm** at ASR receiver

**Interferer:** Marine/coastal surveillance radar (WP 5B)
- Typical ship radar Tx power: 25–30 kW peak (44–45 dBm)
- Pulse duty cycle: 0.1% (average power = 44 − 30 = 14 dBm average)
- Antenna gain: 25–30 dBi

**Critical ITU-R distinction:** ATC radar receivers are pulsed — they integrate energy within the pulse window. Marine radar pulses can appear as **co-channel clutter** if their PRF is not sufficiently different from the ASR PRF.
        """)

        st.subheader("The Sharing Challenge — Why M.1849 Governs")
        st.markdown("""
ITU-R M.1849 is the WP 5B methodology document for radar coexistence. It addresses:

1. **Pulse-on-pulse interference** — radar pulses from one system landing within the receive window of another
2. **Clutter injection** — marine surface returns appearing on ATC radar displays
3. **Desensitization** — strong marine radar sidelobes reducing ATC radar sensitivity

**Key M.1849 parameters FAA defends:**
| Parameter | ASR requirement | ARSR requirement |
|---|---|---|
| I/N threshold | −10 dB | −6 dB |
| Probability of interference | <5% of radar rotations | <1% |
| Minimum separation | Coordination zone | Coordination zone |

**The airborne victim problem (unique to this WP):**
Unlike terrestrial ASR analysis, aircraft at altitude can receive marine radar signals directly (no terrain screening). A ship 200 km at sea with 30 dBi antenna at 20 kW peak:
```
EIRP_peak = 43 dBm + 30 dBi = 73 dBm
FSPL at 200 km, 2800 MHz = 20·log10(200) + 20·log10(2800) + 32.44 = 155 dB
Received by aircraft (at altitude, line-of-sight) = 73 − 155 = −82 dBm peak
ASR noise floor = −100 dBm → I/N = −82 − (−100) = +18 dB >> −10 dB threshold
```
This shows marine radar interference to ASR at 200 km is a real threat for airborne surveillance.
        """)

        st.subheader("WP 5B Study Approach — What FAA Checks")
        st.markdown("""
When reviewing a WP 5B contribution, check:

1. **Is P.452 or P.528 used?**
   - Ground-to-ground path: P.452 ✅
   - Ship-to-aircraft path: P.528 ✅ (Earth-space isn't exactly right but P.528 gives airborne geometry)
   - Ground-to-ground ONLY (ignoring airborne): 🚩 flag this as incomplete

2. **Is average or peak power used?**
   - For pulsed radar interference, peak power × duty cycle → average power
   - But ASR receivers respond to PEAK pulse power, not average → check which metric is used

3. **Is the coordination zone realistic?**
   - M.1849 defines required separation as a function of terrain, power, and geometry
   - A proponent using "urban clutter" to reduce apparent range is underestimating the threat

4. **Are all radar bands treated separately?**
   - 2,700–2,900 MHz (ASR/ARSR) ← direct co-channel concern
   - 9,000–9,500 MHz (airborne weather radar) ← separate WP 5B concern

**US floor language template:**
> *"The United States notes that the study in [Doc X] uses average transmitter power rather than peak pulse power for the interference calculation. ITU-R M.1849 requires that interference analysis use peak power when the victim is a pulsed radar system. The United States requests revision of the analysis using peak power before the Working Party proceeds to CPM text."*
        """)

        with st.expander("✅ Self-Check — WP 5B"):
            st.markdown("""
**Q1:** Why does a ship radar with 0.1% duty cycle still pose an interference threat to ASR despite its low average power?
> **Answer:** ASR receivers use pulse compression and range gating — they integrate power within a narrow receive window. A marine radar pulse that falls within an ASR range gate appears as a full-power return, regardless of the low average duty cycle. The interference is measured by the peak pulse power, not the average.

**Q2:** A WP 5B contribution analyzes only a 20 km coordination zone. What should the US delegation challenge?
> **Answer:** Aircraft at altitude have line-of-sight to marine transmitters at hundreds of km. The analysis must include the airborne victim geometry using P.528. A 20 km zone is only valid for ground-level ASR receivers.

**Q3:** What is the ASR I/N threshold and why is it stricter than the ARSR threshold?
> **Answer:** ASR = −10 dB (tighter), ARSR = −6 dB. Airport surveillance radars operate in terminal environments where traffic density is high and false targets or missed detections have immediate safety consequences. En-route ARSR has more geometric diversity and computational filtering, so the 4 dB relaxation is accepted.
            """)

    # ── LESSON 13 — WP 4C ────────────────────────────────────────────────────
    elif lesson.startswith("Lesson 13"):
        st.header("🛰️ Lesson 13 — WP 4C: MSS Satellite Downlinks vs DME/AMS(R)S (AI 1.13)")
        ex("WP 4C governs Mobile Satellite Service (MSS) and direct satellite-to-device (DC-MSS-IMT) systems. WRC-27 AI 1.13 proposes MSS downlinks in three candidate bands, all immediately adjacent to FAA aviation systems. The key methodological trap is using terrestrial propagation models for satellite geometry.")

        st.subheader("The Three Candidate Bands and Their FAA Victims")
        data = {
            "Candidate Band": ["925–960 MHz", "1,475–1,518 MHz", "2,620–2,690 MHz"],
            "FAA Band at Risk": ["DME/TACAN 960–1,215 MHz", "L-band AMS(R)S 1,525–1,559 MHz", "ASR/ARSR 2,700–2,900 MHz"],
            "Gap": ["0 MHz (touching at 960)", "7 MHz", "10 MHz"],
            "Protection Metric": ["epfd ≤ −121.5 dBW/m²/MHz", "ΔT/T ≤ 6% single-entry", "I/N ≤ −10 dB (ASR)"],
            "Key Standard": ["RR 5.444 + ITU-R S.1586", "ITU-R S.1598", "ITU-R M.1849"],
        }
        import pandas as pd
        st.table(pd.DataFrame(data))

        st.subheader("The Critical Methodology Error — P.452 vs P.619")
        st.markdown("""
**The most common WP 4C error:** Using ITU-R P.452 (terrestrial link model) instead of ITU-R P.619 (Earth-space propagation) for satellite downlink interference analysis.

**Why this matters:**
- P.452 models ground-to-ground propagation — it includes terrain diffraction, clutter, and ducting
- P.619 models satellite slant-path geometry — includes free-space spreading plus atmospheric effects on the slant path
- A satellite downlink arrives from **above** — there is no terrain between the satellite and the aircraft receiver
- Using P.452 adds fictitious terrain losses that make the interference appear smaller than it is

**The numerical difference can be 10–30 dB** — the difference between "compatible" and "harmful" in a protection study.

**FAA floor challenge:** *"The delegation of the United States notes that the interference analysis uses ITU-R P.452, which is not applicable to satellite-to-Earth paths. The correct propagation model for this geometry is ITU-R P.619. The United States requests that the proponent revise the analysis using the appropriate Earth-space propagation model before this Working Party makes any decisions."*
        """)

        st.subheader("The epfd Concept for the 925–960 MHz Band")
        st.latex(r"\text{epfd} = \sum_{i=1}^{N} \frac{P_i \cdot G_i(\theta_i)}{4\pi d_i^2} \cdot \frac{\lambda^2 G_r(\phi_i)}{4\pi}")
        st.markdown("""
**epfd (equivalent power flux density)** is the ITU-R metric for satellite constellation aggregate interference. Unlike single-satellite pfd, epfd sums the contributions of ALL simultaneously visible satellites.

**Why single-satellite pfd is insufficient:**
- A LEO constellation may have 10–20 satellites visible simultaneously
- Each contributes interference to the DME receiver
- The aggregate is 10–20× (10–13 dB) larger than the single-satellite contribution
- Proponents who cite single-satellite pfd are systematically underestimating the interference

**DME protection limit:** epfd ≤ −121.5 dBW/m²/MHz (from RR No. 5.444 and the FAA system protection table)

**Worked check:**
```
Single LEO satellite pfd at aircraft:  −130 dBW/m²/MHz  (proponent's claim)
Number of visible satellites (LEO):     15
Aggregate epfd:  −130 + 10·log10(15) = −130 + 11.8 = −118.2 dBW/m²/MHz
Protection limit:                       −121.5 dBW/m²/MHz
Margin:                                 −121.5 − (−118.2) = −3.3 dB  ← VIOLATED
```
        """)

        st.subheader("The ΔT/T Metric for L-band AMS(R)S")
        st.latex(r"\Delta T/T = \frac{T_{\text{interference}}}{T_{\text{system}}} \times 100\%")
        st.markdown("""
For L-band AMS(R)S satellite receivers, interference raises the system noise temperature. The WP 4C protection criterion is:
- ΔT/T ≤ **6% single-entry** (per system protection table)
- ΔT/T ≤ **20% aggregate** (all systems combined)

**Converting I/N to ΔT/T:**
> ΔT/T (%) ≈ 100 × 10^(I/N_dB / 10)

So I/N = −12.2 dB → ΔT/T = 6%. A 6% ΔT/T corresponds to roughly **I/N ≤ −12.2 dB** — tighter than the generic −6 dB threshold.

**Why the gap at 1,475–1,518 MHz is deceptive:**
The 7 MHz gap between the MSS candidate band and AMS(R)S at 1,525 MHz seems comfortable — but at L-band, OOB emissions from a high-power satellite downlink can easily exceed the ΔT/T limit across a 7 MHz gap. The 250% BN boundary per SM.1541 for a 43 MHz wide MSS carrier extends to 1,525 − (43 × 1.5) = **1,460 MHz** — already inside the candidate band.
        """)

        with st.expander("✅ Self-Check — WP 4C"):
            st.markdown("""
**Q1:** A proponent presents a WP 4C study using P.452 and shows I/N = −15 dB at the DME receiver. Should the US support this result?
> **Answer:** No. P.452 is a terrestrial propagation model and must not be used for satellite-to-Earth geometry. The US should request the study be redone with P.619. The actual interference could be 10–30 dB higher, which could push the result above the epfd limit.

**Q2:** The 925–960 MHz candidate band has epfd = −125 dBW/m²/MHz from a single satellite. The full constellation has 20 visible satellites simultaneously. Does the constellation comply with the −121.5 dBW/m²/MHz epfd limit?
> **Answer:** −125 + 10·log10(20) = −125 + 13 = **−112 dBW/m²/MHz** — exceeds the −121.5 limit by 9.5 dB. The full constellation is non-compliant even though a single satellite appeared compliant.

**Q3:** What is the WP 4C working group title and why does it matter for FAA?
> **Answer:** WP 4C is responsible for Mobile Satellite Service (MSS). Its AI 1.13 specifically addresses new MSS allocations in 694–2700 MHz, which covers bands immediately adjacent to DME (960 MHz), AMS(R)S (1525 MHz), and ASR (2700 MHz) — three FAA safety systems. FAA must actively participate in WP 4C to protect these bands.
            """)

    # ── LESSON 14 — WP 7B & 7C ───────────────────────────────────────────────
    elif lesson.startswith("Lesson 14"):
        st.header("🔭 Lesson 14 — WP 7B & 7C: Lunar SRS and EESS Passive (AI 1.15, 1.17, 1.19)")
        ex("WP 7B (Space Research) and WP 7C (EESS/Science) address novel allocation scenarios where the interference pathway is less obvious than for terrestrial or satellite services. The FAA concern is primarily strategic — allocation precedent in safety bands.")

        st.subheader("WP 7B — Lunar Surface Communications (AI 1.15)")
        st.markdown("""
**What AI 1.15 proposes:** New SRS (Space Research Service) allocations for communications between Earth and the lunar surface, and among systems on the lunar surface.

**Why it matters for FAA:** The candidate bands include frequencies near:
- **2,700–2,900 MHz** — en-route ATC radar (ASR/ARSR)
- **5,350–5,470 MHz** — ARNS and EESS (passive)
- **7,190–7,235 MHz** and **8,450–8,500 MHz** — FAA fixed point-to-point links

**The methodology gap — the key FAA argument:**
No established ITU-R Recommendation covers Earth-Moon propagation interference to terrestrial ARNS receivers. The path geometry involves:
1. A lunar transmitter pointing toward Earth
2. Signal travels ~384,000 km (1.28 light-seconds)
3. Signal arrives at Earth surface and can propagate to ground-based receivers

This is fundamentally different from both terrestrial (P.452) and Earth-satellite (P.619) scenarios.
**If no methodology exists, no allocation should be made until one is established.**

**Worked order-of-magnitude check:**
```
Lunar surface transmitter EIRP (assumed):   +40 dBW  (10W, 10 dBi — small lunar lander)
Free-space path loss (384,000 km, 2800 MHz): 20·log10(384000) + 20·log10(2800) + 32.44
                                           = 111.7 + 68.9 + 32.44 = 213 dB
Received PFD at Earth surface:              40 − 213 = −173 dBW → pfd ≈ −173 − 71 = −244 dBW/m²
ASR sensitivity threshold:                 Approximately −130 dBW/m² (typical)
Margin:                                    −244 − (−130) = −114 dB — not a concern

But: a LARGE lunar relay station with +60 dBW EIRP → received pfd = −224 − 71 = −224 dBW/m²
Still well below ASR threshold. Geometry protects here.
```

**The real concern is not immediate interference but allocation precedent:**
Once SRS gets a primary allocation in a band containing ARNS, it becomes harder to defend exclusivity in future WRCs. The US should argue for **secondary** not primary allocations for lunar SRS in all bands containing ARNS.
        """)

        st.subheader("WP 7C — EESS Passive: AI 1.17 and AI 1.19")
        st.markdown("""
EESS (passive) sensors **do not transmit** — they only receive natural Earth emissions (microwave radiometry). They cannot directly interfere with FAA systems. **The concern is purely allocational and strategic.**

**AI 1.17 — Space Weather Sensors (below 30 MHz):**
- Proposed allocations: 2.1–29.89 MHz (HF), 74.8–75.2 MHz (VHF)
- FAA concern: These bands contain HF SELCAL communications and ILS-related frequencies
- The threat: A EESS passive co-primary allocation in HF/VHF could create coordination obligations requiring FAA to constrain its own transmitters to protect passive space sensors

**AI 1.19 — EESS Passive in 4.2–4.4 GHz:**
- This is the **most strategically significant** WP 7C agenda item for FAA
- The Radio Altimeter band (4,200–4,400 MHz) is the exact band proposed for EESS passive
- If EESS gets a **co-primary** allocation here, the regulatory status of the RA band changes

**Why this weakens FAA's position on AI 1.7:**
```
Current status:    4,200–4,400 MHz  →  ARNS primary (exclusive)
After AI 1.19:     4,200–4,400 MHz  →  ARNS primary + EESS passive co-primary
Consequence for AI 1.7:  IMT proponents argue the band is "already shared" and
                          that ARNS exclusivity has been diluted, weakening FAA's
                          basis for opposing IMT interference at 4,400 MHz
```

**The US policy argument:**
> *"The delegation of the United States opposes the assignment of co-primary status to EESS (passive) in 4,200–4,400 MHz. While passive sensors do not transmit, a co-primary footnote in the Radio Regulations would undermine the exclusive ARNS status that underpins the United States' position on WRC-27 AI 1.7. The United States proposes secondary allocation only, with no coordination obligation on ARNS operations."*

**The I/N calculation for AI 1.19 (academic — passive sensors have no I/N):**
There is no interference metric for EESS passive → ARNS, because EESS passive cannot transmit.
The analysis is purely regulatory — can ARNS transmitters (i.e., Radio Altimeters themselves) interfere with EESS passive sensors? The answer is yes, and a co-primary allocation would require RA manufacturers to consider EESS passive protection — adding regulatory complexity to aviation-certified equipment.
        """)

        st.subheader("Strategic Connections — The Three-Agenda-Item Linkage")
        st.markdown("""
These agenda items are strategically connected:

```
AI 1.19 (WP 7C)     →  weakens RA band exclusivity
     ↓
AI 1.7  (WP 5D)     →  IMT proponents cite "shared" status to push OOB limits
     ↓
AI 1.13 (WP 4C)     →  MSS proponents cite precedent from AI 1.7 for MSS downlinks
```

**US Strategy:**
1. **Oppose AI 1.19 first** — protect the RA band allocation table status
2. **Use AI 1.19 rejection in AI 1.7 debates** — "the band is ARNS-exclusive, no sharing precedent"
3. **Coordinate with ICAO** — ICAO has standing to submit liaison statements on safety-of-life impacts
4. **Build a coalition** — European administrations with active RA programs (UK, France, Germany) share these concerns

**The WP 7C passive allocation argument structure:**
| Proponent claims | US counter-argument |
|---|---|
| "Passive sensors can't interfere" | "Co-primary status creates regulatory precedent and coordination obligations on ARNS" |
| "This is just a footnote" | "Footnotes in the Radio Regulations are binding — they change the allocation table permanently" |
| "We only need a small bandwidth" | "The RA band is only 200 MHz — any sharing of allocation status is significant" |
| "This helps Earth observation science" | "Safety-of-life must take precedence over science in aviation bands — cite RR No. 4.10" |
        """)

        with st.expander("✅ Self-Check — WP 7B & 7C"):
            st.markdown("""
**Q1:** A WP 7B contribution proposes lunar SRS allocation at 2,700–2,900 MHz. The analysis shows the received PFD at Earth from the lunar transmitter is −220 dBW/m². The ASR threshold is approximately −130 dBW/m². Should FAA oppose this?
> **Answer:** On interference grounds alone, −220 dBW/m² is 90 dB below the ASR threshold — not an immediate threat. However, FAA should still oppose **primary allocation** and argue that (a) no ITU-R methodology exists for this geometry, (b) a secondary allocation with no coordination rights for lunar SRS is appropriate, and (c) establishing any SRS primary allocation in ARNS bands sets a dangerous precedent.

**Q2:** Why is AI 1.19 considered a higher strategic priority than AI 1.17 for FAA?
> **Answer:** AI 1.19 targets the Radio Altimeter band (4,200–4,400 MHz) — a safety-of-life band directly adjacent to the AI 1.7 IMT proposal. A co-primary EESS allocation in this band would weaken FAA's regulatory basis for opposing IMT interference under AI 1.7. AI 1.17 affects HF/VHF bands where the coordination obligations are less direct.

**Q3:** What regulatory mechanism does the US use to connect WP 7C AI 1.19 to WP 5D AI 1.7?
> **Answer:** At the CPM (Conference Preparatory Meeting) level, the US argues that the allocation table status of 4,200–4,400 MHz must be considered holistically — any change to ARNS exclusivity under AI 1.19 must be evaluated against the implications for IMT OOB emission limits under AI 1.7. This forces the two Working Parties to coordinate rather than treating the bands as independent questions.
            """)


# ─────────────────────────────────────────────────────────────────────────────
# TAB 10 — MEETING NOTES
# ─────────────────────────────────────────────────────────────────────────────
elif selected_tab == "📓 Meeting Notes":
    st.title("📓 ITU-R Working Party Meeting Notes")
    ex("P1–P5 are the five formal plenary sessions that open, checkpoint, and close a WP meeting. Working Group sessions between plenaries produce the actual technical text — your most important notes are from WG sessions, not plenaries.")

    # ── Initialize session state stores ──────────────────────────────────────
    if "mn_meeting_info" not in st.session_state:
        st.session_state.mn_meeting_info = {
            "meeting_name": "", "location": "", "dates": "", "working_party": "WP 5D"
        }
    if "mn_sessions" not in st.session_state:
        st.session_state.mn_sessions = {}       # key: session_id → session dict
    if "mn_documents" not in st.session_state:
        st.session_state.mn_documents = {}      # key: doc_id → doc dict
    if "mn_ai_items" not in st.session_state:
        st.session_state.mn_ai_items = {}       # key: ai_id → agenda item dict
    if "mn_actions" not in st.session_state:
        st.session_state.mn_actions = []        # list of action items

    # ── Sub-page navigation ───────────────────────────────────────────────────
    sub = st.radio("Section:", [
        "🏁 Meeting Setup",
        "📝 Session Notes",
        "📄 Document Tracker",
        "🗂️ Agenda Item Tracker",
        "✅ Action Items",
        "📊 Meeting Dashboard",
        "📤 Export Full Record",
    ], horizontal=True)

    st.markdown("---")

    # ── MEETING SETUP ─────────────────────────────────────────────────────────
    if sub == "🏁 Meeting Setup":
        st.subheader("🏁 Meeting Setup")
        ex("The US delegation head is the spokesperson who must approve any US floor intervention — always check with them before speaking. The FAA lead is your technical authority for aeronautical positions. Know both before the meeting opens.")

        info = st.session_state.mn_meeting_info
        col1, col2 = st.columns(2)
        with col1:
            info["meeting_name"] = st.text_input("Meeting Name",
                value=info.get("meeting_name", ""),
                placeholder="e.g., ITU-R WP 5D Meeting #45")
            info["working_party"] = st.selectbox("Working Party",
                ["WP 5D", "WP 5B", "WP 5A", "SG 5", "Other"],
                index=["WP 5D","WP 5B","WP 5A","SG 5","Other"].index(info.get("working_party","WP 5D")))
            info["location"] = st.text_input("Location",
                value=info.get("location", ""),
                placeholder="e.g., ITU Headquarters, Geneva")
        with col2:
            info["dates"] = st.text_input("Dates",
                value=info.get("dates", ""),
                placeholder="e.g., 14–25 October 2025")
            info["us_head"] = st.text_input("US Delegation Head",
                value=info.get("us_head", ""),
                placeholder="Name / Organization")
            info["faa_lead"] = st.text_input("FAA Technical Lead",
                value=info.get("faa_lead", ""),
                placeholder="Your name")

        st.markdown("**WRC-27 Agenda Items Under Watch**")
        ex("WRC-27 agenda items are formally adopted at WRC-23. Each item is assigned to one or more Study Groups and Working Parties for technical study. The CPM Report summarizing results is due approximately 12 months before WRC-27.")

        # Live WRC-27 database display
        wrc27_rows = []
        for ai_key, ai in WRC27_AGENDA_ITEMS.items():
            threat_icon = {"HIGH": "🔴", "MEDIUM": "🟡", "LOW-MEDIUM": "🟢"}.get(ai["threat_level"], "⚪")
            wrc27_rows.append({
                "AI": ai["ref"],
                "Title (summary)": ai["title"][:65] + "…",
                "WP": ai["working_party"],
                "Threat": f"{threat_icon} {ai['threat_level']}",
                "FAA Systems": "; ".join(ai["faa_systems_at_risk"])[:80] + "…",
            })
        st.dataframe(pd.DataFrame(wrc27_rows), use_container_width=True, hide_index=True)

        # Expandable detail for each AI
        for ai_key, ai in WRC27_AGENDA_ITEMS.items():
            threat_icon = {"HIGH": "🔴", "MEDIUM": "🟡", "LOW-MEDIUM": "🟢"}.get(ai["threat_level"], "⚪")
            with st.expander(f"{threat_icon} {ai['ref']} — {ai['title'][:70]}"):
                dc1, dc2 = st.columns(2)
                with dc1:
                    st.markdown(f"**Service:** {ai['service']}")
                    st.markdown(f"**Working Party:** {ai['working_party']}")
                    st.markdown(f"**Interference Mechanism:** {ai['mechanism']}")
                    st.markdown(f"**Key Concern:** {ai['key_concern']}")
                with dc2:
                    st.markdown("**FAA Systems at Risk:**")
                    for s in ai["faa_systems_at_risk"]:
                        st.markdown(f"  - {s}")
                    if ai.get("candidate_bands_mhz"):
                        st.markdown(f"**Candidate Bands (MHz):** {ai['candidate_bands_mhz']}")
                    st.markdown(f"**US Position:** {ai['us_position']}")
                    st.markdown("**Key Citations:** " + ", ".join(ai["citations"]))
                if ai.get("notes"):
                    st.caption(f"📝 {ai['notes']}")

        st.markdown("**Additional Agenda Items to Track (free text)**")
        ai_text = st.text_area("One per line:",
            value=info.get("ai_text", ""),
            placeholder="e.g. AI 9.1(b) — Resolution 236 review\nAI 10 — any other business",
            height=80)
        info["ai_text"] = ai_text

        st.session_state.mn_meeting_info = info
        ok("Meeting info saved to session.")

    # ── SESSION NOTES ─────────────────────────────────────────────────────────
    elif sub == "📝 Session Notes":
        st.subheader("📝 Session Notes")
        ex("WP meetings typically last 2 weeks with 50–100+ documents. Working Groups handle specific agenda items and produce draft Recommendations, Reports, and CPM text. Track which WG is handling each of your agenda items — that is where the real work happens.")

        col1, col2, col3 = st.columns(3)
        with col1:
            session_label = st.selectbox("Session", [
                "P1 — Opening Plenary",
                "WG-1 (morning)", "WG-1 (afternoon)",
                "WG-2 (morning)", "WG-2 (afternoon)",
                "WG-3 (morning)", "WG-3 (afternoon)",
                "WG-4 (morning)", "WG-4 (afternoon)",
                "P2 — Mid-week Plenary",
                "WG-5 (morning)", "WG-5 (afternoon)",
                "WG-6 (morning)", "WG-6 (afternoon)",
                "WG-7 (morning)", "WG-7 (afternoon)",
                "P3 — Plenary",
                "WG-8 (morning)", "WG-8 (afternoon)",
                "P4 — Plenary",
                "P5 — Closing Plenary",
                "Coordination Meeting",
                "Bilateral",
            ])
        with col2:
            session_date = st.text_input("Date", placeholder="e.g., Mon 14 Oct")
        with col3:
            session_chair = st.text_input("Chair / Rapporteur", placeholder="Name")

        ai_context = st.text_input("Agenda Item(s) covered", placeholder="e.g., AI 1.2, AI 9.1(b)")

        st.markdown("**Session Notes**")
        ex("Note the document number of any text that was agreed, modified, or deferred — you will need to reference these exactly in your trip report and in any follow-up contributions. Agreed text is marked [AGREED]; contentious text is bracketed [ ] pending resolution.")

        note_text = st.text_area("Notes:", height=200,
            placeholder="""• Doc 5D/123 (China) introduced — proposes IMT identification at 4800-4990 MHz
• US (Smith) intervened citing FAA radio altimeter concerns — I/N analysis shows threshold exceeded
• France/Germany supported further study before any identification
• Rapporteur agreed to hold for additional contributions next meeting
• Japan proposed compromise text with coordination zone — US noted this is insufficient without PFD limits
• Decision: AIs 1.2 and 9.1(b) deferred to next meeting; new contributions due by [date]""")

        faa_outcome = st.selectbox("FAA Outcome This Session", [
            "✅ Favorable — FAA position advanced",
            "⚠️ Mixed — partial progress, concerns remain",
            "🔴 Unfavorable — opposing text gaining traction",
            "⏳ Deferred — no decision, carried forward",
            "ℹ️ Informational — no action required",
        ])

        key_decisions = st.text_area("Key Decisions / Agreed Text:", height=80,
            placeholder="Record any agreed language, decisions, or text that was approved...")

        follow_up = st.text_area("Follow-up Required:", height=80,
            placeholder="What does the US/FAA need to do before the next session?")

        if st.button("💾 Save Session Notes", type="primary"):
            sid = f"{session_label}_{session_date}".replace(" ","_").replace("/","_")
            st.session_state.mn_sessions[sid] = {
                "session": session_label,
                "date": session_date,
                "chair": session_chair,
                "ai_context": ai_context,
                "notes": note_text,
                "faa_outcome": faa_outcome,
                "key_decisions": key_decisions,
                "follow_up": follow_up,
            }
            ok(f"Session notes saved: {session_label} — {session_date}")

        # Show existing sessions
        if st.session_state.mn_sessions:
            st.markdown("---")
            st.subheader("Saved Sessions")
            for sid, s in st.session_state.mn_sessions.items():
                outcome_color = {
                    "✅ Favorable — FAA position advanced": "🟢",
                    "⚠️ Mixed — partial progress, concerns remain": "🟡",
                    "🔴 Unfavorable — opposing text gaining traction": "🔴",
                    "⏳ Deferred — no decision, carried forward": "🔵",
                    "ℹ️ Informational — no action required": "⚪",
                }.get(s["faa_outcome"], "⚪")
                with st.expander(f"{outcome_color} {s['session']} — {s['date']} | {s.get('ai_context','')}"):
                    st.markdown(f"**Chair:** {s.get('chair','—')}")
                    st.markdown(f"**FAA Outcome:** {s['faa_outcome']}")
                    st.markdown("**Notes:**")
                    st.text(s["notes"])
                    if s.get("key_decisions"):
                        st.markdown(f"**Key Decisions:** {s['key_decisions']}")
                    if s.get("follow_up"):
                        st.markdown(f"**Follow-up:** {s['follow_up']}")
                    if st.button(f"🗑️ Delete", key=f"del_sess_{sid}"):
                        del st.session_state.mn_sessions[sid]
                        st.rerun()

    # ── DOCUMENT TRACKER ──────────────────────────────────────────────────────
    elif sub == "📄 Document Tracker":
        st.subheader("📄 Document Tracker")
        ex("Documents are pre-posted to the ITU-R document system (TIES) before each meeting. Flag high-concern documents before the meeting starts so you arrive with prepared interventions — reacting in real time to a complex sharing study is not ideal.")

        with st.expander("➕ Add New Document", expanded=True):
            col1, col2, col3 = st.columns(3)
            with col1:
                doc_num = st.text_input("Document Number", placeholder="e.g., 5D/234-E")
                doc_title = st.text_input("Title / Subject", placeholder="Brief description")
                doc_admin = st.text_input("Submitting Admin(s)", placeholder="e.g., China, Korea")
            with col2:
                doc_ai = st.text_input("WRC Agenda Item", placeholder="e.g., AI 1.2")
                doc_session = st.text_input("Introduced in Session", placeholder="e.g., WG-1 Mon AM")
                doc_type = st.selectbox("Document Type", [
                    "Sharing/compatibility study",
                    "Draft new Recommendation",
                    "Amendment to Recommendation",
                    "Liaison statement",
                    "Information document",
                    "Working document / DT",
                    "Chairman's report",
                ])
            with col3:
                doc_concern = st.selectbox("FAA Concern Level", [
                    "🔴 HIGH — Directly threatens FAA band",
                    "🟡 MEDIUM — Adjacent band / indirect risk",
                    "🟢 LOW — Monitor only",
                    "✅ FAVORABLE — Supports FAA position",
                    "⚪ NEUTRAL — No FAA impact",
                ])
                doc_us_action = st.selectbox("Required US Action", [
                    "Oppose — prepare rebuttal contribution",
                    "Comment — propose amendments",
                    "Support — align with US position",
                    "Monitor — no action this meeting",
                    "Refer to FAA for technical input",
                    "Coordinate with ICAO",
                ])

            doc_summary = st.text_area("Technical Summary / Notes:", height=100,
                placeholder="What does this document propose? What are the FAA implications? Key technical claims?")
            doc_faa_response = st.text_area("FAA Response / Intervention Taken:", height=80,
                placeholder="What did the US say? What text was proposed?")

            if st.button("💾 Save Document", type="primary", key="save_doc"):
                if doc_num:
                    did = doc_num.replace("/","_").replace("-","_").replace(" ","_")
                    st.session_state.mn_documents[did] = {
                        "doc_num": doc_num,
                        "title": doc_title,
                        "admin": doc_admin,
                        "ai": doc_ai,
                        "session": doc_session,
                        "doc_type": doc_type,
                        "concern": doc_concern,
                        "us_action": doc_us_action,
                        "summary": doc_summary,
                        "faa_response": doc_faa_response,
                    }
                    ok(f"Document {doc_num} saved.")
                else:
                    warn("Please enter a document number.")

        # Document table
        if st.session_state.mn_documents:
            st.subheader("Document Index")
            ex("Focus first on documents proposing new or amended Radio Regulations text (footnotes, allocations, Resolutions) — these have direct treaty-level impact. Informational documents and liaison statements are lower priority unless they establish technical precedent.")

            filter_concern = st.multiselect("Filter by concern level:",
                ["🔴 HIGH — Directly threatens FAA band",
                 "🟡 MEDIUM — Adjacent band / indirect risk",
                 "🟢 LOW — Monitor only",
                 "✅ FAVORABLE — Supports FAA position",
                 "⚪ NEUTRAL — No FAA impact"],
                default=["🔴 HIGH — Directly threatens FAA band",
                         "🟡 MEDIUM — Adjacent band / indirect risk"])

            rows = []
            for did, d in st.session_state.mn_documents.items():
                if not filter_concern or d["concern"] in filter_concern:
                    rows.append({
                        "Doc #": d["doc_num"],
                        "Admin": d["admin"],
                        "AI": d["ai"],
                        "Session": d["session"],
                        "Concern": d["concern"].split("—")[0].strip(),
                        "US Action": d["us_action"].split("—")[0].strip(),
                        "Title": d["title"],
                    })
            if rows:
                st.dataframe(pd.DataFrame(rows), use_container_width=True)

            # Detail view
            st.subheader("Document Detail")
            doc_select = st.selectbox("Select document for full details:",
                [d["doc_num"] for d in st.session_state.mn_documents.values()])
            if doc_select:
                did_sel = doc_select.replace("/","_").replace("-","_").replace(" ","_")
                if did_sel in st.session_state.mn_documents:
                    d = st.session_state.mn_documents[did_sel]
                    col1, col2 = st.columns(2)
                    with col1:
                        st.markdown(f"**Submitting Admin:** {d['admin']}")
                        st.markdown(f"**Agenda Item:** {d['ai']}")
                        st.markdown(f"**Type:** {d['doc_type']}")
                        st.markdown(f"**FAA Concern:** {d['concern']}")
                    with col2:
                        st.markdown(f"**Session:** {d['session']}")
                        st.markdown(f"**US Action:** {d['us_action']}")
                    st.markdown(f"**Summary:** {d['summary']}")
                    if d.get("faa_response"):
                        st.markdown(f"**FAA Response:** {d['faa_response']}")
                    if st.button("🗑️ Delete this document", key=f"del_doc_{did_sel}"):
                        del st.session_state.mn_documents[did_sel]
                        st.rerun()

    # ── AGENDA ITEM TRACKER ───────────────────────────────────────────────────
    elif sub == "🗂️ Agenda Item Tracker":
        st.subheader("🗂️ Agenda Item Tracker")
        ex("The US position matrix is the internal coordination document that tracks where each agenda item stands relative to the US goal. Update it at the end of each day — it feeds directly into the delegation head's daily debrief and NTIA post-meeting report.")

        with st.expander("➕ Add / Update Agenda Item", expanded=True):
            col1, col2 = st.columns(2)
            with col1:
                ai_num = st.text_input("Agenda Item Number", placeholder="e.g., 1.2")
                ai_title = st.text_input("Title", placeholder="e.g., IMT identification above 6 GHz")
                ai_faa_bands = st.text_input("FAA Bands at Risk",
                    placeholder="e.g., Radio Altimeter 4200-4400 MHz, ARNS 5000-5150 MHz")
                ai_us_position = st.selectbox("Current US Position", [
                    "🔴 OPPOSE — unacceptable interference risk",
                    "🟡 CONDITIONAL — support with protective measures",
                    "🟢 SUPPORT — compatible with FAA systems",
                    "🔵 STUDYING — analysis underway",
                    "⚪ TBD — position not yet established",
                ])
            with col2:
                ai_status = st.selectbox("Meeting Status", [
                    "📥 Not yet discussed",
                    "🔄 Under discussion — active drafting",
                    "⚠️ Contentious — major disagreement",
                    "🤝 Converging — near agreement",
                    "✅ Agreed — text finalized this meeting",
                    "⏳ Deferred — carried to next meeting",
                ])
                ai_rapporteur = st.text_input("Rapporteur", placeholder="Name / Admin")
                ai_next_steps = st.text_area("Next Steps / FAA Required Actions:", height=80,
                    placeholder="What does the US/FAA need to do before next meeting?")

            ai_current_text = st.text_area("Current Draft Text (paste key proposed RR language):",
                height=100,
                placeholder="Paste the current state of any proposed RR footnote, Resolution, or Recommendation text...")
            ai_faa_concerns = st.text_area("Technical FAA Concerns:", height=80,
                placeholder="What specific interference mechanism or regulatory gap is the FAA concerned about?")
            ai_allies = st.text_input("Aligned Administrations / Organizations",
                placeholder="e.g., EU, Canada, Australia, ICAO")

            if st.button("💾 Save Agenda Item", type="primary", key="save_ai"):
                if ai_num:
                    aid = ai_num.replace(".","_").replace(" ","_")
                    st.session_state.mn_ai_items[aid] = {
                        "num": ai_num, "title": ai_title,
                        "faa_bands": ai_faa_bands,
                        "us_position": ai_us_position,
                        "status": ai_status,
                        "rapporteur": ai_rapporteur,
                        "next_steps": ai_next_steps,
                        "current_text": ai_current_text,
                        "faa_concerns": ai_faa_concerns,
                        "allies": ai_allies,
                    }
                    ok(f"Agenda Item {ai_num} saved.")
                else:
                    warn("Please enter an agenda item number.")

        # Position matrix
        if st.session_state.mn_ai_items:
            st.subheader("US Position Matrix")
            ex("End-of-day updates are critical: overnight, other delegations coordinate bilaterally and positions shift. If you arrive the next morning without an updated picture, you may be caught off-guard by a compromise proposal that moved without you.")
            rows = []
            for aid, a in st.session_state.mn_ai_items.items():
                rows.append({
                    "AI #": a["num"],
                    "Title": a["title"],
                    "FAA Bands at Risk": a["faa_bands"],
                    "US Position": a["us_position"].split("—")[0].strip(),
                    "Status": a["status"].split("—")[0].strip(),
                    "Allies": a["allies"],
                    "Rapporteur": a["rapporteur"],
                })
            st.dataframe(pd.DataFrame(rows), use_container_width=True)

            # Detail expanders
            for aid, a in st.session_state.mn_ai_items.items():
                pos_icon = a["us_position"].split(" ")[0]
                with st.expander(f"{pos_icon} AI {a['num']} — {a['title']}"):
                    col1, col2 = st.columns(2)
                    with col1:
                        st.markdown(f"**FAA Bands:** {a['faa_bands']}")
                        st.markdown(f"**US Position:** {a['us_position']}")
                        st.markdown(f"**Status:** {a['status']}")
                        st.markdown(f"**Rapporteur:** {a['rapporteur']}")
                        st.markdown(f"**Allied Admins:** {a['allies']}")
                    with col2:
                        if a.get("faa_concerns"):
                            st.markdown(f"**FAA Concerns:** {a['faa_concerns']}")
                        if a.get("next_steps"):
                            st.markdown(f"**Next Steps:** {a['next_steps']}")
                    if a.get("current_text"):
                        st.markdown("**Current Draft Text:**")
                        st.code(a["current_text"], language=None)
                    if st.button("🗑️ Delete", key=f"del_ai_{aid}"):
                        del st.session_state.mn_ai_items[aid]
                        st.rerun()

    # ── ACTION ITEMS ──────────────────────────────────────────────────────────
    elif sub == "✅ Action Items":
        st.subheader("✅ Action Items")
        ex("Contribution deadlines are typically 4–6 weeks before each WP meeting. Missing a deadline means your analysis is submitted as 'late' — it may not be formally considered. Action items from the meeting floor (rapporteur requests) often have even shorter turnaround windows.")

        with st.expander("➕ Add Action Item", expanded=True):
            col1, col2, col3 = st.columns(3)
            with col1:
                act_desc = st.text_area("Action Description:", height=80,
                    placeholder="e.g., Prepare interference analysis for AI 1.2 at 4800-4990 MHz using P.528, submit as US contribution")
            with col2:
                act_owner = st.text_input("Owner", placeholder="Name / Organization")
                act_due = st.text_input("Due Date", placeholder="e.g., 3 days before next meeting")
                act_ai = st.text_input("Related Agenda Item", placeholder="e.g., AI 1.2")
            with col3:
                act_priority = st.selectbox("Priority", [
                    "🔴 URGENT — before next session",
                    "🟡 HIGH — before end of meeting",
                    "🟢 NORMAL — before next meeting",
                    "🔵 LOW — informational follow-up",
                ])
                act_status = st.selectbox("Status", [
                    "⬜ Not started",
                    "🔄 In progress",
                    "✅ Complete",
                    "⛔ Blocked",
                ])

            if st.button("💾 Add Action Item", type="primary", key="save_act"):
                if act_desc:
                    st.session_state.mn_actions.append({
                        "desc": act_desc,
                        "owner": act_owner,
                        "due": act_due,
                        "ai": act_ai,
                        "priority": act_priority,
                        "status": act_status,
                    })
                    ok("Action item added.")

        if st.session_state.mn_actions:
            st.subheader("Open Action Items")
            for i, act in enumerate(st.session_state.mn_actions):
                pri_icon = act["priority"].split(" ")[0]
                stat_icon = act["status"].split(" ")[0]
                with st.expander(f"{pri_icon} {stat_icon} AI {act.get('ai','—')} — {act['desc'][:60]}..."):
                    col1, col2 = st.columns(2)
                    with col1:
                        st.markdown(f"**Owner:** {act['owner']}")
                        st.markdown(f"**Due:** {act['due']}")
                        st.markdown(f"**Agenda Item:** {act['ai']}")
                    with col2:
                        st.markdown(f"**Priority:** {act['priority']}")
                        new_status = st.selectbox("Update Status:",
                            ["⬜ Not started","🔄 In progress","✅ Complete","⛔ Blocked"],
                            index=["⬜ Not started","🔄 In progress","✅ Complete","⛔ Blocked"].index(act["status"]),
                            key=f"act_stat_{i}")
                        if new_status != act["status"]:
                            st.session_state.mn_actions[i]["status"] = new_status
                    st.markdown(f"**Description:** {act['desc']}")
                    if st.button("🗑️ Remove", key=f"del_act_{i}"):
                        st.session_state.mn_actions.pop(i)
                        st.rerun()

    # ── DASHBOARD ─────────────────────────────────────────────────────────────
    elif sub == "📊 Meeting Dashboard":
        st.subheader("📊 Meeting Dashboard")
        info = st.session_state.mn_meeting_info
        if info.get("meeting_name"):
            st.markdown(f"### {info.get('meeting_name','Meeting')} — {info.get('working_party','')} | {info.get('location','')} | {info.get('dates','')}")
        ex("The dashboard gives you a rapid red/yellow/green picture of FAA equities across the meeting. A shift from 'Converging' to 'Contentious' on a high-concern agenda item is an escalation trigger — notify NTIA and consider requesting a bilateral with the opposing administration.")

        # Summary metrics
        docs = st.session_state.mn_documents
        ais = st.session_state.mn_ai_items
        sessions = st.session_state.mn_sessions
        actions = st.session_state.mn_actions

        col1, col2, col3, col4, col5 = st.columns(5)
        col1.metric("Sessions Logged", len(sessions))
        col2.metric("Documents Tracked", len(docs))
        col3.metric("Agenda Items", len(ais))
        col4.metric("Action Items", len(actions))
        col5.metric("Open Actions",
            sum(1 for a in actions if "Complete" not in a["status"]))

        # Document concern breakdown
        if docs:
            st.markdown("---")
            col_a, col_b = st.columns(2)
            with col_a:
                st.subheader("Documents by Concern Level")
                concern_counts = {}
                for d in docs.values():
                    c = d["concern"].split("—")[0].strip()
                    concern_counts[c] = concern_counts.get(c, 0) + 1
                fig_dash, ax_dash = plt.subplots(figsize=(5, 4))
                fig_dash.patch.set_facecolor("#0e1117")
                ax_dash.set_facecolor("#0e1117")
                labels = list(concern_counts.keys())
                sizes = list(concern_counts.values())
                colors_pie = ['#ff4444','#ffaa00','#44bb44','#4488ff','#aaaaaa'][:len(labels)]
                wedges, texts, autotexts = ax_dash.pie(sizes, labels=labels,
                    autopct='%1.0f%%', colors=colors_pie, textprops={'color':'white','fontsize':8})
                plt.tight_layout()
                st.pyplot(fig_dash)

            with col_b:
                st.subheader("US Position Summary")
                if ais:
                    pos_counts = {}
                    for a in ais.values():
                        p = a["us_position"].split("—")[0].strip()
                        pos_counts[p] = pos_counts.get(p, 0) + 1
                    fig_pos, ax_pos = plt.subplots(figsize=(5, 4))
                    fig_pos.patch.set_facecolor("#0e1117")
                    ax_pos.set_facecolor("#0e1117")
                    bars = ax_pos.barh(list(pos_counts.keys()),
                        list(pos_counts.values()),
                        color=['#ff4444','#ffaa00','#44bb44','#4488ff','#aaaaaa'][:len(pos_counts)])
                    ax_pos.tick_params(colors='white', labelsize=8)
                    ax_pos.set_xlabel("Count", color='white')
                    for sp in ax_pos.spines.values(): sp.set_color('#444')
                    plt.tight_layout()
                    st.pyplot(fig_pos)
                else:
                    st.info("No agenda items logged yet.")

        # Session outcomes timeline
        if sessions:
            st.markdown("---")
            st.subheader("Session Outcomes")
            for sid, s in sessions.items():
                outcome = s["faa_outcome"]
                icon = "🟢" if "Favorable" in outcome else "🔴" if "Unfavorable" in outcome else "🟡" if "Mixed" in outcome else "🔵"
                st.markdown(f"{icon} **{s['session']}** ({s['date']}) — {s.get('ai_context','—')} — *{outcome.split('—')[-1].strip()}*")

        # High-priority actions
        if actions:
            st.markdown("---")
            st.subheader("🔴 Open High-Priority Actions")
            urgent = [a for a in actions if "URGENT" in a["priority"] and "Complete" not in a["status"]]
            if urgent:
                for a in urgent:
                    st.markdown(f"- **{a['owner']}** | AI {a.get('ai','—')} | Due: {a['due']} — {a['desc'][:80]}...")
            else:
                ok("No urgent open actions.")

    # ── EXPORT ────────────────────────────────────────────────────────────────
    elif sub == "📤 Export Full Record":
        st.subheader("📤 Export Full Meeting Record")
        ex("The trip report is a formal deliverable to NTIA and FAA within 5–10 business days of the meeting. It documents all agreed text affecting FAA interests, US interventions made, outstanding action items, and recommended US positions for the next meeting cycle.")

        info     = st.session_state.mn_meeting_info
        docs     = st.session_state.mn_documents
        ais      = st.session_state.mn_ai_items
        sessions = st.session_state.mn_sessions
        actions  = st.session_state.mn_actions

        # Summary of what will be exported
        col_s1, col_s2, col_s3, col_s4 = st.columns(4)
        col_s1.metric("Sessions logged",   len(sessions))
        col_s2.metric("Documents tracked", len(docs))
        col_s3.metric("Agenda items",       len(ais))
        col_s4.metric("Action items",       len(actions))

        st.markdown("---")

        if st.button("📄 Generate Word Report", type="primary"):
            try:
                import re as _fn_re
                from datetime import date as _fn_date

                docx_bytes = _make_meeting_docx(info, sessions, docs, ais, actions)

                # Build a filename from meeting name, WP, and date
                meeting_name = info.get("meeting_name") or "Meeting"
                wp_short     = info.get("working_party") or "WP"
                safe_meeting = _fn_re.sub(r'[^A-Za-z0-9_-]', '_', meeting_name)[:30].strip('_')
                safe_wp      = _fn_re.sub(r'[^A-Za-z0-9_-]', '_', wp_short)[:10].strip('_')
                fname = f"MeetingRecord_{safe_wp}_{safe_meeting}_{_fn_date.today()}.docx"

                st.download_button(
                    label="⬇️ Download Meeting Record (.docx)",
                    data=docx_bytes,
                    file_name=fname,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    type="primary",
                    use_container_width=True,
                )
                ok(f"Word document ready — {len(docx_bytes)//1024} KB. Click the button above to save to your PC.")

            except Exception as e:
                st.error(f"❌ Word generation error: {e}")
                import traceback; st.code(traceback.format_exc())

# ─────────────────────────────────────────────────────────────────────────────
# TAB — CONTRIBUTION CODE ANALYZER
# ─────────────────────────────────────────────────────────────────────────────
elif selected_tab == "🔬 Contribution Code Analyzer":
    st.title("🔬 Contribution Code Analyzer")
    st.caption("Policy support for **WP 5B · WP 5D · WP 4C · WP 7B · WP 7C** — FAA aeronautical spectrum protection")
    ex("Paste MATLAB or Python code from a WP 5B, 5D, 4C, 7B, or 7C contribution. The critique checks compliance with the ITU-R methodology required for that Working Party — not general code quality.")

    # API key
    api_key = None
    try:
        api_key = st.secrets["ANTHROPIC_API_KEY"]
    except Exception:
        pass

    if not api_key:
        st.warning("⚠️ Anthropic API key not configured — see the Contribution Analyzer module for setup instructions.")
        st.stop()

    st.markdown("---")

    # ── Metadata ─────────────────────────────────────────────────────────────
    st.subheader("📋 Code Context")
    ex("Context determines which WP-specific critique framework is applied — a 5D IMT sharing study is assessed against ARNS/RNSS criteria; a 7C radar study is assessed against radio altimeter and SSR criteria.")

    col1, col2, col3 = st.columns(3)
    with col1:
        code_wp = st.multiselect("Working Party (select all that apply):", [
            "WP 5D — IMT / Mobile",
            "WP 5B — Maritime / Radiodetermination",
            "WP 4C — MSS / DC-MSS-IMT (Satellite)",
            "WP 7B — Space Research / Lunar SRS",
            "WP 7C — EESS / Space Weather (Passive)",
        ], default=["WP 5D — IMT / Mobile"])
        code_doc = st.text_input("Document Number", placeholder="e.g., 5D/567-E")
        code_admin = st.text_input("Submitting Administration", placeholder="e.g., China")

    with col2:
        code_lang = st.selectbox("Code Language", [
            "MATLAB",
            "Python",
            "MATLAB + Python (mixed)",
            "R",
            "Unknown / Pseudocode",
        ])
        code_ai = st.text_input("WRC Agenda Item", placeholder="e.g., AI 1.2")
        code_scenario = st.selectbox("Analysis Type in Code", [
            "Sharing / compatibility study",
            "Aggregate interference (Monte Carlo)",
            "Single-entry interference",
            "Propagation / path loss calculation",
            "Link budget",
            "PFD / EPFD calculation",
            "Receiver sensitivity / selectivity model",
            "Terrain / geographic analysis",
            "Unknown",
        ])

    with col3:
        faa_victim = st.multiselect("FAA Victim System(s) to Assess Against:", [
            "GPS L1 (1559–1610 MHz)",
            "GPS L5 / ARNS (1164–1215 MHz)",
            "Radio Altimeter (4200–4400 MHz)",
            "ADS-B / Mode-S (1085–1095 MHz)",
            "DME / TACAN (960–1215 MHz)",
            "ILS / VOR (108–335 MHz)",
            "En-Route Radar (2700–2900 MHz)",
            "ARNS 5 GHz (5000–5150 MHz)",
            "Airborne Weather Radar (9000–9500 MHz)",
            "WAAS / SBAS (1559–1610 MHz)",
            "All FAA protected bands",
        ], default=["Radio Altimeter (4200–4400 MHz)"])
        critique_depth = st.selectbox("Critique Depth", [
            "Quick scan — flag critical assumption errors",
            "Standard — full line-by-line critique with corrections",
            "Deep — full critique + corrected code + contribution rebuttal text",
        ])

    # ── Code input ────────────────────────────────────────────────────────────
    st.markdown("---")
    st.subheader("📥 Code Input")
    ex("Paste the complete code block. If it spans multiple files, paste all sections concatenated — the AI will identify file/function boundaries. Comments and variable names are critical context.")

    code_tabs = st.tabs(["📝 Paste Code", "📎 Multiple Blocks"])

    with code_tabs[0]:
        code_input = st.text_area(
            "Paste MATLAB / Python code here:",
            height=350,
            placeholder="""% Example MATLAB snippet from a WP 5D sharing study
% ITU-R WP 5D Document 5D/XXX — Sharing study IMT vs ARNS

f_MHz = 4300;           % Interferer center frequency
Pt_dBm = 43;            % Transmit power
Gt_dBi = 15;            % Antenna gain
distances = 1:1:50;     % km

% Free space path loss
FSPL = 20*log10(distances) + 20*log10(f_MHz) + 32.44;

% Received interference
Pr = Pt_dBm + Gt_dBi - FSPL;

% Victim noise floor (assumed)
N_floor = -89;          % dBm

% I/N check
IN_ratio = Pr - N_floor;
threshold = -6;         % dB

compatible = all(IN_ratio < threshold);
fprintf('Compatible: %d\\n', compatible);""",
            key="code_main"
        )

    with code_tabs[1]:
        st.markdown("Paste additional code blocks (e.g., separate functions, helper scripts):")
        code_block2 = st.text_area("Block 2 (optional):", height=150, key="code_b2")
        code_block3 = st.text_area("Block 3 (optional):", height=150, key="code_b3")
        code_input_full = code_input
        if code_block2.strip():
            code_input_full += "\n\n% --- BLOCK 2 ---\n" + code_block2
        if code_block3.strip():
            code_input_full += "\n\n% --- BLOCK 3 ---\n" + code_block3

    # combine blocks — fallback if only tab 0 used
    if not code_block2.strip() and not code_block3.strip():
        code_input_full = code_input

    # ── Optional context ──────────────────────────────────────────────────────
    with st.expander("➕ Additional Context (recommended)"):
        code_summary = st.text_area(
            "What does this code claim to show? (from the contribution text):",
            height=80,
            placeholder="e.g., 'This study demonstrates that IMT base stations operating at 4800 MHz are compatible with radio altimeters at separation distances > 3 km using P.452 urban clutter model.'"
        )
        code_conclusion = st.text_area(
            "What conclusion does the contribution draw from this code?:",
            height=80,
            placeholder="e.g., 'The analysis shows compatibility at 99% of locations, supporting IMT identification at 4800–4990 MHz.'"
        )
        code_prior_concern = st.text_area(
            "Specific FAA technical concern to focus on (optional):",
            height=60,
            placeholder="e.g., Focus on whether receiver blocking / desensitization is modeled, not just in-band interference."
        )

    # ── Run analysis ──────────────────────────────────────────────────────────
    st.markdown("---")
    run_col1, run_col2 = st.columns([2, 1])
    with run_col1:
        run_btn = st.button("🔬 Analyze Code", type="primary",
                           disabled=not code_input.strip())
    with run_col2:
        st.caption("Analysis takes 20–45 seconds for deep critique")

    if run_btn and code_input.strip():

        # Build WP-specific context
        wp_context = {
            "WP 5D — IMT / Mobile": """WP 5D governs IMT (4G/5G/6G) spectrum. Key FAA concerns:
- New IMT allocations adjacent to radio altimeter (4200-4400 MHz), ARNS 5 GHz (5000-5150 MHz), GPS L5 (1164-1215 MHz)
- Sharing studies must use P.528 for airborne victims, not P.452 terrestrial
- Monte Carlo must follow SM.2028; time percentage must be 1% for protection studies
- Out-of-band emissions and receiver blocking must be analyzed, not just in-band sharing
- Protection criteria: I/N ≤ −6 dB (ARNS), I/N ≤ −10 dB (GNSS), per M.1477/M.1642""",

            "WP 5B — Maritime / Radiodetermination": """WP 5B governs maritime mobile, VDES, and radiodetermination (radar).
Key FAA concerns:
- VDES/AIS expansion in VHF maritime band (156-174 MHz) near aeronautical VHF (118-136 MHz)
- Shore-based radar systems near airborne weather radar (9000-9500 MHz) and SSR (1030/1090 MHz)
- Radiodetermination allocations near DME/TACAN (960-1215 MHz)
- Protection criteria: I/N ≤ −6 dB for ARNS radar systems per ITU-R M.1849""",

            "WP 7C — Radiolocation / Radar": """WP 7C governs radiolocation and radar systems.
Key FAA concerns:
- New radiolocation allocations overlapping or adjacent to: radio altimeter (4200-4400 MHz),
  airborne weather radar (9000-9500 MHz), ATC en-route radar (2700-2900 MHz), SSR (1030/1090 MHz)
- Aggregate interference from ground-based radars into airborne radar receivers
- Receiver blocking from high-power ground radars into sensitive airborne LNAs
- Radar pulse characteristics (PRF, peak power, duty cycle) must be accounted for — not just average power""",

            "WP 4C — MSS / DC-MSS-IMT (Satellite)": """WP 4C governs Mobile Satellite Service (MSS) and the new DC-MSS-IMT (direct satellite-to-device) systems.
This is WRC-27 AI 1.13. KEY FAA CONCERN: THREE CANDIDATE BANDS all adjacent to FAA safety systems:
  - 925–960 MHz (adjacent to DME/TACAN at 960 MHz — epfd ≤ −121.5 dBW/m²/MHz required)
  - 1475–1518 MHz (adjacent to L-band AMS(R)S at 1525–1559 MHz — ΔT/T ≤ 6% single-entry)
  - 2620–2690 MHz (adjacent to ASR at 2700–2900 MHz — I/N ≤ −10 dB)

CRITICAL CODE CHECKS FOR WP 4C:
1. PROPAGATION MODEL: Must use ITU-R P.619 (Earth-space), NOT P.452 (terrestrial only).
   If code uses P.452 or FSPL for a satellite downlink: FLAG AS FUNDAMENTAL ERROR — wrong model.
2. INTERFERENCE METRIC: Must use epfd (effective power flux-density) for constellation aggregate,
   NOT single-satellite pfd. If code only calculates pfd for one satellite: FLAG — understates by 10–30 dB.
   epfd formula: sum over all visible satellites of (EIRP_i / (4π R_i²)) × G_victim(θ_i) 
3. VICTIM GEOMETRY: Must analyze airborne aircraft at altitude, not just ground terminals.
   Aircraft at altitude see stronger downlinks (higher elevation angle, less atmosphere).
4. ΔT/T CHECK: For AMS(R)S band — T_noise_rise / T_system_noise ≤ 6% single-entry per protection table.
5. AGGREGATE: All simultaneously visible satellites must contribute — single satellite analysis is insufficient.

Common errors to flag:
- Using P.452 instead of P.619 (wrong model category)
- Single satellite pfd instead of constellation epfd
- Ground-only victim, ignoring airborne geometry
- Average satellite EIRP instead of worst-case nadir-pointing EIRP
- ΔT/T aggregate compliance without showing single-entry compliance
- Failure to analyze all three candidate bands independently""",

            "WP 7B — Space Research / Lunar SRS": """WP 7B governs space research service (SRS), including the novel WRC-27 AI 1.15 — lunar surface communications.
This is a NEW USE CASE with NO established ITU-R methodology for Earth-Moon SRS interference to terrestrial ARNS.

KEY CODE CHECK: If the code attempts to analyze Earth-Moon interference to terrestrial aeronautical systems,
ask: What ITU-R Recommendation is this methodology based on? If the answer is none — FLAG IMMEDIATELY.
The absence of an agreed methodology is a FAA policy argument: allocation before methodology = wrong order.

FAA BANDS AT RISK: 2700–2900 MHz (ASR), 3600–4200 MHz, 5350–5470 MHz (ARNS 5 GHz), 7190–7235 MHz, 8450–8500 MHz.

For Earth-based SRS uplinks (Earth→Moon):
- Uplink EIRP on Earth may be very high — pfd at co-frequency aeronautical receivers must be assessed
- Use FSPL at Earth-Moon distance (~384,000 km) for downlink path to confirm ground receiver impact

Common errors:
- Assuming Earth-Moon path loss makes interference negligible without calculating the actual pfd
- Not analyzing SRS uplink impact on co-frequency aeronautical ground receivers
- Proposing allocation without demonstrating a coordination methodology exists""",

            "WP 7C — EESS / Space Weather (Passive)": """WP 7C (EESS passive) — WRC-27 AI 1.17 and AI 1.19.
PASSIVE SENSORS DO NOT TRANSMIT — do not analyze for interference to FAA systems.
The FAA concern is ALLOCATION POLICY:
- AI 1.19: EESS passive co-primary in 4.2–4.4 GHz (Radio Altimeter band) weakens FAA exclusive ARNS status
- Strategic threat: weakens FAA position against AI 1.7 (IMT at 4.4–4.8 GHz)
- Check if code is being used to argue for co-primary allocation rather than secondary
- Check if any coordination obligations are being placed on aeronautical transmitters""",
        }

        active_wp_context = "\n\n".join([
            wp_context.get(wp, "") for wp in code_wp if wp in wp_context
        ])

        faa_bands_str = "\n".join([
            f"- {name}: {b['f_low_mhz']}–{b['f_high_mhz']} MHz | I/N threshold: {b['in_threshold_db']} dB | Noise floor: {b['noise_floor_dbm']} dBm"
            for name, b in FAA_BANDS.items()
        ])

        depth_map = {
            "Quick scan — flag critical assumption errors":
                "Provide a QUICK SCAN: identify the 3–5 most critical assumption errors or omissions that would invalidate the compatibility finding from an FAA perspective. Be direct and specific. No need to rewrite code.",
            "Standard — full line-by-line critique with corrections":
                "Provide a FULL LINE-BY-LINE CRITIQUE: annotate each significant code block, identify all assumption errors, propose corrected parameter values with regulatory citations, and summarize the overall validity of the analysis.",
            "Deep — full critique + corrected code + contribution rebuttal text":
                "Provide a DEEP ANALYSIS: full line-by-line critique, a corrected version of the code with FAA-appropriate parameters, AND draft rebuttal text suitable for a US contribution opposing or amending the findings.",
        }

        system_prompt = f"""You are a senior RF systems engineer and ITU-R spectrum policy expert supporting the FAA and NTIA.
You audit MATLAB and Python code embedded in WP 5B, 5D, 4C, 7B, and 7C contributions to identify
deviations from the ITU-R methodology required for that specific Working Party.

SCOPE RULE: Flag only code assumptions or parameter choices that violate the methodology
mandated by the applicable ITU-R Recommendations for this WP. Do not critique general
code quality, engineering design choices, or implementation decisions outside the ITU-R mandate.
Every finding must cite the specific Recommendation that the code is violating.

Applicable methodologies by WP:
- WP 5D/5B: M.1642 (IMT/ARNS), P.452 (terrestrial), P.528 (airborne victim), SM.2028 (Monte Carlo)
- WP 4C: P.619 (Earth-space propagation, REQUIRED), S.1586 (epfd), SM.2028 (aggregate)
- WP 7B: No established ITU-R methodology for lunar SRS — flag if code claims one
- WP 7C: Passive sensors — no interference methodology applies; flag allocation policy issues only
- All: M.1477 (6 dB safety margin), FAA protection levels (ARSR −6 dB, ASR −10 dB, DME epfd)

FAA PROTECTED BANDS:
{faa_bands_str}

WORKING PARTY CONTEXT:
{active_wp_context if active_wp_context else "Apply general FAA aeronautical spectrum protection context."}

FAA VICTIM SYSTEMS: {", ".join(faa_victim)}

ACCURACY RULE: Only flag what you can verify is wrong against a specific ITU-R Recommendation.
If you cannot confirm a specific deviation, state "Cannot confirm — requires manual verification."

Common flaws to check:
1. PROPAGATION MODEL: Is P.528 used for airborne victims? Is time % = 1% for protection studies?
   Is urban clutter applied near airports (unjustified)? Is gaseous attenuation claimed below 3 GHz?
2. RECEIVER PARAMETERS: Is the noise floor consistent with RTCA MOPS? Is bandwidth correct?
   Is NF realistic? Is the I/N threshold correct for the system type?
3. EIRP / POWER: Is maximum authorized EIRP used, or median/typical? Are OOBE accounted for?
4. AGGREGATION: If N interferers, are they summed linearly in watts? Is N realistic?
   Is deployment density consistent with actual licensed densities?
5. GEOMETRY: For airborne victims, is the slant distance used (not horizontal)?
   Is aircraft altitude conservative (low altitude = shorter slant path = more interference)?
6. GNSS SPECIFIC: Is c = a − b applied per M.1318? Is the narrowband rule (≤700 Hz) checked?
   Is power normalized to dB(W/Hz) correctly using the right noise bandwidth?
7. RADAR SPECIFIC: For pulsed systems, is peak power vs average power handled correctly?
   Is the pulse duty cycle accounted for in the power density calculation?
8. MONTE CARLO SPECIFIC: Is the spatial distribution uniform in area (not radius)?
   Is the violation probability criterion 5%? Is 99th percentile cited?

{depth_map[critique_depth]}

Format your response with clear section headers. Use technical notation precisely.
Where you cite a correction, give the specific regulatory reference (Recommendation, RR article).
Be direct — this is an adversarial technical review, not a collaborative editorial."""

        user_msg = f"""Please analyze the following {code_lang} code from ITU-R contribution {code_doc or '[unknown doc]'} 
submitted by {code_admin or '[unknown admin]'} to {', '.join(code_wp)}.

ANALYSIS TYPE: {code_scenario}
WRC AGENDA ITEM: {code_ai or 'Not specified'}
FAA VICTIM SYSTEMS: {', '.join(faa_victim)}

{"CONTRIBUTION CLAIM: " + code_summary if code_summary else ""}
{"CONTRIBUTION CONCLUSION: " + code_conclusion if code_conclusion else ""}
{"SPECIFIC FAA CONCERN: " + code_prior_concern if code_prior_concern else ""}

CODE TO ANALYZE:
```{code_lang.lower().split()[0]}
{code_input_full}
```

Please provide your critique structured as follows:

## 1. CODE SUMMARY
What this code actually does (independently of what the contribution claims).

## 2. CRITICAL ASSUMPTION AUDIT
For each significant assumption: parameter name → value used → FAA-appropriate value → impact on I/N result → regulatory citation for the correct value.

## 3. PROPAGATION MODEL ASSESSMENT  
Is the correct ITU-R model used? Is the time percentage appropriate? Are terrain/clutter corrections justified?

## 4. RECEIVER CHARACTERIZATION ERRORS
Noise floor, bandwidth, NF, I/N threshold — are these consistent with RTCA MOPS and ITU-R Recommendations?

## 5. AGGREGATION / MONTE CARLO ERRORS (if applicable)
Spatial distribution, N value, linear summation, violation criterion.

## 6. GNSS-SPECIFIC ISSUES (if GNSS bands involved)
M.1318 c=a−b framework compliance, narrowband rule, power density normalization.

## 7. RADAR/PULSED SYSTEM ISSUES (if applicable)
Peak vs average power, duty cycle, pulse characteristics.

## 8. OVERALL VALIDITY ASSESSMENT
Is the compatibility conclusion defensible? What is the likely true I/N when correct parameters are applied?

## 9. RECOMMENDED US RESPONSE
What should the US delegation say in the Working Party session? What contribution text should be prepared?

{"## 10. CORRECTED CODE" if "Deep" in critique_depth else ""}
{"Provide corrected code with FAA-appropriate parameters annotated with regulatory justifications." if "Deep" in critique_depth else ""}"""

        with st.spinner(f"Running {critique_depth.split('—')[0].strip()} code analysis... 20–45 seconds"):
            try:
                client = anthropic.Anthropic(api_key=api_key)
                response = client.messages.create(
                    model="claude-opus-4-5",
                    max_tokens=6000,
                    messages=[{"role": "user", "content": user_msg}],
                    system=system_prompt,
                )
                critique_text = response.content[0].text

                st.success("✅ Code analysis complete")
                st.markdown("---")

                # Severity banner
                critical_keywords = ["CRITICAL", "violated", "invalid", "incorrect", "wrong",
                                     "understates", "overstates", "unjustified", "missing"]
                n_critical = sum(critique_text.upper().count(kw.upper()) for kw in critical_keywords)
                if n_critical > 15:
                    warn(f"HIGH CONCERN — analysis identified multiple significant methodology errors. Recommend preparing a US rebuttal contribution.")
                elif n_critical > 5:
                    st.warning("⚠️ MODERATE CONCERN — some assumption errors identified. Review carefully before accepting findings.")
                else:
                    ok("LOW CONCERN — code appears methodologically reasonable under initial review. Verify receiver parameters independently.")

                st.markdown("---")
                st.subheader("🔬 Technical Critique")
                st.markdown(critique_text)

                # Export
                st.markdown("---")
                export = f"""FAA RF INTERFERENCE TOOL — CODE ANALYSIS REPORT
Document: {code_doc or 'N/A'} | Admin: {code_admin or 'N/A'}
Working Party: {', '.join(code_wp)}
Agenda Item: {code_ai or 'N/A'}
Language: {code_lang} | Analysis Type: {code_scenario}
FAA Victim Systems: {', '.join(faa_victim)}
Critique Depth: {critique_depth}

{'='*60}
ORIGINAL CODE
{'='*60}
{code_input_full}

{'='*60}
FAA TECHNICAL CRITIQUE
{'='*60}
{critique_text}
"""
                col_dl1, col_dl2 = st.columns(2)
                with col_dl1:
                    st.download_button(
                        "⬇️ Download Full Critique (.txt)",
                        export,
                        file_name=f"code_critique_{(code_doc or 'doc').replace('/','_')}.txt",
                        mime="text/plain"
                    )
                with col_dl2:
                    # Prepare a quick intervention draft
                    intervention = f"""US DELEGATION INTERVENTION — TECHNICAL FLOOR STATEMENT
Document: {code_doc or '[doc number]'} | {', '.join(code_wp)} | AI {code_ai or '[AI]'}

The United States has reviewed the technical analysis in document {code_doc or '[doc number]'} 
and has identified the following technical concerns with the methodology:

[Insert top 3 findings from critique above]

The United States requests that the Working Party:
1. Require reanalysis using [corrected propagation model / receiver parameters / deployment assumptions]
2. Apply the protection criteria established in [M.1642 / M.1318 / M.1477 / SM.2028] 
3. Defer any regulatory conclusion pending submission of a corrected analysis

The United States will submit a formal contribution addressing these technical issues 
before the next meeting.

[US Delegation — {', '.join(code_wp)} | FAA/NTIA]
"""
                    st.download_button(
                        "⬇️ Download Floor Intervention Draft (.txt)",
                        intervention,
                        file_name=f"floor_intervention_{(code_doc or 'doc').replace('/','_')}.txt",
                        mime="text/plain"
                    )

            except anthropic.AuthenticationError:
                st.error("❌ Invalid API key.")
            except anthropic.RateLimitError:
                st.error("❌ Rate limit — wait a moment and retry.")
            except Exception as e:
                st.error(f"❌ Error: {str(e)}")

    elif not code_input.strip():
        st.info("👆 Paste code above and click Analyze.")

    # ── Example code snippets ─────────────────────────────────────────────────
    st.markdown("---")
    with st.expander("📋 Example code snippets to test with"):
        st.markdown("**Example 1 — WP 5D: IMT sharing study with several intentional flaws**")
        st.code("""% WP 5D Sharing Study — IMT at 4800 MHz vs Radio Altimeter
% Assumes: terrestrial P.452 urban, 50% time, low N, good NF

f_MHz = 4800;
Pt_dBm = 43;
Gt_dBi = 15;
cable_dB = 2;
EIRP = Pt_dBm + Gt_dBi - cable_dB;   % = 56 dBm

N_interferers = 3;
deploy_radius_km = 10;

% P.452 path loss with urban clutter
distances = sqrt(rand(N_interferers,1)) * deploy_radius_km;
FSPL = 20*log10(distances) + 20*log10(f_MHz) + 32.44;
urban_clutter_dB = 18;                 % urban correction
PL = FSPL + urban_clutter_dB;

% Received power per interferer
Pr_each = EIRP - PL;                   % dBm

% Aggregate (sum in dB - INCORRECT!)
Pr_agg = 10*log10(sum(10.^(Pr_each/10)));

% Victim receiver (radio altimeter)
BW_MHz = 500;                          % MHz - overstated
NF_dB = 10;                            % dB - pessimistic for victim
N_floor = -174 + 10*log10(BW_MHz*1e6) + NF_dB;

IN_dB = Pr_agg - N_floor;
threshold_dB = -6;
fprintf('I/N = %.1f dB, threshold = %d dB\\n', IN_dB, threshold_dB);
fprintf('Compatible: %d\\n', IN_dB < threshold_dB);
""", language="matlab")

        st.markdown("**Example 2 — WP 7C: Radiolocation study near SSR with pulsed power error**")
        st.code("""# WP 7C radiolocation study - interference to SSR at 1090 MHz
import numpy as np

f_MHz = 1090
# Pulsed radar parameters
peak_power_dBm = 70      # 10 kW peak
pulse_width_us = 1.0     # microseconds
prf_hz = 1000            # pulse repetition frequency

# INCORRECT: using average power without accounting for duty cycle properly
duty_cycle = pulse_width_us * 1e-6 * prf_hz
avg_power_dBm = peak_power_dBm + 10*np.log10(duty_cycle)

Gt_dBi = 20
EIRP_avg = avg_power_dBm + Gt_dBi    # Using average - wrong for pulse desensitization

dist_km = 10
FSPL = 20*np.log10(dist_km) + 20*np.log10(f_MHz) + 32.44

Pr = EIRP_avg - FSPL    # dBm at SSR receiver

# SSR noise floor - using wrong bandwidth
BW_MHz = 10             # SSR reply bandwidth ~1-2 MHz, not 10
NF_dB = 8
N_floor = -174 + 10*np.log10(BW_MHz*1e6) + NF_dB

IN_dB = Pr - N_floor
print(f"I/N = {IN_dB:.1f} dB")
print(f"Compatible: {IN_dB < -10}")
""", language="python")

        st.markdown("**Example 3 — WP 5D: GNSS L5 study missing M.1318 framework**")
        st.code("""% GPS L5 interference study - missing c = a-b framework
% WP 5D document - IMT at 1200 MHz vs GPS L5

f_MHz = 1200;
EIRP_dBm = 46;
dist_km = 5;

FSPL = 20*log10(dist_km) + 20*log10(f_MHz) + 32.44;
Pr_dBm = EIRP_dBm - FSPL;

% Using generic noise floor instead of M.1318 framework
BW_MHz = 20;
NF_dB = 2;
N_floor = -174 + 10*log10(BW_MHz*1e6) + NF_dB;

% Missing: conversion to dB(W/Hz), comparison to M.1318 'a' parameter
% Missing: application of 6 dB aeronautical safety margin (M.1477)
% Missing: narrowband rule check (M.1477 Annex 5 - 700 Hz threshold)
% Using simple I/N instead of c = a - b methodology

IN_dB = Pr_dBm - N_floor;
fprintf('I/N = %.1f dB vs threshold -10 dB\\n', IN_dB);
fprintf('Result: %s\\n', IN_dB < -10 ? 'Compatible' : 'Incompatible');
""", language="matlab")

# ─────────────────────────────────────────────────────────────────────────────
# TAB 12 — GLOSSARY
# ─────────────────────────────────────────────────────────────────────────────
elif selected_tab == "📖 Glossary":
    st.title("📖 Glossary — Terms, Acronyms & Definitions")
    st.markdown("*Every acronym and technical term used in this tool and in ITU-R aeronautical spectrum work, in plain language.*")
    ex("Use Ctrl+F (or Cmd+F on Mac) to search this page, or use the category filter below to jump to a section.")

    # Search box
    search = st.text_input("🔍 Search glossary:", placeholder="Type any term or acronym...").strip().lower()

    # Category filter
    categories = [
        "All",
        "🔤 Acronyms — Systems & Organizations",
        "📡 RF & Signal Theory",
        "🌐 ITU-R & Regulatory",
        "✈️ Aeronautical Systems",
        "📐 Measurements & Units",
        "💻 Analysis Methods",
    ]
    cat_filter = st.selectbox("Filter by category:", categories)

    st.markdown("---")

    # ── Glossary data ─────────────────────────────────────────────────────────
    GLOSSARY = [

        # ── ACRONYMS — Systems & Organizations ───────────────────────────────
        ("ADS-B", "Automatic Dependent Surveillance — Broadcast",
         "A surveillance technology where aircraft determine their own position using GPS and periodically broadcast it. Air Traffic Control and other aircraft receive these broadcasts to track traffic. Operates at 1090 MHz. Mandated in US airspace since 2020.",
         "🔤 Acronyms — Systems & Organizations"),

        ("AGL", "Above Ground Level",
         "Height measured from the ground directly below, not from sea level. Antenna heights in interference analyses are typically specified in meters Above Ground Level (AGL).",
         "🔤 Acronyms — Systems & Organizations"),

        ("ARNS", "Aeronautical Radionavigation Service",
         "An ITU Radio Regulations allocation category covering radio systems used for aircraft navigation. ARNS is a primary allocation in most aeronautical frequency bands and carries safety-of-life status under Radio Regulations No. 4.10.",
         "🔤 Acronyms — Systems & Organizations"),

        ("ANS", "Aeronautical Navigation Service",
         "An ITU Radio Regulations service category covering navigation aids for aviation, including both ground-based and satellite-based systems.",
         "🔤 Acronyms — Systems & Organizations"),

        ("ATC", "Air Traffic Control",
         "The ground-based service that directs aircraft in controlled airspace. ATC uses radar (SSR, PSR), voice radio (VHF), and data links (ADS-B, CPDLC) to manage aircraft separation and flow.",
         "🔤 Acronyms — Systems & Organizations"),

        ("CAT I / II / III", "Instrument Landing System Categories I, II, III",
         "Categories of precision instrument approach defined by visibility and decision height. CAT I: visibility ≥ 550m. CAT II: visibility ≥ 300m. CAT III: visibility < 300m down to zero. Higher categories require more precise ILS signals and more stringent interference protection.",
         "🔤 Acronyms — Systems & Organizations"),

        ("CEPT", "European Conference of Postal and Telecommunications Administrations",
         "The European regional body that coordinates spectrum policy and produces shared European positions for ITU-R meetings. CEPT produces SEAMCAT (Spectrum Engineering Advanced Monte Carlo Analysis Tool) used in interference studies.",
         "🔤 Acronyms — Systems & Organizations"),

        ("CPM", "Conference Preparatory Meeting",
         "An ITU-R meeting held approximately 12 months before each World Radiocommunication Conference (WRC) that produces the CPM Report — a summary of all technical study results and proposed methods for each WRC agenda item.",
         "🔤 Acronyms — Systems & Organizations"),

        ("DME", "Distance Measuring Equipment",
         "A radio navigation system that measures the slant-range distance between an aircraft and a ground transponder. Aircraft interrogate the ground station at 1025–1150 MHz; the station responds at 962–1213 MHz. Used alongside VOR for position fixing and with ILS for approach guidance.",
         "🔤 Acronyms — Systems & Organizations"),

        ("EASA", "European Union Aviation Safety Agency",
         "The European counterpart to the FAA. EASA certifies aircraft and avionics in Europe and participates in ICAO and EUROCAE standards work.",
         "🔤 Acronyms — Systems & Organizations"),

        ("EGPWS", "Enhanced Ground Proximity Warning System",
         "An avionics system that alerts pilots when an aircraft is in dangerous proximity to terrain. Uses GPS position and a terrain database. Relies on accurate radio altimeter readings for low-altitude warnings.",
         "🔤 Acronyms — Systems & Organizations"),

        ("EPFD", "Equivalent Power Flux Density",
         "A measure of aggregate interference from non-geostationary satellite constellations at a point on the Earth's surface. EPFD limits are specified in ITU Radio Regulations Appendix 7 to protect terrestrial and geostationary services.",
         "🔤 Acronyms — Systems & Organizations"),

        ("ESA", "European Space Agency",
         "The intergovernmental organization responsible for Europe's space program, including the Galileo Global Navigation Satellite System (GNSS). ESA participates in ITU-R coordination meetings affecting GNSS spectrum.",
         "🔤 Acronyms — Systems & Organizations"),

        ("FAA", "Federal Aviation Administration",
         "The United States government agency responsible for regulating civil aviation, managing US airspace, and certifying aircraft and avionics. The FAA is the primary stakeholder defending aeronautical spectrum in US ITU-R proceedings through NTIA.",
         "🔤 Acronyms — Systems & Organizations"),

        ("FMG", "Frequency Management Group",
         "An ICAO working group that manages the frequency assignments for aeronautical services globally and prepares ICAO positions for ITU-R World Radiocommunication Conferences.",
         "🔤 Acronyms — Systems & Organizations"),

        ("FSMP", "Frequency Spectrum Management Panel",
         "An ICAO body responsible for managing the frequency spectrum used by international civil aviation and coordinating with ITU-R on spectrum issues affecting aviation.",
         "🔤 Acronyms — Systems & Organizations"),

        ("Galileo", "Galileo Global Navigation Satellite System",
         "The European Union's Global Navigation Satellite System (GNSS), operated by the European Space Agency (ESA) and European Union Agency for the Space Programme (EUSPA). Provides civil and safety-of-life positioning services on L1 (1575.42 MHz) and E5 (1164–1215 MHz) bands.",
         "🔤 Acronyms — Systems & Organizations"),

        ("GLONASS", "Global Navigation Satellite System (Russian)",
         "Russia's Global Navigation Satellite System, operated by the Russian Aerospace Defence Forces. Provides positioning on L1 (1602 MHz region) and L2 (1246 MHz region). Protected under ITU-R M.1904.",
         "🔤 Acronyms — Systems & Organizations"),

        ("GMDSS", "Global Maritime Distress and Safety System",
         "An internationally agreed-upon set of safety procedures, types of equipment, and communication protocols used to increase safety and make it easier to rescue distressed ships, boats, and aircraft. Governed by ITU-R Working Party 5B (WP 5B).",
         "🔤 Acronyms — Systems & Organizations"),

        ("GNSS", "Global Navigation Satellite System",
         "The generic term for satellite-based positioning systems including GPS (United States), GLONASS (Russia), Galileo (European Union), and BeiDou (China). GNSS signals arrive at Earth at approximately −130 dBm — roughly 20 dB above the receiver noise floor — making them extremely vulnerable to interference.",
         "🔤 Acronyms — Systems & Organizations"),

        ("GPS", "Global Positioning System",
         "The United States satellite-based navigation system operated by the US Space Force. Provides positioning, navigation, and timing (PNT) services on L1 (1575.42 MHz), L2 (1227.60 MHz), and L5 (1176.45 MHz). GPS L1 and L5 are used for aviation safety-of-life applications.",
         "🔤 Acronyms — Systems & Organizations"),

        ("GPWS", "Ground Proximity Warning System",
         "An avionics system that warns pilots when an aircraft is in immediate danger of controlled flight into terrain (CFIT). The enhanced version is called EGPWS (Enhanced Ground Proximity Warning System) and uses GPS and terrain databases.",
         "🔤 Acronyms — Systems & Organizations"),

        ("IALA", "International Association of Marine Aids to Navigation and Lighthouse Authorities",
         "The international organization that coordinates maritime navigation aids worldwide. IALA participates in ITU-R Working Party 5B (WP 5B) on maritime spectrum issues.",
         "🔤 Acronyms — Systems & Organizations"),

        ("IATA", "International Air Transport Association",
         "The trade association representing airlines worldwide. IATA participates in ITU-R as a Sector Member to protect airline operational spectrum interests.",
         "🔤 Acronyms — Systems & Organizations"),

        ("ICAO", "International Civil Aviation Organization",
         "A United Nations specialized agency that codifies the principles and techniques of international air navigation and promotes the planning of international air transport. ICAO has Sector Member status at ITU-R and submits influential contributions defending aeronautical spectrum allocations.",
         "🔤 Acronyms — Systems & Organizations"),

        ("ILS", "Instrument Landing System",
         "A precision radio navigation system that provides aircraft with horizontal (Localizer, 108–117.975 MHz) and vertical (Glide Slope, 328.6–335.4 MHz) guidance for landing in low visibility. ILS enables Category I, II, and III approaches. It is a safety-of-life system protected under Radio Regulations No. 4.10.",
         "🔤 Acronyms — Systems & Organizations"),

        ("IMO", "International Maritime Organization",
         "A United Nations agency responsible for regulating shipping. IMO coordinates with ITU-R Working Party 5B (WP 5B) on maritime radio spectrum, including GMDSS and VDES.",
         "🔤 Acronyms — Systems & Organizations"),

        ("IMT", "International Mobile Telecommunications",
         "The ITU-R family of standards for mobile broadband, including 4G (Long Term Evolution / LTE) and 5G (New Radio / NR). IMT spectrum studies are conducted in ITU-R Working Party 5D (WP 5D). New IMT allocations near aeronautical bands are a primary FAA concern.",
         "🔤 Acronyms — Systems & Organizations"),

        ("JAXA", "Japan Aerospace Exploration Agency",
         "Japan's national aerospace and space agency. JAXA operates the Multi-functional Satellite Augmentation System (MSAS), Japan's Satellite-Based Augmentation System (SBAS) for GPS. Participates in ITU-R coordination affecting GNSS spectrum.",
         "🔤 Acronyms — Systems & Organizations"),

        ("LNA", "Low Noise Amplifier",
         "The first active component in a receiver chain, designed to amplify weak signals while adding minimal noise. The LNA sets the receiver's noise figure. A strong out-of-band interferer can drive the LNA into compression (saturation), causing it to generate intermodulation products and lose sensitivity — this is called receiver desensitization or blocking.",
         "🔤 Acronyms — Systems & Organizations"),

        ("MOPS", "Minimum Operational Performance Standards",
         "RTCA documents that specify the minimum performance requirements for avionics systems, including receiver sensitivity, selectivity, and interference immunity. MOPS are the primary source for protection criteria cited in ITU-R contributions. Examples: DO-235B (GNSS), DO-260B (ADS-B), DO-155 (Radio Altimeter).",
         "🔤 Acronyms — Systems & Organizations"),

        ("MSAS", "Multi-functional Satellite Augmentation System",
         "Japan's Satellite-Based Augmentation System (SBAS) for GPS, operated by JAXA. Transmits GPS correction signals on L1 (1575.42 MHz) from geostationary satellites to improve positioning accuracy for aviation.",
         "🔤 Acronyms — Systems & Organizations"),

        ("NASA", "National Aeronautics and Space Administration",
         "The United States government agency responsible for civilian space program and aeronautics research. NASA operates GNSS reference stations and participates in ITU-R coordination affecting GNSS spectrum.",
         "🔤 Acronyms — Systems & Organizations"),

        ("NTIA", "National Telecommunications and Information Administration",
         "The US government agency, part of the Department of Commerce, responsible for managing the federal government's use of radio spectrum. NTIA coordinates all US government positions for ITU-R meetings and submits official US contributions. FAA submits contributions directly on technical aviation matters and coordinates its spectrum positions through NTIA for policy alignment.",
         "🔤 Acronyms — Systems & Organizations"),

        ("OOBE", "Out-of-Band Emissions",
         "Radio frequency energy emitted by a transmitter outside its assigned channel or band. OOBE can cause in-band interference to a victim receiver even when the interferer is nominally operating in a different frequency band. OOBE limits are specified in ITU-R Radio Regulations and system-specific Recommendations.",
         "🔤 Acronyms — Systems & Organizations"),

        ("PSR", "Primary Surveillance Radar",
         "A radar system that detects aircraft by reflecting radio waves off the aircraft skin — no cooperation from the aircraft is required. PSR provides range and bearing but no altitude. Used by Air Traffic Control alongside Secondary Surveillance Radar (SSR).",
         "🔤 Acronyms — Systems & Organizations"),

        ("RNSS", "Radionavigation Satellite Service",
         "An ITU Radio Regulations allocation category covering satellite-based navigation systems (GPS, GLONASS, Galileo, BeiDou). RNSS allocations in the L-band are protected under Resolution 233 and the GNSS protection methodology in ITU-R M.1318.",
         "🔤 Acronyms — Systems & Organizations"),

        ("RTCA", "Radio Technical Commission for Aeronautics",
         "A US standards organization that develops aviation standards, including Minimum Operational Performance Standards (MOPS) for avionics. RTCA documents (called DO-xxx) define the receiver sensitivity, selectivity, and interference immunity requirements cited in ITU-R contributions to defend aeronautical frequencies.",
         "🔤 Acronyms — Systems & Organizations"),

        ("SBAS", "Satellite-Based Augmentation System",
         "A system that improves GPS accuracy and integrity for aviation by broadcasting correction signals from geostationary satellites. Examples include WAAS (Wide Area Augmentation System, United States), EGNOS (Europe), MSAS (Japan), and GAGAN (India). SBAS signals occupy the GPS L1 band (1575.42 MHz).",
         "🔤 Acronyms — Systems & Organizations"),

        ("SSR", "Secondary Surveillance Radar",
         "An Air Traffic Control radar system that interrogates aircraft transponders (Mode A, C, S) to obtain identification and altitude. SSR operates at 1030 MHz (interrogation) and 1090 MHz (reply). ADS-B (Automatic Dependent Surveillance — Broadcast) uses the 1090 MHz SSR reply frequency.",
         "🔤 Acronyms — Systems & Organizations"),

        ("TACAN", "Tactical Air Navigation",
         "A military radio navigation system providing both bearing and distance information, operating in the 960–1215 MHz band. Civilian DME (Distance Measuring Equipment) and TACAN are co-located at VORTAC stations, providing both VOR bearing and distance to civilian and military aircraft.",
         "🔤 Acronyms — Systems & Organizations"),

        ("TAWS", "Terrain Awareness and Warning System",
         "An avionics system that provides pilots with alerts about terrain, obstacles, and excessive descent rates. TAWS uses GPS position, a terrain database, and radio altimeter data to generate alerts. Radio altimeter accuracy is critical to TAWS performance at low altitude.",
         "🔤 Acronyms — Systems & Organizations"),

        ("TCAS", "Traffic Collision Avoidance System",
         "An airborne system that monitors the airspace around an aircraft for other transponder-equipped aircraft and provides Resolution Advisories (RA) to avoid collisions. TCAS operates at 1030 MHz (interrogation) and 1090 MHz (reply), sharing the SSR/ADS-B frequencies.",
         "🔤 Acronyms — Systems & Organizations"),

        ("USTR", "United States Trade Representative",
         "The US government office that coordinates international trade and telecommunications negotiations. The USTR, together with the State Department and NTIA, represents the US at World Radiocommunication Conferences (WRC).",
         "🔤 Acronyms — Systems & Organizations"),

        ("VDES", "VHF Data Exchange System",
         "A maritime communication system operating in the Very High Frequency (VHF) maritime band (156–174 MHz), intended to succeed the Automatic Identification System (AIS). VDES adds data exchange capabilities for e-navigation. Governed by ITU-R Working Party 5B (WP 5B).",
         "🔤 Acronyms — Systems & Organizations"),

        ("VHF", "Very High Frequency",
         "The radio frequency range from 30 to 300 MHz, with wavelengths of 1 to 10 meters. Aeronautical VHF communications (118–136 MHz) and VOR/ILS navigation (108–117.975 MHz) operate in this band. VHF propagation is primarily line-of-sight.",
         "🔤 Acronyms — Systems & Organizations"),

        ("VOR", "VHF Omnidirectional Range",
         "A Very High Frequency (VHF) radio navigation beacon operating at 108–117.975 MHz that transmits bearing information in all directions simultaneously. Aircraft use VOR to determine their magnetic bearing (called a radial) from the station. VOR stations form the backbone of the US and international airway system. Co-located with DME at VORTAC stations.",
         "🔤 Acronyms — Systems & Organizations"),

        ("VORTAC", "VOR and TACAN Co-located Station",
         "A navigation facility combining a VHF Omnidirectional Range (VOR) beacon with a Tactical Air Navigation (TACAN) transponder. Provides civilian aircraft with VOR bearing and DME distance, and military aircraft with TACAN bearing and distance.",
         "🔤 Acronyms — Systems & Organizations"),

        ("WAAS", "Wide Area Augmentation System",
         "The United States Satellite-Based Augmentation System (SBAS), operated by the FAA. WAAS broadcasts GPS correction and integrity signals from geostationary satellites on L1 (1575.42 MHz), enabling GPS approaches to as low as 200 feet Decision Height (Category I equivalent). WAAS reference stations are a specific FAA concern in L-band spectrum discussions.",
         "🔤 Acronyms — Systems & Organizations"),

        ("WP 5D", "Working Party 5D — IMT Systems",
         "The ITU-R Working Party responsible for International Mobile Telecommunications (IMT), including 4G (LTE), 5G (NR), and IMT-2030 (6G). WP 5D conducts sharing studies for new IMT spectrum that may threaten aeronautical bands including the Radio Altimeter band (4200–4400 MHz), ARNS at 5 GHz, and GNSS L-band.",
         "🔤 Acronyms — Systems & Organizations"),

        ("WP 5B", "Working Party 5B — Maritime Mobile and Radiodetermination",
         "The ITU-R Working Party responsible for maritime mobile communications (GMDSS, VDES, AIS) and radiodetermination (radar) services. WP 5B proposals near aeronautical VHF and radar bands require FAA monitoring.",
         "🔤 Acronyms — Systems & Organizations"),

        ("WP 7C", "Working Party 7C — Radiolocation",
         "The ITU-R Working Party responsible for radiolocation services including ground-based radar systems. WP 7C proposals can affect the Radio Altimeter band (4200–4400 MHz), Airborne Weather Radar band (9000–9500 MHz), and ATC radar bands (2700–2900 MHz).",
         "🔤 Acronyms — Systems & Organizations"),

        ("WP 7D", "Working Party 7D — Radio Astronomy",
         "The ITU-R Working Party responsible for radio astronomy and passive sensing services. WP 7D protection zones and emission limit proposals adjacent to aeronautical bands can indirectly constrain FAA system parameters.",
         "🔤 Acronyms — Systems & Organizations"),

        ("WRC", "World Radiocommunication Conference",
         "An international treaty conference held every 3–4 years under the International Telecommunication Union (ITU) that reviews and revises the Radio Regulations — the international treaty governing use of the radio frequency spectrum. WRC decisions are legally binding on all 193 ITU member states. WRC-27 is the next conference, scheduled for 2027.",
         "🔤 Acronyms — Systems & Organizations"),

        ("WRC-27", "World Radiocommunication Conference 2027",
         "The next WRC scheduled for 2027. Five agenda items directly threaten FAA aeronautical bands: AI 1.7 (IMT near Radio Altimeter band 4.2–4.4 GHz), AI 1.13 (MSS near DME/ASR bands), AI 1.15 (lunar SRS near ASR and radar bands), AI 1.17 (EESS passive sensors near HF/VHF), and AI 1.19 (EESS passive in RA band). AI 1.7 is the highest priority — IMT identification at 4.4–4.8 GHz is only 200 MHz above the Radio Altimeter band and repeats the pattern of the 5G C-band / RA controversy from 2019–2022.",
         "🌐 ITU-R & Regulatory"),

        ("WRC-27 AI 1.7", "WRC-27 Agenda Item 1.7 — IMT in 4.4–4.8 GHz, 7.125–8.4 GHz, 14.8–15.35 GHz",
         "Proposes IMT (5G/6G) identification in three frequency ranges, all of which threaten FAA systems. The 4.4–4.8 GHz range is the highest priority — it sits immediately above the Radio Altimeter band (4.2–4.4 GHz), and given 5G OOB emission characteristics and RA receiver blocking susceptibility, this directly parallels the C-band / Radio Altimeter controversy that caused FAA to restrict 5G deployment near airports in 2022. The 7.125–8.4 GHz range overlaps FAA microwave backbone links. FAA systems at risk: RA and WAICS (4.2–4.4 GHz), fixed microwave links (7.125–8.4 GHz and 14.8–15.35 GHz). US position: oppose without OOB compliance per SM.1541 and coordination zones. 6 dB aviation safety factor applies.",
         "🌐 ITU-R & Regulatory"),

        ("WRC-27 AI 1.13", "WRC-27 Agenda Item 1.13 — MSS/IMT Connectivity (DC-MSS-IMT) 694–2700 MHz",
         "Proposes new Mobile Satellite Service (MSS) allocations for direct satellite-to-device (s-E) connectivity linking space stations with IMT users. Three candidate frequency bands: 925–960 MHz (immediately adjacent to DME/TACAN at 960 MHz), 1475–1518 MHz (adjacent to L-band AMS(R)S at 1525–1559 MHz), and 2620–2690 MHz (adjacent to ASR at 2700–2900 MHz). All three threaten critical FAA safety systems via OOB and spurious emissions. Aggregate interference from multiple satellite downlinks must be assessed per SM.2028. US position: require aggregate analysis per SM.2028 for all candidate bands before any allocation.",
         "🌐 ITU-R & Regulatory"),

        ("WRC-27 AI 1.15", "WRC-27 Agenda Item 1.15 — Space Research Service for Lunar Surface Communications",
         "Proposes Space Research Service (SRS) space-to-space allocations to support lunar surface communications — a novel use case with no established interference methodology for effects on terrestrial aeronautical systems. FAA bands potentially affected: 2700–2900 MHz (ASR), 3600–4200 MHz, 5350–5470 MHz (ARNS), 7190–7235 MHz, and 8450–8500 MHz. Threat level medium — primary concern is setting precedents for high-power SRS links without adequate coordination requirements for terrestrial ARNS.",
         "🌐 ITU-R & Regulatory"),

        ("WRC-27 AI 1.19", "WRC-27 Agenda Item 1.19 — EESS (passive) in 4.2–4.4 GHz and 8.4–8.5 GHz",
         "Proposes Earth Exploration Satellite Service (EESS) passive allocation co-primary in the 4.2–4.4 GHz Radio Altimeter band. Passive sensors don't transmit, so the direct interference risk is low. However, the strategic threat is significant: adding a co-primary EESS allocation in the RA band dilutes the exclusive ARNS status that is FAA's strongest argument against AI 1.7 (IMT identification at 4.4–4.8 GHz). US position: oppose — the RA band must remain exclusively ARNS to maintain the strongest regulatory protection against IMT encroachment.",
         "🌐 ITU-R & Regulatory"),

        # ── RF & Signal Theory ────────────────────────────────────────────────
        ("Blocking / Desensitization", "Receiver Front-End Overload",
         "A condition where a strong out-of-band signal drives the receiver's Low Noise Amplifier (LNA) into compression (saturation), raising the effective noise floor and reducing sensitivity to the desired signal. The interferer does not need to be in-band to cause this — it only needs to be strong enough at the antenna input. This was the core mechanism in the 5G / Radio Altimeter interference problem of 2019–2022.",
         "📡 RF & Signal Theory"),

        ("C/I", "Carrier-to-Interference Ratio",
         "The ratio of the desired signal power to the interference power, expressed in decibels (dB). C/I = C(dBm) − I(dBm). Used in communications system design to assess whether the desired signal can be detected in the presence of interference. A higher C/I indicates less interference impact.",
         "📡 RF & Signal Theory"),

        ("C/N", "Carrier-to-Noise Ratio",
         "The ratio of the desired signal power to the thermal noise floor power, expressed in decibels (dB). C/N = C(dBm) − N(dBm). Determines whether a receiver can successfully decode a signal. GPS L1 operates at C/N of approximately 40–45 dB-Hz under normal conditions.",
         "📡 RF & Signal Theory"),

        ("dB", "Decibel",
         "A logarithmic unit expressing the ratio of two power levels: dB = 10 × log10(P2/P1). Because RF signal levels span many orders of magnitude (from kilowatts to femtowatts), the decibel scale makes calculations manageable. In decibels, multiplication of powers becomes addition, and division becomes subtraction.",
         "📡 RF & Signal Theory"),

        ("dBm", "Decibels Relative to One Milliwatt",
         "An absolute power level expressed in decibels referenced to 1 milliwatt. Formula: P(dBm) = 10 × log10(P(mW)). Common reference points: 0 dBm = 1 mW; 30 dBm = 1 W; 43 dBm = 20 W; −130 dBm = GPS signal at Earth's surface.",
         "📡 RF & Signal Theory"),

        ("dBi", "Decibels Relative to an Isotropic Antenna",
         "A unit of antenna gain relative to a theoretical isotropic antenna that radiates equally in all directions. 0 dBi = isotropic; 2 dBi = simple dipole; 15 dBi = typical cellular sector antenna; 30+ dBi = high-gain dish or phased array.",
         "📡 RF & Signal Theory"),

        ("dBW", "Decibels Relative to One Watt",
         "An absolute power level in decibels referenced to 1 watt. Relationship: 0 dBW = 30 dBm; 27 dBW = 57 dBm = 500 W (typical GPS satellite transmit power).",
         "📡 RF & Signal Theory"),

        ("EIRP", "Effective Isotropic Radiated Power",
         "The product of transmitter output power and antenna gain in a given direction, less cable losses. EIRP = Pt(dBm) + Gt(dBi) − Lcable(dB). EIRP represents the power that would need to be fed to an isotropic antenna to produce the same signal strength in the direction of interest. It is the standard metric for interference analysis and is used in ITU Radio Regulations power flux density limits.",
         "📡 RF & Signal Theory"),

        ("FSPL", "Free Space Path Loss",
         "The signal power loss between a transmitting and receiving antenna in free space (vacuum, no obstructions). Formula: FSPL(dB) = 20·log10(d_km) + 20·log10(f_MHz) + 32.44. FSPL is the most optimistic (least loss) propagation model and represents the worst-case interference scenario — if compatibility holds under FSPL, any realistic deployment will also be compatible.",
         "📡 RF & Signal Theory"),

        ("I/N", "Interference-to-Noise Ratio",
         "The ratio of interference power to thermal noise floor power at a receiver input, expressed in decibels (dB). I/N = I(dBm) − N(dBm). The ITU-R standard compatibility metric. At I/N = −6 dB, the noise floor rises by 0.97 dB. At I/N = −10 dB (GNSS standard), the rise is 0.41 dB. Exceeding the I/N threshold is the technical basis for citing harmful interference under Radio Regulations No. 4.10.",
         "📡 RF & Signal Theory"),

        ("Noise Figure (NF)", "Receiver Noise Figure",
         "A measure of how much noise a receiver adds to the signal beyond the fundamental thermal noise limit, expressed in decibels (dB). A 0 dB noise figure is theoretically perfect. Real avionics receivers typically have noise figures of 2–8 dB. Every 1 dB increase in noise figure raises the noise floor by 1 dB and reduces sensitivity by 1 dB.",
         "📡 RF & Signal Theory"),

        ("Noise Floor", "Receiver Thermal Noise Floor",
         "The minimum signal power a receiver can detect, set by thermal noise. Formula: N(dBm) = −174 + 10·log10(B_Hz) + NF(dB), where −174 dBm/Hz is thermal noise at 290 Kelvin, B is bandwidth in Hertz, and NF is the receiver noise figure in dB. Interference must remain below the I/N threshold relative to this level.",
         "📡 RF & Signal Theory"),

        ("PFD", "Power Flux Density",
         "The radio frequency power arriving per unit area at a given location, expressed in dB(W/m²) or dB(W/m²/MHz). PFD limits are specified in the ITU Radio Regulations to protect Earth-based receivers from satellite transmissions. RR Appendix 7 specifies PFD limits for non-geostationary satellite systems.",
         "📡 RF & Signal Theory"),

        ("PRF", "Pulse Repetition Frequency",
         "The number of radar pulses transmitted per second, expressed in Hertz (Hz) or pulses per second (pps). Used in pulsed radar systems such as ATC radar and airborne weather radar. The average power of a pulsed transmitter = Peak Power × Pulse Width × PRF (the duty cycle).",
         "📡 RF & Signal Theory"),

        ("SNR", "Signal-to-Noise Ratio",
         "The ratio of the desired signal power to the noise floor power, expressed in decibels (dB). SNR = S(dBm) − N(dBm). Interference raises the effective noise floor, thereby reducing SNR and degrading receiver performance. GPS requires SNR above approximately 25 dB for reliable tracking.",
         "📡 RF & Signal Theory"),

        # ── ITU-R & Regulatory ────────────────────────────────────────────────
        ("AI", "Agenda Item (WRC)",
         "A specific topic adopted by a World Radiocommunication Conference (WRC) for study and potential regulatory action at the following WRC. Agenda items are numbered (e.g., AI 1.2) and assigned to ITU-R Study Groups and Working Parties for technical study over the 3–4 year inter-conference period.",
         "🌐 ITU-R & Regulatory"),

        ("CPM Report", "Conference Preparatory Meeting Report",
         "The document produced by the ITU-R Conference Preparatory Meeting (CPM) summarizing all technical studies and presenting one or more 'methods' for each WRC Agenda Item. The CPM Report is the primary input to the WRC negotiation process.",
         "🌐 ITU-R & Regulatory"),

        ("Harmful Interference", "Harmful Interference — RR Article 1.169",
         "Interference which endangers the functioning of a radionavigation service or other safety service, OR seriously degrades, obstructs, or repeatedly interrupts a radiocommunication service operating in accordance with the Radio Regulations. This is the legally defined threshold for regulatory action under RR No. 4.10. When interference rises to this level, affected administrations have the right to require cessation of the interfering transmission. For FAA aeronautical systems, demonstrating that a new proposal causes harmful interference under RR 1.169 is the strongest possible policy lever.",
         "🌐 ITU-R & Regulatory"),

        ("Permissible Interference", "Permissible Interference — RR Article 1.167",
         "Observed or predicted interference which complies with quantitative interference and sharing criteria contained in the Radio Regulations or in ITU-R Recommendations or in special agreements as provided for in these Regulations. Permissible interference is below the harmful threshold — it is legally acceptable under the treaty framework. Both permissible interference and accepted interference are used in the coordination of frequency assignments between administrations.",
         "🌐 ITU-R & Regulatory"),

        ("Accepted Interference", "Accepted Interference — RR Article 1.168",
         "Interference at a higher level than that defined as permissible interference, which has been agreed upon between two or more administrations without prejudice to other administrations. Accepted interference is used in bilateral frequency coordination agreements where administrations choose to tolerate higher interference levels than the standard criteria in exchange for operational flexibility. It cannot be imposed on a third administration — only agreed bilaterally.",
         "🌐 ITU-R & Regulatory"),

        ("Interference (RR 1.166)", "Interference — RR Article 1.166",
         "The effect of unwanted energy due to one or a combination of emissions, radiations, or inductions upon reception in a radiocommunication system, manifested by any performance degradation, misinterpretation, or loss of information which could be extracted in the absence of such unwanted energy. The three regulatory categories — harmful (1.169), permissible (1.167), and accepted (1.168) — all fall under this parent definition.",
         "🌐 ITU-R & Regulatory"),

        ("Spurious Emissions", "Spurious Emissions (ITU-R definition)",
         "Emissions at frequencies outside the necessary bandwidth whose level can be reduced without affecting the corresponding transmission of information. Spurious emissions include harmonic emissions, parasitic emissions, intermodulation products, and frequency conversion products, but explicitly EXCLUDE out-of-band (OOB) emissions. Spurious emissions arise from hardware non-idealities (non-linearities, oscillator leakage) rather than the modulation process itself. They are regulated by ITU-R Recommendation SM.329 and are subject to limits in the Radio Regulations Appendix 3.",
         "🌐 ITU-R & Regulatory"),

        ("Out-of-Band (OOB) Emissions", "Out-of-Band Emissions (ITU-R definition)",
         "Emissions at frequencies immediately outside the necessary bandwidth which result from the modulation process, but excluding spurious emissions. OOB emissions are an unavoidable consequence of the modulation waveform — they arise from the spectral sidebands of the modulated signal. Unlike spurious emissions, OOB emissions cannot be eliminated without affecting the information content of the transmission. They can be reduced by filtering or by reducing modulation bandwidth. OOB emissions are the primary mechanism for interference into adjacent aeronautical bands — for example, 5G New Radio OOB emissions near the Radio Altimeter band (4200–4400 MHz). The OOB domain boundary is defined as 250% of the occupied bandwidth on each side of the channel (per ITU-R SM.1540). Within that zone, SM.1540 and SM.1541 apply. Beyond it, RR Appendix 3 spurious limits take over.",
         "🌐 ITU-R & Regulatory"),

        ("SM.1540", "ITU-R Recommendation SM.1540 — Unwanted Emissions in OOB Domain",
         "Governs unwanted emissions falling into adjacent allocated bands from the OOB domain. Establishes the framework for evaluating OOB emissions that land in neighboring services. The OOB domain is defined as extending 250% of the occupied bandwidth on each side of the channel edge. Within this zone, SM.1540 provides the general methodology and Figure 1 of Annex 1 defines the evaluation framework showing the OOB spectrum mask against the spurious limit floor. Key concept: the OOB mask at the allocated band edge must be at least 23 dB down (per RR No. 1.153 — the 0.5% power criterion). Cite SM.1540 when a new service proposes to operate in a band adjacent to an FAA protected band and its OOB emissions cross the boundary.",
         "🌐 ITU-R & Regulatory"),

        ("SM.1541", "ITU-R Recommendation SM.1541 — OOB Emission Mask 23 dB Rule",
         "Specifies that the maximum value of 99% power occupied bandwidth permitted by a particular emission mask can be determined from the 23 dB attenuation levels. In plain terms: the OOB emission mask must attenuate the signal by at least 23 dB at the edge of the allocated band. This comes from RR No. 1.153 which defines occupied bandwidth as the band containing 99% of total mean power, meaning 0.5% (β/2) sits outside each edge — and the mask at that 0.5% point must be 23 dB down. Practical impact: if a proponent's emission mask doesn't show 23 dB attenuation at the boundary of the FAA protected band, SM.1541 gives you the regulatory basis to demand a tighter mask or a larger guard band.",
         "🌐 ITU-R & Regulatory"),

        ("RR No. 1.153", "Radio Regulations No. 1.153 — Occupied Bandwidth Definition",
         "Defines occupied bandwidth as the frequency band such that below its lower frequency limit and above its upper frequency limit, the mean powers emitted are each equal to a specified percentage β/2 of the total mean power of the emission. Unless otherwise specified in an ITU-R Recommendation, β/2 = 0.5% — meaning the occupied bandwidth contains 99% of the total transmitted power, with 0.5% sitting outside each edge. This definition is the anchor for both the OOB domain boundary (250% of occupied BW per SM.1540) and the 23 dB OOB mask rule (SM.1541). When a proponent claims a narrow occupied bandwidth to minimize their OOB footprint, challenge them by demanding a 99% power measurement per RR 1.153.",
         "🌐 ITU-R & Regulatory"),

        ("RR Appendix 3", "Radio Regulations Appendix 3 (Rev. WRC-12) — Spurious Emission Limits",
         "Defines maximum permitted power levels for unwanted emissions in the spurious domain — the region beyond 250% of occupied bandwidth. The limits are: All services (general): 43 + 10·log(P) dB below carrier, or 70 dBc, whichever is less stringent. Space services (earth stations): same formula or 70 dBc. Space services (space stations): 43 + 10·log(P) dB or 60 dBc — whichever is less stringent. The formula scales with transmitter power P in watts. Example (ROSE-L satellite, 39 dBW peak): 43 + 39 = 82 dBc, but 60 dBc cap applies, so limit = 60 dBc below 39 dBW = −21 dBW spurious power. When a satellite's spurious product falls inside an FAA protected band, RR Appendix 3 is the quantitative limit you cite.",
         "🌐 ITU-R & Regulatory"),

        ("SM.329", "ITU-R Recommendation SM.329 — Spurious Domain Measurement and Limits",
         "Companion to RR Appendix 3. Provides the detailed measurement methodology for characterizing and verifying spurious emission levels. Defines the measurement procedures, test conditions, and how to apply the limits from RR Appendix 3. Cite SM.329 alongside Appendix 3 when challenging a proponent's spurious emission measurements or methodology — Appendix 3 gives the limit, SM.329 gives the method for verifying compliance.",
         "🌐 ITU-R & Regulatory"),

        ("Unwanted Emissions", "Unwanted Emissions",
         "The combination of spurious emissions and out-of-band (OOB) emissions. This is the umbrella category under which both emission types fall. In interference analysis for ITU-R contributions, you must characterize which type of unwanted emission is causing the problem, as they are subject to different regulatory limits and different mitigation approaches: OOB can be reduced by transmitter filtering or guard bands; spurious can be reduced by improving transmitter linearity or adding cavity filters.",
         "🌐 ITU-R & Regulatory"),

        ("In-Band Interference", "In-Band Interference (technical mechanism)",
         "Interference where the interfering signal's carrier or main lobe falls within the allocated frequency band of the victim service. In-band interference is the most direct form — the interferer and victim are co-frequency or co-channel. For FAA systems, in-band interference means the new service is proposed for allocation within a band already assigned to ARNS or RNSS. This is the most straightforward case for citing harmful interference under RR No. 4.10.",
         "🌐 ITU-R & Regulatory"),

        ("Adjacent Band Interference", "Adjacent Band / Near-Band Interference (technical mechanism)",
         "Interference where the interfering system operates in a different band but its out-of-band (OOB) or spurious emissions, or the victim receiver's finite selectivity, allow coupling of interfering energy into the victim receiver. The 5G C-band (3.7–3.98 GHz) / Radio Altimeter (4.2–4.4 GHz) case is a classic adjacent-band interference scenario: the 5G signal was outside the altimeter band, but OOB emissions and receiver front-end selectivity limitations allowed destructive coupling.",
         "🌐 ITU-R & Regulatory"),

        ("Intermodulation", "Intermodulation Interference (technical mechanism)",
         "Interference produced when two or more signals are applied to a non-linear device (amplifier, mixer, or antenna), producing new signals at frequencies that are sums and differences of the input frequencies and their harmonics. The most troublesome are third-order intermodulation products (2f1−f2 and 2f2−f1) which can fall in-band even when the original signals are outside the protected band. Intermodulation is particularly important at sites with multiple co-located transmitters.",
         "🌐 ITU-R & Regulatory"),

        ("Receiver Blocking", "Receiver Blocking / Desensitization (technical mechanism)",
         "A form of adjacent-band or out-of-band interference where a strong signal overloads the victim receiver's Low Noise Amplifier (LNA) or front-end stages, compressing their gain and raising the effective noise floor. Unlike intermodulation, blocking does not require the interferer to be at a specific frequency relationship to the victim — it only requires the interferer to be sufficiently powerful. Blocking is characterized by the receiver's 1 dB compression point (P1dB). The Radio Altimeter / 5G interference problem was primarily a blocking mechanism, not an in-band interference problem.",
         "🌐 ITU-R & Regulatory"),

        ("ITU", "International Telecommunication Union",
         "A United Nations specialized agency for information and communication technologies, headquartered in Geneva, Switzerland. The ITU-R (Radiocommunication Sector) manages the global radio frequency spectrum and satellite orbits through international treaty — the Radio Regulations.",
         "🌐 ITU-R & Regulatory"),

        ("ITU-R", "International Telecommunication Union — Radiocommunication Sector",
         "The ITU sector responsible for managing the radio frequency spectrum and satellite orbits. ITU-R produces international standards (Recommendations, Reports) and maintains the Radio Regulations through the World Radiocommunication Conference (WRC) process.",
         "🌐 ITU-R & Regulatory"),

        ("M.1318", "ITU-R Recommendation M.1318",
         "ITU-R Recommendation M.1318 — 'Model for the evaluation of continuous interference levels to radionavigation-satellite service receivers.' Defines the c = a − b calculation framework for assessing aggregate non-RNSS interference to GNSS receivers in the 1164–1215, 1215–1300, 1559–1610, and 5010–5030 MHz bands.",
         "🌐 ITU-R & Regulatory"),

        ("M.1477", "ITU-R Recommendation M.1477",
         "ITU-R Recommendation M.1477 — Annex 5 establishes the aeronautical safety margin for GNSS: at least 6 dB protection margin for safety-of-life applications, plus an additional 10 dB margin when the interfering signal bandwidth is 700 Hz or less (the narrowband rule).",
         "🌐 ITU-R & Regulatory"),

        ("M.1642", "ITU-R Recommendation M.1642",
         "ITU-R Recommendation M.1642 — 'Methodology for assessing the interference from IMT-2000 to aeronautical radionavigation systems.' The standard methodology for compatibility studies between IMT systems and ARNS, including radio altimeters and other aeronautical navigation aids.",
         "🌐 ITU-R & Regulatory"),

        ("M.1904", "ITU-R Recommendation M.1904",
         "ITU-R Recommendation M.1904 — Establishes protection criteria for GLONASS spaceborne receivers. Annex 1 Table 1 Note 3 specifies a safety margin of 6 dB. Part of the citation stack defending the 6 dB aeronautical safety margin doctrine alongside M.1477 and M.1905.",
         "🌐 ITU-R & Regulatory"),

        ("M.1905", "ITU-R Recommendation M.1905",
         "ITU-R Recommendation M.1905 — Recommends that a safety margin be applied for protection of the safety aspects and applications of RNSS when performing interference analyses. Note 1 specifies an aeronautical safety margin of 6 dB. Applies to ALL RNSS systems, making it the broadest authority for the 6 dB doctrine.",
         "🌐 ITU-R & Regulatory"),

        ("P.452", "ITU-R Recommendation P.452",
         "ITU-R Recommendation P.452 — 'Prediction procedure for the evaluation of interference between stations on the surface of the Earth at frequencies above about 0.1 GHz.' The standard terrestrial propagation model for point-to-area interference prediction. Used for ground-to-ground interference scenarios.",
         "🌐 ITU-R & Regulatory"),

        ("P.528", "ITU-R Recommendation P.528",
         "ITU-R Recommendation P.528 — 'Propagation curves for aeronautical mobile and radionavigation services using the VHF, UHF and SHF bands.' The standard propagation model for aeronautical scenarios. Required when the victim receiver is airborne. Accounts for slant-path geometry, atmospheric refraction, and radio horizon effects.",
         "🌐 ITU-R & Regulatory"),

        ("P.676", "ITU-R Recommendation P.676",
         "ITU-R Recommendation P.676 — 'Attenuation by atmospheric gases and related effects.' Specifies gaseous attenuation from oxygen and water vapor. Implemented in this tool via the itur Python library. Relevant above approximately 3 GHz; negligible for L-band and VHF/UHF paths.",
         "🌐 ITU-R & Regulatory"),

        ("Radio Regulations (RR)", "ITU Radio Regulations",
         "The international treaty governing use of the radio frequency spectrum and satellite orbits, binding on all 193 ITU member states. Key articles for aeronautical spectrum defense: RR No. 4.10 (no harmful interference to safety-of-life services), RR No. 5.444 (ARNS protection at 960–1215 MHz), RR No. 5.328 (ARNS at 108–137 MHz).",
         "🌐 ITU-R & Regulatory"),

        ("RR No. 4.10", "Radio Regulations Article No. 4.10",
         "The Radio Regulations provision stating that all stations must be established and operated so as not to cause harmful interference to safety-of-life services. This is the strongest regulatory instrument for protecting aeronautical frequencies — it applies regardless of allocation status and cannot be waived.",
         "🌐 ITU-R & Regulatory"),

        ("Resolution 233", "WRC Resolution 233",
         "A World Radiocommunication Conference Resolution requiring that new spectrum allocations and uses not degrade the performance of Radionavigation Satellite Service (RNSS) systems including GPS, GLONASS, and Galileo.",
         "🌐 ITU-R & Regulatory"),

        ("SM.2028", "ITU-R Recommendation SM.2028",
         "ITU-R Recommendation SM.2028 — 'Monte Carlo simulation methodology for use in sharing and compatibility studies between different radio services or systems.' The authoritative methodology for aggregate interference Monte Carlo analysis. Specifies that time percentage should be 1% for protection studies and that violation probability must be less than 5%.",
         "🌐 ITU-R & Regulatory"),

        ("SG 5", "Study Group 5 — Terrestrial Services",
         "An ITU-R Study Group responsible for terrestrial services including mobile, fixed, amateur, and broadcasting. Working Parties 5A, 5B, and 5D operate under SG 5. SG 5 plenary sessions approve and coordinate the work of these Working Parties.",
         "🌐 ITU-R & Regulatory"),

        ("TIES", "ITU Electronic Document System",
         "The ITU-R electronic document management system where all Working Party contributions, chairman's reports, and meeting documents are posted. Access to TIES requires ITU membership or delegation credentials.",
         "🌐 ITU-R & Regulatory"),

        # ── Aeronautical Systems ──────────────────────────────────────────────
        ("AIS", "Automatic Identification System",
         "A maritime tracking system that broadcasts vessel identity, position, speed, and heading on VHF frequencies (161.975 MHz and 162.025 MHz). AIS is being succeeded by VDES (VHF Data Exchange System) for expanded maritime data communications. Governed by ITU-R Working Party 5B (WP 5B).",
         "✈️ Aeronautical Systems"),

        ("ARSR", "Air Route Surveillance Radar",
         "Long-range Air Traffic Control radar used to track aircraft along airways. Operates in the 2700–2900 MHz band. The FAA operates the ARSR-4 system. Also refers to the L-band Air Route Surveillance Radar variant of WAAS reference station receivers.",
         "✈️ Aeronautical Systems"),

        ("ASR", "Airport Surveillance Radar",
         "Short-range Air Traffic Control radar used to track aircraft in the terminal area around an airport. Typically operates in the 2700–2900 MHz (S-band) or 9000–9500 MHz (X-band) range.",
         "✈️ Aeronautical Systems"),

        ("Radio Altimeter", "Radio Altimeter (Rad Alt)",
         "An aircraft instrument that measures the exact height above ground by transmitting radio pulses downward and measuring the time for the echo to return. Operates in 4200–4400 MHz. Critical for Category III instrument landings (near-zero visibility), Terrain Awareness and Warning System (TAWS), Enhanced Ground Proximity Warning System (EGPWS), and helicopter operations. A safety-of-life system protected under Radio Regulations No. 4.10.",
         "✈️ Aeronautical Systems"),

        # ── Measurements & Units ─────────────────────────────────────────────
        ("GHz", "Gigahertz",
         "A unit of frequency equal to one billion (10⁹) cycles per second, or 1000 Megahertz (MHz). Used for microwave and millimeter-wave frequencies. Example: Radio Altimeter at 4.2–4.4 GHz; 5G NR at 3.5 GHz or 26 GHz.",
         "📐 Measurements & Units"),

        ("Hz", "Hertz",
         "The base unit of frequency, equal to one cycle per second. Radio frequencies are typically expressed in kilohertz (kHz = 10³ Hz), Megahertz (MHz = 10⁶ Hz), or Gigahertz (GHz = 10⁹ Hz).",
         "📐 Measurements & Units"),

        ("kHz", "Kilohertz",
         "A unit of frequency equal to one thousand (10³) cycles per second. Used for low-frequency and medium-frequency radio systems. The 700 Hz narrowband threshold in ITU-R M.1477 Annex 5 is specified in Hertz — approximately 0.0007 MHz.",
         "📐 Measurements & Units"),

        ("kT", "Boltzmann's Constant × Temperature",
         "The product of Boltzmann's constant (k = 1.38 × 10⁻²³ Joules per Kelvin) and temperature in Kelvin (T). At 290 Kelvin (room temperature), kT = −174 dBm/Hz — the fundamental thermal noise spectral density. This is the starting point for all receiver noise floor calculations.",
         "📐 Measurements & Units"),

        ("MHz", "Megahertz",
         "A unit of frequency equal to one million (10⁶) cycles per second. The standard unit for specifying radio frequencies in the VHF, UHF, and microwave ranges. All FAA protected bands are specified in MHz in this tool.",
         "📐 Measurements & Units"),

        ("W", "Watt",
         "The SI unit of power. In radio engineering, transmitter power is often expressed in Watts or converted to decibels relative to one milliwatt (dBm): P(dBm) = 10 × log10(P(mW)). Common reference: 1 W = 30 dBm; 20 W = 43 dBm (typical LTE base station).",
         "📐 Measurements & Units"),

        # ── Analysis Methods ─────────────────────────────────────────────────
        ("CCDF", "Complementary Cumulative Distribution Function",
         "A statistical function showing the probability that a variable exceeds a given value. In Monte Carlo interference analysis, the CCDF of I/N is plotted to show the probability that aggregate interference exceeds the protection threshold. Reading the CCDF at the I/N threshold directly gives the violation probability — which must be less than 5% per ITU-R SM.2028.",
         "💻 Analysis Methods"),

        ("Monte Carlo", "Monte Carlo Simulation",
         "A statistical simulation method that uses random sampling to model uncertain or variable inputs. In ITU-R interference analysis (SM.2028), Monte Carlo simulates thousands of random interferer deployments, computing aggregate I/N for each trial, to produce a statistical distribution of interference outcomes and a violation probability.",
         "💻 Analysis Methods"),

        ("SEAMCAT", "Spectrum Engineering Advanced Monte Carlo Analysis Tool",
         "Free software developed by CEPT (European Conference of Postal and Telecommunications Administrations) for Monte Carlo aggregate interference analysis. Widely used in ITU-R Working Party contributions. Implements the methodology of ITU-R SM.2028.",
         "💻 Analysis Methods"),
    ]

    # ── Filter ────────────────────────────────────────────────────────────────
    filtered = []
    for term, full, definition, cat in GLOSSARY:
        if cat_filter != "All" and cat != cat_filter:
            continue
        if search and search not in term.lower() and search not in full.lower() and search not in definition.lower():
            continue
        filtered.append((term, full, definition, cat))

    st.markdown(f"**Showing {len(filtered)} of {len(GLOSSARY)} entries**")
    st.markdown("---")

    if not filtered:
        st.info("No entries match your search. Try a broader term.")
    else:
        # Group by category
        current_cat = None
        for term, full, definition, cat in filtered:
            if cat != current_cat:
                current_cat = cat
                st.subheader(cat)
            with st.expander(f"**{term}** — {full}"):
                st.markdown(definition)

    st.markdown("---")
    st.caption(f"FAA RF Interference Analysis Tool Glossary — {len(GLOSSARY)} terms across 6 categories.")

# ─────────────────────────────────────────────────────────────────────────────
# TAB 14 — MICROWAVE LINK BUDGET CALCULATOR
# ─────────────────────────────────────────────────────────────────────────────
elif selected_tab == "📻 Microwave Link Budget":
    st.title("📻 Microwave Link Budget Calculator")
    st.markdown("*Enter your site parameters and the tool instantly calculates EIRP, Free Space Path Loss, atmospheric absorption, and Received Signal Level (RSL) in both directions.*")
    ex("RSL = EIRP − Free Space Path Loss − Absorption + Rx Antenna Gain − Rx Losses. Enter any combination of sites — this calculator is not tied to any specific report format.")

    # ── Session state for saved links ─────────────────────────────────────────
    if "mw_saved_links" not in st.session_state:
        st.session_state.mw_saved_links = {}

    # ── Load saved link ───────────────────────────────────────────────────────
    saved_names = list(st.session_state.mw_saved_links.keys())
    load_choice = "(new link)"
    if saved_names:
        load_choice = st.selectbox("📂 Load saved link:", ["(new link)"] + saved_names)
    else:
        st.caption("No saved links yet — fill in parameters below and save.")

    # Defaults — load from saved if chosen
    D = {
        "site_a": "", "site_b": "",
        "location_a": "", "location_b": "",
        "lat_a": "", "lon_a": "", "lat_b": "", "lon_b": "",
        "elev_a": 0.0, "elev_b": 0.0,
        "struct_a": 0.0, "struct_b": 0.0,
        "path_override": False, "path_km_manual": 10.0,
        "band_ghz": 7.0, "freq_a": 7902.5, "freq_b": 8262.5,
        "tx_pwr_a": 30.0, "tx_pwr_b": 30.0,
        "ant_gain_a": 40.0, "ant_gain_b": 40.0,
        "ant_model_a": "", "ant_model_b": "",
        "other_loss_a": 2.0, "other_loss_b": 2.0,
        "branch_a": 0.0, "branch_b": 0.0,
        "radio_model": "", "bandwidth_mhz": 30.0,
        "modulation": "256QAM", "bitrate_mbps": 100.0,
        "spec_atten": 0.0102, "link_id": "",
    }
    if load_choice != "(new link)" and load_choice in st.session_state.mw_saved_links:
        for k, v in st.session_state.mw_saved_links[load_choice].items():
            D[k] = v

    # ── Coordinate helpers ────────────────────────────────────────────────────
    def parse_latlon(s):
        if not s or not s.strip(): return None
        s = s.strip().upper()
        hemi = None
        for suffix, sign in [(' N','N','N'),(' S','S','S'),(' E','E','E'),(' W','W','W')]:
            pass
        for ch, sign in [('N',1),('S',-1),('E',1),('W',-1)]:
            if s.endswith(' '+ch) or s.endswith(ch):
                hemi = sign; s = s.rstrip(ch).strip(); break
        s = s.replace('-',' ').replace('°',' ').replace("'",' ').replace('"',' ')
        parts = s.split()
        try:
            if len(parts)==1:   val = float(parts[0])
            elif len(parts)==2: val = float(parts[0]) + float(parts[1])/60
            else:               val = float(parts[0]) + float(parts[1])/60 + float(parts[2])/3600
            return -val if hemi==-1 else val
        except: return None

    def haversine_km(la, lo, lb, lb2):
        R = 6371.0
        p1,p2 = math.radians(la), math.radians(lb)
        dp,dl = math.radians(lb-la), math.radians(lb2-lo)
        a = math.sin(dp/2)**2 + math.cos(p1)*math.cos(p2)*math.sin(dl/2)**2
        return R*2*math.atan2(math.sqrt(a), math.sqrt(1-a))

    def dms(dd, is_lat):
        h = ("N" if dd>=0 else "S") if is_lat else ("E" if dd>=0 else "W")
        dd=abs(dd); d=int(dd); m=int((dd-d)*60); s=((dd-d)*60-m)*60
        return f"{d}°{m}'{s:.1f}\"{h}"

    st.markdown("---")

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 1 — SITE IDENTITY & LOCATION
    # ══════════════════════════════════════════════════════════════════════════
    st.subheader("📍 Section 1 — Site Identity & Location")
    ex("Enter any site names and optional coordinates. Coordinates auto-calculate path length via the Haversine great-circle formula. All fields are free-form — not tied to any specific report.")

    c1, c2 = st.columns(2)
    with c1:
        st.markdown("**Site A**")
        site_a     = st.text_input("Site A — Name / Identifier", value=D["site_a"], placeholder="e.g. BEDFORD or TX Site")
        location_a = st.text_input("Site A — Location Description", value=D["location_a"], placeholder="e.g. Bedford, Virginia, USA")
        lat_a_s    = st.text_input("Latitude A (optional)", value=D["lat_a"], placeholder="e.g. 37-31-21.9 N  or  37.5261")
        lon_a_s    = st.text_input("Longitude A (optional)", value=D["lon_a"], placeholder="e.g. 79-30-36.5 W  or  -79.5101")
        elev_a     = st.number_input("Ground Elevation A (m)", value=float(D["elev_a"]), step=1.0, help="Meters above mean sea level")
        struct_a   = st.number_input("Structure Height A (m AGL)", value=float(D["struct_a"]), step=0.5, help="Tower height above ground")
    with c2:
        st.markdown("**Site B**")
        site_b     = st.text_input("Site B — Name / Identifier", value=D["site_b"], placeholder="e.g. ROANOKE or RX Site")
        location_b = st.text_input("Site B — Location Description", value=D["location_b"], placeholder="e.g. Roanoke, Virginia, USA")
        lat_b_s    = st.text_input("Latitude B (optional)", value=D["lat_b"], placeholder="e.g. 37-18-32.7 N  or  37.3091")
        lon_b_s    = st.text_input("Longitude B (optional)", value=D["lon_b"], placeholder="e.g. 80-09-36.5 W  or  -80.1601")
        elev_b     = st.number_input("Ground Elevation B (m)", value=float(D["elev_b"]), step=1.0)
        struct_b   = st.number_input("Structure Height B (m AGL)", value=float(D["struct_b"]), step=0.5)

    # Parse coords
    la = parse_latlon(lat_a_s); lo = parse_latlon(lon_a_s)
    lb = parse_latlon(lat_b_s); lb2= parse_latlon(lon_b_s)
    coords_ok = all(v is not None for v in [la, lo, lb, lb2])
    auto_km = round(haversine_km(la, lo, lb, lb2), 3) if coords_ok else None

    if any([lat_a_s, lon_a_s, lat_b_s, lon_b_s]):
        st.markdown("**Coordinate Status**")
        cs1, cs2, cs3 = st.columns(3)
        with cs1:
            if la is not None and lo is not None:
                ok(f"Site A: {dms(la,True)}, {dms(lo,False)}")
            else:
                warn("Site A coordinates not parseable")
        with cs2:
            if lb is not None and lb2 is not None:
                ok(f"Site B: {dms(lb,True)}, {dms(lb2,False)}")
            else:
                warn("Site B coordinates not parseable")
        with cs3:
            if auto_km:
                ok(f"Path (Haversine): **{auto_km:.3f} km**")
            else:
                st.info("Enter both sites for auto path length")

    # Path length entry
    path_override = st.checkbox("Enter path length manually (override coordinate calculation)",
                                value=D["path_override"])
    if path_override or not auto_km:
        path_km = st.number_input("Path Length (km)",
                                  value=float(D["path_km_manual"]) if D["path_km_manual"] else (auto_km or 10.0),
                                  min_value=0.001, step=0.001, format="%.3f")
        if auto_km and abs(path_km - auto_km) > 0.1:
            warn(f"Manual path ({path_km:.3f} km) differs from coordinate calculation ({auto_km:.3f} km) by {abs(path_km-auto_km):.3f} km")
    else:
        path_km = auto_km
        st.metric("Path Length", f"{path_km:.3f} km", help="Haversine great-circle distance from coordinates")

    # Map
    if coords_ok:
        st.markdown("**🗺️ Link Map**")
        st.map(pd.DataFrame({'lat':[la,lb],'lon':[lo,lb2]}), zoom=7, use_container_width=True)
        st.caption(f"📍 {site_a or 'Site A'}: {la:.5f}°, {lo:.5f}°   |   📍 {site_b or 'Site B'}: {lb:.5f}°, {lb2:.5f}°   |   Path: {path_km:.3f} km")

    st.markdown("---")

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 2 — RF PARAMETERS
    # ══════════════════════════════════════════════════════════════════════════
    st.subheader("📡 Section 2 — RF & Antenna Parameters")
    ex("Enter parameters for each end independently — supports asymmetric links where sites have different antennas, power levels, or losses.")

    c3, c4 = st.columns(2)
    with c3:
        st.markdown(f"**{site_a or 'Site A'} — Transmitter**")
        ant_model_a  = st.text_input("Antenna Model A", value=D["ant_model_a"], placeholder="e.g. Andrew HP6-71W")
        ant_gain_a   = st.number_input("Antenna Gain A (dBi)", value=float(D["ant_gain_a"]), step=0.1,
                                        help="Peak gain toward Site B. Use 0 dBi for isotropic worst-case.")
        tx_pwr_a     = st.number_input("Tx Power A (dBm)", value=float(D["tx_pwr_a"]), step=0.5,
                                        help="Radio output power in dBm. 30 dBm = 1 W, 43 dBm = 20 W")
        other_loss_a = st.number_input("Feeder / Connector Loss A (dB)", value=float(D["other_loss_a"]), step=0.1,
                                        help="Cable, connectors, jumpers between radio and antenna")
        branch_a     = st.number_input("Branching / Hybrid Loss A (dB)", value=float(D["branch_a"]), step=0.1,
                                        help="Combiner or branching losses. Zero for direct antenna connection.")
    with c4:
        st.markdown(f"**{site_b or 'Site B'} — Transmitter**")
        ant_model_b  = st.text_input("Antenna Model B", value=D["ant_model_b"], placeholder="e.g. Andrew HP10-71W")
        ant_gain_b   = st.number_input("Antenna Gain B (dBi)", value=float(D["ant_gain_b"]), step=0.1,
                                        help="Peak gain toward Site A")
        tx_pwr_b     = st.number_input("Tx Power B (dBm)", value=float(D["tx_pwr_b"]), step=0.5)
        other_loss_b = st.number_input("Feeder / Connector Loss B (dB)", value=float(D["other_loss_b"]), step=0.1)
        branch_b     = st.number_input("Branching / Hybrid Loss B (dB)", value=float(D["branch_b"]), step=0.1)

    st.markdown("---")

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 3 — PATH & SYSTEM PARAMETERS
    # ══════════════════════════════════════════════════════════════════════════
    st.subheader("🌐 Section 3 — Path & System Parameters")
    ex("Band and path length drive the dominant Free Space Path Loss term. Specific attenuation is from ITU-R P.676 — 0.010 dB/km is a good starting value for 7 GHz under standard conditions.")

    c5, c6, c7 = st.columns(3)
    with c5:
        band_ghz   = st.number_input("Band (GHz)", value=float(D["band_ghz"]), step=0.5, min_value=0.1,
                                      help="Nominal band center frequency used for FSPL calculation")
        freq_a     = st.number_input("Tx Frequency A→B (MHz)", value=float(D["freq_a"]), step=0.5)
        freq_b     = st.number_input("Tx Frequency B→A (MHz)", value=float(D["freq_b"]), step=0.5)
    with c6:
        spec_atten = st.number_input("P.676 Attenuation (dB/km)", value=float(D["spec_atten"]),
                                      step=0.0001, format="%.4f",
                                      help="ITU-R P.676 specific gaseous attenuation. Typical values: 7 GHz ≈ 0.010, 15 GHz ≈ 0.015, 23 GHz ≈ 0.18, 60 GHz ≈ 15")
        radio_model= st.text_input("Radio Model", value=D["radio_model"], placeholder="e.g. Aviat ODU600v2")
        link_id    = st.text_input("Link ID / Reference", value=D["link_id"], placeholder="e.g. MW-001 or any identifier")
    with c7:
        bandwidth_mhz = st.number_input("Channel Bandwidth (MHz)", value=float(D["bandwidth_mhz"]), step=1.0)
        modulation    = st.text_input("Modulation", value=D["modulation"], placeholder="e.g. 256QAM, QPSK, 4096QAM")
        bitrate_mbps  = st.number_input("Bit Rate (Mb/s)", value=float(D["bitrate_mbps"]), step=1.0)

    st.markdown("---")

    # ══════════════════════════════════════════════════════════════════════════
    # CALCULATIONS — run when path_km > 0
    # ══════════════════════════════════════════════════════════════════════════
    if path_km and path_km > 0 and band_ghz > 0:

        # Core link budget calculations
        eirp_a   = round(tx_pwr_a + ant_gain_a - other_loss_a - branch_a, 2)
        eirp_b   = round(tx_pwr_b + ant_gain_b - other_loss_b - branch_b, 2)
        fspl     = round(20*np.log10(path_km) + 20*np.log10(band_ghz*1000) + 32.44, 2)
        abs_db   = round(spec_atten * path_km, 3)
        tpl      = round(fspl + abs_db, 2)
        rsl_ab   = round(eirp_a - tpl + ant_gain_b - other_loss_b - branch_b, 2)
        rsl_ba   = round(eirp_b - tpl + ant_gain_a - other_loss_a - branch_a, 2)
        balanced = abs(rsl_ab - rsl_ba) < 0.05
        sA = site_a or "Site A"
        sB = site_b or "Site B"

        # ── TOP METRICS ───────────────────────────────────────────────────────
        st.subheader("⚡ Results")
        ex("All values recalculate instantly when any input changes. EIRP is the effective radiated power; RSL is the power arriving at the receiver input.")

        m1, m2, m3, m4, m5, m6 = st.columns(6)
        m1.metric("EIRP — "+sA, f"{eirp_a:.2f} dBm",
            help=f"Pt({tx_pwr_a}) + Gt({ant_gain_a}) − Feeder({other_loss_a}) − Branching({branch_a}) = {eirp_a}")
        m2.metric("EIRP — "+sB, f"{eirp_b:.2f} dBm",
            help=f"Pt({tx_pwr_b}) + Gt({ant_gain_b}) − Feeder({other_loss_b}) − Branching({branch_b}) = {eirp_b}")
        m3.metric("Free Space Path Loss", f"{fspl:.2f} dB",
            help=f"20·log({path_km}) + 20·log({band_ghz*1000:.0f}) + 32.44")
        m4.metric("Absorption (P.676)", f"{abs_db:.3f} dB",
            help=f"{spec_atten} dB/km × {path_km} km")
        m5.metric(f"RSL at {sB}", f"{rsl_ab:.2f} dBm",
            delta_color="normal" if rsl_ab > -80 else "inverse",
            help=f"Power arriving at {sB} receiver input")
        m6.metric(f"RSL at {sA}", f"{rsl_ba:.2f} dBm",
            delta_color="normal" if rsl_ba > -80 else "inverse",
            help=f"Power arriving at {sA} receiver input (reverse path)")

        # Balance / capacity callouts
        if balanced:
            ok(f"Link is balanced — RSL is equal in both directions ({rsl_ab:.2f} dBm). Both sites use equal effective EIRP.")
        else:
            diff = abs(rsl_ab - rsl_ba)
            if diff < 1.0:
                warn(f"Link is slightly unbalanced — {diff:.2f} dB difference between forward and reverse RSL. Usually acceptable for ACM systems.")
            else:
                warn(f"Link imbalance of {diff:.2f} dB — check antenna gains and losses at each end. Significant imbalance reduces overall link availability.")

        # Capacity sanity
        mod_bits = {"4096QAM":12,"2048QAM":11,"1024QAM":10,"512QAM":9,"256QAM":8,
                    "128QAM":7,"64QAM":6,"32QAM":5,"16QAM":4,"QPSK":2,"BPSK":1}
        mod_key = modulation.strip().upper().split()[0] if modulation else ""
        bits = mod_bits.get(mod_key)
        if bits and bandwidth_mhz > 0:
            theo_max = bandwidth_mhz * bits
            if bitrate_mbps > theo_max:
                warn(f"Bit rate {bitrate_mbps:.0f} Mb/s exceeds theoretical {mod_key} maximum of {theo_max:.0f} Mb/s over {bandwidth_mhz:.0f} MHz — check inputs.")

        # ── STEP BY STEP WATERFALL ─────────────────────────────────────────────
        st.markdown("---")
        st.subheader("📊 Step-by-Step Link Budget")
        ex("Running total tracks cumulative signal power from transmitter output through to the receiver. Blue highlighted rows are EIRP and RSL — the two engineering milestones.")

        dir_tab_a, dir_tab_b = st.tabs([
            f"➡️  {sA} → {sB}",
            f"⬅️  {sB} → {sA}",
        ])

        def make_waterfall(tx_pwr, tx_gain, feeder, branch, fspl_v, abs_v,
                           rx_gain, rx_feeder, rx_branch, rsl_v, tx_name, rx_name):
            running = tx_pwr
            rows = []
            steps = [
                ("1", f"Transmit Power — {tx_name}",       f"+{tx_pwr:.2f}",     tx_pwr,       f"{running:.2f} dBm",   "Radio PA output"),
                ("2", "＋ Tx Antenna Gain",                  f"+{tx_gain:.2f}",    tx_gain,      None,                   "Peak gain toward far end (dBi)"),
                ("3", "－ Feeder / Connector Loss (Tx)",     f"−{feeder:.2f}",     -feeder,      None,                   "Cable, connectors, jumpers"),
                ("4", "－ Branching / Hybrid Loss (Tx)",     f"−{branch:.2f}",     -branch,      None,                   "Combiner losses (0 if direct)"),
                ("★", f"EIRP",                               f"{tx_pwr+tx_gain-feeder-branch:.2f} dBm", None, f"{tx_pwr+tx_gain-feeder-branch:.2f} dBm", "Effective Isotropic Radiated Power"),
                ("5", "－ Free Space Path Loss (FSPL)",      f"−{fspl_v:.2f}",     -fspl_v,      None,                   f"20·log({path_km})+20·log({band_ghz*1000:.0f})+32.44"),
                ("6", "－ Atmospheric Absorption (P.676)",   f"−{abs_v:.3f}",      -abs_v,       None,                   f"{spec_atten} dB/km × {path_km} km"),
                ("7", "Total Propagation Loss",              f"−{fspl_v+abs_v:.2f} dB", None,   "—",                    "FSPL + Absorption"),
                ("8", "＋ Rx Antenna Gain",                   f"+{rx_gain:.2f}",    rx_gain,      None,                   f"Peak gain toward {tx_name} (dBi)"),
                ("9", "－ Feeder / Connector Loss (Rx)",     f"−{rx_feeder:.2f}",  -rx_feeder,   None,                   "Rx side cable and connectors"),
                ("10","－ Branching / Hybrid Loss (Rx)",     f"−{rx_branch:.2f}",  -rx_branch,   None,                   "Rx combiner losses"),
                ("★", f"RSL at {rx_name}",                   f"{rsl_v:.2f} dBm",   None,         f"{rsl_v:.2f} dBm",     "Power at receiver input"),
            ]
            for snum, label, value, delta, run_override, note in steps:
                if delta is not None:
                    running = round(running + delta, 2)
                run_str = run_override if run_override is not None else f"{running:.2f} dBm"
                rows.append({"Step": snum, "Description": label,
                             "Value": value, "Running Total": run_str, "Notes": note})

            df = pd.DataFrame(rows)

            def style_row(row):
                if row["Step"] == "★" and "EIRP" in row["Description"]:
                    return ["background-color:#1a3560;font-weight:bold;color:#aaddff"]*len(row)
                if row["Step"] == "★" and "RSL" in row["Description"]:
                    return ["background-color:#1a4a1a;font-weight:bold;color:#aaffaa"]*len(row)
                if row["Step"] == "7":
                    return ["background-color:#3a1a1a;color:#ffaaaa"]*len(row)
                return [""]*len(row)

            st.dataframe(df.style.apply(style_row, axis=1),
                         use_container_width=True, hide_index=True)

        with dir_tab_a:
            make_waterfall(tx_pwr_a, ant_gain_a, other_loss_a, branch_a,
                           fspl, abs_db, ant_gain_b, other_loss_b, branch_b,
                           rsl_ab, sA, sB)
        with dir_tab_b:
            make_waterfall(tx_pwr_b, ant_gain_b, other_loss_b, branch_b,
                           fspl, abs_db, ant_gain_a, other_loss_a, branch_a,
                           rsl_ba, sB, sA)

        # ── P.676 SECTION ─────────────────────────────────────────────────────
        st.markdown("---")
        st.subheader("🌡️ Atmospheric Absorption (ITU-R P.676)")
        ex("Specific attenuation × path length = total absorption loss. At 7 GHz this is small (≈0.6 dB over 60 km). At 23 GHz or 60 GHz it becomes the dominant loss term.")

        p1, p2, p3, p4 = st.columns(4)
        p1.metric("Specific Attenuation", f"{spec_atten:.4f} dB/km",
            help="Change B56 to model different frequencies, temperatures, or humidity")
        p2.metric("Path Length", f"{path_km:.3f} km")
        p3.metric("Total Absorption", f"{abs_db:.3f} dB")
        p4.metric("% of Total Path Loss", f"{100*abs_db/tpl:.1f}%",
            help="Absorption as a fraction of total propagation loss")

        # P.676 reference table
        with st.expander("📋 P.676 Reference — Specific Attenuation at Common Frequencies"):
            st.markdown("""
| Frequency | Specific Attenuation | Notes |
|---|---|---|
| 2 GHz | ~0.003 dB/km | Negligible |
| 7 GHz | ~0.010 dB/km | Standard microwave backbone |
| 11 GHz | ~0.012 dB/km | Common microwave band |
| 15 GHz | ~0.015 dB/km | Upper microwave |
| 23 GHz | ~0.18 dB/km | Rain scatter starts to matter |
| 38 GHz | ~0.12 dB/km | Short-haul millimeter wave |
| 60 GHz | ~14 dB/km | Oxygen absorption peak — very short range |
| 80 GHz | ~0.5 dB/km | E-band license-free |

*Values at standard atmosphere (290K, 50% relative humidity, sea level)*
            """)

        # ── RSL SWEEP CHART ───────────────────────────────────────────────────
        st.markdown("---")
        st.subheader("📉 RSL vs Path Length")
        ex("Shows how RSL degrades as distance increases at the current settings. The vertical marker shows your current path. Use this to find the maximum path length before RSL falls below a minimum threshold.")

        max_sweep = max(path_km * 3, 50.0)
        d_sw = np.linspace(0.1, max_sweep, 500)
        rsl_ab_sw = eirp_a - (20*np.log10(d_sw)+20*np.log10(band_ghz*1000)+32.44) \
                    - spec_atten*d_sw + ant_gain_b - other_loss_b - branch_b
        rsl_ba_sw = eirp_b - (20*np.log10(d_sw)+20*np.log10(band_ghz*1000)+32.44) \
                    - spec_atten*d_sw + ant_gain_a - other_loss_a - branch_a

        min_thresh = st.number_input("Minimum RSL threshold (dBm)",
                                      value=-75.0, step=1.0,
                                      help="Typical minimum receive threshold for operation. Enter your radio's threshold for the chosen modulation.")

        fig_sw, ax_sw = plt.subplots(figsize=(12, 4))
        fig_sw.patch.set_facecolor("#0e1117"); ax_sw.set_facecolor("#0e1117")
        ax_sw.plot(d_sw, rsl_ab_sw, color="#4488ff", lw=2, label=f"RSL at {sB} ({sA}→{sB})")
        if not balanced:
            ax_sw.plot(d_sw, rsl_ba_sw, color="#ff8844", lw=2, ls="--",
                       label=f"RSL at {sA} ({sB}→{sA})")
        ax_sw.axvline(path_km, color="white", lw=1.2, ls=":", label=f"Current: {path_km:.1f} km")
        ax_sw.axhline(rsl_ab,  color="#4488ff", lw=0.8, ls=":", alpha=0.5)
        ax_sw.axhline(min_thresh, color="red", lw=1.2, ls="--", alpha=0.8,
                      label=f"Min threshold: {min_thresh:.0f} dBm")

        # Mark max range at threshold
        try:
            cross_idx = np.where(rsl_ab_sw < min_thresh)[0][0]
            max_range = d_sw[cross_idx]
            ax_sw.axvline(max_range, color="red", lw=1, ls=":", alpha=0.6)
            ax_sw.annotate(f"  Max range\n  {max_range:.1f} km",
                           xy=(max_range, min_thresh+3), color="red", fontsize=8)
        except IndexError:
            pass

        ax_sw.annotate(f"  {rsl_ab:.2f} dBm",
                       xy=(path_km, rsl_ab), color="white", fontsize=8, va="bottom")
        ax_sw.set_xlabel("Path Length (km)", color="white", fontsize=10)
        ax_sw.set_ylabel("RSL (dBm)", color="white", fontsize=10)
        ax_sw.set_title(f"RSL vs Distance — {sA} ↔ {sB}  |  {band_ghz} GHz  |  EIRP={eirp_a:.1f}/{eirp_b:.1f} dBm",
                        color="white", fontsize=11)
        ax_sw.legend(facecolor="#1a1a2e", labelcolor="white", fontsize=9)
        ax_sw.tick_params(colors="white"); ax_sw.grid(color="#333", alpha=0.4)
        for sp in ax_sw.spines.values(): sp.set_color("#444")
        plt.tight_layout()
        st.pyplot(fig_sw)

        # ── SYSTEM SUMMARY CARD ───────────────────────────────────────────────
        st.markdown("---")
        st.subheader("📋 Link Summary")
        ex("Quick-reference card for this link showing all key parameters and results.")

        summary_data = {
            "Link ID": link_id or "—",
            f"Site A": f"{sA}{' — '+location_a if location_a else ''}",
            f"Site B": f"{sB}{' — '+location_b if location_b else ''}",
            "Path Length": f"{path_km:.3f} km",
            "Band / Frequency": f"{band_ghz} GHz  ({freq_a:.1f}/{freq_b:.1f} MHz)",
            "Free Space Path Loss": f"{fspl:.2f} dB",
            "Atmospheric Absorption": f"{abs_db:.3f} dB",
            "Total Propagation Loss": f"{tpl:.2f} dB",
            f"EIRP — {sA}": f"{eirp_a:.2f} dBm  ({ant_model_a or 'antenna'} {ant_gain_a} dBi)",
            f"EIRP — {sB}": f"{eirp_b:.2f} dBm  ({ant_model_b or 'antenna'} {ant_gain_b} dBi)",
            f"RSL at {sB} (A→B)": f"{rsl_ab:.2f} dBm",
            f"RSL at {sA} (B→A)": f"{rsl_ba:.2f} dBm",
            "Link Balance": f"{'Balanced ✓' if balanced else f'Unbalanced — {abs(rsl_ab-rsl_ba):.2f} dB difference'}",
            "Radio": radio_model or "—",
            "Bandwidth / Modulation": f"{bandwidth_mhz:.0f} MHz  /  {modulation}",
            "Bit Rate": f"{bitrate_mbps:.0f} Mb/s",
        }
        if bits:
            summary_data["Spectral Efficiency"] = f"{bitrate_mbps/bandwidth_mhz:.2f} b/s/Hz  (max for {mod_key}: {bits:.0f})"

        for k, v in summary_data.items():
            highlight = "RSL" in k or "EIRP" in k
            bg = "#1a3560" if "EIRP" in k else ("#1a4a1a" if "RSL" in k else "#1a1a2e")
            st.markdown(
                f"<div style='display:flex;justify-content:space-between;"
                f"padding:5px 12px;margin:2px 0;border-radius:4px;"
                f"background:{bg};border-left:3px solid {'#4488ff' if 'EIRP' in k else '#44bb44' if 'RSL' in k else '#2a2a3a'}'>"
                f"<span style='color:#aaa;font-size:0.85em'>{k}</span>"
                f"<span style='color:white;font-weight:{'bold' if highlight else 'normal'};font-size:0.9em'>{v}</span>"
                f"</div>", unsafe_allow_html=True
            )

        # ── FORMULA REFERENCE ─────────────────────────────────────────────────
        with st.expander("📐 Formula Reference"):
            st.latex(r"\text{EIRP (dBm)} = P_t + G_t - L_{feeder} - L_{branch}")
            st.latex(r"\text{FSPL (dB)} = 20\log_{10}(d_{km}) + 20\log_{10}(f_{MHz}) + 32.44")
            st.latex(r"\text{Absorption (dB)} = \gamma \times d_{km} \quad (\gamma = \text{specific attenuation, dB/km})")
            st.latex(r"\text{RSL (dBm)} = \text{EIRP} - \text{FSPL} - \text{Absorption} + G_r - L_{feeder,Rx} - L_{branch,Rx}")

    else:
        st.info("👆 Enter your path length and band in Section 3 above to see results.")

    # ── SAVE / DELETE ─────────────────────────────────────────────────────────
    st.markdown("---")
    st.subheader("💾 Save This Link")
    ex("Save all current parameters to session memory. Reload any saved link from the dropdown at the top.")
    sv1, sv2 = st.columns([3, 1])
    with sv1:
        save_name = st.text_input("Save as:",
            value=f"{(site_a or 'Site A')}–{(site_b or 'Site B')}",
            placeholder="Give this link a name")
    with sv2:
        st.markdown("<br>", unsafe_allow_html=True)
        if st.button("💾 Save", type="primary", use_container_width=True):
            st.session_state.mw_saved_links[save_name] = {
                "site_a": site_a, "site_b": site_b,
                "location_a": location_a, "location_b": location_b,
                "lat_a": lat_a_s, "lon_a": lon_a_s,
                "lat_b": lat_b_s, "lon_b": lon_b_s,
                "elev_a": elev_a, "elev_b": elev_b,
                "struct_a": struct_a, "struct_b": struct_b,
                "path_override": path_override,
                "path_km_manual": float(path_km),
                "band_ghz": band_ghz, "freq_a": freq_a, "freq_b": freq_b,
                "tx_pwr_a": tx_pwr_a, "tx_pwr_b": tx_pwr_b,
                "ant_gain_a": ant_gain_a, "ant_gain_b": ant_gain_b,
                "ant_model_a": ant_model_a, "ant_model_b": ant_model_b,
                "other_loss_a": other_loss_a, "other_loss_b": other_loss_b,
                "branch_a": branch_a, "branch_b": branch_b,
                "radio_model": radio_model, "bandwidth_mhz": bandwidth_mhz,
                "modulation": modulation, "bitrate_mbps": bitrate_mbps,
                "spec_atten": spec_atten, "link_id": link_id,
            }
            ok(f"'{save_name}' saved. Select from the dropdown at top to reload.")

    if saved_names and load_choice != "(new link)":
        if st.button(f"🗑️ Delete '{load_choice}'"):
            del st.session_state.mw_saved_links[load_choice]
            st.rerun()



# ─────────────────────────────────────────────────────────────────────────────
# TAB 13 — ADMIN PANEL (admin only)
# ─────────────────────────────────────────────────────────────────────────────
elif selected_tab == "⚙️ Admin Panel":
    if is_admin():
        show_admin_panel()
    else:
        st.error("⛔ Access denied — admin role required.")

# ─────────────────────────────────────────────────────────────────────────────
# FOOTER
# ─────────────────────────────────────────────────────────────────────────────
st.sidebar.markdown("---")
st.sidebar.caption(
    "Tool uses itur library (P.676, P.618) for ITU-R propagation models. "
    "P.452 and P.528 implementations are simplified analytical approximations. "
    "For full regulatory submissions, validate with ITU-R SoftTools and SEAMCAT."
)
